"""Baubesprechungsprotokoll: Fortschreibung, Pruefschritt, Erzeugung.

DER PUNKT, DEN DIESER TEST BEWACHT
=================================
Ein Protokoll ist kein Behaelter fuer die Themen einer Sitzung, sondern ein
Schnappschuss der fortlaufenden Themenliste des Projekts. Die naheliegende
Fehlimplementierung — "je Protokoll eine frische Liste" — sieht in der
Oberflaeche voellig richtig aus und faellt erst Wochen spaeter auf, wenn
niemand mehr nachvollziehen kann, seit wann ein Punkt offen ist.

Deshalb pruefen die Faelle hier vor allem eines: Was passiert mit einem Punkt,
ueber den in einer Sitzung NICHT gesprochen wurde?

  * Er bleibt stehen.
  * Er behaelt seine alte BB-Nummer (die dritte Zahl von "02. 08. 16").
  * Er verschwindet erst, wenn ein Mensch ihn auf "e" setzt.

Dazu der zweite Grundsatz des Moduls: Ohne Pruefung kein Dokument.
"""
from __future__ import annotations

import asyncio
import os
import shutil
import sys
import tempfile
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

STORAGE = Path(tempfile.gettempdir()) / "hpp-test-besprechung"
if STORAGE.exists():
    shutil.rmtree(STORAGE)
STORAGE.mkdir(parents=True)

WIN = str(STORAGE).replace("\\", "/")
os.environ["BTB_DATABASE_URL"] = f"sqlite:///{WIN}/test.db"
os.environ["BTB_UPLOAD_DIR"] = f"{WIN}/uploads"
os.environ["BTB_OUTPUT_DIR"] = f"{WIN}/output"

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from fastapi.testclient import TestClient  # noqa: E402

from app.database import init_db  # noqa: E402
from app.main import app  # noqa: E402
from app.services import besprechung_analyse as analyse  # noqa: E402

# Ohne den Lifespan-Kontext des TestClient legt niemand die Tabellen an.
init_db()

ok = 0
fehler: list[str] = []


def pruefe(bedingung, text):
    global ok
    if bedingung:
        ok += 1
    else:
        fehler.append(text)


def j(antwort, erlaubt=(200, 201)):
    erlaubt = (erlaubt,) if isinstance(erlaubt, int) else erlaubt
    if antwort.status_code not in erlaubt:
        raise SystemExit(
            f"{antwort.request.url} -> {antwort.status_code}: {antwort.text[:400]}"
        )
    return antwort.json() if antwort.content else None


c = TestClient(app)


# ─────────────────────────────────────────────────────────────────────────────
# Stammdaten
# ─────────────────────────────────────────────────────────────────────────────

projekt = j(c.post("/api/projekte", json={
    "name": "Neubau Institutsgebaeude Weidenstieg",
    "adresse": "Weidenstieg 29, 20259 Hamburg",
}))
pid = projekt["id"]
c.patch(f"/api/projekte/{pid}", json={
    "projekt_nummer": "225100", "bauherr": "SBH | Schulbau Hamburg",
})

bearbeiter = j(c.post("/api/mangel-stammdaten/bearbeiter", json={
    "name": "Katharina Blanck", "email": "katharina.blanck@hpp.com",
    "kuerzel": "kbl", "durchwahl": "22",
}))
pruefe(bearbeiter["kuerzel"] == "kbl" and bearbeiter["durchwahl"] == "22",
       "Bearbeiter traegt Kuerzel und Durchwahl")

for firma, code, bez in (
    ("Rolfes Bau", "VE300.01", "VE01 Erweiterte Rohbauarbeiten"),
    ("Heinze-Stockfisch-Grabis+Partner", "VE400", "VE02 Technische Anlagen"),
):
    j(c.post("/api/gewerke", json={
        "projekt_id": pid, "firma_name": firma,
        "vergabeeinheit_code": code, "vergabeeinheit_bezeichnung": bez,
    }))

kapitel = j(c.post(f"/api/besprechungsprotokolle/kapitel/aus-gewerken?projekt_id={pid}"))
pruefe(len(kapitel) == 3, f"Allgemein + zwei Vergabeeinheiten: {len(kapitel)} Kapitel")
pruefe(kapitel[0]["gewerk_id"] is None and "Allgemein" in kapitel[0]["titel"],
       "Kapitel 1 ist Allgemein und haengt an keinem Gewerk")
pruefe("VE300.01" in kapitel[1]["titel"],
       f"Vergabeeinheit im Kapiteltitel: {kapitel[1]['titel']}")
# Zweiter Aufruf darf nichts verdoppeln.
pruefe(len(j(c.post(
    f"/api/besprechungsprotokolle/kapitel/aus-gewerken?projekt_id={pid}"))) == 3,
    "Erneutes Erzeugen legt keine Dubletten an")
k_allg, k_rol = kapitel[0]["id"], kapitel[1]["id"]

for kz, name, rolle, ap, tel in (
    ("SBH", "Schulbau Hamburg", "Bauherr", "Herr R. Melms", "+49 40 42823-6111"),
    ("HPP", "HPP Architekten", "Objektueberwachung", "Frau K. Blanck", "+49 173 5489021"),
    ("ROL", "Rolfes Bau", "Rohbauer", "Frau R. Stark", "+49 159 04974373"),
    ("MOR", "Mo Re Architekten", "Objektplanung", "Herr T. Reinhardt", "+49 40 73084428"),
):
    j(c.post("/api/besprechungsprotokolle/beteiligte", json={
        "projekt_id": pid, "kuerzel": kz, "name": name, "rolle": rolle,
        "ansprechpartner": ap, "telefon": tel,
    }))


# ─────────────────────────────────────────────────────────────────────────────
# Protokoll 1 — die Themenliste entsteht
# ─────────────────────────────────────────────────────────────────────────────

p1 = j(c.post("/api/besprechungsprotokolle", json={
    "projekt_id": pid, "besprechungsdatum": "2026-08-18",
    "besprechungsort": "Weidenstieg 29, Baufeld",
    "ersteller_id": bearbeiter["id"],
}))
pruefe(p1["nummer"] == 1, "Erstes Protokoll bekommt Nummer 1")
pruefe(p1["themen_updates"] == [], "Ohne Vorgaenger startet es leer")
pruefe(p1["ersteller_kuerzel"] == "kbl" and p1["ersteller_durchwahl"] == "22",
       "Kopfzeile wird aus den Bearbeiter-Stammdaten vorbelegt")

THEMEN = [
    (k_allg, "Baubesprechung Hochbau im woechentlichen Turnus,\ndienstags um 14:00 Uhr", "", "", "i"),
    (k_allg, "Parkplatzflaeche Baederland als Rangierflaeche", "", "", "i"),
    (k_rol, "Uebermittlung eines Detailterminplans", "ROL", "KW 35'26", "n"),
    (k_rol, "Die Lieferung des Buerocontainers steht aus", "ROL", "KW 41'26", "n"),
    (k_rol, "Baumschutz nachbessern und herstellen", "ROL", "12.08.26", "n"),
    (k_rol, "Geruestaufstockung", "ROL", "20.08.26", "n"),
]
angelegt = [
    j(c.post(f"/api/besprechungsprotokolle/{p1['id']}/themen", json={
        "kapitel_id": kid, "thema_text": text, "zustaendig": zu,
        "bearb_bis": frist, "status": st,
    }))
    for kid, text, zu, frist, st in THEMEN
]
pruefe(all(z["bb_nr"] == "01" for z in angelegt), "Alle Zeilen tragen BB 01")
pruefe([z["nummer"] for z in angelegt][:3] == ["01. 01. 01", "01. 02. 01", "02. 01. 01"],
       f"Nummern wie im Original: {[z['nummer'] for z in angelegt][:3]}")

j(c.post(f"/api/besprechungsprotokolle/{p1['id']}/teilnehmer/aus-beteiligten"))
pruefe(len(j(c.get(f"/api/besprechungsprotokolle/{p1['id']}"))["teilnehmer"]) == 4,
       "Ansprechpartner werden als Teilnehmer uebernommen")

j(c.post(f"/api/besprechungsprotokolle/{p1['id']}/freigeben", json={}))
p1 = j(c.get(f"/api/besprechungsprotokolle/{p1['id']}"))
pruefe(p1["status"] == "freigegeben" and p1["hat_dokument"],
       "Freigabe erzeugt das Dokument")
pruefe(c.get(f"/api/besprechungsprotokolle/{p1['id']}/dokument").status_code == 200,
       "Dokument laesst sich herunterladen")

# Ein freigegebenes Protokoll ist ein Dokument, kein Entwurf mehr.
gesperrt = c.patch(f"/api/besprechungsprotokolle/{p1['id']}", json={"leistung": "X"})
pruefe(gesperrt.status_code == 409, "Freigegebenes Protokoll ist gesperrt")


# ─────────────────────────────────────────────────────────────────────────────
# Protokoll 2 — Fortschreibung
# ─────────────────────────────────────────────────────────────────────────────

p2 = j(c.post("/api/besprechungsprotokolle", json={
    "projekt_id": pid, "besprechungsdatum": "2026-08-25",
    "ersteller_id": bearbeiter["id"],
}))
pruefe(p2["nummer"] == 2, "Protokoll 2 zaehlt hoch")
pruefe(len(p2["themen_updates"]) == 6,
       f"Es erbt alle 6 offenen Punkte (hat {len(p2['themen_updates'])})")
pruefe(all(u["bb_nr"] == "01" for u in p2["themen_updates"]),
       "Uebernommene Zeilen behalten BB 01 — sie stammen aus Sitzung 1")
pruefe(all(u["uebernommen"] for u in p2["themen_updates"]),
       "…und sind als Fortschreibung gekennzeichnet")
pruefe(p2["anzahl_ungeprueft"] == 0,
       "Unveraenderte Uebernahmen gelten als geprueft")

nach_text = {u["thema_text"].split("\n")[0]: u for u in p2["themen_updates"]}

geaendert = j(c.patch(
    f"/api/besprechungsprotokolle/{p2['id']}/themen/"
    f"{nach_text['Uebermittlung eines Detailterminplans']['id']}",
    json={"bearb_bis": "KW 36'26", "status": "b", "hervorheben": True},
))
pruefe(geaendert["bb_nr"] == "02",
       f"Wer heute besprochen wurde, rueckt auf BB 02 (ist {geaendert['bb_nr']})")
pruefe(geaendert["vorher_bb"] == 1 and geaendert["vorher_status"] == "n",
       "Der Stand aus Protokoll 1 liegt zum Vergleich bei")

j(c.patch(f"/api/besprechungsprotokolle/{p2['id']}/themen/"
          f"{nach_text['Baumschutz nachbessern und herstellen']['id']}",
          json={"status": "e"}))

neu = j(c.post(f"/api/besprechungsprotokolle/{p2['id']}/themen", json={
    "kapitel_id": k_rol, "thema_text": "Ausbessern Geruestarbeiten",
    "zustaendig": "ROL", "bearb_bis": "KW 35'26", "status": "n",
}))
pruefe(neu["bb_nr"] == "02", "Neuer Punkt bekommt BB 02")
pruefe(neu["inhalt_nr"] == "05" and neu["nummer"] == "02. 05. 02",
       f"…und die naechste freie laufende Nummer ({neu['nummer']})")

# Dasselbe Thema zweimal im selben Protokoll waere ein Fehler.
doppelt = c.post(f"/api/besprechungsprotokolle/{p2['id']}/themen",
                 json={"thema_id": neu["thema_id"], "thema_text": "nochmal"})
pruefe(doppelt.status_code == 409, "Ein Thema kann nicht zweimal im Protokoll stehen")

j(c.post(f"/api/besprechungsprotokolle/{p2['id']}/teilnehmer/aus-beteiligten"))
j(c.post(f"/api/besprechungsprotokolle/{p2['id']}/freigeben", json={}))


# ─────────────────────────────────────────────────────────────────────────────
# Protokoll 3 — was bleibt, was geht
# ─────────────────────────────────────────────────────────────────────────────

p3 = j(c.post("/api/besprechungsprotokolle", json={
    "projekt_id": pid, "besprechungsdatum": "2026-09-01",
}))
texte = {u["thema_text"].split("\n")[0]: u for u in p3["themen_updates"]}

pruefe("Baumschutz nachbessern und herstellen" not in texte,
       "Der auf 'e' gesetzte Punkt taucht nicht mehr auf")
pruefe("Geruestaufstockung" in texte,
       "Der nicht erwaehnte Punkt bleibt offen stehen")
pruefe(texte["Geruestaufstockung"]["bb_nr"] == "01",
       "…unveraendert mit seiner alten BB-Nummer 01 — Schweigen ist keine Aenderung")
pruefe(texte["Uebermittlung eines Detailterminplans"]["bb_nr"] == "02",
       "Der in Sitzung 2 besprochene Punkt steht jetzt mit BB 02 da")
pruefe(texte["Uebermittlung eines Detailterminplans"]["bearb_bis"] == "KW 36'26",
       "…mit der dort gesetzten Frist")

liste = j(c.get(f"/api/besprechungsprotokolle/themen?projekt_id={pid}"))
pruefe(len(liste) == 6, f"Die laufende Themenliste hat 6 offene Themen ({len(liste)})")
alle = j(c.get(f"/api/besprechungsprotokolle/themen?projekt_id={pid}&nur_offen=false"))
pruefe(len(alle) == 7, "Mit den erledigten sind es 7")


# ─────────────────────────────────────────────────────────────────────────────
# Ohne Pruefung kein Dokument
# ─────────────────────────────────────────────────────────────────────────────

ungeprueft = j(c.post(f"/api/besprechungsprotokolle/{p3['id']}/themen", json={
    "kapitel_id": k_allg, "thema_text": "Ungeprueft", "status": "n",
}))
c.patch(f"/api/besprechungsprotokolle/{p3['id']}/themen/{ungeprueft['id']}",
        json={"bestaetigt": False})
abgelehnt = c.post(f"/api/besprechungsprotokolle/{p3['id']}/freigeben", json={})
pruefe(abgelehnt.status_code == 409, "Freigabe mit ungeprueften Zeilen wird abgelehnt")
pruefe(not j(c.get(f"/api/besprechungsprotokolle/{p3['id']}"))["hat_dokument"],
       "…und es entsteht kein Dokument")
mit_zwang = c.post(f"/api/besprechungsprotokolle/{p3['id']}/freigeben",
                   json={"trotz_ungeprueft": True})
pruefe(mit_zwang.status_code == 200,
       "Mit ausdruecklicher Bestaetigung geht es doch")


# ─────────────────────────────────────────────────────────────────────────────
# Analyse: Zuordnung und Absicherung gegen erfundene Angaben
# ─────────────────────────────────────────────────────────────────────────────

OFFEN = [analyse.OffenesThema(id=11, kennung="02. 01.", kapitel="VE01",
                              text="Uebermittlung eines Detailterminplans",
                              zustaendig="ROL", bearb_bis="KW 35'26", status="b")]
KAP = [analyse.KapitelInfo(id=1, nummer="01.", titel="Allgemein")]
BET = [analyse.BeteiligterInfo(kuerzel="ROL", name="Rolfes Bau")]

prompt, hinweise = analyse.baue_prompt(
    transkript="[00:02] Terminplan kommt KW 36.", notizen="", offene_themen=OFFEN,
    kapitel=KAP, beteiligte=BET, projektname="P", besprechungsdatum="01.09.2026")
pruefe("id=11" in prompt and "02. 01." in prompt,
       "Die offenen Themen stehen mit id und Nummer im Prompt")
pruefe("ROL" in prompt, "Die gueltigen Firmenkuerzel stehen im Prompt")

_, gekuerzt = analyse.baue_prompt(
    transkript="x" * (analyse.MAX_TRANSKRIPT + 100), notizen="", offene_themen=[],
    kapitel=KAP, beteiligte=[], projektname="P", besprechungsdatum="01.09.2026")
pruefe(gekuerzt and "gekuerzt" in gekuerzt[0].replace("ü", "ue"),
       "Ein zu langes Transkript wird gekuerzt und das steht als Hinweis drin")

ergebnis = analyse._zu_ergebnis({"punkte": [
    {"bestehendes_thema_id": 11, "text": "Detailterminplan", "zustaendig": "ROL",
     "bearb_bis": "KW 36'26", "status": "b"},
    {"bestehendes_thema_id": 999, "text": "Erfunden", "zustaendig": "XYZ", "status": "n"},
    {"bestehendes_thema_id": None, "kapitel_id": 42, "text": "Kapitel unbekannt",
     "status": "quatsch"},
    {"bestehendes_thema_id": None, "kapitel_id": 1, "text": "   "},
], "teilnehmer": [{"name": "Herr A. Eberz", "firma_kuerzel": "ROL"},
                  {"name": "", "firma_kuerzel": "ROL"}]},
    offene_themen=OFFEN, kapitel=KAP, beteiligte=BET)

pruefe(len(ergebnis.themen) == 3, "Leerer Text wird verworfen")
pruefe(ergebnis.themen[0].thema_id == 11, "Fortschreibung wird uebernommen")
pruefe(ergebnis.themen[1].thema_id is None,
       "Eine erfundene Thema-id wird NICHT uebernommen")
pruefe(ergebnis.themen[1].zustaendig == "",
       "Ein unbekanntes Firmenkuerzel landet nicht im Protokoll")
pruefe(any("XYZ" in h for h in ergebnis.hinweise), "…sondern wird gemeldet")
pruefe(ergebnis.themen[2].status == "n", "Ein ungueltiger Status faellt auf 'n' zurueck")
pruefe(ergebnis.themen[2].kapitel_id == 1,
       "Eine unbekannte Kapitel-id faellt auf das erste Kapitel zurueck")
pruefe([t.name for t in ergebnis.teilnehmer] == ["Herr A. Eberz"],
       "Namenlose Teilnehmer werden verworfen")

wert, unbekannt = analyse._pruefe_kuerzel("ROL/ SBH", {"ROL", "SBH"})
pruefe(wert == "ROL/\nSBH" and not unbekannt,
       f"Mehrere Zustaendige werden gestapelt ({wert!r})")
pruefe(analyse._pruefe_kuerzel("ALL", {"ROL"})[0] == "ALL",
       "'ALL' ist kein Firmenkuerzel und bleibt stehen")

from app.config import settings  # noqa: E402

alt = settings.anthropic_api_key
settings.anthropic_api_key = ""
try:
    asyncio.run(analyse.analysiere(
        transkript="x", notizen="", offene_themen=[], kapitel=KAP,
        beteiligte=[], projektname="P", besprechungsdatum="01.09.2026"))
    pruefe(False, "Ohne Schluessel muss ein AnalyseFehler kommen")
except analyse.AnalyseFehler as f:
    pruefe("sk-ant-" in str(f), "Ohne Schluessel sagt die Meldung, was zu tun ist")
finally:
    settings.anthropic_api_key = alt


# ─────────────────────────────────────────────────────────────────────────────
# Die Analyse im Router — mit gestelltem Modell
# ─────────────────────────────────────────────────────────────────────────────

p4 = j(c.post("/api/besprechungsprotokolle", json={
    "projekt_id": pid, "besprechungsdatum": "2026-09-08",
}))
laufend = next(u for u in p4["themen_updates"]
               if u["thema_text"].startswith("Geruestaufstockung"))


async def gestellt(**kw):
    pruefe(any(t.id == laufend["thema_id"] for t in kw["offene_themen"]),
           "Die offenen Themen werden dem Modell wirklich vorgelegt")
    return analyse.AnalyseErgebnis(
        themen=[
            analyse.ThemenVorschlag(
                thema_id=laufend["thema_id"], kapitel_id=None,
                text="Geruestaufstockung abgeschlossen", zustaendig="ROL",
                bearb_bis="", status="e", begruendung="Heute abgenommen."),
            analyse.ThemenVorschlag(
                thema_id=None, kapitel_id=k_allg,
                text="Baustellenbesichtigung Studierende am 31.08.26",
                zustaendig="MOR", bearb_bis="31.08.26", status="n",
                begruendung="Neuer Punkt."),
        ],
        teilnehmer=[analyse.TeilnehmerVorschlag(name="Herr A. Eberz",
                                                firma_kuerzel="ROL")],
        hinweise=["Zur Sockelabdichtung wurde nichts Konkretes gesagt."],
    )


echt = analyse.analysiere
analyse.analysiere = gestellt
try:
    bericht = j(c.post(f"/api/besprechungsprotokolle/{p4['id']}/tldv-import", json={
        "transkript": "[00:02] ROL: Geruest ist fertig.",
        "notizen": "- Geruestaufstockung abgeschlossen",
    }))
finally:
    analyse.analysiere = echt

pruefe(bericht["fortschreibungen"] == 1 and bericht["neue_themen"] == 1,
       f"1 Fortschreibung + 1 neues Thema: {bericht}")
pruefe(bericht["hinweise"], "Hinweise des Modells werden durchgereicht")

p4 = j(c.get(f"/api/besprechungsprotokolle/{p4['id']}"))
nach_id = {u["thema_id"]: u for u in p4["themen_updates"]}
fort = nach_id[laufend["thema_id"]]
pruefe(fort["bb_nr"] == "04", f"Fortgeschriebene Zeile rueckt auf BB 04 ({fort['bb_nr']})")
pruefe(fort["herkunft"] == "ki" and not fort["bestaetigt"],
       "Ein KI-Vorschlag ist NIE bestaetigt")
pruefe(any(t["aus_transkript"] for t in p4["teilnehmer"]),
       "Erkannte Teilnehmer sind als 'aus Transkript' markiert")
pruefe(c.post(f"/api/besprechungsprotokolle/{p4['id']}/freigeben",
              json={}).status_code == 409,
       "Ungeprueft kommt kein Dokument heraus")

for u in p4["themen_updates"]:
    j(c.patch(f"/api/besprechungsprotokolle/{p4['id']}/themen/{u['id']}",
              json={"bestaetigt": True}))
fertig = j(c.post(f"/api/besprechungsprotokolle/{p4['id']}/freigeben", json={}))
pruefe(fertig["hat_dokument"], "Nach der Pruefung entsteht das Dokument")

nachher = j(c.get(f"/api/besprechungsprotokolle/themen?projekt_id={pid}"))
pruefe(not any(t["id"] == laufend["thema_id"] for t in nachher),
       "Das abgehakte Thema ist aus den offenen Punkten verschwunden")


# ─────────────────────────────────────────────────────────────────────────────
# Anlage: die unterschriebene Teilnehmerliste
# ─────────────────────────────────────────────────────────────────────────────

import io  # noqa: E402

from PIL import Image  # noqa: E402
from docx import Document  # noqa: E402

p5 = j(c.post("/api/besprechungsprotokolle", json={
    "projekt_id": pid, "besprechungsdatum": "2026-09-15",
}))
bild = io.BytesIO()
Image.new("RGB", (1240, 1754), (245, 245, 245)).save(bild, format="JPEG")
anlage = j(c.post(
    f"/api/besprechungsprotokolle/{p5['id']}/anlagen",
    files={"datei": ("scan.jpg", bild.getvalue(), "image/jpeg")},
    data={"bezeichnung": "Teilnehmerliste, unterschrieben"},
))
pruefe(anlage["bezeichnung"] == "Teilnehmerliste, unterschrieben",
       "Anlage wird mit Bezeichnung gespeichert")

j(c.post(f"/api/besprechungsprotokolle/{p5['id']}/freigeben",
         json={"trotz_ungeprueft": True}))
p5 = j(c.get(f"/api/besprechungsprotokolle/{p5['id']}"))
pruefe(p5["hat_dokument"], "Protokoll mit Anlage wurde erzeugt")
datei = max((STORAGE / "output").glob("Protokoll_05_*.docx"),
            key=lambda p: p.stat().st_mtime)
dok = Document(str(datei))
pruefe(len(dok.sections) == 2,
       f"Die Anlagen bekommen einen eigenen Abschnitt ({len(dok.sections)})")
kopf_anlage = dok.sections[1].header
pruefe(not kopf_anlage.paragraphs[0].runs,
       "…ohne Briefkopf, damit er nicht ueber den Scan stempelt")
pruefe(dok.sections[1].footer.is_linked_to_previous,
       "…aber mit durchlaufender Fusszeile")
pruefe(any("Teilnehmerliste, unterschrieben" in a.text for a in dok.paragraphs),
       "Die Bezeichnung steht ueber der Abbildung")

gesperrt = c.post(
    f"/api/besprechungsprotokolle/{p5['id']}/anlagen",
    files={"datei": ("x.jpg", bild.getvalue(), "image/jpeg")}, data={})
pruefe(gesperrt.status_code == 409,
       "An ein freigegebenes Protokoll kommt keine Anlage mehr")


print(f"\n{ok} Pruefungen ok, {len(fehler)} Fehler")
for f in fehler:
    print("  FEHLER:", f)
sys.exit(1 if fehler else 0)
