"""Der ganze Weg: ein Wochenpaket hochladen, fuenf Tagesberichte bekommen.

Geprueft wird das, was der Bauleiter tatsaechlich tut — Dateien einer Woche
waehlen, kurz auf die erkannten Tage sehen, erzeugen lassen. Und vor allem:
dass am Ende in Dienstags Bericht Dienstags Arbeit steht.
"""
from __future__ import annotations

import os
import shutil
import sys
import tempfile
from datetime import date
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

STORAGE = Path(tempfile.gettempdir()) / "hpp-test-wocheneinreichung"
if STORAGE.exists():
    shutil.rmtree(STORAGE)
STORAGE.mkdir(parents=True)

WIN = str(STORAGE).replace("\\", "/")
os.environ["BTB_DATABASE_URL"] = f"sqlite:///{WIN}/test.db"
os.environ["BTB_UPLOAD_DIR"] = f"{WIN}/uploads"
os.environ["BTB_OUTPUT_DIR"] = f"{WIN}/output"

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from fastapi.testclient import TestClient  # noqa: E402

from app.database import init_db  # noqa: E402
from app.main import app  # noqa: E402
from app.services import wochenpaket_ablage as ablage  # noqa: E402
from app.services.wochenaufteilung import woche_um  # noqa: E402
from app.services.word_pdf import nach_pdf, word_vorhanden  # noqa: E402

ok = 0
fehler: list[str] = []


def pruefe(bedingung, text):
    global ok
    if bedingung:
        ok += 1
    else:
        fehler.append(text)


def gleich(ist, soll, text):
    pruefe(ist == soll, f"{text}: erwartet {soll!r}, war {ist!r}")


if not word_vorhanden():
    print("Word nicht erreichbar — dieser Test braucht es zum Bauen der "
          "Beispiel-PDFs. Uebersprungen.")
    raise SystemExit(0)

# Tabellen anlegen wie beim Start der App. Ohne das laeuft der TestClient
# gegen eine leere Datenbank (die Startup-Ereignisse feuern nur im
# with-Block, und der wuerde den ganzen Test eine Ebene tiefer schieben).
init_db()

client = TestClient(app)
TAGE = woche_um(date(2026, 8, 5))          # Mo 03.08. bis Fr 07.08.2026


def baue_pdf(name: str, eintraege: dict[date, list[tuple[str, int, str]]]) -> Path:
    """Baut ein Firmen-Wochenpaket: je Datum eine Seite mit Firmen und Arbeiten."""
    from docx import Document

    dok = Document()
    for i, (tag, firmen) in enumerate(sorted(eintraege.items())):
        if i:
            dok.add_page_break()
        dok.add_paragraph("Bautagesbericht")
        dok.add_paragraph(f"Datum: {tag.strftime('%d.%m.%Y')}")
        for firma, personen, arbeit in firmen:
            dok.add_paragraph(f"Fa. {firma} ({personen} Mann)")
            dok.add_paragraph(f"- {arbeit}")
        dok.add_paragraph("Fertigstellung laut Vertrag bis 30.11.2026")
    docx_pfad = STORAGE / (name + ".docx")
    dok.save(docx_pfad)
    pdf_pfad = STORAGE / (name + ".pdf")
    pdf_pfad.write_bytes(nach_pdf(docx_pfad.read_bytes()))
    return pdf_pfad


print("─── Beispielpaket bauen ───")

# Rohbau liefert alle fuenf Tage, Elektro nur Mittwoch bis Freitag — genau die
# ungleiche Verteilung, die es in der Praxis gibt.
rohbau = baue_pdf("rohbau_woche", {
    TAGE[0]: [("Meyer Bau", 4, "Schalung Achse A")],
    TAGE[1]: [("Meyer Bau", 5, "Bewehrung Decke UG")],
    TAGE[2]: [("Meyer Bau", 6, "Betonage Decke UG")],
    TAGE[3]: [("Meyer Bau", 3, "Ausschalen UG")],
    TAGE[4]: [("Meyer Bau", 4, "Schalung Treppenhaus")],
})
elektro = baue_pdf("elektro_woche", {
    TAGE[2]: [("Elektro Schulz", 2, "Leerrohre UG")],
    TAGE[3]: [("Elektro Schulz", 2, "Leerrohre EG")],
    TAGE[4]: [("Elektro Schulz", 1, "Verteiler setzen")],
})
print(f"   {rohbau.name} (5 Tage), {elektro.name} (3 Tage)")


print("─── Stammdaten ───")

# Ohne Koordinaten: Der Wetterabruf entfaellt und meldet eine Warnung. Genau
# das ist hier erwuenscht — der Test braucht kein Internet.
antwort = client.post("/api/projekte", json={
    "name": "Wochentest Baustelle", "adresse": "Teststrasse 1",
})
gleich(antwort.status_code, 201, "Projekt angelegt")
projekt_id = antwort.json()["id"]

antwort = client.post("/api/empfaenger", json={
    "label": "Bauleitung", "email": "bauleitung@example.org",
})
gleich(antwort.status_code, 201, "Empfaenger angelegt")
empfaenger_id = antwort.json()["id"]


print("─── Schritt 1: Paket hochladen und ansehen ───")

with rohbau.open("rb") as a, elektro.open("rb") as b:
    antwort = client.post(
        "/api/einreichungen/woche/analyse",
        files=[("dateien", (rohbau.name, a, "application/pdf")),
               ("dateien", (elektro.name, b, "application/pdf"))],
        data={"woche_von": TAGE[0].isoformat()},
    )
gleich(antwort.status_code, 200, "Analyse beantwortet")
analyse = antwort.json()

gleich(len(analyse["tage"]), 5, "fuenf Tage erkannt")
gleich([t["datum"] for t in analyse["tage"]],
       [t.isoformat() for t in TAGE], "und zwar Montag bis Freitag")
gleich(analyse["ohne_datum"], None, "nichts blieb unzugeordnet")
gleich(sorted(analyse["dateien"]), sorted([rohbau.name, elektro.name]),
       "beide Dateien im Paket")

montag, mittwoch = analyse["tage"][0], analyse["tage"][2]
gleich(len(montag["quellen"]), 1, "Montag kommt aus einer Datei")
gleich(montag["quellen"][0]["seiten"], [1], "Montag ist Seite 1 des Rohbaus")
gleich(len(mittwoch["quellen"]), 2, "Mittwoch kommt aus beiden Dateien")
pruefe(any("erkannt" in h for h in analyse["hinweise"]), "Hinweis nennt die Tage")

kennung = analyse["kennung"]
pruefe(ablage.ordner(kennung).is_dir(), "Zwischenablage liegt auf der Platte")

# Die Kennung ist eine Eingabe von aussen und wird auch so behandelt.
antwort = client.post("/api/einreichungen/woche", json={
    "kennung": "../../etc", "projekt_id": projekt_id,
    "empfaenger_id": empfaenger_id, "tage": [],
})
gleich(antwort.status_code, 400, "erfundene Kennung wird abgewiesen")


print("─── Schritt 2: fuenf Berichte erzeugen ───")

tage_eingabe = []
for i, tag in enumerate(analyse["tage"]):
    eintrag = dict(tag)
    if i == 0:
        eintrag["ergaenzende_angaben"] = "Kranmontage am Vormittag"
    tage_eingabe.append(eintrag)

antwort = client.post("/api/einreichungen/woche", json={
    "kennung": kennung,
    "projekt_id": projekt_id,
    "empfaenger_id": empfaenger_id,
    "tage": tage_eingabe,
})
gleich(antwort.status_code, 201, "Berichte angelegt")
ergebnis = antwort.json()
gleich(len(ergebnis["einreichungen"]), 5, "fuenf Einreichungen")

pruefe(not ablage.ordner(kennung).is_dir(),
       "Zwischenablage danach aufgeraeumt")

for eintrag, tag in zip(ergebnis["einreichungen"], TAGE):
    gleich(eintrag["datum"], tag.isoformat(), f"Datum {tag}")
    pruefe(len(eintrag["quelle_dateien"]) >= 1,
           f"{tag}: mindestens eine Quelldatei")

gleich(len(ergebnis["einreichungen"][0]["quelle_dateien"]), 1,
       "Montag: nur der Rohbau")
gleich(len(ergebnis["einreichungen"][2]["quelle_dateien"]), 2,
       "Mittwoch: Rohbau und Elektro")
gleich(ergebnis["einreichungen"][0]["ergaenzende_angaben"],
       "Kranmontage am Vormittag", "eigener Text nur am Montag")
gleich(ergebnis["einreichungen"][1]["ergaenzende_angaben"], "",
       "andere Tage ohne Zusatztext")


print("─── Steht in jedem Bericht der richtige Tag? ───")

ids = [e["id"] for e in ergebnis["einreichungen"]]

# Ohne Koordinaten meldet die Verarbeitung eine Wetterwarnung und wartet.
# Ob Warnungen auftreten, haengt davon ab, ob Geocoding und Wetterdienst
# erreichbar waren. Beides ist fuer diesen Test unerheblich — hier zaehlt, dass
# am Ende je Tag ein fertiges Dokument steht. Deshalb wird bestaetigt, falls
# die Verarbeitung darauf wartet, und sonst nichts getan.
for kennnummer, tag in zip(ids, TAGE):
    zustand = client.get(f"/api/einreichungen/{kennnummer}").json()
    pruefe(zustand["status"] in ("wartet_auf_bestaetigung", "abgeschlossen"),
           f"{tag}: Verarbeitung gelaufen (Status {zustand['status']})")
    if zustand["status"] == "wartet_auf_bestaetigung":
        antwort = client.post(f"/api/einreichungen/{kennnummer}/bestaetigen")
        gleich(antwort.status_code, 200, f"{tag}: bestaetigt")

import re  # noqa: E402
import zipfile as zf  # noqa: E402


def dokument_text(pfad: Path) -> str:
    """Der gesamte sichtbare Text einer Word-Datei.

    Bewusst ueber das XML statt ueber python-docx: Das HPP-Formular verschachtelt
    Tabellen, und ``Document(...).tables`` steigt in Untertabellen nicht hinab —
    genau dort stehen aber die Firmenzeilen.
    """
    with zf.ZipFile(pfad) as archiv:
        xml = archiv.read("word/document.xml").decode("utf-8")
    return "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml))


erwartet = {
    TAGE[0]: ("Schalung Achse A", ["Bewehrung", "Betonage", "Leerrohre"]),
    TAGE[1]: ("Bewehrung Decke UG", ["Schalung Achse A", "Betonage", "Leerrohre"]),
    TAGE[2]: ("Betonage Decke UG", ["Schalung Achse A", "Ausschalen"]),
    TAGE[3]: ("Leerrohre EG", ["Leerrohre UG", "Betonage"]),
    TAGE[4]: ("Verteiler setzen", ["Leerrohre EG", "Betonage"]),
}

for kennnummer, tag in zip(ids, TAGE):
    zustand = client.get(f"/api/einreichungen/{kennnummer}").json()
    gleich(zustand["status"], "abgeschlossen", f"{tag}: fertig")

    antwort = client.get(f"/api/einreichungen/{kennnummer}/dokument")
    gleich(antwort.status_code, 200, f"{tag}: Word abrufbar")

    pfad = STORAGE / f"bericht_{tag.isoformat()}.docx"
    pfad.write_bytes(antwort.content)
    text = dokument_text(pfad)

    soll, nicht_drin = erwartet[tag]
    pruefe(soll in text, f"{tag}: '{soll}' steht im Bericht")
    for fremd in nicht_drin:
        pruefe(fremd not in text, f"{tag}: '{fremd}' steht NICHT drin")
    pruefe(tag.strftime("%d.%m.%Y") in text, f"{tag}: Datum im Dokument")

gleich(len(erwartet), 5, "alle fuenf Tage geprueft")


print("─── Sammel-Download ───")

antwort = client.get("/api/einreichungen/dokumente.zip",
                     params={"ids": ",".join(str(i) for i in ids)})
gleich(antwort.status_code, 200, "ZIP abrufbar")
pruefe("Bautagesberichte" in antwort.headers.get("content-disposition", ""),
       "ZIP hat einen sprechenden Namen")

import io  # noqa: E402
import zipfile  # noqa: E402

with zipfile.ZipFile(io.BytesIO(antwort.content)) as archiv:
    namen = archiv.namelist()
    gleich(len(namen), 5, "fuenf Dokumente im Archiv")
    pruefe(all(n.endswith(".docx") for n in namen), "alle sind Word-Dateien")
    gleich(len(set(namen)), 5, "keine doppelten Namen")

antwort = client.get("/api/einreichungen/dokumente.zip", params={"ids": "999999"})
gleich(antwort.status_code, 404, "unbekannte Kennung -> klare Absage")

antwort = client.get("/api/einreichungen/dokumente.zip", params={"ids": ""})
gleich(antwort.status_code, 400, "leere Auswahl -> 400")


print("─── Mehrere Tage auf EINEM Blatt, durch den ganzen Ablauf ───")

# So schicken es viele Firmen: die ganze Woche als fortlaufender Text auf
# einem Blatt, nicht eine Seite je Tag.
BLATT_TAGE = [date(2026, 8, 10), date(2026, 8, 11), date(2026, 8, 12)]


def baue_wochenblatt(name: str) -> Path:
    from docx import Document

    dok = Document()
    dok.add_paragraph("Wochenbericht Fa. Sammelmeier")
    for tag, personen, arbeit in zip(
        BLATT_TAGE, (7, 8, 9),
        ("Estrich EG eingebracht", "Estrich OG1 eingebracht",
         "Trennschnitte gesetzt"),
    ):
        dok.add_paragraph(f"Datum: {tag.strftime('%d.%m.%Y')}")
        dok.add_paragraph(f"Fa. Sammelmeier ({personen} Mann)")
        dok.add_paragraph(f"- {arbeit}")
    dok.add_paragraph("Fertigstellung laut Vertrag bis 30.11.2026")
    docx_pfad = STORAGE / (name + ".docx")
    dok.save(docx_pfad)
    pdf_pfad = STORAGE / (name + ".pdf")
    pdf_pfad.write_bytes(nach_pdf(docx_pfad.read_bytes()))
    return pdf_pfad


blatt = baue_wochenblatt("sammelblatt")

with blatt.open("rb") as f:
    antwort = client.post(
        "/api/einreichungen/woche/analyse",
        files=[("dateien", (blatt.name, f, "application/pdf"))],
        data={"woche_von": BLATT_TAGE[0].isoformat(),
              "woche_bis": BLATT_TAGE[-1].isoformat()},
    )
gleich(antwort.status_code, 200, "Sammelblatt analysiert")
blatt_analyse = antwort.json()

gleich(len(blatt_analyse["tage"]), 3, "drei Tage aus einem Blatt erkannt")
gleich([t["datum"] for t in blatt_analyse["tage"]],
       [t.isoformat() for t in BLATT_TAGE], "und zwar die richtigen")
pruefe(any("Blatt" in h for h in blatt_analyse["hinweise"]),
       "Hinweis nennt die Trennung")
# Jeder Tag verweist auf eine eigene Textdatei, nicht auf dieselbe Seite.
quellen = [t["quellen"][0]["datei"] for t in blatt_analyse["tage"]]
gleich(len(set(quellen)), 3, "drei getrennte Quellen statt einer Seite")
pruefe(all(q.endswith(".txt") for q in quellen),
       "Abschnitte liegen als Textdatei vor")

antwort = client.post("/api/einreichungen/woche", json={
    "kennung": blatt_analyse["kennung"],
    "projekt_id": projekt_id,
    "empfaenger_id": empfaenger_id,
    "tage": blatt_analyse["tage"],
})
gleich(antwort.status_code, 201, "drei Berichte aus einem Blatt")
blatt_ergebnis = antwort.json()["einreichungen"]
gleich(len(blatt_ergebnis), 3, "drei Einreichungen")

erwartet_blatt = {
    BLATT_TAGE[0]: ("Estrich EG eingebracht", ["OG1", "Trennschnitte"]),
    BLATT_TAGE[1]: ("Estrich OG1 eingebracht", ["Trennschnitte"]),
    BLATT_TAGE[2]: ("Trennschnitte gesetzt", ["Estrich EG", "Estrich OG1"]),
}

for eintrag, tag in zip(blatt_ergebnis, BLATT_TAGE):
    kennnummer = eintrag["id"]
    zustand = client.get(f"/api/einreichungen/{kennnummer}").json()
    if zustand["status"] == "wartet_auf_bestaetigung":
        client.post(f"/api/einreichungen/{kennnummer}/bestaetigen")
    zustand = client.get(f"/api/einreichungen/{kennnummer}").json()
    gleich(zustand["status"], "abgeschlossen", f"{tag}: fertig")

    antwort = client.get(f"/api/einreichungen/{kennnummer}/dokument")
    gleich(antwort.status_code, 200, f"{tag}: Word abrufbar")
    pfad = STORAGE / f"blatt_{tag.isoformat()}.docx"
    pfad.write_bytes(antwort.content)
    text = dokument_text(pfad)

    soll, nicht_drin = erwartet_blatt[tag]
    pruefe(soll in text, f"{tag}: '{soll}' steht im Bericht")
    for fremd in nicht_drin:
        pruefe(fremd not in text, f"{tag}: '{fremd}' steht NICHT drin")
    pruefe("Sammelmeier" in text, f"{tag}: Firma uebernommen")


print("─── Zeitraum wird geprueft ───")

with blatt.open("rb") as f:
    antwort = client.post(
        "/api/einreichungen/woche/analyse",
        files=[("dateien", (blatt.name, f, "application/pdf"))],
        data={"woche_von": "2026-08-12", "woche_bis": "2026-08-10"},
    )
gleich(antwort.status_code, 400, "Ende vor Beginn wird abgewiesen")

with blatt.open("rb") as f:
    antwort = client.post(
        "/api/einreichungen/woche/analyse",
        files=[("dateien", (blatt.name, f, "application/pdf"))],
        data={"woche_von": "2026-01-01", "woche_bis": "2026-12-31"},
    )
gleich(antwort.status_code, 400, "zu langer Zeitraum wird abgewiesen")

with blatt.open("rb") as f:
    antwort = client.post(
        "/api/einreichungen/woche/analyse",
        files=[("dateien", (blatt.name, f, "application/pdf"))],
        data={"woche_von": "kein Datum"},
    )
gleich(antwort.status_code, 400, "unsinniges Datum wird abgewiesen")


print("─── Faehigkeiten ───")

antwort = client.get("/api/einreichungen/faehigkeiten")
gleich(antwort.status_code, 200, "Faehigkeiten abrufbar")
faehig = antwort.json()
pruefe(isinstance(faehig.get("handschrift"), bool), "Handschrift-Flag ist ein Ja/Nein")
pruefe(len(faehig.get("hinweis", "")) > 30, "Hinweis erklaert den Zustand")
if not faehig["handschrift"]:
    pruefe("anthropic_key" in faehig["hinweis"],
           "ohne Schluessel wird gesagt, wo er hingehoert")


print("─── Einzelner Tag geht unveraendert weiter ───")

with rohbau.open("rb") as a:
    antwort = client.post("/api/einreichungen", files=[
        ("dateien", (rohbau.name, a, "application/pdf")),
    ], data={
        "projekt_id": projekt_id, "empfaenger_id": empfaenger_id,
        "datum": TAGE[1].isoformat(), "ergaenzende_angaben": "",
    })
gleich(antwort.status_code, 201, "alter Weg funktioniert weiterhin")


print("─── Was gar nicht auszulesen ist, wird gleich abgewiesen ───")

# Vorher lief eine Tabelle durch den ganzen Ablauf und endete in einem leeren
# Bericht mit der Warnung "Keine Firmendaten extrahiert" — der Grund stand
# dort nicht.
tabelle = STORAGE / "kalkulation.xlsx"
tabelle.write_bytes(b"PK\x03\x04 keine echte Tabelle, aber die Endung zaehlt")

with tabelle.open("rb") as f:
    antwort = client.post("/api/einreichungen", files=[
        ("dateien", (tabelle.name, f, "application/vnd.ms-excel")),
    ], data={
        "projekt_id": projekt_id, "empfaenger_id": empfaenger_id,
        "datum": TAGE[1].isoformat(), "ergaenzende_angaben": "",
    })
gleich(antwort.status_code, 400, "Tabelle wird abgewiesen")
pruefe("kalkulation.xlsx" in antwort.text,
       f"die Meldung nennt die Datei: {antwort.text[:120]}")
pruefe("PDF" in antwort.text or "Foto" in antwort.text,
       "…und sagt, was stattdessen gebraucht wird")

with tabelle.open("rb") as f:
    antwort = client.post(
        "/api/einreichungen/woche/analyse",
        files=[("dateien", (tabelle.name, f, "application/vnd.ms-excel"))],
    )
gleich(antwort.status_code, 400, "auch im Wochenweg wird sie abgewiesen")

# Ein Handyfoto muss durchgehen, auch als HEIC — die Oberflaeche laesst es zu,
# und die Auswertung hat es vorher stillschweigend verworfen.
heic = STORAGE / "bericht.heic"
heic.write_bytes(b"kein echtes HEIC, aber die Endung zaehlt")
with heic.open("rb") as f:
    antwort = client.post("/api/einreichungen", files=[
        ("dateien", (heic.name, f, "image/heic")),
    ], data={
        "projekt_id": projekt_id, "empfaenger_id": empfaenger_id,
        "datum": TAGE[2].isoformat(), "ergaenzende_angaben": "",
    })
gleich(antwort.status_code, 201, "HEIC vom iPhone wird angenommen")

# Und die Groessengrenze, die es beim Maengelmodul und den Baufotos schon
# gibt. Direkt gegen die Pruefung, nicht ueber einen 50-MB-Upload: Ein
# still wirkungsloser Grenzwert waere schlimmer als keiner, und genau das
# passiert, wenn UploadFile eines Tages kein ".size" mehr liefert.
from fastapi import HTTPException as _HTTPException  # noqa: E402

from app.config import settings as _einst  # noqa: E402
from app.routers.einreichungen import _pruefe_dateien  # noqa: E402


class _Upload:
    def __init__(self, filename, size):
        self.filename = filename
        self.size = size


grenze_bytes = _einst.max_file_size_mb * 1024 * 1024
try:
    _pruefe_dateien([_Upload("plan.pdf", grenze_bytes - 1)])
    pruefe(True, "eine Datei unter der Grenze geht durch")
except _HTTPException as exc:
    fehler.append(f"Datei unter der Grenze wurde abgewiesen: {exc.detail}")

try:
    _pruefe_dateien([_Upload("riesig.pdf", grenze_bytes + 1)])
    fehler.append("eine Datei ueber der Grenze muesste abgewiesen werden")
except _HTTPException as exc:
    pruefe(exc.status_code == 400 and "MB" in str(exc.detail),
           f"zu grosse Datei wird abgewiesen: {exc.detail}")
    pruefe("riesig.pdf" in str(exc.detail),
           "die Meldung nennt die Datei")

# Fehlt die Groessenangabe, wird nicht geraten — die Endung entscheidet dann.
try:
    _pruefe_dateien([_Upload("ohne_groesse.pdf", None)])
    pruefe(True, "ohne Groessenangabe greift die Grenze nicht")
except _HTTPException as exc:
    fehler.append(f"ohne Groessenangabe faelschlich abgewiesen: {exc.detail}")


print("─── Ein fehlgeschlagener Bericht ist keine Sackgasse ───")

# Vorher: rote Plakette "fehlgeschlagen", kein Grund, kein Knopf — man konnte
# nur alles noch einmal hochladen. Die Dateien liegen aber weiterhin da.
from app.database import SessionLocal  # noqa: E402
from app.models import Einreichung  # noqa: E402

with rohbau.open("rb") as a:
    antwort = client.post("/api/einreichungen", files=[
        ("dateien", (rohbau.name, a, "application/pdf")),
    ], data={
        "projekt_id": projekt_id, "empfaenger_id": empfaenger_id,
        "datum": TAGE[3].isoformat(), "ergaenzende_angaben": "",
    })
gleich(antwort.status_code, 201, "Bericht fuer den Versuch angelegt")
versuchs_id = antwort.json()["id"]

# So sieht ein Bericht aus, dessen Word-Erzeugung gescheitert ist.
sitzung = SessionLocal()
try:
    eintrag = sitzung.get(Einreichung, versuchs_id)
    eintrag.status = "fehlgeschlagen"
    eintrag.warnungen = [{
        "feld": "dokument",
        "problem": "Das Word-Dokument konnte nicht erzeugt werden: Testfall.",
        "quelle_datei": "",
    }]
    sitzung.commit()
finally:
    sitzung.close()

zustand = client.get(f"/api/einreichungen/{versuchs_id}").json()
gleich(zustand["status"], "fehlgeschlagen", "Ausgangslage steht")
pruefe(any(w["feld"] == "dokument" for w in zustand["warnungen"]),
       "der Grund steht am Bericht und nicht nur im Protokoll")

antwort = client.post(f"/api/einreichungen/{versuchs_id}/bestaetigen")
gleich(antwort.status_code, 200, "ein zweiter Versuch ist moeglich")
zustand = client.get(f"/api/einreichungen/{versuchs_id}").json()
pruefe(zustand["status"] in ("abgeschlossen", "wird_verarbeitet"),
       f"der Bericht laeuft wieder: {zustand['status']}")
if zustand["status"] == "abgeschlossen":
    pruefe(not any(w["feld"] == "dokument" for w in zustand["warnungen"]),
           "die alte Fehlermeldung klebt nicht am gelungenen Bericht")

# Ein fertiger Bericht laesst sich nicht erneut anstossen.
antwort = client.post(f"/api/einreichungen/{versuchs_id}/bestaetigen")
gleich(antwort.status_code, 400, "ein fertiger Bericht wird nicht neu erzeugt")

print()
if fehler:
    print(f"{ok} Pruefungen ok, {len(fehler)} Fehler:")
    for f in fehler:
        print("  -", f)
    raise SystemExit(1)
print(f"{ok} Pruefungen ok, 0 Fehler")
