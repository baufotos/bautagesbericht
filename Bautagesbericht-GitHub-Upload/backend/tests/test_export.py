"""Rauchtest: Mängelliste erzeugen."""
import sys
from datetime import date
from pathlib import Path

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from PIL import Image  # noqa: E402

from app.schemas import MaengellisteJSON, MangelExportEintrag  # noqa: E402
from app.services.maengelliste_generation import generate_maengelliste  # noqa: E402

tmp = Path(__file__).parent
foto1 = tmp / "foto1.jpg"
foto2 = tmp / "foto2.jpg"
Image.new("RGB", (1200, 900), (180, 120, 90)).save(foto1)
Image.new("RGB", (900, 1200), (100, 140, 180)).save(foto2)

daten = MaengellisteJSON(
    projekt="Neubau Verwaltungsgebäude Süd",
    stand=date(2026, 8, 20),
    filter_beschreibung="Status: offen · Firma: Rolfes Bau GmbH",
    intern=True,
    maengel=[
        MangelExportEintrag(
            nummer="00012",
            kurzbezeichnung="Stahlbeton",
            typ="Hinweis",
            status="offen",
            prioritaet="mittel",
            firma="Rolfes Bau GmbH | VE300-01 Erweiterter Rohbau",
            ort="EG",
            raumnummer="E.014",
            beschreibung="Sichtbetonoberflaeche weist Lunker auf.\nNacharbeit erforderlich.",
            erstellt_am=date(2026, 8, 12),
            frist_bis=date(2026, 8, 26),
            rueckmeldung_status="keine Rückmeldung",
            ist_ueberfaellig=False,
            plan_markierung="Grundriss EG, Seite 1 (42 % / 68 %)",
            foto_pfade=[str(foto1), str(foto2), str(foto1), str(foto2)],
            interne_bemerkung="Bauleiter telefonisch informiert, Frist nicht verlaengern.",
        ),
        MangelExportEintrag(
            nummer="00012.1",
            kurzbezeichnung="Stahlbeton (Duplikat NU)",
            typ="Mangel",
            status="Nachfrist",
            prioritaet="hoch",
            firma="Muster Bau GmbH | VE300-02 Rohbau Ergaenzung",
            ort="1. OG",
            beschreibung="Gleicher Sachverhalt, anderer Nachunternehmer.",
            erstellt_am=date(2026, 8, 13),
            frist_bis=date(2026, 8, 20),
            nachfrist_bis=date(2026, 8, 15),
            ist_ueberfaellig=True,
            foto_pfade=[],
        ),
        MangelExportEintrag(
            nummer="00013",
            kurzbezeichnung="Fehlende Brandschottung",
            status="erledigt",
            firma="Elektro Nord GmbH",
            erledigt_am=date(2026, 8, 18),
            foto_pfade=[str(tmp / "gibtsnicht.jpg")],
        ),
    ],
)

pfad = generate_maengelliste(daten)
print("erzeugt:", pfad, pfad.stat().st_size, "Bytes")

# Gegenprobe: Firmen-Fassung darf die interne Bemerkung nicht enthalten.
import docx  # noqa: E402

firmen_daten = daten.model_copy(update={"intern": False})
pfad2 = generate_maengelliste(firmen_daten)
text = "\n".join(
    p.text for p in docx.Document(str(pfad2)).paragraphs
) + "\n".join(
    z.text
    for t in docx.Document(str(pfad2)).tables
    for r in t.rows
    for z in r.cells
)
print("Firmen-Fassung:", pfad2)
print("interne Bemerkung enthalten:", "telefonisch" in text)
print("Platzhalter uebrig:", [p for p in ("{{PROJEKT}}", "{{STAND}}", "{{MAENGEL}}", "{{ANZAHL}}", "{{FILTER}}", "{{HINWEIS}}") if p in text])
print("Tabellen:", len(docx.Document(str(pfad2)).tables))
