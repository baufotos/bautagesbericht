"""Layout des Baubesprechungsprotokolls gegen die Bueromasse pruefen.

Die Werte hier stammen nicht aus dem Gefuehl, sondern aus zwei Quellen:

  * dem Ausdruck von Protokoll 16 vom 25.08.2026 (Spaltengrenzen, Hoehen,
    Schriftgroessen — am PDF gemessen),
  * der Excel-Vorlage ``260825_BB_16.xlsm`` selbst (Fuellfarben und
    Spaltenbreiten stehen dort im XML).

Der Test schreibt sie fest. Wer eine Zahl aendert, aendert das Aussehen eines
Dokuments, das an Bauherrn und Firmen geht — dann soll er es bewusst tun und
diesen Test mit anpassen.

Besonders bewacht wird die Statusfaerbung. Sie sieht im PDF nach Handarbeit
aus, ist aber eine bedingte Formatierung der Excel-Vorlage: "n" faerbt die
ganze Zeile, "e"/"b"/"k" nur die Status-Zelle, "i" gar nichts.
"""
from __future__ import annotations

import os
import sys
import tempfile
from datetime import date
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ARBEIT = Path(tempfile.gettempdir()) / "hpp-test-besprechung-layout"
ARBEIT.mkdir(exist_ok=True)
os.environ.setdefault("BTB_OUTPUT_DIR", str(ARBEIT / "output"))
os.environ.setdefault("BTB_UPLOAD_DIR", str(ARBEIT / "uploads"))
os.environ.setdefault("BTB_DATABASE_URL", f"sqlite:///{(ARBEIT / 'x.db').as_posix()}")

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt  # noqa: E402

from app.services import besprechungsprotokoll_generation as g  # noqa: E402

ok = 0
fehler: list[str] = []


def pruefe(bedingung, text):
    global ok
    if bedingung:
        ok += 1
    else:
        fehler.append(text)


def nah(wert, soll, toleranz=0.6):
    return abs(float(wert) - float(soll)) <= toleranz


# ─────────────────────────────────────────────────────────────────────────────
# Die gemessenen Sollwerte
# ─────────────────────────────────────────────────────────────────────────────

# Senkrechte Linien der Themen-Tabelle im Original-PDF:
# 58.56 | 79.1 | 100.1 | 121.2 | 400.8 | 441.2 | 508.0 | 535.92
SPALTEN_SOLL = (20.5, 21.0, 21.1, 279.6, 40.4, 66.8, 27.4)

pruefe(all(nah(Pt(soll), breite, 1000)  # EMU-Vergleich unten genauer
           for soll, breite in zip(SPALTEN_SOLL, g.SPALTEN_BREITEN)),
       "Spaltenbreiten der Themen-Tabelle")
pruefe(all(nah(b.pt, s, 0.1) for b, s in zip(g.SPALTEN_BREITEN, SPALTEN_SOLL)),
       f"Spaltenbreiten exakt: {[round(b.pt, 1) for b in g.SPALTEN_BREITEN]}")
pruefe(nah(sum(b.pt for b in g.SPALTEN_BREITEN), 476.8, 1.0),
       "Summe der Spalten entspricht der Tabellenbreite 477.4 pt")

pruefe(nah(g.THEMEN_KOPF_HOEHE.pt, 60.4), "Schwarze Kopfzeile 60.4 pt hoch")
pruefe(nah(g.THEMEN_KAPITEL_HOEHE.pt, 19.4), "Grauer Kapitelbalken 19.4 pt hoch")
pruefe(nah(g.THEMEN_ZEILE.pt, 12.36, 0.05), "Eine Textzeile ist 12.36 pt hoch")

pruefe(g.FUELLUNG_KOPF == "000000", "Kopfzeile schwarz")
pruefe(g.FUELLUNG_KAPITEL == "F2F2F2",
       "Kapitelbalken hellgrau F2F2F2 — nicht schwarz")
pruefe(g.FUELLUNG_NEU == "DDEBF7", "Zeilenfuellung fuer Status n: DDEBF7")

pruefe(g.STATUS_FARBEN["e"][0] == "C6EFCE", "Status e gruen")
pruefe(g.STATUS_FARBEN["b"][0] == "FFEB9C", "Status b gelb")
pruefe(g.STATUS_FARBEN["k"][0] == "FFC7CE", "Status k rot")
pruefe(g.STATUS_FARBEN["i"][0] is None, "Status i ohne Fuellung")
pruefe(g.STATUS_FARBEN["n"][0] is None,
       "Status n faerbt die Zelle nicht — die ganze Zeile ist blau")
pruefe(g.STATUS_ZEILE_BLAU == "n", "Nur n faerbt die ganze Zeile")

pruefe(g.SCHRIFT == "Arial", "Hausschrift Arial")
pruefe(nah(g.GR_TABELLE.pt, 10), "Tabellentext 10 pt")
pruefe(nah(g.GR_DECKBLATT.pt, 11), "Deckblatt 11 pt")
pruefe(nah(g.GR_FLIESSTEXT.pt, 9), "Verteilungshinweis 9 pt")
pruefe(nah(g.GR_FUSS.pt, 7), "Fusszeile 7 pt")

pruefe(nah(g.RAND_LINKS.pt, 50.4), "Linker Rand 0.7 Zoll wie in Excel")
pruefe(nah(g.RAND_RECHTS.pt, 50.4), "Rechter Rand 0.7 Zoll")
pruefe(nah(g.TEXT_BREITE.pt, 494.4), "Textbreite 494.4 pt")

pruefe(len(g.LEGENDE) == 5 and [k for k, _ in g.LEGENDE] == ["k", "b", "e", "n", "i"],
       "Legende in der Reihenfolge des Originals")
pruefe(g.SPALTEN_THEMEN[0] == "Kapitel" and g.SPALTEN_THEMEN[-1] == "Status",
       "Spaltenkoepfe wie im Original")
pruefe("Die Verteilung dieses Protokolls" in g.VERTEILUNGSHINWEIS
       and "umgehend mitzuteilen." in g.VERTEILUNGSHINWEIS,
       "Verteilungshinweis im Wortlaut des Originals")


# ─────────────────────────────────────────────────────────────────────────────
# Ein erzeugtes Dokument nachmessen
# ─────────────────────────────────────────────────────────────────────────────

DATEN = g.ProtokollDaten(
    nummer=16,
    projektname="Neubau Institutsgebaeude Weidenstieg",
    projekt_adresse_zeilen=["Weidenstieg 29", "20259 Hamburg"],
    projekt_nummer="225100",
    leistung="Baubesprechung",
    bauherr="SBH | Schulbau Hamburg",
    besprechungsort="Weidenstieg 29, Baufeld",
    besprechungsdatum=date(2026, 8, 25),
    erstellt_datum=date(2026, 8, 25),
    ersteller_name="Fr. K. Blanck",
    ersteller_kuerzel="kbl",
    ersteller_durchwahl="22",
    ersteller_email="katharina.blanck@hpp.com",
    kapitel=[
        g.KapitelBlock("01.", "Allgemein/ Projektorganisation", [
            g.ThemaZeile("01.", "01.", "12", "Turnus", "", "", "i"),
            g.ThemaZeile("01.", "08.", "16", "Besichtigung", "MOR", "31.08.26", "n"),
        ]),
        g.KapitelBlock("2.", "VE01 Rohbau", [
            g.ThemaZeile("02.", "08.", "10", "Musterflaeche", "ROL", "19.08.26", "e"),
            g.ThemaZeile("02.", "09.", "02", "Terminplan", "ROL", "KW 35'26", "b",
                         hervorheben=True),
            g.ThemaZeile("02.", "10.", "16", "Kritisch", "ROL", "", "k"),
        ]),
    ],
    teilnehmer=[g.TeilnehmerZeile("Frau K. Blanck", "HPP", "+49 173 5489021")],
    beteiligte=[g.BeteiligterZeile("SBH", "Schulbau Hamburg", "Bauherr")],
)

pfad, probleme = g.generate_besprechungsprotokoll(DATEN)
pruefe(pfad.is_file() and not probleme, f"Dokument erzeugt: {pfad.name}")
pruefe(pfad.name == "Protokoll_16_2026-08-25_Neubau_Institutsgebaeude_Weidenstieg.docx",
       f"Dateiname nach Konvention: {pfad.name}")

dok = Document(str(pfad))
abschnitt = dok.sections[0]
pruefe(nah(abschnitt.left_margin.pt, 50.4) and nah(abschnitt.right_margin.pt, 50.4),
       "Seitenraender im Dokument")
pruefe(abschnitt.different_first_page_header_footer,
       "Seite 1 hat einen eigenen Kopf (voller Briefkopf)")
pruefe(len(abschnitt.first_page_header.paragraphs[0].runs) == 1,
       "Der Briefkopf der ersten Seite ist genau eine Grafik")
pruefe(len(abschnitt.header.paragraphs[0].runs) == 1,
       "Die Folgeseiten tragen nur die Wortmarke")

# Die Kopfzeilen-Absaetze muessen winzig bleiben: Sonst schiebt Word den
# gesamten Seiteninhalt nach unten und keine gemessene Hoehe stimmt mehr.
for name, teil in (("erste Seite", abschnitt.first_page_header),
                   ("Folgeseiten", abschnitt.header)):
    hoehe = teil.paragraphs[0].paragraph_format.line_spacing
    pruefe(hoehe is not None and nah(hoehe.pt, 1),
           f"Kopfzeilenabsatz ({name}) bleibt 1 pt hoch")

tabellen = dok.tables
themen = next(t for t in tabellen if len(t.columns) == 7)
breiten = [c.width.pt for c in themen.rows[0].cells]
pruefe(all(nah(b, s, 0.3) for b, s in zip(breiten, SPALTEN_SOLL)),
       f"Spaltenbreiten im Dokument: {[round(b, 1) for b in breiten]}")

kopfzeile = themen.rows[0]
tc_pr = kopfzeile.cells[0]._tc.get_or_add_tcPr()
pruefe(tc_pr.find(qn("w:textDirection")) is not None,
       "Die Spaltenkoepfe sind gedreht")
shd = tc_pr.find(qn("w:shd"))
pruefe(shd is not None and shd.get(qn("w:fill")) == "000000",
       "Die Kopfzeile ist schwarz gefuellt")


def fuellung(zelle):
    el = zelle._tc.get_or_add_tcPr().find(qn("w:shd"))
    return el.get(qn("w:fill")) if el is not None else None


# Zeile 1 Kopf, 2 Kapitelbalken, 3 Status i, 4 Status n,
# 5 Kapitelbalken, 6 Status e, 7 Status b, 8 Status k
zeilen = themen.rows
pruefe(fuellung(zeilen[1].cells[0]) == "F2F2F2", "Kapitelbalken grau")
pruefe(fuellung(zeilen[2].cells[3]) is None, "Status i: Zeile ohne Fuellung")
pruefe(fuellung(zeilen[2].cells[6]) is None, "Status i: auch die Status-Zelle leer")
pruefe(fuellung(zeilen[3].cells[3]) == "DDEBF7", "Status n: ganze Zeile blau")
pruefe(fuellung(zeilen[5].cells[3]) is None, "Status e: Zeile bleibt weiss")
pruefe(fuellung(zeilen[5].cells[6]) == "C6EFCE", "Status e: Status-Zelle gruen")
pruefe(fuellung(zeilen[6].cells[6]) == "FFEB9C", "Status b: Status-Zelle gelb")
pruefe(fuellung(zeilen[6].cells[5]) == "DDEBF7",
       "hervorheben faerbt die Zelle 'Bearb. bis'")
pruefe(fuellung(zeilen[7].cells[6]) == "FFC7CE", "Status k: Status-Zelle rot")

texte = [a.text for a in dok.paragraphs]
pruefe("protokoll" in texte, "Titel 'protokoll' steht im Text, nicht als Grafik")
pruefe(any("teilnehmer" in t for t in texte), "Die Teilnehmerliste hat ihren Titel")
pruefe(any(g.UEBERSCHRIFT_LEGENDE in t for t in texte), "Legende ist enthalten")
pruefe(any(g.UEBERSCHRIFT_ABKUERZUNGEN in t for t in texte),
       "Abkuerzungen Projektbeteiligte sind enthalten")
pruefe(sum(1 for a in dok.paragraphs
           if a._p.findall(".//" + qn("w:br"))) >= 3,
       "Mindestens drei Seitenumbrueche (Deckblatt, Themen, Legende, Teilnehmer)")

fuss = abschnitt.footer.paragraphs[0]
pruefe(pfad.name in fuss.text, "Die Fusszeile nennt den Dateinamen")
pruefe(len(fuss.paragraph_format.tab_stops) == 1
       and nah(fuss.paragraph_format.tab_stops[0].position.pt, 493.4),
       "Rechter Tabstopp der Fusszeile fuer die Seitenzahl")
pruefe(fuss.style.name == "Normal",
       "Die Fusszeile nutzt nicht die Formatvorlage 'Footer' — deren eigene "
       "Tabstopps wuerden die Seitenzahl in die Seitenmitte ziehen")

print(f"\n{ok} Pruefungen ok, {len(fehler)} Fehler")
for f in fehler:
    print("  FEHLER:", f)
sys.exit(1 if fehler else 0)
