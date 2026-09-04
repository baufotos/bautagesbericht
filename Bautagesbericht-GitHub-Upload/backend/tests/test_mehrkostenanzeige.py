"""Rauchtest: Anzeigen beantworten — Leser, Erzeuger und die fuenf Endpunkte.

Kein pytest, sondern ein Skript wie die uebrigen Tests hier:

    PY="C:/Users/ben.gagelmann/Desktop/Claude App/laufzeit/python/python.exe"
    export PYTHONPATH="...\\laufzeit\\pakete;...\\backend"
    "$PY" tests/test_mehrkostenanzeige.py

Die Pruefungen laufen ohne Fremddateien: Das eingehende Schreiben wird als
Word-Datei erzeugt und dann eingelesen. Liegen die echten Referenz-PDFs auf
diesem Rechner (Bens Ordner "Mehrkostenanzeige"), werden sie zusaetzlich
geprueft — sonst wird dieser Teil uebersprungen und gemeldet.
"""
from __future__ import annotations

import asyncio
import email
import os
import shutil
import sys
import tempfile
from datetime import date
from email import policy
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# Eigene Ablage im Temp-Ordner — die echte storage/ bleibt unberuehrt.
STORAGE = Path(tempfile.gettempdir()) / "hpp-anzeigetest"
if STORAGE.exists():
    shutil.rmtree(STORAGE)
STORAGE.mkdir(parents=True)

WIN = str(STORAGE).replace("\\", "/")
os.environ["BTB_DATABASE_URL"] = f"sqlite:///{WIN}/test.db"
os.environ["BTB_UPLOAD_DIR"] = f"{WIN}/uploads"
os.environ["BTB_OUTPUT_DIR"] = f"{WIN}/output"

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from docx import Document                                        # noqa: E402
from docx.oxml.ns import qn                                      # noqa: E402
from fastapi.testclient import TestClient                        # noqa: E402

from app.main import app                                         # noqa: E402
from app.services import anzeige_bausteine as bau                # noqa: E402
from app.services import anzeige_formulierung as form            # noqa: E402
from app.services import mehrkostenanzeige_generation as gen     # noqa: E402
from app.services import mehrkostenanzeige_lesen as les          # noqa: E402

ok = 0
fehler: list[str] = []


def pruefe(bedingung, text):
    global ok
    if bedingung:
        ok += 1
    else:
        fehler.append(text)


def abschnitt(titel: str) -> None:
    print(f"\n── {titel} " + "─" * max(0, 66 - len(titel)))


# ─────────────────────────────────────────────────────────────────────────────
# Eine eingehende Anzeige als Word-Datei
# ─────────────────────────────────────────────────────────────────────────────

ANZEIGE_TEXT = """Muster Rueckbau & Sanierung GmbH
Attilastrasse 24
12105 Berlin
Telefon: +49 (0)30 / 565 555 600
E-Mail: info@muster-rueckbau.de

Bauherr Immobilien GmbH
Schildhornstrasse 5
10163 Berlin

Datum
17. August 2026

Ansprechpartner
Qasem Ashgarzada
Qasem.Ashgarzada@muster-rueckbau.de

Leistungsort: Boulevard Berlin - 12163 Berlin, Schlossstrasse 10
Gewerk: Abbruch-, Rueckbau- und Entsorgungsmassnahme
Hier: Mehrkostenanzeige Nr. 03, Zusaetzlicher Rueckbauaufwand aufgrund
abweichender Bodenaufbauten

Sehr geehrte Damen und Herren,

hiermit zeigen wir gemaess § 2 Abs. 6 VOB/B die Entstehung zusaetzlicher
Verguetungsansprueche aufgrund abweichender Bodenaufbauten an.

1. Doppelboeden
Es wurden aufgestaenderte Doppelboeden vorgefunden. Gemaess Pos. 03.0110 ist
nur der Rueckbau von Bodenbelaegen vorgesehen.

2. Mehrlagige Belaege
Unterhalb der oberen Lage befinden sich weitere Belagslagen. Diese sind in
Pos. 03.0120 nicht beschrieben.

3. Getrennte Entsorgung
Die getrennte Aufnahme und Entsorgung verursacht zusaetzlichen Aufwand.

Nach derzeitigem Stand fuehren die hieraus resultierenden Mehraufwendungen
voraussichtlich zu einer Verlaengerung der Ausfuehrungsdauer um eine Woche.

Wir bitten um Kenntnisnahme und entsprechende Bestaetigung.

Freundliche Gruesse

Qasem Ashgarzada
-Bauleitung-
"""


def schreibe_anzeige(ziel: Path) -> Path:
    dok = Document()
    for zeile in ANZEIGE_TEXT.split("\n"):
        dok.add_paragraph(zeile)
    dok.save(str(ziel))
    return ziel


# ─────────────────────────────────────────────────────────────────────────────
# 1. Bausteine des Lesers
# ─────────────────────────────────────────────────────────────────────────────

abschnitt("Leser: einzelne Bausteine")

pruefe(les._erstes_datum("17. August 2026") == date(2026, 8, 17),
       "langes Datum nicht erkannt")
pruefe(les._erstes_datum("Stand 06.08.2024, gez.") == date(2024, 8, 6),
       "kurzes Datum nicht erkannt")
pruefe(les._erstes_datum("vom 20.08.26") == date(2026, 8, 20),
       "zweistelliges Jahr muesste 2026 werden")
pruefe(les._erstes_datum("kein Datum hier") is None,
       "ohne Datum muesste None kommen")
pruefe(les._erstes_datum("32.13.2026") is None,
       "unmoegliches Datum muesste None sein")

pruefe(les._ist_strasse("Attilastrasse 24"), "Strasse nicht erkannt")
pruefe(les._ist_strasse("Krumbaeken Kaempe 2"),
       "Strasse ohne Strassenwort nicht erkannt")
pruefe(not les._ist_strasse("Riedel Bau GmbH & Co. KG"),
       "Firmenzeile faelschlich als Strasse erkannt")

pruefe(les._ist_umgedreht("etieS"), "gespiegeltes Wort nicht erkannt")
pruefe(les._ist_umgedreht("1 etieS"), "gespiegelte Zeile nicht erkannt")
pruefe(not les._ist_umgedreht("Seite 1 von 2"),
       "lesbare Zeile faelschlich als gespiegelt verworfen")

anschrift = les._anschrift_aus_zeilen(
    ["Muster Bau GmbH & Co. KG", "Herrn Steffen Wegner",
     "Silbersteinstrasse 4", "97424 Schweinfurt", "Deutschland"]
)
pruefe(anschrift.firma == "Muster Bau GmbH & Co. KG", f"Firma: {anschrift.firma!r}")
pruefe(anschrift.zusatz == "Herrn Steffen Wegner", f"Zusatz: {anschrift.zusatz!r}")
pruefe(anschrift.strasse == "Silbersteinstrasse 4", f"Strasse: {anschrift.strasse!r}")
pruefe((anschrift.plz, anschrift.ort) == ("97424", "Schweinfurt"),
       f"PLZ/Ort: {anschrift.plz!r} {anschrift.ort!r}")
pruefe("Deutschland" not in anschrift.zeilen(),
       "Inlandspost braucht kein Land im Adressblock")

impressum = les._anschrift_aus_zeilen(
    ["Muster GmbH · Attilastrasse 24 · 12105 Berlin · Deutschland"]
)
pruefe(impressum.plz == "12105" and impressum.strasse == "Attilastrasse 24",
       f"Impressumszeile nicht zerlegt: {impressum}")

# Fortlaufende Gliederung: ein Briefdatum "22. Juli 2026" darf keine
# Ueberschrift Nr. 22 werden.
lauf = les._fortlaufend([
    les.Punkt("22", "Juli 2026"), les.Punkt("1", "Erstens"),
    les.Punkt("2", "Zweitens"), les.Punkt("3", "Drittens"),
])
pruefe([p.nummer for p in lauf] == ["1", "2", "3"],
       f"Lauf falsch: {[p.nummer for p in lauf]}")

# ─────────────────────────────────────────────────────────────────────────────
# 2. Leser auf einer echten Datei
# ─────────────────────────────────────────────────────────────────────────────

abschnitt("Leser: Word-Anzeige komplett")

quelle = schreibe_anzeige(STORAGE / "MKA 03.docx")
gelesen = les.lies(quelle)

pruefe(gelesen.art == "Mehrkostenanzeige", f"Art: {gelesen.art!r}")
pruefe(gelesen.nummer == "03", f"Nummer: {gelesen.nummer!r}")
pruefe(gelesen.datum == date(2026, 8, 17), f"Datum: {gelesen.datum}")
pruefe(gelesen.betreff.startswith("Mehrkostenanzeige Nr. 03"),
       f"Betreff: {gelesen.betreff!r}")
pruefe("Bodenaufbauten" in gelesen.betreff,
       f"Fortsetzungszeile des Betreffs fehlt: {gelesen.betreff!r}")
pruefe(gelesen.kurzbezeichnung.startswith("Zusaetzlicher"),
       f"Kurzbezeichnung: {gelesen.kurzbezeichnung!r}")
pruefe(gelesen.absender.firma == "Muster Rueckbau & Sanierung GmbH",
       f"Absenderfirma: {gelesen.absender.firma!r}")
pruefe(gelesen.absender.strasse == "Attilastrasse 24",
       f"Absenderstrasse: {gelesen.absender.strasse!r}")
pruefe((gelesen.absender.plz, gelesen.absender.ort) == ("12105", "Berlin"),
       f"Absender PLZ/Ort: {gelesen.absender.plz!r} {gelesen.absender.ort!r}")
pruefe(gelesen.ansprechpartner_email == "Qasem.Ashgarzada@muster-rueckbau.de",
       f"Mail des Ansprechpartners: {gelesen.ansprechpartner_email!r}")
pruefe(gelesen.leistungsort.startswith("Boulevard Berlin"),
       f"Leistungsort: {gelesen.leistungsort!r}")
pruefe(gelesen.gewerk.startswith("Abbruch-"), f"Gewerk: {gelesen.gewerk!r}")
pruefe("§ 2 Abs. 6 VOB/B" in gelesen.rechtsgrundlage,
       f"Rechtsgrundlage: {gelesen.rechtsgrundlage!r}")
pruefe([p.nummer for p in gelesen.punkte] == ["1", "2", "3"],
       f"Punkte: {[(p.nummer, p.titel) for p in gelesen.punkte]}")
pruefe(gelesen.punkte[0].titel == "Doppelboeden",
       f"Titel Punkt 1: {gelesen.punkte[0].titel!r}")
pruefe("Doppelboeden" in gelesen.punkte[0].text or gelesen.punkte[0].text,
       "Text zu Punkt 1 fehlt")
pruefe(gelesen.lv_positionen == ["03.0110", "03.0120"],
       f"LV-Positionen: {gelesen.lv_positionen}")
pruefe("eine Woche" in gelesen.bauzeit, f"Bauzeit: {gelesen.bauzeit!r}")
pruefe(gelesen.forderung.startswith("Wir bitten um"),
       f"Forderung: {gelesen.forderung!r}")
pruefe(gelesen.unterzeichner == "Qasem Ashgarzada",
       f"Unterzeichner: {gelesen.unterzeichner!r}")
pruefe(gelesen.unterzeichner_funktion == "Bauleitung",
       f"Funktion: {gelesen.unterzeichner_funktion!r}")

# Die Art muss auch bei den anderen Anzeigearten stimmen — darum geht es dem
# Buero: Es sind nicht nur Mehrkostenanzeigen.
for text, erwartet in [
    ("Behinderungsanzeige Nr. 01 - Erdabfuhr", "Behinderungsanzeige"),
    ("Ihr Schreiben BMA 07 Behinderungs- und Mehrkostenanzeige",
     "Behinderungs- und Mehrkostenanzeige"),
    ("Verguetung zusaetzlicher Leistungen VZL 03", "Vergütung zusätzlicher Leistungen"),
    ("unsere Bedenkenanmeldung zum Untergrund", "Bedenkenanmeldung"),
    ("Nachtragsangebot NT 12 zur Position", "Nachtragsangebot"),
    ("MEKO 11 - Bewehrung fuer Filigrandecken", "Mehrkostenanzeige"),
    ("BEH 01 Erdabfuhr", "Behinderungsanzeige"),
]:
    erkannt = les._aus_text(text, "probe.txt", 1)
    pruefe(erkannt.art == erwartet,
           f"Art aus {text!r}: {erkannt.art!r}, erwartet {erwartet!r}")

# Die Kennung fuer den Dateinamen: Kuerzel der Firma, sonst das des Bueros.
for text, erwartet in [
    ("Behinderungsanzeige Nr. 02, Erdabfuhr", "BEH 02"),
    ("Mehrkostenanzeige Nr. 03, Bodenaufbauten", "MKA 03"),
    ("Behinderungs- und Mehrkostenanzeige Nr. 07", "BMA 07"),
    # Fuer diese Art gibt es in den Bueroschreiben kein Kuerzel — dann bleibt
    # der ausgeschriebene Name stehen, statt eines zu erfinden.
    ("Bedenkenanmeldung Nr. 04 zum Untergrund", "Bedenkenanmeldung Nr. 04"),
    # Ein eigenes Kuerzel der Firma gewinnt immer.
    ("Ihr Schreiben MEKO 11 - Bewehrung", "MEKO 11"),
]:
    erkannt = les._aus_text(text, "probe.txt", 1)
    pruefe(erkannt.kennung == erwartet,
           f"Kennung aus {text!r}: {erkannt.kennung!r}, erwartet {erwartet!r}")

# Warnung, wenn Firma und Maildomaene nicht zusammenpassen — der Fall, in dem
# die Antwort an den Bauherrn statt an die Firma ginge.
falsch = les.GelesenesSchreiben(
    absender=les.Anschrift(firma="Bauherr Immobilien GmbH", strasse="Weg 1",
                           plz="10163", ort="Berlin"),
    ansprechpartner_email="info@muster-rueckbau.de",
)
pruefe(not les._passt_zur_maildomain(falsch),
       "Firma und Maildomaene passen nicht — haette auffallen muessen")
richtig = les.GelesenesSchreiben(
    absender=les.Anschrift(firma="Muster Rueckbau & Sanierung GmbH"),
    ansprechpartner_email="info@muster-rueckbau.de",
)
pruefe(les._passt_zur_maildomain(richtig),
       "passende Firma faelschlich beanstandet")

pruefe(les.lies(STORAGE / "MKA 03.docx").art == "Mehrkostenanzeige",
       "zweiter Lesevorgang derselben Datei schlug fehl")

try:
    les.lies(STORAGE / "gibtsnicht.xlsx")
    pruefe(False, "unbekannte Endung muesste einen Lesefehler geben")
except les.Lesefehler as f:
    pruefe("PDF" in str(f), f"Meldung ohne Hinweis auf erlaubte Formate: {f}")

# ─────────────────────────────────────────────────────────────────────────────
# 3. Erzeuger
# ─────────────────────────────────────────────────────────────────────────────

abschnitt("Erzeuger: Textbausteine")


def daten(**abweichung) -> gen.Schreiben:
    grund = dict(
        empfaenger=gen.Empfaenger(
            firma="Muster Rueckbau & Sanierung GmbH", anrede="Herr",
            ansprechpartner="Qasem Ashgarzada", strasse="Attilastrasse 24",
            plz="12105", ort="Berlin",
            email="Qasem.Ashgarzada@muster-rueckbau.de",
        ),
        sachbearbeiter=gen.Sachbearbeiter(
            name="Katharina Blanck", funktion="-Baumanagement-", zeichen="kb",
            durchwahl="22", email="w29@hpp.com",
        ),
        anzeige=gen.Anzeige(
            art="Mehrkostenanzeige", nummer="03", kennung="MKA 03",
            datum=date(2026, 8, 17),
            kurzbezeichnung="Zusaetzlicher Rueckbauaufwand",
            bauzeit="Verlaengerung der Ausfuehrungsdauer um eine Woche.",
        ),
        projektzeile="BOB_Boulevard Berlin",
        vergabeeinheit="VE100- Abbrucharbeiten",
        briefdatum=date(2026, 9, 3),
        stellungnahme="Wir sehen hier keinen Mehrverguetungsanspruch.",
        dateikuerzel="BOB",
    )
    grund.update(abweichung)
    return gen.Schreiben(**grund)


pruefe(gen.anredezeile(daten().empfaenger) == "Sehr geehrter Herr Ashgarzada,",
       f"Anrede Herr: {gen.anredezeile(daten().empfaenger)!r}")
frau = gen.Empfaenger(firma="X", anrede="Frau", ansprechpartner="Regina Braun")
pruefe(gen.anredezeile(frau) == "Sehr geehrte Frau Braun,",
       f"Anrede Frau: {gen.anredezeile(frau)!r}")
ohne = gen.Empfaenger(firma="X", ansprechpartner="Q. Ashgarzada")
pruefe(gen.anredezeile(ohne) == "Sehr geehrte Damen und Herren,",
       "ohne Angabe von Herr/Frau muss die Anrede neutral bleiben")

betreff = gen.betreffzeile(daten())
pruefe(betreff == "Ihre Mehrkostenanzeige Nr. 03 vom 17.08.2026 – "
                  "Zusaetzlicher Rueckbauaufwand", f"Betreff: {betreff!r}")
beh = gen.betreffzeile(daten(anzeige=gen.Anzeige(
    art="Behinderungsanzeige", nummer="01", datum=date(2026, 5, 7),
    kurzbezeichnung="Erdabfuhr")))
pruefe(beh == "Ihre Behinderungsanzeige Nr. 01 vom 07.05.2026 – Erdabfuhr",
       f"Betreff Behinderungsanzeige: {beh!r}")
pruefe(gen.betreffzeile(daten(betreff="Von Hand")) == "Von Hand",
       "eigener Betreff muss gewinnen")

pruefe(gen.dateiname(daten()) == "260903 BOB-VE100-MKA 03.docx",
       f"Dateiname: {gen.dateiname(daten())!r}")
pruefe(gen._ve_code("VE300.01- Erweiterter Rohbau") == "VE300.01",
       f"VE-Kennung: {gen._ve_code('VE300.01- Erweiterter Rohbau')!r}")

pruefe(gen.mail_betreff(daten()) ==
       "Boulevard Berlin_Mehrkostenanzeige vom 17.08.26",
       f"Mailbetreff: {gen.mail_betreff(daten())!r}")
pruefe(gen.mail_betreff(daten(anzeige=gen.Anzeige(
    art="Behinderungsanzeige", datum=date(2026, 8, 20)))) ==
       "Boulevard Berlin_Behinderungsanzeige vom 20.08.26",
       "Mailbetreff muss die Art der Anzeige tragen")
pruefe(gen.mail_text(daten()) ==
       "Sehr geehrte Damen und Herren,\n\n"
       "in der Anlage senden wir unsere Schreiben vom 03.09.26.\n",
       f"Mailtext: {gen.mail_text(daten())!r}")
pruefe(gen.mail_text(daten()).rstrip().endswith("26."),
       "der Satz muss mit einem Punkt enden")

# Die Mail bekommt Arial 10 — eine reine Textmail hat keine Schrift, und
# Outlook nimmt sonst die aus den Einstellungen des jeweiligen Absenders.
html = gen.mail_html("Sehr geehrte Damen und Herren,\n\nZeile A\nZeile B")
pruefe("font-family:Arial" in html.replace(", ", ""),
       f"Arial fehlt im HTML-Teil: {html[:120]}")
pruefe("font-size:10pt" in html, f"10 pt fehlt im HTML-Teil: {html[:120]}")
pruefe("<br>" in html, "einfache Zeilenumbrueche muessen <br> werden")
pruefe(html.count("<p ") == 2, f"zwei Absaetze erwartet: {html}")
# Sonderzeichen duerfen die Mail nicht zerlegen.
pruefe("&amp;" in gen.mail_html("Meyer & Sohn") and
       "&lt;" in gen.mail_html("a < b"),
       "im HTML-Teil muss escaped werden")

adressblock = gen.adressblock(daten())
pruefe(adressblock == [
    "Muster Rueckbau & Sanierung GmbH", "Herrn Qasem Ashgarzada",
    "Attilastrasse 24", "12105 Berlin", "", "per E-Mail:",
    "Qasem.Ashgarzada@muster-rueckbau.de",
], f"Adressblock: {adressblock}")
pruefe("per E-Mail:" not in gen.adressblock(daten(
    empfaenger=gen.Empfaenger(firma="X", strasse="Weg 1", plz="1", ort="Y"))),
    "ohne Mailadresse darf keine leere Zeile 'per E-Mail:' entstehen")

# Haltung bestimmt den Schlusssatz.
pruefe(gen.schlusssaetze(daten(haltung="kenntnisnahme")) ==
       ["Wir bitten um Kenntnisnahme."], "Schlusssatz Kenntnisnahme")
pruefe("lehnen wir Ihre Mehrkostenanzeige ab" in
       " ".join(gen.schlusssaetze(daten(haltung="ablehnung"))),
       "Schlusssatz Ablehnung nennt die Art der Anzeige nicht")
pruefe(gen.schlusssaetze(daten(schlusssatz="Eigener Schluss.")) ==
       ["Eigener Schluss."], "eigener Schlusssatz muss gewinnen")

# Infofeld: Aufzaehlung bleibt Absatz, LV-Auszug wird eingerueckt.
bloecke = gen.koerper(daten(stellungnahme=(
    "Vorbemerkung Satz eins.\nGleicher Absatz Satz zwei.\n"
    "\n1) Erster Punkt.\n2) Zweiter Punkt.\n"
    "\nAuszug LV:\nZeile A des Auszugs\nZeile B des Auszugs\n"
    "\nSchlussgedanke."
)))
texte = [b.text for b in bloecke]
pruefe("Vorbemerkung Satz eins. Gleicher Absatz Satz zwei." in texte,
       f"Absatzbildung falsch: {texte}")
pruefe("1) Erster Punkt." in texte and "2) Zweiter Punkt." in texte,
       f"Aufzaehlung nicht als eigene Absaetze: {texte}")
zitate = [b.text for b in bloecke if b.zitat]
pruefe(zitate == ["Zeile A des Auszugs", "Zeile B des Auszugs"],
       f"LV-Auszug nicht eingerueckt: {zitate}")
pruefe(any(b.marke for b in bloecke), "Auszug-Marke nicht gesetzt")
pruefe(texte[0].startswith("wir haben Ihre Mehrkostenanzeige Nr. 03"),
       f"Einleitung falsch: {texte[0]!r}")
pruefe(texte[-1] == "Wir bitten um Kenntnisnahme.",
       f"letzter Absatz: {texte[-1]!r}")
pruefe(gen.BAUZEIT_ABLEHNUNG in
       [b.text for b in gen.koerper(daten(bauzeit_ablehnen=True))],
       "Bauzeit-Ablehnung fehlt im Koerper")

# Pruefung meldet fehlende Angaben.
for feld, wert, stichwort in [
    ("stellungnahme", "", "Stellungnahme"),
    ("projektzeile", "", "Projektzeile"),
]:
    try:
        gen.pruefe(daten(**{feld: wert}))
        pruefe(False, f"fehlendes Feld {feld} haette auffallen muessen")
    except gen.MehrkostenanzeigeFehler as f:
        pruefe(stichwort in str(f), f"Meldung zu {feld}: {f}")

abschnitt("Erzeuger: das Word-Dokument")

inhalt = gen.erzeuge(daten(
    stellungnahme="Wir sehen hier keinen Mehrverguetungsanspruch.\n"
                  "\nAuszug LV:\nBodenbelaege sind vollstaendig zurueckzubauen.",
    haltung="ablehnung", bauzeit_ablehnen=True,
    anlagen="Auszug LV Pos. 03.0110",
    verteiler="Bauherr Immobilien GmbH, Frau Eschelbach\nsowie Projektadressen",
))
pruefe(len(inhalt) > 20000, f"Dokument zu klein: {len(inhalt)} Bytes")

pfad = STORAGE / "brief.docx"
pfad.write_bytes(inhalt)
brief = Document(str(pfad))
alle = [p.text for p in brief.paragraphs]
volltext = "\n".join(alle)

pruefe("Muster Rueckbau & Sanierung GmbH" in alle, "Firma fehlt im Adressblock")
pruefe("per E-Mail:" in alle, "Zeile 'per E-Mail:' fehlt")
pruefe("BOB_Boulevard Berlin" in alle, "Projektzeile fehlt")
pruefe("VE100- Abbrucharbeiten" in alle, "Zeile der Vergabeeinheit fehlt")
pruefe(any(z.startswith("Ihre Mehrkostenanzeige Nr. 03") for z in alle),
       "Betreffzeile fehlt")
pruefe("Sehr geehrter Herr Ashgarzada," in alle, "Anrede fehlt")
pruefe("Katharina Blanck" in alle, "Unterzeichner fehlt")
pruefe("-Baumanagement-" in alle, "Funktion fehlt")
pruefe("Götz Gagelmann" not in volltext,
       "Der Name aus der Vorlage steht noch im Schreiben")
pruefe("Projektnr_Projektname" not in volltext and "Betreff…." not in volltext,
       "Platzhalter der Vorlage sind stehen geblieben")
pruefe("Anlage:" in alle and "Verteiler:" in alle,
       "Anlagen-/Verteilerblock fehlt")
pruefe(gen.BAUZEIT_ABLEHNUNG in alle, "Bauzeit-Ablehnung fehlt im Dokument")

# Der Adressblock muss im Textrahmen der Vorlage stehen — sonst rutscht die
# Anschrift mitten in den Brief.
rahmen = [
    p for p in brief.paragraphs
    if p.style.name == "Adressblock"
    and p._p.find(qn("w:pPr")) is not None
    and p._p.find(qn("w:pPr")).find(qn("w:framePr")) is not None
]
pruefe(len(rahmen) == 7, f"Adressblock-Absaetze im Rahmen: {len(rahmen)}")

# Datumszeile
zellen = [z.text for z in brief.tables[0].rows[0].cells]
pruefe(zellen == ["03.09.2026", "Ze: kb", "T - 22", "w29@hpp.com"],
       f"Datumszeile: {zellen}")

# Der LV-Auszug ist eingerueckt.
eingerueckt = [
    p.text for p in brief.paragraphs
    if p._p.find(qn("w:pPr")) is not None
    and p._p.find(qn("w:pPr")).find(qn("w:ind")) is not None
    and p._p.find(qn("w:pPr")).find(qn("w:ind")).get(qn("w:left")) ==
    str(gen.ZITAT_EINZUG)
]
pruefe(eingerueckt == ["Bodenbelaege sind vollstaendig zurueckzubauen."],
       f"eingerueckte Absaetze: {eingerueckt}")

# Erste Seite voller Briefkopf, Folgeseiten der verkleinerte.
brief_sectpr = brief.sections[0]._sectPr
pruefe(brief_sectpr.find(qn("w:titlePg")) is not None,
       "titlePg fehlt — Seite 2 truege den vollen Briefkopf")
bezuege = {
    b.get(qn("w:type")): b.get(qn("r:id"))
    for b in brief_sectpr.findall(qn("w:headerReference"))
}
pruefe(bezuege.get("first") == gen.KOPF_VOLL,
       f"Kopf der ersten Seite: {bezuege}")
pruefe(bezuege.get("default") == gen.KOPF_FOLGE,
       f"Kopf der Folgeseiten: {bezuege}")
fuesse = {
    b.get(qn("w:type")): b.get(qn("r:id"))
    for b in brief_sectpr.findall(qn("w:footerReference"))
}
pruefe(fuesse.get("first") == gen.FUSS_SEITEN,
       f"Fuss der ersten Seite (Seitenzahl): {fuesse}")
pruefe(len(brief.sections) == 2, "Verteiler braucht den zweiten Abschnitt")
pruefe(brief.sections[1]._sectPr.find(qn("w:type")).get(qn("w:val")) ==
       "nextPage", "Verteiler muss auf einer neuen Seite beginnen")

# Dateiname in der Fusszeile
fusstext = "".join(p.text for p in brief.sections[0].footer.paragraphs)
pruefe("260903 BOB-VE100-MKA 03.docx" in fusstext,
       f"Dateiname fehlt in der Fusszeile: {fusstext!r}")

# Ohne Verteiler bleibt kein leerer Abschnitt uebrig.

(STORAGE / "ohne.docx").write_bytes(gen.erzeuge(daten()))
ohne_verteiler = Document(str(STORAGE / "ohne.docx"))
pruefe(len(ohne_verteiler.sections) == 1,
       f"ohne Verteiler muesste ein Abschnitt bleiben: "
       f"{len(ohne_verteiler.sections)}")
pruefe(ohne_verteiler.sections[0]._sectPr.find(qn("w:titlePg")) is not None,
       "titlePg fehlt im einabschnittigen Brief")

abschnitt("Erzeuger: keine eingescannte Unterschrift")


def hat_unterschrift(pfad: Path) -> bool:
    """Steckt im Textkoerper ein eingebettetes Bild?

    Geprueft wird auf das Element ``w:drawing``, nicht auf die Zeichenkette
    "drawing" im XML: lxml schreibt beim Serialisieren alle geerbten
    Namensraeume mit, und einer davon heisst "wordprocessingDrawing" — eine
    Textsuche findet ihn in jedem Absatz und meldet immer "Bild vorhanden".
    """
    for absatz in Document(str(pfad)).paragraphs:
        for lauf in absatz._p.findall(qn("w:r")):
            if lauf.findall(qn("w:drawing")):
                return True
    return False


# Das Unterschriftsbild der Vorlage kommt in kein Schreiben — jeder Brief
# geht zum Unterschreiben aus dem Haus.
ohne = daten()                                        # Name ist Katharina Blanck
ohne_dok = STORAGE / "ohne_unterschrift.docx"
ohne_dok.write_bytes(gen.erzeuge(ohne))
pruefe(not hat_unterschrift(ohne_dok),
       "im Schreiben darf keine eingescannte Unterschrift stehen")
pruefe(not ohne.hinweise, f"unerwartete Hinweise: {ohne.hinweise}")

# Auch unter dem Namen dessen, dessen Unterschrift in der Vorlage steckt.
eigen = daten(sachbearbeiter=gen.Sachbearbeiter(
    name="Götz Gagelmann", funktion="-Partner-", zeichen="gg",
    durchwahl="25", email="goetz.gagelmann@hpp.com"))
eigen_dok = STORAGE / "eigen.docx"
eigen_dok.write_bytes(gen.erzeuge(eigen))
pruefe(not hat_unterschrift(eigen_dok),
       "auch unter dem eigenen Namen bleibt das Unterschriftsbild weg")

# ─────────────────────────────────────────────────────────────────────────────
# 3b. Stichpunkte ausformulieren
#
# Der Aufruf an die Schnittstelle wird NICHT gemacht — er kostet Geld und
# braucht einen Schluessel. Geprueft wird alles davor und danach: der Aufbau
# der Anfrage, die Uebersetzung der Antwort ins Infofeld und dass der Erzeuger
# daraus wieder genau die gemeinten Bloecke macht.
# ─────────────────────────────────────────────────────────────────────────────

abschnitt("Angaben im Infofeld")

# Eine Zeile mit Beschriftung und Doppelpunkt geht an ihren Platz im
# Schreiben, statt Brieftext zu werden.
mit_angaben = daten(
    empfaenger=gen.Empfaenger(firma="SERVISA GmbH"),
    sachbearbeiter=gen.Sachbearbeiter(),
    projektzeile="", vergabeeinheit="", dateikuerzel="", briefdatum=None,
    haltung="kenntnisnahme",
    stellungnahme=(
        "Projekt: BOB_Boulevard Berlin\n"
        "VE: VE100- Abbrucharbeiten\n"
        "Kuerzel: BOB\n"
        "Anrede: Herr\n"
        "Ansprechpartner: Qasem Ashgarzada\n"
        "Strasse: Attilastrasse 24\n"
        "PLZ Ort: 12105 Berlin\n"
        "Mail: Qasem@servisa-gruppe.de\n"
        "Unterzeichner: Katharina Blanck\n"
        "Zeichen: kb\n"
        "Durchwahl: 22\n"
        "E-Mail HPP: w29@hpp.com\n"
        "Datum: 03.09.2026\n"
        "Haltung: ablehnen\n"
        "Bauzeit ablehnen\n"
        "Betreff: Ihre MKA Nr. 03\n"
        "\n"
        "Wir sehen hier keinen Mehrverguetungsanspruch.\n"
        "Die Anlage 3 haben wir geprueft und fuer richtig befunden.\n"
        "\n"
        "Anlage: Auszug LV Pos. 03.0110\n"
        "Anlage: Bestandsschnitt Ebene 0\n"
        "Verteiler: sowie Projektadressen"
    ),
)
gemeldet = gen.vorbereiten(mit_angaben)

pruefe(mit_angaben.projektzeile == "BOB_Boulevard Berlin",
       f"Projektzeile: {mit_angaben.projektzeile!r}")
pruefe(mit_angaben.vergabeeinheit == "VE100- Abbrucharbeiten",
       f"Vergabeeinheit: {mit_angaben.vergabeeinheit!r}")
pruefe(mit_angaben.dateikuerzel == "BOB", f"Kuerzel: {mit_angaben.dateikuerzel!r}")
pruefe(mit_angaben.betreff == "Ihre MKA Nr. 03", f"Betreff: {mit_angaben.betreff!r}")
pruefe(mit_angaben.briefdatum == date(2026, 9, 3),
       f"Briefdatum: {mit_angaben.briefdatum}")
pruefe(mit_angaben.haltung == "ablehnung", f"Haltung: {mit_angaben.haltung!r}")
pruefe(mit_angaben.bauzeit_ablehnen is True, "Schalter Bauzeit nicht gesetzt")
pruefe(mit_angaben.empfaenger.anrede == "Herr",
       f"Anrede: {mit_angaben.empfaenger.anrede!r}")
pruefe(mit_angaben.empfaenger.ansprechpartner == "Qasem Ashgarzada",
       f"Ansprechpartner: {mit_angaben.empfaenger.ansprechpartner!r}")
pruefe(mit_angaben.empfaenger.strasse == "Attilastrasse 24",
       f"Strasse: {mit_angaben.empfaenger.strasse!r}")
pruefe((mit_angaben.empfaenger.plz, mit_angaben.empfaenger.ort)
       == ("12105", "Berlin"),
       f"PLZ/Ort: {mit_angaben.empfaenger.plz!r} {mit_angaben.empfaenger.ort!r}")
pruefe(mit_angaben.anlagen.split("\n")
       == ["Auszug LV Pos. 03.0110", "Bestandsschnitt Ebene 0"],
       f"Anlagen sammeln sich nicht: {mit_angaben.anlagen!r}")
pruefe(mit_angaben.verteiler == "sowie Projektadressen",
       f"Verteiler: {mit_angaben.verteiler!r}")

# Der schlimmste denkbare Fehler dieser Liste: die beiden Mailadressen
# verwechseln. Die der Firma steht im Adressfeld und ist Mailempfaenger, die
# des Bueros in der Datumszeile.
pruefe(mit_angaben.empfaenger.email == "Qasem@servisa-gruppe.de",
       f"Mail der Firma: {mit_angaben.empfaenger.email!r}")
pruefe(mit_angaben.sachbearbeiter.email == "w29@hpp.com",
       f"Mail HPP: {mit_angaben.sachbearbeiter.email!r}")
pruefe(mit_angaben.sachbearbeiter.name == "Katharina Blanck",
       f"Unterzeichner: {mit_angaben.sachbearbeiter.name!r}")
pruefe((mit_angaben.sachbearbeiter.zeichen, mit_angaben.sachbearbeiter.durchwahl)
       == ("kb", "22"),
       f"Zeichen/Durchwahl: {mit_angaben.sachbearbeiter.zeichen!r} "
       f"{mit_angaben.sachbearbeiter.durchwahl!r}")

# Ein gewoehnlicher Satz bleibt Brieftext, auch wenn er dieselben Woerter hat.
pruefe("Die Anlage 3 haben wir geprueft" in mit_angaben.stellungnahme,
       f"ein Satz ohne Doppelpunkt darf nicht verschwinden: "
       f"{mit_angaben.stellungnahme!r}")
pruefe("Projekt:" not in mit_angaben.stellungnahme
       and "Verteiler:" not in mit_angaben.stellungnahme,
       f"Angabenzeilen stehen noch im Brieftext: {mit_angaben.stellungnahme!r}")
pruefe(mit_angaben.stellungnahme.startswith("Wir sehen hier keinen"),
       f"Brieftext faengt falsch an: {mit_angaben.stellungnahme[:60]!r}")

# Es muss gemeldet werden, was gewandert ist — sonst ist es Zauberei.
gesamtmeldung = " | ".join(gemeldet)
for stichwort in ("Projektzeile", "Anlage", "Verteiler", "E-Mail der Firma",
                  "E-Mail HPP", "Bauzeit"):
    pruefe(stichwort in gesamtmeldung,
           f"in der Rueckmeldung fehlt {stichwort!r}: {gesamtmeldung[:200]}")

# Idempotent: Vorschau und Erzeugen rufen beide auf, ohne sich abzustimmen.
vorher_text = mit_angaben.stellungnahme
pruefe(gen.vorbereiten(mit_angaben) == [],
       "der zweite Aufruf darf nichts mehr finden")
pruefe(mit_angaben.stellungnahme == vorher_text,
       "der zweite Aufruf darf den Brieftext nicht veraendern")

# Zweimal dieselbe Angabe: Der letzte Wert gewinnt, und es wird gesagt.
doppelt = daten(betreff="", stellungnahme=(
    "Betreff: Erster Versuch\nBetreff: Zweiter Versuch\n\nText."))
doppelt_meldung = " | ".join(gen.vorbereiten(doppelt))
pruefe(doppelt.betreff == "Zweiter Versuch", f"Betreff: {doppelt.betreff!r}")
pruefe("zweimal" in doppelt_meldung,
       f"doppelte Angabe nicht gemeldet: {doppelt_meldung}")

# Unbekannte Werte werden nicht stillschweigend geraten.
krumm = daten(stellungnahme="Haltung: irgendwas\nAnrede: Firma\nDatum: gestern\n\nText.")
krumm_meldung = " | ".join(gen.vorbereiten(krumm))
pruefe(krumm.haltung == "kenntnisnahme",
       f"unbekannte Haltung haette bleiben muessen: {krumm.haltung!r}")
pruefe("nicht erkannt" in krumm_meldung and "nicht lesbar" in krumm_meldung,
       f"krumme Angaben nicht gemeldet: {krumm_meldung}")

# "Anlage:" ohne Wert ist ein Tippfehler, kein Befehl — und bleibt stehen.
leer_angabe = daten(stellungnahme="Anlage:\n\nText.")
gen.vorbereiten(leer_angabe)
pruefe("Anlage:" in leer_angabe.stellungnahme,
       "eine leere Angabe darf nicht verschwinden")
pruefe(not leer_angabe.anlagen, f"Anlagen: {leer_angabe.anlagen!r}")

abschnitt("Bausteine und Glaetten (ohne Schluessel)")

# Der Katalog muss zur Haltung passen. Das ist keine Kosmetik: Ein Fehlklick
# auf "Dies ist eine Zusatzleistung" in einem Ablehnungsschreiben sagt das
# Gegenteil des Gemeinten.
alle_gruppen = bau.fuer_haltung()
pruefe(len(alle_gruppen) == len(bau.KATALOG),
       f"ohne Haltung muss der ganze Katalog kommen: {len(alle_gruppen)}")

ablehnung_texte = [
    b.kennung for g in bau.fuer_haltung("ablehnung") for b in g.bausteine
]
anerkennung_texte = [
    b.kennung for g in bau.fuer_haltung("anerkennung") for b in g.bausteine
]
pruefe("ist_zusatz" not in ablehnung_texte,
       "„Dies ist eine Zusatzleistung“ darf bei einer Ablehnung nicht "
       "angeboten werden")
pruefe("ist_zusatz" in anerkennung_texte,
       "bei einer Anerkennung muss der Zusatzleistungs-Satz da sein")
pruefe("kein_anspruch" in ablehnung_texte,
       "bei einer Ablehnung fehlt „Kein Mehrverguetungsanspruch“")
pruefe("kein_anspruch" not in anerkennung_texte,
       "„Kein Mehrverguetungsanspruch“ passt nicht zu einer Anerkennung")
# Bausteine ohne Haltungsbindung sind immer dabei.
for haltung in gen.HALTUNGEN:
    kennungen = [b.kennung for g in bau.fuer_haltung(haltung) for b in g.bausteine]
    pruefe("lv_marke" in kennungen,
           f"der LV-Auszug muss bei „{haltung}“ angeboten werden")
    pruefe(kennungen, f"bei „{haltung}“ kam kein einziger Baustein")

# Jeder Baustein muss ein brauchbarer Satz sein — oder eine Zitatzeile.
for gruppe in bau.KATALOG:
    for baustein in gruppe.bausteine:
        pruefe(baustein.titel and len(baustein.titel) <= 45,
               f"Beschriftung unbrauchbar: {baustein.titel!r}")
        if baustein.zitat or baustein.text.endswith(":"):
            continue
        pruefe(baustein.text.rstrip().endswith((".", "!", "?")),
               f"Baustein ohne Satzende: {baustein.text!r}")
        pruefe(baustein.text[0].isupper(),
               f"Baustein faengt klein an: {baustein.text!r}")

# Glaetten: Form, nicht Inhalt.
geglaettet, glaett_hinweise = bau.glaette(
    "- pos 3.11 ausgeschrieben, ep abgegolten\n"
    "- be-kosten sind einzukalkulieren\n"
    "- entsorgung nach din 18459"
)
pruefe(geglaettet.split("\n") == [
    "1) Pos. 3.11 ausgeschrieben, EP abgegolten.",
    "2) BE-Kosten sind einzukalkulieren.",
    "3) Entsorgung nach DIN 18459.",
], f"Glaetten falsch:\n{geglaettet}")
pruefe(not glaett_hinweise, f"unerwartete Hinweise: {glaett_hinweise}")

pruefe(bau.glaette("wir sehen keinen anspruch nach §2 abs 6 vob/b")[0]
       == "Wir sehen keinen anspruch nach § 2 Abs. 6 VOB/B.",
       f"Paragraf nicht geglaettet: {bau.glaette('§2 abs 6 vob/b')[0]!r}")

# Gewoehnliche Woerter duerfen nicht zu Abkuerzungen werden.
pruefe(bau.glaette("an der wand war nichts zu sehen")[0]
       == "An der wand war nichts zu sehen.",
       "„an“ mitten im Satz darf nicht zu „AN“ werden")
pruefe("BE-Kosten" in bau.glaette("die be-kosten sind hoch")[0],
       "„be-kosten“ muesste „BE-Kosten“ werden")

# Die Nummerierung laeuft weiter, wenn ein eingesetzter Baustein zwischen zwei
# Punkten steht — sonst stuenden zwei Punkte mit derselben Nummer im Brief.
durchlauf, _ = bau.glaette(
    "- pos 03.0110 abgegolten\n"
    "Hier gibt es eine Position im LV.\n"
    "- bestandsschnitte lagen bei"
)
pruefe(durchlauf.split("\n") == [
    "1) Pos. 03.0110 abgegolten.",
    "Hier gibt es eine Position im LV.",
    "2) Bestandsschnitte lagen bei.",
], f"Nummerierung nach einem Satz dazwischen:\n{durchlauf}")

# Eine Leerzeile beginnt dagegen eine neue Aufzaehlung.
zwei_bloecke, _ = bau.glaette("- erstens\n- zweitens\n\n- neu erstens")
pruefe(
    zwei_bloecke.split("\n")
    == ["1) Erstens.", "2) Zweitens.", "", "1) Neu erstens."],
    f"neue Liste nach Leerzeile:\n{zwei_bloecke}",
)

# Ein LV-Zitat wird NICHT geglaettet — sonst stimmt das Zitat nicht mehr.
zitat_ein = "Auszug LV:\nzeitliche unterbrechungen sind einzukalkulieren"
zitat_aus, _ = bau.glaette(zitat_ein)
pruefe(zitat_aus.split("\n")[1] == "zeitliche unterbrechungen sind einzukalkulieren",
       f"das Zitat wurde veraendert: {zitat_aus!r}")

# Zwei Saetze in einer Zeile: beide Anfaenge gross.
pruefe(bau.glaette("erstens gilt das. zweitens gilt jenes")[0]
       == "Erstens gilt das. Zweitens gilt jenes.",
       f"zweiter Satzanfang nicht gross: "
       f"{bau.glaette('erstens gilt das. zweitens gilt jenes')[0]!r}")

# Die Grenze wird ausgesprochen, nicht verschwiegen.
_, klein_hinweise = bau.glaette(
    "wir sehen keinen anspruch. die leistung ist im leistungsverzeichnis "
    "enthalten und mit den einheitspreisen abgegolten"
)
pruefe(any("Großschreibung" in h for h in klein_hinweise),
       f"bei durchgehend kleiner Schreibung fehlt der Hinweis: {klein_hinweise}")
_, luecken_hinweise = bau.glaette(f"Die Abrechnung erfolgt über Pos. {bau.LUECKE}.")
pruefe(any("Lücken" in h for h in luecken_hinweise),
       f"offene Luecke nicht gemeldet: {luecken_hinweise}")

# Ein eingesetzter Baustein muss den Erzeuger sauber durchlaufen.
mit_baustein = "\n".join([
    b.text for g in bau.fuer_haltung("ablehnung") for b in g.bausteine
    if b.kennung in ("kein_anspruch", "im_lv_enthalten")
])
bloecke_baustein = gen._stellungnahme_bloecke(mit_baustein)
pruefe(len(bloecke_baustein) == 1,
       "zwei Bausteine ohne Leerzeile gehoeren in einen Absatz: "
       f"{[b.text[:30] for b in bloecke_baustein]}")

abschnitt("Ausformulieren: Anfrage und Rueckweg")

auftrag = form.Auftrag(
    stichpunkte="Pos 03.0110 ausgeschrieben, EP abgegolten\nBestandsschnitte lagen bei",
    art="Mehrkostenanzeige", nummer="03", datum="17.08.2026",
    kurzbezeichnung="Abweichende Bodenaufbauten",
    punkte=["1. Doppelboeden", "2. Mehrlagige Belaege"],
    lv_positionen=["03.0110", "03.0120"],
    bauzeit="Verlaengerung um eine Woche.",
    rechtsgrundlage="§ 2 Abs. 6 VOB/B",
    anzeigetext="Wortlaut der Anzeige mit den Feststellungen der Firma.",
    haltung="ablehnung",
    projekt="BOB_Boulevard Berlin", vergabeeinheit="VE100- Abbrucharbeiten",
)
anfragetext, anfragehinweise = form.baue_anfrage(auftrag)

pruefe("Pos 03.0110 ausgeschrieben" in anfragetext,
       "die Stichpunkte fehlen in der Anfrage")
pruefe("1. Doppelboeden" in anfragetext, "die Punkte der Firma fehlen")
pruefe("03.0110, 03.0120" in anfragetext, "die LV-Positionen fehlen")
pruefe("Wortlaut der Anzeige" in anfragetext, "der Anzeigetext fehlt")
pruefe("§ 2 Abs. 6 VOB/B" in anfragetext, "die Rechtsgrundlage fehlt")
pruefe(form.HALTUNG_TEXT["ablehnung"] in anfragetext,
       "die Haltung steht nicht in der Anfrage")
pruefe(not anfragehinweise, f"unerwartete Hinweise: {anfragehinweise}")

# Zu lange Eingaben werden gekuerzt — und das wird gesagt, nicht verschwiegen.
lang = form.Auftrag(stichpunkte="x" * (form.MAX_STICHPUNKTE_ZEICHEN + 500),
                    anzeigetext="y" * (form.MAX_ANZEIGE_ZEICHEN + 500))
_, gekuerzt = form.baue_anfrage(lang)
pruefe(len(gekuerzt) == 2, f"Kuerzungen nicht gemeldet: {gekuerzt}")

# Die harten Regeln muessen in der Anweisung stehen — daran haengt, dass das
# Modell keine LV-Positionen und keine Termine erfindet.
for stichwort in ("Erfinde nichts", "offene_fragen", "KEINE Anrede",
                  "Auszug LV:", "Entscheide nicht"):
    pruefe(stichwort in form.ANWEISUNG,
           f"Regel fehlt in der Anweisung: {stichwort!r}")

# Rueckweg: Absaetze -> Infofeld -> Erzeuger.
feld = form.zu_infofeld([
    ("Wir sehen hier keinen Mehrverguetungsanspruch.", False),
    ("1) Doppelboeden — Der Rueckbau ist mit Pos. 03.0110 abgegolten.", False),
    ("Auszug LV:", False),
    ("Bodenbelaege sind vollstaendig zurueckzubauen.", True),
    ("Mehrlagige Aufbauten sind einzukalkulieren.", True),
    ("Eine Mengenmehrung liegt nicht vor.", False),
])
pruefe("Auszug LV:\nBodenbelaege" in feld,
       f"nach der Zitat-Einleitung darf keine Leerzeile stehen:\n{feld!r}")
pruefe("zurueckzubauen.\nMehrlagige" in feld,
       f"innerhalb des Zitats darf keine Leerzeile stehen:\n{feld!r}")
pruefe("einzukalkulieren.\n\nEine Mengenmehrung" in feld,
       f"nach dem Zitat fehlt die Leerzeile:\n{feld!r}")

# Und der Beweis, dass der Erzeuger daraus dasselbe macht — es gibt nur einen
# Weg ins Dokument, und der fuehrt durch das Infofeld.
bloecke_rueck = gen._stellungnahme_bloecke(feld)
pruefe([b.zitat for b in bloecke_rueck] ==
       [False, False, False, True, True, False],
       f"Einzuege nach dem Rueckweg falsch: "
       f"{[(b.text[:25], b.zitat) for b in bloecke_rueck]}")
pruefe(bloecke_rueck[3].text == "Bodenbelaege sind vollstaendig zurueckzubauen.",
       f"Zitatzeile verfaelscht: {bloecke_rueck[3].text!r}")

# Ohne Schluessel: klare Meldung, kein Absturz, und der Grund steht drin.
pruefe("einstellungen.txt" in form.warum_nicht()
       and "anthropic_key" in form.warum_nicht(),
       f"Meldung sagt nicht, wo der Schluessel hingehoert: {form.warum_nicht()}")

alter_schluessel = form.settings.anthropic_api_key
form.settings.anthropic_api_key = ""
try:
    asyncio.run(form.formuliere(form.Auftrag(stichpunkte="etwas")))
    pruefe(False, "ohne Schluessel muesste formuliere() werfen")
except form.FormulierungFehler as f:
    pruefe("anthropic_key" in str(f), f"falsche Meldung: {f}")

# Mit Schluessel, aber ohne Stichpunkte, gibt es nichts zu tun.
form.settings.anthropic_api_key = "sk-ant-test"
try:
    asyncio.run(form.formuliere(form.Auftrag(stichpunkte="   ")))
    pruefe(False, "ohne Stichpunkte muesste formuliere() werfen")
except form.FormulierungFehler as f:
    pruefe("Stichpunkte" in str(f), f"falsche Meldung: {f}")
form.settings.anthropic_api_key = alter_schluessel

# Die Modellantwort wird gesaeubert: leere Absaetze und Doppelte fliegen raus.
absaetze, fragen = form._aus_antwort({
    "absaetze": [{"text": "Erster Absatz."}, {"text": "   "},
                 {"text": "Zitat", "zitat": True}, "kaputt"],
    "offene_fragen": ["Datum fehlt", "Datum fehlt", ""],
})
pruefe([a[0] for a in absaetze] == ["Erster Absatz.", "Zitat"],
       f"Saeuberung der Absaetze: {absaetze}")
pruefe(absaetze[1][1] is True, "Zitatkennzeichen verloren")
pruefe(fragen == ["Datum fehlt"], f"offene Fragen nicht entdoppelt: {fragen}")

# Zeitgrenze und Wiederholung. Ohne Zeitgrenze wartet das Anthropic-Paket zehn
# Minuten — im Buero ist die Schnittstelle je nach Firewall nicht erreichbar,
# und die Oberflaeche stuende so lange auf "Wird formuliert...".
form.settings.anthropic_api_key = "sk-ant-test"
try:
    zugang = form._client()
    pruefe(float(zugang.timeout) == form.ZEITGRENZE_SEKUNDEN,
           f"Zeitgrenze nicht gesetzt: {zugang.timeout!r}")
    pruefe(zugang.timeout is not None and float(zugang.timeout) <= 120,
           f"Zeitgrenze zu gross: {zugang.timeout!r}")
    # Innen NICHT wiederholen: aussen tut es schnittstelle.mit_wiederholung.
    # Beides zusammen waeren neun Anfragen und eine Viertelstunde Warten.
    pruefe(zugang.max_retries == 0,
           f"das Paket wiederholt selbst mit: max_retries={zugang.max_retries}")
finally:
    form.settings.anthropic_api_key = alter_schluessel

# Ein falscher Schluessel muss als Satz ankommen, nicht als Statuscode: Wer
# ihn vertippt hat, soll das lesen koennen.
class _Streik:
    def __init__(self, fehler):
        self._fehler = fehler
        self.messages = self

    def create(self, **_):
        raise self._fehler


for rohfehler, erwartet_stichwort in [
    (Exception("Error code: 401 - invalid x-api-key"), "chl"),
    (Exception("Error code: 429 - rate_limit_error"), "berlast"),
]:
    echter = form._client
    form._client = lambda f=rohfehler: _Streik(f)
    form.settings.anthropic_api_key = "sk-ant-test"
    try:
        asyncio.run(form.formuliere(form.Auftrag(
            stichpunkte="EP abgegolten", haltung="ablehnung")))
        pruefe(False, f"{rohfehler} haette eine Meldung geben muessen")
    except form.FormulierungFehler as f:
        pruefe("Error code" not in str(f),
               f"der rohe Statuscode steht in der Meldung: {f}")
        pruefe(erwartet_stichwort in str(f),
               f"Meldung zu {rohfehler} unklar: {f}")
    finally:
        form._client = echter
        form.settings.anthropic_api_key = alter_schluessel


abschnitt("Ausformulieren: die ganze Kette mit Attrappe")

# Ohne Schluessel laesst sich der echte Aufruf nicht machen. Die Kette
# dahinter ist aber genau die, die im Buero zaehlt: Antwort -> Infofeld ->
# Erzeuger -> Word. Deshalb wird hier nur der Client ersetzt und der Rest
# wirklich durchlaufen. So faellt auf, wenn sich das Format zwischen
# Formulierung und Erzeuger auseinanderentwickelt.


class _Block:
    def __init__(self, name, eingabe):
        self.type = "tool_use"
        self.name = name
        self.input = eingabe


class _Antwort:
    def __init__(self, eingabe):
        self.content = [_Block(form.WERKZEUG["name"], eingabe)]


class _Nachrichten:
    def __init__(self, eingabe, protokoll):
        self._eingabe = eingabe
        self._protokoll = protokoll

    def create(self, **kwargen):
        self._protokoll.append(kwargen)
        return _Antwort(self._eingabe)


class _Attrappe:
    def __init__(self, eingabe, protokoll):
        self.messages = _Nachrichten(eingabe, protokoll)


ANTWORT_ATTRAPPE = {
    "absaetze": [
        {"text": "Wir sehen hier keinen Mehrverguetungsanspruch. In Ihrem "
                 "Auftrags-LV ist Ihr Vertragssoll dargelegt."},
        {"text": "1) Doppelboeden — Der Rueckbau aufgestaenderter Doppelboeden "
                 "ist mit Pos. 03.0110 ausgeschrieben und mit den "
                 "Einheitspreisen abgegolten."},
        {"text": "2) Mehrlagige Belaege — Die Bestandsschnitte lagen der "
                 "Ausschreibung bei."},
        {"text": "Auszug LV:"},
        {"text": "Bodenbelaege einschliesslich Kleber sind vollstaendig "
                 "zurueckzubauen.", "zitat": True},
        {"text": "Mehrlagige Aufbauten sind einzukalkulieren.", "zitat": True},
    ],
    "offene_fragen": ["Zur getrennten Entsorgung fehlt die LV-Position."],
}

protokoll: list[dict] = []
echter_client = form._client
form._client = lambda: _Attrappe(ANTWORT_ATTRAPPE, protokoll)
alter_schluessel = form.settings.anthropic_api_key
form.settings.anthropic_api_key = "sk-ant-attrappe"
try:
    ergebnis = asyncio.run(form.formuliere(auftrag))
finally:
    form._client = echter_client
    form.settings.anthropic_api_key = alter_schluessel

pruefe(len(protokoll) == 1, f"genau eine Anfrage erwartet: {len(protokoll)}")
gestellt = protokoll[0]
pruefe(gestellt["model"] == form.CLAUDE_MODELL,
       f"falsches Modell: {gestellt['model']!r}")
pruefe(gestellt["tool_choice"]["name"] == form.WERKZEUG["name"],
       "die Antwort muss ins Werkzeug erzwungen werden, sonst kommt Prosa")
pruefe("Erfinde nichts" in gestellt["system"],
       "die Anweisung wurde nicht mitgeschickt")
pruefe("Pos 03.0110 ausgeschrieben" in gestellt["messages"][0]["content"],
       "die Stichpunkte wurden nicht mitgeschickt")

pruefe(ergebnis.offene_fragen ==
       ["Zur getrennten Entsorgung fehlt die LV-Position."],
       f"offene Fragen verloren: {ergebnis.offene_fragen}")
pruefe(ergebnis.stellungnahme.startswith("Wir sehen hier keinen"),
       f"Stellungnahme faengt falsch an: {ergebnis.stellungnahme[:60]!r}")

# Jetzt der eigentliche Punkt: Aus dem Feld muss ein richtiges Schreiben werden.
mit_text = daten(stellungnahme=ergebnis.stellungnahme, haltung="ablehnung")
docx_attrappe = STORAGE / "formuliert.docx"
docx_attrappe.write_bytes(gen.erzeuge(mit_text))
gebaut = Document(str(docx_attrappe))
zeilen_im_brief = [p.text for p in gebaut.paragraphs]

pruefe(any(z.startswith("1) Doppelboeden") for z in zeilen_im_brief),
       "der formulierte Punkt 1 steht nicht im Schreiben")
pruefe("Auszug LV:" in zeilen_im_brief,
       "die Zitat-Einleitung fehlt im Schreiben")
eingezogen = [
    p.text for p in gebaut.paragraphs
    if p._p.find(qn("w:pPr")) is not None
    and p._p.find(qn("w:pPr")).find(qn("w:ind")) is not None
    and p._p.find(qn("w:pPr")).find(qn("w:ind")).get(qn("w:left"))
    == str(gen.ZITAT_EINZUG)
]
pruefe(len(eingezogen) == 2,
       f"die zwei zitierten LV-Zeilen muessten eingerueckt sein: {eingezogen}")
# Die offene Frage darf NICHT im Schreiben stehen — sie ist eine Rueckfrage
# an den Menschen, kein Satz fuer die Baufirma.
pruefe(not any("fehlt die LV-Position" in z for z in zeilen_im_brief),
       "eine offene Rueckfrage ist im Schreiben an die Firma gelandet")


# ─────────────────────────────────────────────────────────────────────────────
# 4. Endpunkte
# ─────────────────────────────────────────────────────────────────────────────

abschnitt("Endpunkte")

with TestClient(app) as c:
    # ── Auslesen ──
    with open(quelle, "rb") as f:
        antwort = c.post(
            "/api/anzeigen/auslesen",
            files={"dateien": ("MKA 03.docx", f.read(), "application/vnd."
                               "openxmlformats-officedocument.wordprocessingml."
                               "document")},
        )
    pruefe(antwort.status_code == 200, f"auslesen: {antwort.status_code} "
                                       f"{antwort.text[:200]}")
    ergebnis = antwort.json()
    pruefe(len(ergebnis["schreiben"]) == 1,
           f"ein Schreiben erwartet: {ergebnis}")
    eines = ergebnis["schreiben"][0]
    pruefe(eines["art"] == "Mehrkostenanzeige", f"Art: {eines['art']!r}")
    pruefe(eines["nummer"] == "03", f"Nummer: {eines['nummer']!r}")
    pruefe(eines["absender"]["firma"] == "Muster Rueckbau & Sanierung GmbH",
           f"Absender: {eines['absender']}")
    pruefe(eines["datum"] == "2026-08-17", f"Datum: {eines['datum']!r}")
    pruefe(len(eines["punkte"]) == 3, f"Punkte: {eines['punkte']}")
    pruefe("Mehrkostenanzeige Nr. 03" in eines["volltext"],
           "der Volltext fehlt in der Antwort — ohne ihn kann die "
           "Stellungnahme nicht ausformuliert werden")

    # Eine unlesbare Datei darf die lesbare nicht mitreissen.
    with open(quelle, "rb") as f:
        gemischt = c.post(
            "/api/anzeigen/auslesen",
            files=[
                ("dateien", ("MKA 03.docx", f.read(), "application/octet-stream")),
                ("dateien", ("tabelle.xlsx", b"nichts", "application/octet-stream")),
            ],
        )
    pruefe(gemischt.status_code == 200, f"gemischter Upload: {gemischt.text[:200]}")
    pruefe(len(gemischt.json()["schreiben"]) == 1 and
           len(gemischt.json()["fehlgeschlagen"]) == 1,
           f"gemischter Upload falsch aufgeteilt: {gemischt.json()}")

    nur_kaputt = c.post("/api/anzeigen/auslesen",
                        files={"dateien": ("x.xlsx", b"nichts",
                                           "application/octet-stream")})
    pruefe(nur_kaputt.status_code == 422,
           f"nur unlesbare Dateien muessten 422 sein: {nur_kaputt.status_code}")

    # ── Vorbelegung ──
    leer = c.get("/api/anzeigen/vorbelegung")
    pruefe(leer.status_code == 200, f"Vorbelegung ohne Projekt: {leer.text[:200]}")
    pruefe("Mehrkostenanzeige" in leer.json()["arten"] and
           "Behinderungsanzeige" in leer.json()["arten"],
           f"Arten fehlen: {leer.json().get('arten')}")
    pruefe("ablehnung" in leer.json()["haltungen"],
           f"Haltungen fehlen: {leer.json().get('haltungen')}")

    # Die Oberflaeche muss wissen, ob der Formulieren-Knopf etwas tun kann.
    pruefe("formulieren_verfuegbar" in leer.json(),
           "die Vorbelegung sagt nicht, ob formuliert werden kann")
    if not leer.json()["formulieren_verfuegbar"]:
        pruefe("anthropic_key" in leer.json()["formulieren_hinweis"],
               f"Hinweis ohne Schluessel unklar: "
               f"{leer.json()['formulieren_hinweis']!r}")
        ohne = c.post("/api/anzeigen/formulieren",
                      json={"stichpunkte": "EP abgegolten"})
        pruefe(ohne.status_code == 422,
               f"ohne Schluessel muesste 422 kommen: {ohne.status_code}")
        pruefe("anthropic_key" in ohne.text,
               f"Meldung nennt den Schluessel nicht: {ohne.text[:200]}")
    pruefe(c.post("/api/anzeigen/formulieren", json={}).status_code == 422,
           "fehlende Stichpunkte muessten 422 sein")

    # ── Bausteine und Glaetten brauchen keinen Schluessel und muessen daher
    #    unter allen Umstaenden antworten ──
    kat = c.get("/api/anzeigen/bausteine?haltung=ablehnung")
    pruefe(kat.status_code == 200, f"bausteine: {kat.status_code} {kat.text[:150]}")
    katalog = kat.json()
    pruefe(katalog["luecke"] == bau.LUECKE,
           f"Lueckenzeichen: {katalog['luecke']!r}")
    kennungen = [b["kennung"] for g in katalog["gruppen"] for b in g["bausteine"]]
    pruefe("kein_anspruch" in kennungen and "ist_zusatz" not in kennungen,
           f"Katalog passt nicht zur Ablehnung: {kennungen[:6]}")
    pruefe(all(g["titel"] and g["bausteine"] for g in katalog["gruppen"]),
           "leere Gruppe im Katalog")
    pruefe(c.get("/api/anzeigen/bausteine").status_code == 200,
           "ohne Haltung muss der ganze Katalog kommen")
    pruefe(c.get("/api/anzeigen/bausteine?haltung=quatsch").status_code == 400,
           "unbekannte Haltung muesste 400 sein")

    gl = c.post("/api/anzeigen/glaetten",
                json={"text": "- pos 3.11 ep abgegolten\n- be-kosten dazu"})
    pruefe(gl.status_code == 200, f"glaetten: {gl.status_code} {gl.text[:150]}")
    pruefe(gl.json()["text"].startswith("1) Pos. 3.11"),
           f"Glaetten ueber die Schnittstelle: {gl.json()['text']!r}")
    pruefe(c.post("/api/anzeigen/glaetten", json={"text": "   "}).status_code == 422,
           "leerer Text muesste 422 sein")

    pid = c.post("/api/projekte", json={"name": "Boulevard Berlin",
                                        "adresse": ""}).json()["id"]
    c.patch(f"/api/projekte/{pid}", json={"name": "Boulevard Berlin"})
    gid = c.post("/api/gewerke", json={
        "projekt_id": pid, "firma_name": "Muster Rueckbau & Sanierung GmbH",
        "vergabeeinheit_code": "VE100", "vergabeeinheit_bezeichnung": "Abbruch",
        "email": "Qasem.Ashgarzada@muster-rueckbau.de",
        "ansprechpartner": "Qasem Ashgarzada", "strasse": "Attilastrasse 24",
        "plz": "12105", "ort": "Berlin",
    }).json()["id"]

    voll = c.get(f"/api/anzeigen/vorbelegung?projekt_id={pid}&gewerk_id={gid}")
    pruefe(voll.status_code == 200, f"Vorbelegung: {voll.text[:200]}")
    daten_voll = voll.json()
    pruefe(daten_voll["empfaenger"]["firma"] ==
           "Muster Rueckbau & Sanierung GmbH",
           f"Empfaenger aus Stammdaten: {daten_voll['empfaenger']}")
    pruefe(daten_voll["vergabeeinheit"] == "VE100 Abbruch",
           f"Vergabeeinheit: {daten_voll['vergabeeinheit']!r}")
    pruefe(any(g["id"] == gid for g in daten_voll["gewerke"]),
           "Gewerkliste fehlt in der Vorbelegung")
    pruefe(c.get("/api/anzeigen/vorbelegung?projekt_id=99999").status_code == 404,
           "unbekanntes Projekt muesste 404 sein")

    # ── Vorschau, Dokument, Mail ──
    anfrage = {
        "empfaenger": {
            "firma": "Muster Rueckbau & Sanierung GmbH", "anrede": "Herr",
            "ansprechpartner": "Qasem Ashgarzada",
            "strasse": "Attilastrasse 24", "plz": "12105", "ort": "Berlin",
            "email": "Qasem.Ashgarzada@muster-rueckbau.de",
        },
        "sachbearbeiter": {
            "name": "Katharina Blanck", "funktion": "-Baumanagement-",
            "zeichen": "kb", "durchwahl": "22", "email": "w29@hpp.com",
        },
        "anzeige": {
            "art": "Mehrkostenanzeige", "nummer": "03", "kennung": "MKA 03",
            "datum": "2026-08-17",
            "kurzbezeichnung": "Zusaetzlicher Rueckbauaufwand",
            "bauzeit": "Verlaengerung um eine Woche.",
        },
        "projektzeile": "BOB_Boulevard Berlin",
        "vergabeeinheit": "VE100- Abbrucharbeiten",
        "briefdatum": "2026-09-03",
        "stellungnahme": "Wir sehen hier keinen Mehrverguetungsanspruch.",
        "haltung": "ablehnung",
        "bauzeit_ablehnen": True,
        "anlagen": "Auszug LV Pos. 03.0110",
        "verteiler": "sowie Projektadressen",
        "dateikuerzel": "BOB",
    }

    vor = c.post("/api/anzeigen/vorschau", json=anfrage)
    pruefe(vor.status_code == 200, f"vorschau: {vor.status_code} {vor.text[:250]}")
    v = vor.json()
    pruefe(v["dateiname"] == "260903 BOB-VE100-MKA 03.docx",
           f"Dateiname: {v['dateiname']!r}")
    pruefe(v["anrede"] == "Sehr geehrter Herr Ashgarzada,", f"Anrede: {v['anrede']!r}")
    pruefe(v["adressblock"][0] == "Muster Rueckbau & Sanierung GmbH",
           f"Adressblock: {v['adressblock']}")
    pruefe(v["datumszeile"] == ["03.09.2026", "Ze: kb", "T - 22", "w29@hpp.com"],
           f"Datumszeile: {v['datumszeile']}")
    pruefe(v["mail_betreff"] ==
           "Boulevard Berlin_Mehrkostenanzeige vom 17.08.26",
           f"Mailbetreff: {v['mail_betreff']!r}")
    pruefe("in der Anlage senden wir unsere Schreiben vom 03.09.26"
           in v["mail_text"], f"Mailtext: {v['mail_text']!r}")
    pruefe(v["mail_an"] == "Qasem.Ashgarzada@muster-rueckbau.de",
           f"Mailempfaenger: {v['mail_an']!r}")
    pruefe(v["verteilerseite"][:1] == ["Anlage:"],
           f"Verteilerseite: {v['verteilerseite']}")
    pruefe(any("Bauzeit" in a["text"] for a in v["absaetze"]),
           "Bauzeit-Absatz fehlt in der Vorschau")

    pruefe(c.post("/api/anzeigen/vorschau",
                  json={**anfrage, "stellungnahme": ""}).status_code == 422,
           "leere Stellungnahme muesste 422 sein")
    pruefe(c.post("/api/anzeigen/vorschau",
                  json={**anfrage, "haltung": "erfunden"}).status_code == 422,
           "unbekannte Haltung muesste 422 sein")

    # Der geglaettete Text muss durch die restliche Kette gehen — sonst waere
    # das Glaetten eine Sackgasse.
    glatt = c.post("/api/anzeigen/glaetten",
                   json={"text": "- pos 3.11 ausgeschrieben, ep abgegolten"})
    mit_glatt = c.post("/api/anzeigen/vorschau",
                       json={**anfrage, "stellungnahme": glatt.json()["text"]})
    pruefe(mit_glatt.status_code == 200,
           f"geglaetteter Text in der Vorschau: {mit_glatt.text[:150]}")
    pruefe(any(a["text"].startswith("1) Pos. 3.11")
               for a in mit_glatt.json()["absaetze"]),
           f"der geglaettete Punkt fehlt in der Vorschau: "
           f"{[a['text'][:40] for a in mit_glatt.json()['absaetze']]}")

    # Und ein eingesetzter Baustein ebenfalls.
    satz = katalog["gruppen"][0]["bausteine"][0]["text"]
    mit_baustein_ep = c.post("/api/anzeigen/vorschau",
                             json={**anfrage, "stellungnahme": satz})
    pruefe(mit_baustein_ep.status_code == 200,
           f"Baustein in der Vorschau: {mit_baustein_ep.text[:150]}")
    pruefe(any(satz[:40] in a["text"] for a in mit_baustein_ep.json()["absaetze"]),
           "der eingesetzte Baustein fehlt in der Vorschau")

    dok = c.post("/api/anzeigen/dokument", json=anfrage)
    pruefe(dok.status_code == 200, f"dokument: {dok.status_code} {dok.text[:200]}")
    pruefe(dok.content[:2] == b"PK", "Antwort ist keine Word-Datei")
    pruefe("MKA 03.docx" in dok.headers.get("content-disposition", ""),
           f"Dateiname im Kopf: {dok.headers.get('content-disposition')!r}")
    pruefe(c.post("/api/anzeigen/dokument?format=xps",
                  json=anfrage).status_code == 400,
           "unbekanntes Format muesste 400 sein")

    mail = c.post("/api/anzeigen/mail/entwurf",
                  json={"antwort": anfrage, "kopie": ["desy.bl@hpp.com"]})
    pruefe(mail.status_code == 200, f"mail/entwurf: {mail.status_code} "
                                    f"{mail.text[:200]}")
    nachricht = email.message_from_bytes(mail.content, policy=policy.default)
    pruefe(nachricht["To"] == "Qasem.Ashgarzada@muster-rueckbau.de",
           f"Mailempfaenger: {nachricht['To']!r}")
    pruefe(nachricht["Cc"] == "desy.bl@hpp.com", f"Kopie: {nachricht['Cc']!r}")
    pruefe(nachricht["Subject"] ==
           "Boulevard Berlin_Mehrkostenanzeige vom 17.08.26",
           f"Mailbetreff: {nachricht['Subject']!r}")
    pruefe(nachricht["X-Unsent"] == "1",
           "ohne X-Unsent zeigt Outlook die Datei als empfangene Mail")
    pruefe(nachricht["From"] is None,
           "der Absender muss leer bleiben, damit Outlook das Konto waehlt")
    anhaenge = [t.get_filename() for t in nachricht.iter_attachments()]
    pruefe(anhaenge == ["260903 BOB-VE100-MKA 03.docx"],
           f"Anhaenge: {anhaenge}")
    anhang = next(iter(nachricht.iter_attachments()))
    pruefe(anhang.get_payload(decode=True)[:2] == b"PK",
           "der Anhang ist keine gueltige Word-Datei")
    pruefe("wordprocessingml" in anhang.get_content_type(),
           f"falscher Anhangtyp: {anhang.get_content_type()}")

    # Zwei Fassungen desselben Textes: Der HTML-Teil legt Arial 10 fest, der
    # Textteil bleibt als Rueckfallebene daneben. Ohne HTML nimmt Outlook die
    # Schrift aus den Einstellungen des jeweiligen Absenders.
    typen = [teil.get_content_type() for teil in nachricht.walk()]
    pruefe("text/plain" in typen and "text/html" in typen,
           f"Mail braucht Text- und HTML-Teil: {typen}")
    pruefe("multipart/alternative" in typen,
           f"die beiden Fassungen muessen Alternativen sein: {typen}")
    html_teil = nachricht.get_body(preferencelist=("html",))
    pruefe(html_teil is not None, "der HTML-Teil fehlt")
    html_text = html_teil.get_content()
    pruefe("Arial" in html_text and "10pt" in html_text,
           f"Arial 10 fehlt im HTML-Teil: {html_text[:160]}")
    pruefe("in der Anlage senden wir unsere Schreiben vom 03.09.26."
           in html_text, f"Mailtext im HTML-Teil: {html_text[:250]}")
    koerpertext = nachricht.get_body(preferencelist=("plain",)).get_content()
    pruefe("in der Anlage senden wir unsere Schreiben vom 03.09.26"
           in koerpertext, f"Mailtext: {koerpertext!r}")

    ohne_anhang = c.post("/api/anzeigen/mail/entwurf",
                         json={"antwort": anfrage, "dokument_anhaengen": False})
    pruefe(not list(email.message_from_bytes(
        ohne_anhang.content, policy=policy.default).iter_attachments()),
        "ohne dokument_anhaengen darf kein Anhang drin sein")

    ohne_adresse = c.post("/api/anzeigen/mail/entwurf", json={
        "antwort": {**anfrage,
                    "empfaenger": {**anfrage["empfaenger"], "email": ""}}})
    pruefe(ohne_adresse.status_code == 422,
           f"ohne Empfaengeradresse muesste 422 sein: "
           f"{ohne_adresse.status_code}")

# ─────────────────────────────────────────────────────────────────────────────
# 5. Die echten Referenzschreiben, falls vorhanden
# ─────────────────────────────────────────────────────────────────────────────

abschnitt("Referenz-PDFs (nur auf Bens Rechner)")

REFERENZ = Path.home() / "Desktop" / "Mehrkostenanzeige"
ERWARTET = {
    "BOB-260722-MKA 01-EG Deckenöffnung.pdf": {
        "art": "Mehrkostenanzeige", "nummer": "01", "datum": date(2026, 7, 22),
        "firma": "SERVISA Rückbau & Sanierung GmbH", "punkte": 6,
    },
    "BOB-260812-MKA 02-GK-Ständerwände.pdf": {
        "art": "Mehrkostenanzeige", "nummer": "02", "datum": date(2026, 8, 12),
        "firma": "SERVISA Rückbau & Sanierung GmbH", "punkte": 0,
    },
    "BOB-260817-MKA 03-abweichende Bodenaufbauten.pdf": {
        "art": "Mehrkostenanzeige", "nummer": "03", "datum": date(2026, 8, 17),
        "firma": "SERVISA Rückbau & Sanierung GmbH", "punkte": 0,
    },
}

if not REFERENZ.is_dir():
    print("   uebersprungen — Ordner nicht vorhanden:", REFERENZ)
else:
    for name, soll in ERWARTET.items():
        pfad = REFERENZ / name
        if not pfad.is_file():
            print("   uebersprungen — fehlt:", name)
            continue
        ist = les.lies(pfad)
        pruefe(ist.art == soll["art"], f"{name}: Art {ist.art!r}")
        pruefe(ist.nummer == soll["nummer"], f"{name}: Nummer {ist.nummer!r}")
        pruefe(ist.datum == soll["datum"], f"{name}: Datum {ist.datum}")
        pruefe(ist.absender.firma == soll["firma"],
               f"{name}: Firma {ist.absender.firma!r}")
        pruefe(ist.absender.strasse == "Attilastraße 24",
               f"{name}: Strasse {ist.absender.strasse!r}")
        pruefe(ist.ansprechpartner == "Qasem Ashgarzada",
               f"{name}: Ansprechpartner {ist.ansprechpartner!r}")
        pruefe(ist.ansprechpartner_email ==
               "Qasem.Ashgarzada@servisa-gruppe.de",
               f"{name}: Mail {ist.ansprechpartner_email!r}")
        pruefe(ist.betreff.startswith(f"Mehrkostenanzeige Nr. {soll['nummer']}"),
               f"{name}: Betreff {ist.betreff!r}")
        pruefe(ist.gewerk.startswith("Abbruch-"), f"{name}: Gewerk {ist.gewerk!r}")
        pruefe(ist.leistungsort.startswith("Boulevard Berlin"),
               f"{name}: Leistungsort {ist.leistungsort!r}")
        pruefe(len(ist.punkte) == soll["punkte"],
               f"{name}: Punkte {len(ist.punkte)}, erwartet {soll['punkte']}")
        pruefe(ist.kennung == f"MKA {soll['nummer']}",
               f"{name}: Kennung {ist.kennung!r} — im Dateinamen des Bueros "
               f"steht 'MKA {soll['nummer']}'")
        pruefe(not ist.hinweise, f"{name}: Hinweise {ist.hinweise}")
        # Der Bauherr darf nie im Adressfeld der Antwort landen.
        pruefe("BoB" not in ist.absender.firma,
               f"{name}: Bauherr als Absender erkannt — {ist.absender.firma!r}")

print(f"\n{ok} Pruefungen ok, {len(fehler)} Fehler")
for f in fehler:
    print("  FEHLER:", f)
sys.exit(1 if fehler else 0)
