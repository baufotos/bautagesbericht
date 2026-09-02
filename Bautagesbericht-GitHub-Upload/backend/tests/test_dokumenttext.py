"""Steuerzeichen duerfen kein Dokument aufhalten.

WARUM DIESE REIHE
=================
Word-Dokumente sind XML, und XML 1.0 verbietet die meisten Steuerzeichen. Ein
Seitenvorschub aus einem kopierten PDF oder ein vertikaler Tabulator aus einer
Texterkennung reichte, und das Erzeugen brach ab:

    XMLSyntaxError: PCDATA invalid Char value 12
    ValueError: All strings must be XML compatible

Der Anwender sah davon nur "fehlgeschlagen" — ohne Grund, und mit
hochgeladenen Blaettern, an denen nichts falsch war.

Betroffen waren alle fuenf Erzeuger, weil sie dieselben drei Textquellen
haben: Texterkennung, PDF-Textebene und Eingaben aus der Oberflaeche. Geprueft
wird hier deshalb beides: die Aufbereitung selbst und dass jeder Erzeuger sie
wirklich benutzt.
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ARBEIT = Path(tempfile.gettempdir()) / "hpp-test-dokumenttext"
ARBEIT.mkdir(exist_ok=True)
os.environ.setdefault("BTB_OUTPUT_DIR", str(ARBEIT / "output"))
os.environ.setdefault("BTB_UPLOAD_DIR", str(ARBEIT / "uploads"))
os.environ.setdefault("BTB_DATABASE_URL", f"sqlite:///{(ARBEIT / 'x.db').as_posix()}")

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from docx import Document  # noqa: E402

from app.services import dokumenttext  # noqa: E402

ok = 0
fehler: list[str] = []


def pruefe(bedingung, text):
    global ok
    if bedingung:
        ok += 1
    else:
        fehler.append(text)


def gleich(ist, soll, text):
    pruefe(ist == soll, f"{text}: erwartet {soll!r}, war {ist!r}")


def abschnitt(text):
    print(f"─── {text} ───")


abschnitt("Was ersetzt wird und was bleibt")

# Zeichen, die einen Umbruch meinten, werden einer. Der vertikale Tabulator
# IST in Words eigenem Textmodell der Zeilenumbruch, der Seitenvorschub
# trennte Zeilen.
gleich(dokumenttext.xml_sicher("Wand\x0bgestellt"), "Wand\ngestellt",
       "vertikaler Tabulator wird ein Umbruch")
gleich(dokumenttext.xml_sicher("Seite\x0cZwei"), "Seite\nZwei",
       "Seitenvorschub wird ein Umbruch")
gleich(dokumenttext.xml_sicher("a b"), "a\nb",
       "Unicode-Zeilentrenner wird ein Umbruch")

# Zeichen ohne Bedeutung verschwinden, ohne den Text zu zerreissen.
gleich(dokumenttext.xml_sicher("a\x01b"), "ab", "Steuerzeichen fliegt raus")
gleich(dokumenttext.xml_sicher("a\x00b"), "ab", "Nullbyte fliegt raus")
gleich(dokumenttext.xml_sicher("Fu​sse"), "Fusse",
       "Zeichen ohne Breite fliegt raus")

# Zeilenenden vereinheitlichen — sonst haengt an jeder Zeile ein
# Wagenruecklauf, der spaeter in Ortsangaben wieder auftaucht.
gleich(dokumenttext.xml_sicher("a\r\nb\rc"), "a\nb\nc", "CRLF und CR")

# Und das, was bleiben MUSS. Das weiche Trennzeichen steht in der HPP-Vorlage
# als Zierstrich neben der Seitenzahl; wuerde es hier weggeworfen, aenderte
# sich die Fusszeile jedes Berichts.
for text, name in [
    ("Ried­el", "weiches Trennzeichen (Zierstrich der Vorlage)"),
    ("4.OG – Ost", "Gedankenstrich"),
    ("„Riedel Bau“", "typografische Anfuehrungszeichen"),
    ("Spalte\tWert", "Tabulator"),
    ("3 °C", "Gradzeichen"),
    ("Grünanlagen GmbH & Co. KG", "Umlaute und Kaufmanns-Und"),
]:
    gleich(dokumenttext.xml_sicher(text), text, f"{name} bleibt")

gleich(dokumenttext.xml_sicher(""), "", "leerer Text")
gleich(dokumenttext.xml_sicher(None), "", "None ergibt leeren Text")


abschnitt("Einzeilige Werte")

# Fuer Werte, die in eine Tabellenzelle mit knapper Hoehe gehen: Ein
# Firmenname darf nicht zwei Zeilen hoch werden, bloss weil in der
# Texterkennung ein Umbruch steckte.
gleich(dokumenttext.einzeilig("Riedel\nBau   GmbH"), "Riedel Bau GmbH",
       "Umbruch und doppelte Leerzeichen werden eines")
gleich(dokumenttext.einzeilig("  Kita Nord \x0c "), "Kita Nord",
       "Rand wird abgeschnitten")


abschnitt("Zeilen fuer echte Umbrueche")

gleich(dokumenttext.zeilen("a\nb"), ["a", "b"], "zwei Zeilen")
gleich(dokumenttext.zeilen("\n a \n b \n\n"), [" a ", " b "],
       "leere Zeilen am Rand fallen weg, innen bleibt der Text")
gleich(dokumenttext.zeilen(""), [], "leerer Text ergibt keine Zeile")


abschnitt("Jeder Erzeuger benutzt die Aufbereitung wirklich")

# Ohne diese Pruefung waere die Aufbereitung ein Modul, das niemand aufruft.
# Angesteuert wird je Erzeuger die zentrale Stelle, durch die sein Text laeuft.
STOERTEXT = "Wand\x0bgestellt\x0cund\x01geprueft"
#: Die beiden Umbruchzeichen werden Umbrueche, das bedeutungslose \x01 fliegt
#: weg, ohne den Text zu zerreissen.
SAUBER = "Wand\ngestellt\nundgeprueft"


def leerer_absatz():
    dokument = Document()
    return dokument.add_paragraph()


from app.services import (  # noqa: E402
    besprechungsprotokoll_generation as bp,
    docx_generation as btb,
    maengelanzeige_generation as ma,
    maengelliste_generation as ml,
    projektbericht_generation as pb,
)

erzeuger = [
    ("Maengelliste", lambda: ml._schreibe(leerer_absatz(), STOERTEXT)),
    ("Projektbericht", lambda: pb._lauf(leerer_absatz(), STOERTEXT)),
    ("Besprechungsprotokoll", lambda: bp._schreibe(leerer_absatz(), STOERTEXT)),
]

for name, ruf in erzeuger:
    try:
        lauf = ruf()
        pruefe("\x0b" not in lauf.text and "\x01" not in lauf.text,
               f"{name}: Steuerzeichen ist weg ({lauf.text!r})")
        pruefe("Wand" in lauf.text and "geprueft" in lauf.text,
               f"{name}: der Text ist noch da ({lauf.text!r})")
    except Exception as exc:
        fehler.append(f"{name}: bricht an einem Steuerzeichen ab ({exc})")

# Die Maengelanzeige baut Absatz und Lauf in einem Schritt.
try:
    dokument = Document()
    absatz = ma._absatz(dokument, STOERTEXT)
    pruefe("\x0b" not in absatz.text and "\x01" not in absatz.text,
           f"Maengelanzeige: Steuerzeichen ist weg ({absatz.text!r})")
    pruefe("Wand" in absatz.text,
           f"Maengelanzeige: der Text ist noch da ({absatz.text!r})")
except Exception as exc:
    fehler.append(f"Maengelanzeige: bricht an einem Steuerzeichen ab ({exc})")

# Der Bautagesbericht baut sein XML selbst und setzt echte Umbrueche.
gleich(dokumenttext.xml_sicher(STOERTEXT), SAUBER,
       "Umbruchzeichen werden Umbrueche, der Rest fliegt weg")

xml = btb._run(STOERTEXT)
pruefe("\x0b" not in xml and "\x01" not in xml,
       "Bautagesbericht: kein Steuerzeichen im XML")
gleich(xml.count("<w:br/>"), SAUBER.count("\n"),
       "Bautagesbericht: je Umbruch ein <w:br/>")


print()
if fehler:
    print(f"{ok} Pruefungen ok, {len(fehler)} Fehler:")
    for f in fehler:
        print("  -", f)
    raise SystemExit(1)
print(f"{ok} Pruefungen ok, 0 Fehler")
