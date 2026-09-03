"""Rauchtest der Mängel-API gegen eine frische SQLite-Datenbank."""
import io
import os
import shutil
import sys
import tempfile
from pathlib import Path

# Eigene Ablage im Temp-Ordner — die echte storage/ bleibt unberuehrt.
STORAGE = Path(tempfile.gettempdir()) / "hpp-apitest"
if STORAGE.exists():
    shutil.rmtree(STORAGE)
STORAGE.mkdir(parents=True)

os.environ["BTB_DATABASE_URL"] = f"sqlite:///{STORAGE / 'test.db'}"
os.environ["BTB_UPLOAD_DIR"] = str(STORAGE / "uploads")
os.environ["BTB_OUTPUT_DIR"] = str(STORAGE / "output")

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from fastapi.testclient import TestClient  # noqa: E402
from PIL import Image  # noqa: E402

from app.main import app  # noqa: E402

# Fristen relativ zu heute, nicht als festes Datum: Ein Mangel mit fester
# Frist wird irgendwann von selbst ueberfaellig, und dann scheitert der Test
# am Kalender statt an einem echten Fehler. Genau das ist passiert.
from datetime import date as _date, timedelta as _timedelta  # noqa: E402

VOR_ZWEI_WOCHEN = (_date.today() - _timedelta(days=14)).isoformat()
IN_ZWEI_WOCHEN = (_date.today() + _timedelta(days=14)).isoformat()

ok = 0
fehler = []


def pruefe(bedingung, text):
    global ok
    if bedingung:
        ok += 1
    else:
        fehler.append(text)


def foto_bytes(groesse=(1600, 1200), farbe=(200, 100, 50)) -> bytes:
    puffer = io.BytesIO()
    Image.new("RGB", groesse, farbe).save(puffer, format="JPEG")
    return puffer.getvalue()


with TestClient(app) as c:
    # ── Stammdaten wurden geseedet ──
    sd = c.get("/api/mangel-stammdaten").json()
    pruefe([t["bezeichnung"] for t in sd["typen"]] ==
           ["Mangel", "Hinweis", "Gefahr", "Frage", "Sonstiges"], f"Typen: {sd['typen']}")
    pruefe(len(sd["status"]) == 6, f"Status: {len(sd['status'])}")
    pruefe(any(s["ist_abgeschlossen"] for s in sd["status"]), "kein Abschluss-Status")

    # ── Bearbeiter ──
    b = c.post("/api/mangel-stammdaten/bearbeiter",
               json={"name": "B. Gagelmann", "email": ""}).json()
    pruefe(b["email"] is None, f"Bearbeiter-Mail: {b}")

    # ── Projekt + Gewerke ──
    p = c.post("/api/projekte", json={"name": "Testprojekt Nord", "adresse": ""}).json()
    pid = p["id"]
    g1 = c.post("/api/gewerke", json={
        "projekt_id": pid,
        "firma_name": "Rolfes Bau GmbH",
        "vergabeeinheit_code": "VE300-01",
        "vergabeeinheit_bezeichnung": "Erweiterter Rohbau",
        "email": "bau@example.com",
    }).json()
    pruefe(g1["anzeige_name"] == "Rolfes Bau GmbH | VE300-01 Erweiterter Rohbau",
           f"Anzeigename: {g1['anzeige_name']}")

    g2 = c.post("/api/gewerke", json={
        "projekt_id": pid, "firma_name": "Ohne Mail GmbH", "email": "",
    }).json()
    pruefe(g2["email"] is None, f"g2 mail: {g2['email']}")

    # ── Mangel anlegen: Nummernkreis ──
    m1 = c.post("/api/maengel", json={
        "projekt_id": pid, "kurzbezeichnung": "Stahlbeton", "typ": "Hinweis",
        "gewerk_id": g1["id"], "hinweis_ort": "EG", "raumnummer": "E.014",
        "prioritaet": "mittel", "beschreibung": "Lunker in Sichtbeton.",
        "erstellt_am": VOR_ZWEI_WOCHEN, "erste_frist_bis": IN_ZWEI_WOCHEN,
        "interne_bemerkung": "intern: Bauleiter informiert",
        "zustaendiger_user_id": b["id"],
    })
    pruefe(m1.status_code == 201, f"create: {m1.status_code} {m1.text[:300]}")
    m1 = m1.json()
    pruefe(m1["nummer"] == "00001", f"Nummer: {m1['nummer']}")
    pruefe(m1["mail_fehler"] is None, f"mail_fehler: {m1['mail_fehler']}")
    pruefe(m1["aktuelle_frist"] == IN_ZWEI_WOCHEN, f"frist: {m1['aktuelle_frist']}")

    m2 = c.post("/api/maengel", json={
        "projekt_id": pid, "kurzbezeichnung": "Fehlende Brandschottung",
        "gewerk_id": g2["id"], "mail_autosend": True,
        "mail_versendemodus": "automatisch",
        "erste_frist_bis": "2020-01-01", "prioritaet": "hoch",
    }).json()
    pruefe(m2["nummer"] == "00002", f"Nummer2: {m2['nummer']}")
    pruefe(m2["mail_autosend"] is False, "Autosend haette blockiert werden muessen")
    pruefe(m2["mail_versendemodus"] == "manuell", f"Modus: {m2['mail_versendemodus']}")
    pruefe("keine Email-Adresse" in (m2["mail_fehler"] or ""),
           f"mail_fehler2: {m2['mail_fehler']}")
    pruefe(m2["ist_ueberfaellig"] is True, "m2 muesste ueberfaellig sein")

    # ── Fotos hochladen ──
    up = c.post(f"/api/maengel/{m1['id']}/fotos", files=[
        ("dateien", ("IMG_1.jpg", foto_bytes(), "image/jpeg")),
        ("dateien", ("IMG_2.jpg", foto_bytes((900, 1600), (60, 120, 200)), "image/jpeg")),
    ], data={"bildunterschrift": "Nordwand"})
    pruefe(up.status_code == 201, f"foto upload: {up.status_code} {up.text[:300]}")
    fotos = up.json()
    pruefe(len(fotos) == 2 and fotos[0]["reihenfolge"] == 1, f"fotos: {fotos}")

    bild = c.get(f"/api/maengel/fotos/{fotos[0]['id']}/bild")
    thumb = c.get(f"/api/maengel/fotos/{fotos[0]['id']}/bild?thumb=true")
    pruefe(bild.status_code == 200 and bild.headers["content-type"].startswith("image"),
           f"bild: {bild.status_code} {bild.headers.get('content-type')}")
    pruefe(0 < len(thumb.content) < len(bild.content),
           f"thumb {len(thumb.content)} vs bild {len(bild.content)}")

    # ── Anhang ──
    anh = c.post(f"/api/maengel/{m1['id']}/dateien", files=[
        ("dateien", ("Protokoll.txt", b"Pruefprotokoll", "text/plain")),
    ])
    pruefe(anh.status_code == 201, f"anhang: {anh.status_code} {anh.text[:200]}")

    # ── Plan hochladen + Markierung ──
    plan_bytes = io.BytesIO()
    Image.new("RGB", (2000, 1400), (245, 245, 240)).save(plan_bytes, format="PNG")
    plan = c.post("/api/plaene", data={"projekt_id": pid},
                  files={"datei": ("Grundriss_EG.png", plan_bytes.getvalue(), "image/png")})
    pruefe(plan.status_code == 201, f"plan: {plan.status_code} {plan.text[:300]}")
    plan = plan.json()
    pruefe(plan["seiten"] == 1, f"seiten: {plan}")

    vor = c.get(f"/api/plaene/{plan['id']}/vorschau")
    pruefe(vor.status_code == 200, f"vorschau: {vor.status_code} {vor.text[:200]}")

    mark = c.put(f"/api/maengel/{m1['id']}/markierung", json={
        "plan_datei_id": plan["id"], "x_prozent": 42.5, "y_prozent": 68.0, "seite": 1,
    })
    pruefe(mark.status_code == 200, f"markierung: {mark.status_code} {mark.text[:300]}")
    pruefe(mark.json()["plan_dateiname"] == "Grundriss_EG.png", f"{mark.json()}")

    schlecht = c.put(f"/api/maengel/{m1['id']}/markierung", json={
        "plan_datei_id": plan["id"], "x_prozent": 120, "y_prozent": 5, "seite": 1,
    })
    pruefe(schlecht.status_code == 422, f"ungueltige Position: {schlecht.status_code}")

    seite_falsch = c.put(f"/api/maengel/{m1['id']}/markierung", json={
        "plan_datei_id": plan["id"], "x_prozent": 10, "y_prozent": 5, "seite": 4,
    })
    pruefe(seite_falsch.status_code == 400, f"falsche Seite: {seite_falsch.status_code}")

    # ── Duplikat NU ──
    dup = c.post(f"/api/maengel/{m1['id']}/duplizieren")
    pruefe(dup.status_code == 201, f"dup: {dup.status_code} {dup.text[:300]}")
    dup = dup.json()
    pruefe(dup["nummer"] == "00001.1", f"dup nummer: {dup['nummer']}")
    pruefe(dup["eltern_nummer"] == "00001" and dup["eltern_kurzbezeichnung"] == "Stahlbeton",
           f"Ist Kopie von: {dup['eltern_nummer']} {dup['eltern_kurzbezeichnung']}")
    pruefe(len(dup["fotos"]) == 2, f"dup fotos: {len(dup['fotos'])}")
    pruefe(dup["markierung"] is not None, "dup markierung fehlt")
    pruefe(dup["mail_autosend"] is False, "dup autosend")

    dup2 = c.post(f"/api/maengel/{m1['id']}/duplizieren").json()
    pruefe(dup2["nummer"] == "00001.2", f"dup2 nummer: {dup2['nummer']}")
    dup3 = c.post(f"/api/maengel/{dup2['id']}/duplizieren").json()
    pruefe(dup3["nummer"] == "00001.3", f"dup3 (Dup vom Dup): {dup3['nummer']}")

    # ── Update: Nachfrist setzt Datum automatisch, Status wechselt ──
    upd = c.patch(f"/api/maengel/{m1['id']}", json={
        "erste_nachfrist_bis": "2026-09-10", "status": "Nachfrist",
    })
    pruefe(upd.status_code == 200, f"update: {upd.status_code} {upd.text[:300]}")
    upd = upd.json()
    pruefe(upd["erste_nachfrist_gesetzt_am"] is not None,
           "Nachfrist-gesetzt-am nicht automatisch gefuellt")
    pruefe(upd["aktuelle_frist"] == "2026-09-10", f"Nachfrist zaehlt: {upd['aktuelle_frist']}")

    leeren = c.patch(f"/api/maengel/{m1['id']}", json={"raumnummer": None}).json()
    pruefe(leeren["raumnummer"] is None, f"raumnummer leeren: {leeren['raumnummer']}")

    # ── Filter ──
    alle = c.get(f"/api/maengel?projekt_id={pid}").json()
    pruefe(len(alle) == 5, f"alle: {len(alle)} -> {[m['nummer'] for m in alle]}")
    pruefe([m["nummer"] for m in alle] ==
           ["00001", "00001.1", "00001.2", "00001.3", "00002"],
           f"Sortierung: {[m['nummer'] for m in alle]}")
    ueber = c.get(f"/api/maengel?projekt_id={pid}&ueberfaellig=true").json()
    pruefe([m["nummer"] for m in ueber] == ["00002"], f"ueberfaellig: {[m['nummer'] for m in ueber]}")
    nicht = c.get(f"/api/maengel?projekt_id={pid}&ueberfaellig=false").json()
    pruefe(len(nicht) == 4, f"nicht ueberfaellig: {len(nicht)}")
    firma = c.get(f"/api/maengel?projekt_id={pid}&gewerk_id={g2['id']}").json()
    pruefe(len(firma) == 1, f"firma-filter: {len(firma)}")
    suche = c.get(f"/api/maengel?projekt_id={pid}&suche=brandschott").json()
    pruefe(len(suche) == 1, f"suche: {len(suche)}")
    prio = c.get(f"/api/maengel?projekt_id={pid}&prioritaet=hoch").json()
    pruefe(len(prio) == 1, f"prio: {len(prio)}")
    with_thumb = [m for m in alle if m["titel_foto_id"]]
    pruefe(len(with_thumb) == 4, f"titelfotos: {len(with_thumb)}")

    # ── Erledigt -> nicht mehr ueberfaellig ──
    c.patch(f"/api/maengel/{m2['id']}", json={"status": "erledigt",
                                             "erledigt_am": "2026-08-19"})
    ueber2 = c.get(f"/api/maengel?projekt_id={pid}&ueberfaellig=true").json()
    pruefe(ueber2 == [], f"nach Erledigung noch ueberfaellig: {ueber2}")
    offen = c.get(f"/api/maengel?projekt_id={pid}&abgeschlossen=false").json()
    pruefe(len(offen) == 4, f"offene: {len(offen)}")

    # ── Senden ohne Mailadresse ──
    senden_fehler = c.post(f"/api/maengel/{m2['id']}/senden").json()
    pruefe(senden_fehler["versendet"] is False and
           "keine Email-Adresse" in senden_fehler["nachricht"],
           f"senden ohne Mail: {senden_fehler}")
    senden_ohne_kanal = c.post(f"/api/maengel/{m1['id']}/senden").json()
    pruefe(senden_ohne_kanal["versendet"] is False and
           senden_ohne_kanal["kanal"] == "keiner",
           f"senden ohne Webhook: {senden_ohne_kanal}")

    # ── Export ──
    exp = c.get(f"/api/maengel/export?projekt_id={pid}")
    pruefe(exp.status_code == 200 and len(exp.content) > 10000,
           f"export: {exp.status_code} {len(exp.content)}")
    pruefe("wordprocessingml" in exp.headers["content-type"], exp.headers["content-type"])

    exp_intern = c.get(f"/api/maengel/export?projekt_id={pid}&intern=true&status=Nachfrist")
    pruefe(exp_intern.status_code == 200, f"export intern: {exp_intern.status_code}")

    ziel = STORAGE / "export_firma.docx"
    ziel.write_bytes(exp.content)
    ziel_intern = STORAGE / "export_intern.docx"
    ziel_intern.write_bytes(exp_intern.content)

    import docx  # noqa: E402

    def volltext(pfad):
        d = docx.Document(str(pfad))
        teile = [p.text for p in d.paragraphs]
        teile += [z.text for t in d.tables for r in t.rows for z in r.cells]
        return "\n".join(teile)

    firmen_text = volltext(ziel)
    intern_text = volltext(ziel_intern)
    pruefe("intern: Bauleiter" not in firmen_text,
           "INTERNE BEMERKUNG IM FIRMEN-EXPORT!")
    pruefe("intern: Bauleiter" in intern_text,
           "interne Bemerkung fehlt im internen Export")
    pruefe("Grundriss_EG.png" in intern_text, "Plan-Markierung fehlt im Export")
    pruefe("{{" not in firmen_text, f"Platzhalter uebrig: {firmen_text[:200]}")

    # ── Loeschen ──
    del_foto = c.delete(f"/api/maengel/fotos/{fotos[1]['id']}")
    pruefe(del_foto.status_code == 204, f"foto loeschen: {del_foto.status_code}")

    del_gewerk = c.delete(f"/api/gewerke/{g1['id']}")
    pruefe(del_gewerk.status_code == 409, f"gewerk konflikt: {del_gewerk.status_code}")
    pruefe(del_gewerk.json()["detail"]["anzahl_maengel"] >= 1, del_gewerk.text)

    # Ein Kapitel der Besprechungsprotokolle, das auf dieselbe Firma zeigt.
    # Genau daran scheiterte das Loeschen: "FOREIGN KEY constraint failed",
    # fuer den Anwender ein Serverfehler ohne Erklaerung — die Firma blieb in
    # den Stammdaten stehen. Ein Kapitel soll das Gewerk ausdruecklich
    # ueberleben (siehe models.BesprechungsKapitel).
    kapitel = c.post("/api/besprechungsprotokolle/kapitel", json={
        "projekt_id": pid, "nummer": "2.", "titel": "VE01 Rohbau",
        "sortierung": 1, "gewerk_id": g1["id"],
    })
    pruefe(kapitel.status_code in (200, 201),
           f"Kapitel mit Gewerk angelegt: {kapitel.status_code} {kapitel.text[:120]}")

    weg = c.delete(f"/api/gewerke/{g1['id']}?force=true")
    pruefe(weg.status_code == 204,
           f"gewerk force loeschen: {weg.status_code} {weg.text[:200]}")
    pruefe(c.get(f"/api/maengel/{m1['id']}").json()["gewerk_id"] is None,
           "Mangel muesste Firma verloren haben, aber erhalten bleiben")

    if kapitel.status_code in (200, 201):
        kid = kapitel.json()["id"]
        uebrig = [k for k in c.get(f"/api/besprechungsprotokolle/kapitel?projekt_id={pid}").json()
                  if k["id"] == kid]
        pruefe(len(uebrig) == 1,
               "das Kapitel ueberlebt das Loeschen der Firma")
        pruefe(uebrig and uebrig[0]["gewerk_id"] is None,
               f"…und hat nur den Verweis verloren: {uebrig[0]['gewerk_id'] if uebrig else '?'}")

    del_plan = c.delete(f"/api/plaene/{plan['id']}")
    pruefe(del_plan.status_code == 409, f"plan konflikt: {del_plan.status_code}")
    pruefe(c.delete(f"/api/plaene/{plan['id']}?force=true").status_code == 204,
           "plan force loeschen")

    del_mangel = c.delete(f"/api/maengel/{dup3['id']}")
    pruefe(del_mangel.status_code == 204, f"mangel loeschen: {del_mangel.status_code}")
    pruefe(c.get(f"/api/maengel/{dup3['id']}").status_code == 404, "mangel noch da")

    # Original loeschen -> Duplikate bleiben, verlieren nur den Verweis
    pruefe(c.delete(f"/api/maengel/{m1['id']}").status_code == 204, "original loeschen")
    rest = c.get(f"/api/maengel/{dup['id']}").json()
    pruefe(rest["eltern_mangel_id"] is None and rest["eltern_nummer"] == "",
           f"Duplikat nach Loeschen des Originals: {rest['eltern_mangel_id']}")

    # Projekt loeschen: 409, dann force
    konflikt = c.delete(f"/api/projekte/{pid}")
    pruefe(konflikt.status_code == 409, f"projekt konflikt: {konflikt.status_code}")
    pruefe(konflikt.json()["detail"]["anzahl_maengel"] == 3, konflikt.text)
    pruefe(c.delete(f"/api/projekte/{pid}?force=true").status_code == 204,
           "projekt force loeschen")
    pruefe(c.get(f"/api/maengel?projekt_id={pid}").json() == [], "Maengel noch da")
    pruefe(c.get(f"/api/gewerke?projekt_id={pid}").json() == [], "Gewerke noch da")
    pruefe(c.get(f"/api/plaene?projekt_id={pid}").json() == [], "Plaene noch da")

print(f"\n{ok} Pruefungen ok, {len(fehler)} Fehler")
for f in fehler:
    print("  FEHLER:", f)
sys.exit(1 if fehler else 0)
