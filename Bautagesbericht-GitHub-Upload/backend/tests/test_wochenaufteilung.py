"""Prueft die Zerlegung eines Wochenpakets in Tagesberichte.

Der Kern ist die Datumserkennung: Ein falsch erkannter Tag schreibt die Arbeit
eines Tages in den Bericht eines anderen, und das faellt beim Durchsehen kaum
auf. Deshalb wird hier vor allem geprueft, wann NICHT geraten wird.
"""
from __future__ import annotations

import sys
import tempfile
from datetime import date, timedelta
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BACKEND = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND))

from app.services.wochenaufteilung import (  # noqa: E402
    Seitenfund,
    tagesabschnitte,
    daten_in_text,
    datum_der_seite,
    finde_seitendaten,
    gruppiere_nach_tag,
    schreibe_teil_pdf,
    woche_um,
    wochentag,
)

ok = 0
fehler: list[str] = []


def pruefe(bedingung, text):
    global ok
    if bedingung:
        ok += 1
    else:
        fehler.append(text)


def gleich(ist, soll, text):
    pruefe(ist == soll, f"{text}: erwartet {soll!r}, war {ist!r}")


print("─── Datumsschreibweisen ───")

gleich(daten_in_text("Datum: 03.08.2026"), [date(2026, 8, 3)], "TT.MM.JJJJ")
gleich(daten_in_text("Datum: 3.8.2026"), [date(2026, 8, 3)], "einstellig")
gleich(daten_in_text("Datum: 03.08.26"), [date(2026, 8, 3)], "zweistelliges Jahr")
gleich(daten_in_text("Stand 2026-08-03"), [date(2026, 8, 3)], "ISO")
gleich(daten_in_text("03-08-2026"), [date(2026, 8, 3)], "Bindestriche")
gleich(daten_in_text("03/08/2026"), [date(2026, 8, 3)], "Schraegstriche")

# Unsinn darf kein Datum ergeben.
gleich(daten_in_text("32.13.2026"), [], "unmoeglicher Tag")
gleich(daten_in_text("Pos. 1.2.3"), [], "Gliederungsnummer ist kein Datum")
gleich(daten_in_text("Baujahr 1975, Halle 4"), [], "Jahreszahl allein")
gleich(daten_in_text("01.01.1899"), [], "Jahr vor 2000")

# Reihenfolge und Entdopplung
gleich(daten_in_text("03.08.2026 ... 04.08.2026 ... 03.08.2026"),
       [date(2026, 8, 3), date(2026, 8, 4)], "Reihenfolge, ohne Doppel")


print("─── Datumsformen aus der Texterkennung ───")

# Alles echte Ausgaben der Windows-Texterkennung an einem gescannten
# Firmen-Formblatt: mal fehlt ein Punkt, mal alle, mal wird er zum Leerzeichen.
gleich(daten_in_text("Sa. 0604.2024"), [date(2024, 4, 6)],
       "TTMM.JJJJ (ein Punkt verloren)")
gleich(daten_in_text("Mi. 03042024"), [date(2024, 4, 3)], "TTMMJJJJ (alle Punkte weg)")
gleich(daten_in_text("Di. 02 04.2024"), [date(2024, 4, 2)], "Punkt als Leerzeichen")
gleich(daten_in_text("erstellt 31.012020 CZ"), [date(2020, 1, 31)], "TT.MMJJJJ")
gleich(daten_in_text("ersten: 31 01 2020 CZ"), [date(2020, 1, 31)], "zwei Leerzeichen")
gleich(daten_in_text("06,04.2024"), [date(2024, 4, 6)], "Punkt als Komma")

# Achtstellige Zahlen ohne plausibles Jahr sind kein Datum.
gleich(daten_in_text("Auftrag 12345678"), [], "Auftragsnummer ist kein Datum")
gleich(daten_in_text("Tel 040 12345678"), [], "Telefonnummer auch nicht")

# Abgekuerzte Wochentage zaehlen als Datumswort.
gleich(datum_der_seite("Datum:\nSa. 0604.2024\nFa. X"), date(2024, 4, 6),
       "Wochentagskuerzel erkannt")

print("─── Verlesene Jahreszahl ueber den Zeitraum retten ───")

april = {date(2024, 4, 2) + timedelta(days=i) for i in range(5)}

# Echter Fall: die Erkennung las "9024" statt "2024".
gleich(datum_der_seite("Datum:\nDo. 04.04 9024\nFa. X (2 Mann)", april),
       date(2024, 4, 4), "Tag und Monat retten den Tag")

# Ohne Zeitraum wird NICHT geraten.
gleich(datum_der_seite("Datum:\nDo. 04.04 9024\nFa. X (2 Mann)"), None,
       "ohne Zeitraum kein Rateversuch")

# Ohne Datumswort in der Zeile ebenfalls nicht.
gleich(datum_der_seite("Achse 04.04 irgendwas", april), None,
       "ohne Datumswort kein Rateversuch")

# Kaeme mehr als ein Tag in Frage, bleibt es beim Nichtstun.
zwei_monate = {date(2024, 4, 4), date(2024, 5, 4)}
gleich(datum_der_seite("Datum:\nDo. 04.04 9024", zwei_monate), date(2024, 4, 4),
       "eindeutiger Tag im Zeitraum wird genommen")


print("─── Welches Datum ist das Berichtsdatum ───")

seite_einfach = """Bautagesbericht
Datum: 05.08.2026
Fa. Meyer Bau (4 Mann)
- Schalung OG2
"""
gleich(datum_der_seite(seite_einfach), date(2026, 8, 5), "Datumszeile im Kopf")

# Der Klassiker: Vertragsfrist steht weiter unten und darf nicht gewinnen.
seite_mit_frist = """Bautagesbericht
Datum: 05.08.2026
Fa. Meyer Bau (4 Mann)
- Fertigstellung laut Vertrag bis 30.11.2026
- Nachtrag vom 12.01.2026 beauftragt
"""
gleich(datum_der_seite(seite_mit_frist), date(2026, 8, 5),
       "Kopfdatum schlaegt Frist im Text")

# Wochentag statt des Wortes "Datum"
seite_wochentag = """Tagesbericht Nr. 42
Mittwoch, 05.08.2026
Fa. Schulz (2 Mann)
"""
gleich(datum_der_seite(seite_wochentag), date(2026, 8, 5), "Wochentagszeile")

# Kein Datumswort, aber ein Datum -> zweitbeste Regel greift
gleich(datum_der_seite("Irgendein Kopf\n05.08.2026\nFa. X (1 Mann)"),
       date(2026, 8, 5), "erstes Datum ohne Schluesselwort")

# Zeitraum in einer Zeile ist kein Tag -> nicht raten
seite_zeitraum = """Wochenbericht
Berichtszeitraum vom 03.08.2026 bis 07.08.2026
Fa. Meyer Bau (4 Mann)
"""
gleich(datum_der_seite(seite_zeitraum), date(2026, 8, 3),
       "Zeitraumzeile faellt auf 'erstes Datum' zurueck")

# Gar kein Datum -> None, kein Raten
gleich(datum_der_seite("Fa. Meyer Bau (4 Mann)\n- Schalung"), None,
       "ohne Datum wird nicht geraten")

print("─── Zeitraum-Filter ───")

woche = set(woche_um(date(2026, 8, 5)))          # Mo 03.08. - Fr 07.08.
gleich(sorted(woche)[0], date(2026, 8, 3), "Woche beginnt am Montag")
gleich(len(woche), 5, "fuenf Arbeitstage")
gleich(wochentag(date(2026, 8, 5)), "Mittwoch", "Wochentagsname")

# Mit Zeitraum: die Vertragsfrist ist ausgeschlossen, auch wenn sie zuerst kaeme.
seite_frist_oben = """Fertigstellung bis 30.11.2026
Bautagesbericht
Datum: 06.08.2026
"""
gleich(datum_der_seite(seite_frist_oben, woche), date(2026, 8, 6),
       "Zeitraum schliesst fremde Daten aus")
gleich(datum_der_seite("Nur 30.11.2026 steht hier", woche), None,
       "kein Datum im Zeitraum -> None")


print("─── Seiten zu Tagen gruppieren ───")

funde = [
    Seitenfund("woche.pdf", 1, date(2026, 8, 3), "kopf"),
    Seitenfund("woche.pdf", 2, date(2026, 8, 4), "kopf"),
    Seitenfund("woche.pdf", 3, date(2026, 8, 4), "fortsetzung"),
    Seitenfund("woche.pdf", 4, date(2026, 8, 5), "kopf"),
    Seitenfund("elektro.pdf", 1, date(2026, 8, 3), "kopf"),
    Seitenfund("foto.jpg", 1, None, ""),
]
bloecke = gruppiere_nach_tag(funde)

gleich(len(bloecke), 4, "drei Tage plus ein Block ohne Datum")
gleich([b.datum for b in bloecke[:3]],
       [date(2026, 8, 3), date(2026, 8, 4), date(2026, 8, 5)],
       "aufsteigend sortiert")
gleich(bloecke[-1].datum, None, "unbekannter Tag steht hinten")

montag = bloecke[0]
gleich(montag.seiten_je_datei["woche.pdf"], [1], "Montag: Seite 1")
gleich(montag.seiten_je_datei["elektro.pdf"], [1], "Montag: zweite Firma dabei")
gleich(montag.anzahl_seiten, 2, "Montag hat zwei Seiten")
gleich(sorted(montag.dateien), ["elektro.pdf", "woche.pdf"], "beide Dateien")

dienstag = bloecke[1]
gleich(dienstag.seiten_je_datei["woche.pdf"], [2, 3],
       "Fortsetzungsseite bleibt beim Dienstag")


print("─── Echte PDFs: lesen, trennen, wiederfinden ───")

ARBEIT = Path(tempfile.gettempdir()) / "hpp-test-wochenaufteilung"
if ARBEIT.exists():
    import shutil
    shutil.rmtree(ARBEIT)
ARBEIT.mkdir(parents=True)


def baue_wochen_pdf(ziel: Path, tage: list[date]) -> Path:
    """Erzeugt ein PDF mit einer Tagesseite je Datum — wie ein Firmenpaket."""
    from docx import Document

    from app.services.word_pdf import nach_pdf

    dok = Document()
    for i, tag in enumerate(tage):
        if i:
            dok.add_page_break()
        dok.add_paragraph("Bautagesbericht")
        dok.add_paragraph(f"Datum: {tag.strftime('%d.%m.%Y')}")
        dok.add_paragraph(f"Fa. Meyer Bau ({3 + i} Mann)")
        dok.add_paragraph("- Schalarbeiten")
        dok.add_paragraph("Fertigstellung laut Vertrag bis 30.11.2026")
    docx_pfad = ARBEIT / (ziel.stem + ".docx")
    dok.save(docx_pfad)
    ziel.write_bytes(nach_pdf(docx_pfad.read_bytes()))
    return ziel


try:
    from app.services.word_pdf import word_vorhanden
    hat_word = word_vorhanden()
except Exception:
    hat_word = False

if not hat_word:
    print("   Word nicht erreichbar — PDF-Teil uebersprungen.")
else:
    tage = woche_um(date(2026, 8, 5))
    paket = baue_wochen_pdf(ARBEIT / "wochenpaket.pdf", tage)

    funde = finde_seitendaten([paket], set(tage))
    gleich(len(funde), 5, "fuenf Seiten gelesen")
    gleich([f.datum for f in funde], tage, "jede Seite dem richtigen Tag")
    pruefe(all(f.herkunft == "kopf" for f in funde),
           "alle Tage im Seitenkopf gefunden")

    bloecke = gruppiere_nach_tag(funde)
    gleich(len(bloecke), 5, "fuenf Tagesbloecke")

    # Mittwoch herausschneiden und nachsehen, ob wirklich nur er drinsteht.
    mittwoch = bloecke[2]
    gleich(mittwoch.datum, date(2026, 8, 5), "dritter Block ist Mittwoch")
    teil = schreibe_teil_pdf(paket, mittwoch.seiten_je_datei[str(paket)],
                             ARBEIT / "mittwoch.pdf")
    pruefe(teil.is_file() and teil.stat().st_size > 0, "Teil-PDF geschrieben")

    nachgelesen = finde_seitendaten([teil], set(tage))
    gleich(len(nachgelesen), 1, "Teil-PDF hat genau eine Seite")
    gleich(nachgelesen[0].datum, date(2026, 8, 5), "und es ist der Mittwoch")

    from app.services.wochenaufteilung import seiten_lesen
    # Mittwoch ist der dritte Tag: 3 + 2 = 5 Mann. Die Personenzahl steigt je
    # Tag, deshalb taugt sie hier als Fingerabdruck der richtigen Seite.
    text = seiten_lesen(teil)[0]
    pruefe("5 Mann" in text, "Inhalt des Mittwochs erhalten")
    pruefe("3 Mann" not in text and "4 Mann" not in text,
           "keine fremden Tage im Teil-PDF")

    # Ganze Auswahl -> unveraenderte Kopie
    alles = schreibe_teil_pdf(paket, [1, 2, 3, 4, 5], ARBEIT / "alles.pdf")
    gleich(alles.stat().st_size, paket.stat().st_size,
           "vollstaendige Auswahl wird nur kopiert")

print("─── Mehrere Tage auf EINEM Blatt ───")

woche = set(woche_um(date(2026, 8, 5)))

wochenblatt = """Wochenbericht Fa. Meyer Bau
Datum: 03.08.2026
Fa. Meyer Bau (4 Mann)
- Schalung Achse A
Fertigstellung laut Vertrag bis 30.11.2026
Datum: 04.08.2026
Fa. Meyer Bau (5 Mann)
- Bewehrung Decke UG
Datum: 05.08.2026
Fa. Meyer Bau (6 Mann)
- Betonage Decke UG"""

teile = dict(tagesabschnitte(wochenblatt, woche))
gleich(sorted(teile), [date(2026, 8, 3), date(2026, 8, 4), date(2026, 8, 5)],
       "drei Tage aus einem Blatt")
pruefe("Schalung Achse A" in teile[date(2026, 8, 3)], "Montag: eigener Text")
pruefe("Bewehrung" not in teile[date(2026, 8, 3)], "Montag ohne Dienstags Arbeit")
pruefe("4 Mann" in teile[date(2026, 8, 3)] and "5 Mann" not in teile[date(2026, 8, 3)],
       "Montag: richtige Personenzahl")
pruefe("Betonage Decke UG" in teile[date(2026, 8, 5)], "Mittwoch: eigener Text")
pruefe("30.11.2026" in teile[date(2026, 8, 3)],
       "Vertragsfrist bleibt im Abschnitt, wo sie steht")

# Eine Seite mit nur einem Tag darf NICHT geschnitten werden.
gleich(tagesabschnitte("Datum: 03.08.2026\nFa. X (2 Mann)", woche), [],
       "ein Tag pro Blatt -> kein Schnitt")
gleich(tagesabschnitte("Datum: 03.08.2026\nFa. X (2 Mann)\nFrist bis 30.11.2026",
                       woche), [],
       "Vertragsfrist startet keinen zweiten Abschnitt")

# Ohne Zeitraum zaehlen nur Zeilen mit Datumswort.
ohne_zeitraum = dict(tagesabschnitte(wochenblatt))
gleich(sorted(ohne_zeitraum), [date(2026, 8, 3), date(2026, 8, 4), date(2026, 8, 5)],
       "auch ohne Zeitraum: drei Tage")

# Kopf- und Fusszeile mit demselben Datum sind kein neuer Abschnitt.
doppelkopf = """Datum: 03.08.2026
Fa. X (2 Mann)
- Arbeiten
Seite 1 - Bericht vom 03.08.2026
Datum: 04.08.2026
Fa. X (3 Mann)"""
gleich(sorted(dict(tagesabschnitte(doppelkopf, woche))),
       [date(2026, 8, 3), date(2026, 8, 4)],
       "wiederholtes Datum erzeugt keinen dritten Abschnitt")

# Der Fund traegt den Abschnittstext mit.
if hat_word:
    blatt_docx = ARBEIT / "wochenblatt.docx"
    from docx import Document as Dok
    d = Dok()
    for zeile in wochenblatt.splitlines():
        d.add_paragraph(zeile)
    d.save(blatt_docx)
    blatt_pdf = ARBEIT / "wochenblatt.pdf"
    from app.services.word_pdf import nach_pdf as _nach_pdf
    blatt_pdf.write_bytes(_nach_pdf(blatt_docx.read_bytes()))

    funde = finde_seitendaten([blatt_pdf], woche)
    gleich(len(funde), 3, "ein Blatt -> drei Funde")
    pruefe(all(f.herkunft == "abschnitt" for f in funde), "als Abschnitt markiert")
    pruefe(all(f.abschnitt for f in funde), "jeder Fund traegt seinen Text")
    gleich([f.seite for f in funde], [1, 1, 1], "alle von Seite 1")
    montag_fund = next(f for f in funde if f.datum == date(2026, 8, 3))
    pruefe("Schalung" in montag_fund.abschnitt, "Montagstext im Fund")
    pruefe("Betonage" not in montag_fund.abschnitt, "ohne fremde Tage")

print("─── Robustheit ───")

kaputt = ARBEIT / "kaputt.pdf"
kaputt.write_bytes(b"kein PDF")
funde = finde_seitendaten([kaputt])
gleich(len(funde), 1, "kaputte Datei ergibt einen Fund")
gleich(funde[0].datum, None, "und zwar ohne Datum statt eines Absturzes")

bild = ARBEIT / "handyfoto.jpg"
bild.write_bytes(b"\xff\xd8\xff")
funde = finde_seitendaten([bild])
gleich(funde[0].datum, None, "Foto hat kein lesbares Datum")

print()
if fehler:
    print(f"{ok} Pruefungen ok, {len(fehler)} Fehler:")
    for f in fehler:
        print("  -", f)
    raise SystemExit(1)
print(f"{ok} Pruefungen ok, 0 Fehler")
