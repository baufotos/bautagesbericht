"""Erzeugt den Projektbericht (Monatsbericht) als Word-Dokument.

WAS HIER PASSIERT
=================
Eingabe sind Stammdaten plus ein Wörterbuch „Kapitelschlüssel → Inhalt“.
Welche Kapitel erscheinen und welche Nummer sie tragen, entscheidet
``projektbericht_gliederung.nummeriere`` — hier wird nur noch gesetzt. Diese
Trennung ist Absicht: Die Nummerierung ist die fehleranfällige Stelle (siehe
dortigen Modulkommentar) und muss ohne Word prüfbar bleiben.

WOHER DIE MASSE KOMMEN
======================
Aus der Referenz ``BoB- Projektbericht Nr.3 20260731_1``, und zwar zweimal
gelesen: die Seiteneinrichtung aus der ``.doc`` selbst (über Word), die
Positionen und Schriftgrößen aus dem PDF (Zeichenkoordinaten):

    Seite          A4, Ränder 2.50 / 2.00 / 2.50 / 3.25 cm
    Kopf / Fuß     1.27 cm / 0.95 cm Abstand, erste Seite anders
    Kopfzeile      Projektname links, „HPP“ rechts, darunter eine Linie —
                   nur auf Seite 1; ab Seite 2 bleibt die Linie allein
    Titel          12 pt, eingerückt auf 1 cm
    Hauptkapitel   Nummer fett 14 pt, Titel fett 10 pt
    Unterkapitel   Nummer fett 13 pt, Titel fett 10 pt
    Fließtext      10 pt, eingerückt auf 1 cm
    Verzeichnis    10 pt, Punktführung, Seitenzahl rechts bei 16.6 cm
    Fußzeile       Dateiname fett 6 pt links, „Seite“ + Seitenzahl in
                   Impact 16 pt rechts

SCHRIFTART
==========
Die Vorlage nennt als Standardschrift ``MetaPlusNormal-Roman`` — die
Hausschrift. Sie ist auf den Rechnern nicht installiert; schon das
Referenz-PDF ist deshalb in **Arial** gesetzt. Genau das macht dieses Modul
auch. Wer die Hausschrift ausrollt, ändert ``SCHRIFT`` an einer Stelle.

SEITENZAHLEN IM VERZEICHNIS
===========================
Das Verzeichnis kann die Seitenzahlen nicht selbst wissen — sie stehen erst
fest, wenn Word umbricht. Deshalb bekommt jede Überschrift eine Textmarke und
das Verzeichnis ein ``PAGEREF``-Feld darauf, und das Dokument ist auf
„Felder beim Öffnen aktualisieren“ gestellt. Word trägt die Zahlen dann selbst
ein — auch nachdem jemand im Dokument noch etwas ergänzt hat.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor
from PIL import Image, ImageOps

from app.services import projektbericht_gliederung as gliederung
from app.services import dokumenttext

# ─────────────────────────────────────────────────────────────────────────────
# Maße und Schrift
# ─────────────────────────────────────────────────────────────────────────────

SCHRIFT = "Arial"
#: Schmale Schnitte im Original (ArialNarrow) — für die Fußzeile.
SCHRIFT_SCHMAL = "Arial Narrow"
#: Die Seitenzahl steht im Original in Impact 16 pt. Sieht ungewöhnlich aus,
#: ist aber das Erkennungszeichen dieser Berichtsreihe.
SCHRIFT_SEITENZAHL = "Impact"

GR_TEXT = Pt(10)
GR_TITEL = Pt(12)
GR_NUMMER_H1 = Pt(14)
GR_NUMMER_H2 = Pt(13)
GR_KOPF = Pt(9.5)
GR_FUSS_DATEI = Pt(6)
GR_FUSS_TEXT = Pt(10)
GR_SEITENZAHL = Pt(16)

RAND_LINKS = Cm(2.5)
RAND_RECHTS = Cm(2.0)
RAND_OBEN = Cm(2.5)
RAND_UNTEN = Cm(3.25)
KOPF_ABSTAND = Cm(1.27)
FUSS_ABSTAND = Cm(0.95)

#: Einzug von Text und Kapiteltiteln: gemessen 99.3 pt = 1.0 cm hinter dem Rand.
EINZUG = Cm(1.0)
#: Textbreite (16.5 cm) — der rechte Tabstopp des Verzeichnisses.
TEXTBREITE = Cm(16.5)

#: Fotos: größter Rahmen, in den ein Bild gesetzt wird. Im Original sind es
#: 323–428 pt breit und bis 503 pt hoch, jeweils ein Foto je Seite.
FOTO_MAX_BREITE = Cm(15.1)
FOTO_MAX_HOEHE = Cm(17.7)
FOTO_MAX_KANTE = 1600
FOTO_QUALITAET = 82

ROT = RGBColor(0xC0, 0x00, 0x00)
#: Zeilen, die so beginnen, werden rot gesetzt (terminkritisch).
ROT_MARKE = "!"

TITEL_VORLAGE = "Zusammenfassende Bewertung / Monatsbericht Nr:{nummer}"
VERZEICHNIS_UEBERSCHRIFT = "Weiterer Inhalt:"


# ─────────────────────────────────────────────────────────────────────────────
# Datenmodell (ohne Web- und Datenbankbezug, damit es prüfbar bleibt)
# ─────────────────────────────────────────────────────────────────────────────


@dataclass
class Baubegehung:
    datum: str = ""
    teilnehmer: str = ""
    firma: str = ""

    def zeile(self) -> str:
        teile = [t for t in (self.datum, self.teilnehmer, self.firma) if t.strip()]
        return " — ".join(teile)


@dataclass
class Besprechung:
    bezeichnung: str = ""
    rhythmus: str = ""
    uhrzeit: str = ""


@dataclass
class SollIstZeile:
    bezeichnung: str = ""
    soll: str = ""
    ist: str = ""
    verzug: str = ""


@dataclass
class Berichtsfoto:
    daten: bytes
    bildunterschrift: str = ""


@dataclass
class Projektbericht:
    """Alles, was ein Bericht braucht."""

    projektname: str
    #: Kürzel für Dateiname und Fußzeile, z. B. „BoB“.
    projektkuerzel: str
    nummer: int
    berichtsdatum: date
    ersteller: str = ""
    #: Kürzel oben rechts in der Kopfzeile.
    buero: str = "HPP"
    zeitraum_von: date | None = None
    zeitraum_bis: date | None = None
    #: Kapitelschlüssel → Text (siehe projektbericht_gliederung.GLIEDERUNG).
    kapitel: dict[str, str] = field(default_factory=dict)
    baubegehungen: list[Baubegehung] = field(default_factory=list)
    besprechungen: list[Besprechung] = field(default_factory=list)
    soll_ist: list[SollIstZeile] = field(default_factory=list)
    fotos: list[Berichtsfoto] = field(default_factory=list)

    def inhalte(self) -> dict[str, object]:
        """Kapitelschlüssel → Inhalt, wie die Nummerierung ihn erwartet."""
        werte: dict[str, object] = dict(self.kapitel)
        werte["baubegehungen"] = self.baubegehungen
        werte["besprechungen"] = self.besprechungen
        werte["soll_ist"] = self.soll_ist
        werte["fotos"] = self.fotos
        return werte


class ProjektberichtFehler(ValueError):
    """Eingaben unbrauchbar — Text ist für die Oberfläche gedacht."""


def pruefe(bericht: Projektbericht) -> None:
    """Pflichtangaben prüfen. Kapitelinhalte sind alle freiwillig."""
    fehlt: list[str] = []
    if not bericht.projektname.strip():
        fehlt.append("Projektname")
    if not bericht.projektkuerzel.strip():
        fehlt.append("Projektkürzel (steht in der Fußzeile und im Dateinamen)")
    if bericht.nummer is None or bericht.nummer < 1:
        fehlt.append("Berichtsnummer (mindestens 1)")
    if bericht.berichtsdatum is None:
        fehlt.append("Berichtsdatum")
    if fehlt:
        raise ProjektberichtFehler(
            "Für den Projektbericht fehlen: " + ", ".join(fehlt) + "."
        )
    if (bericht.zeitraum_von and bericht.zeitraum_bis
            and bericht.zeitraum_bis < bericht.zeitraum_von):
        raise ProjektberichtFehler(
            "Der Berichtszeitraum endet vor seinem Beginn."
        )


def dateiname(bericht: Projektbericht, endung: str = "docx") -> str:
    """``{Kürzel}-Projektbericht_Nr.{n}_{JJJJMMTT}.docx``.

    Das Original heißt ``BoB- Projektbericht Nr.3 20260731_1.doc``. Hier ohne
    Leerzeichen: Der Name wandert durch Mailanhänge und Netzlaufwerke, und
    Leerzeichen darin sind auf Dauer Ärger. Das Datumsformat des Originals
    (JJJJMMTT) bleibt.
    """
    kuerzel = "".join(
        z for z in bericht.projektkuerzel.strip().replace(" ", "-")
        if z.isalnum() or z in "-_."
    ) or "Projekt"
    return (f"{kuerzel}-Projektbericht_Nr.{bericht.nummer}_"
            f"{bericht.berichtsdatum:%Y%m%d}.{endung}")


# ─────────────────────────────────────────────────────────────────────────────
# Word-Werkzeuge
# ─────────────────────────────────────────────────────────────────────────────


def _lauf(absatz, text: str, *, fett=False, groesse=GR_TEXT, schrift=SCHRIFT,
          farbe=None):
    lauf = absatz.add_run(dokumenttext.xml_sicher(text))
    lauf.font.name = schrift
    lauf.font.size = groesse
    lauf.bold = fett
    if farbe is not None:
        lauf.font.color.rgb = farbe
    return lauf


def _absatz(behaelter, *, einzug=None, abstand_vor=Pt(0), abstand_nach=Pt(0),
            zusammenhalten=False, ausrichtung=None):
    absatz = behaelter.add_paragraph()
    form = absatz.paragraph_format
    form.space_before = abstand_vor
    form.space_after = abstand_nach
    if einzug is not None:
        form.left_indent = einzug
    if zusammenhalten:
        form.keep_with_next = True
    if ausrichtung is not None:
        absatz.alignment = ausrichtung
    return absatz


def _ohne_stilstopps(absatz) -> None:
    """Absatz auf den Normal-Stil setzen.

    Die mitgelieferten Stile „Kopfzeile“ und „Fußzeile“ bringen eigene
    Tabstopps mit (Mitte und rechts). Ein Tabulator sprang deshalb in die
    Seitenmitte statt an den rechten Rand — im Kopf sah das aus, als stünde
    das Bürokürzel mitten auf der Seite.
    """
    eigenschaften = absatz._p.get_or_add_pPr()
    for vorhandener in eigenschaften.findall(qn("w:pStyle")):
        eigenschaften.remove(vorhandener)
    stil = OxmlElement("w:pStyle")
    stil.set(qn("w:val"), "Normal")
    eigenschaften.insert(0, stil)


def _linie(absatz, *, oben=False) -> None:
    """Waagerechte Linie als Absatzrahmen — so macht es die Word-Vorlage.

    Eine echte Linie (statt einer Reihe Unterstriche) verschiebt sich nicht,
    wenn jemand den Text in der Kopfzeile ändert.
    """
    eigenschaften = absatz._p.get_or_add_pPr()
    rahmen = OxmlElement("w:pBdr")
    kante = OxmlElement("w:top" if oben else "w:bottom")
    kante.set(qn("w:val"), "single")
    kante.set(qn("w:sz"), "6")          # 0.75 pt
    kante.set(qn("w:space"), "1")
    kante.set(qn("w:color"), "000000")
    rahmen.append(kante)
    eigenschaften.append(rahmen)


def _feld(absatz, anweisung: str, vorschau: str, *, groesse=GR_TEXT,
          schrift=SCHRIFT, fett=False) -> None:
    """Word-Feld (PAGE, NUMPAGES, PAGEREF) mit fünf Läufen.

    Nur als echtes Feld stimmen Seitenzahlen auch nach Änderungen im Dokument.
    """
    def teil(inhalt):
        element = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        schriften = OxmlElement("w:rFonts")
        schriften.set(qn("w:ascii"), schrift)
        schriften.set(qn("w:hAnsi"), schrift)
        rpr.append(schriften)
        if fett:
            rpr.append(OxmlElement("w:b"))
        groessen = OxmlElement("w:sz")
        groessen.set(qn("w:val"), str(int(groesse.pt * 2)))
        rpr.append(groessen)
        element.append(rpr)
        element.append(inhalt)
        absatz._p.append(element)

    anfang = OxmlElement("w:fldChar")
    anfang.set(qn("w:fldCharType"), "begin")
    teil(anfang)

    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = anweisung
    teil(text)

    trennung = OxmlElement("w:fldChar")
    trennung.set(qn("w:fldCharType"), "separate")
    teil(trennung)

    inhalt = OxmlElement("w:t")
    inhalt.text = vorschau
    teil(inhalt)

    ende = OxmlElement("w:fldChar")
    ende.set(qn("w:fldCharType"), "end")
    teil(ende)


def _textmarke(absatz, name: str, nummer: int) -> None:
    """Textmarke um einen Absatz — Ziel der PAGEREF-Felder im Verzeichnis."""
    anfang = OxmlElement("w:bookmarkStart")
    anfang.set(qn("w:id"), str(nummer))
    anfang.set(qn("w:name"), name)
    ende = OxmlElement("w:bookmarkEnd")
    ende.set(qn("w:id"), str(nummer))
    absatz._p.insert(0, anfang)
    absatz._p.append(ende)


def _felder_aktualisieren(dokument) -> None:
    """Word beim Öffnen fragen lassen, ob es die Felder aktualisiert.

    Ohne das stünden im Verzeichnis die Platzhalterzahlen aus der Erzeugung.
    """
    einstellungen = dokument.settings.element
    if einstellungen.find(qn("w:updateFields")) is not None:
        return
    element = OxmlElement("w:updateFields")
    element.set(qn("w:val"), "true")
    vorgaenger = einstellungen.find(qn("w:defaultTabStop"))
    if vorgaenger is not None:
        vorgaenger.addnext(element)
    else:
        einstellungen.insert(0, element)


def _grundschrift(dokument) -> None:
    stil = dokument.styles["Normal"]
    stil.font.name = SCHRIFT
    stil.font.size = GR_TEXT
    rpr = stil.element.get_or_add_rPr()
    schriften = rpr.find(qn("w:rFonts"))
    if schriften is None:
        schriften = OxmlElement("w:rFonts")
        rpr.append(schriften)
    for schluessel in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        schriften.set(qn(schluessel), SCHRIFT)
    sprache = rpr.find(qn("w:lang"))
    if sprache is None:
        sprache = OxmlElement("w:lang")
        rpr.append(sprache)
    sprache.set(qn("w:val"), "de-DE")
    stil.paragraph_format.space_before = Pt(0)
    stil.paragraph_format.space_after = Pt(0)


def _rahmen_aus(tabelle) -> None:
    eigenschaften = tabelle._tbl.tblPr
    rahmen = OxmlElement("w:tblBorders")
    for kante in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{kante}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        rahmen.append(element)
    eigenschaften.append(rahmen)


def _foto_aufbereiten(daten: bytes) -> tuple[bytes, Emu, Emu]:
    """Dreht nach EXIF, verkleinert und passt ins Rahmenmaß.

    Handyfotos liegen quer in der Datei und tragen die Drehung nur im
    EXIF-Vermerk; ohne ``exif_transpose`` läge jedes Hochformatfoto auf der
    Seite. Das Verkleinern hält den Bericht handlich — die Originalvorlage ist
    mit eingebetteten Fotos über 12 MB groß.
    """
    try:
        with Image.open(io.BytesIO(daten)) as bild:
            gedreht = ImageOps.exif_transpose(bild)
            if gedreht.mode not in ("RGB", "L"):
                gedreht = gedreht.convert("RGB")
            gedreht.thumbnail((FOTO_MAX_KANTE, FOTO_MAX_KANTE), Image.LANCZOS)
            breite_px, hoehe_px = gedreht.size
            puffer = io.BytesIO()
            gedreht.save(puffer, format="JPEG", quality=FOTO_QUALITAET,
                         optimize=True)
    except Exception as fehler:                       # noqa: BLE001
        raise ProjektberichtFehler(
            f"Ein Foto lässt sich nicht lesen ({fehler})."
        ) from fehler

    breite = int(FOTO_MAX_BREITE)
    hoehe = int(breite * hoehe_px / breite_px)
    if hoehe > int(FOTO_MAX_HOEHE):
        hoehe = int(FOTO_MAX_HOEHE)
        breite = int(hoehe * breite_px / hoehe_px)
    return puffer.getvalue(), Emu(breite), Emu(hoehe)


# ─────────────────────────────────────────────────────────────────────────────
# Bausteine des Berichts
# ─────────────────────────────────────────────────────────────────────────────


def _kopfzeilen(abschnitt, bericht: Projektbericht) -> None:
    """Erste Seite mit Projektname und Bürokürzel, Folgeseiten nur mit Linie."""
    abschnitt.different_first_page_header_footer = True
    abschnitt.header_distance = KOPF_ABSTAND

    erste = abschnitt.first_page_header
    erste.is_linked_to_previous = False
    absatz = erste.paragraphs[0]
    _ohne_stilstopps(absatz)
    absatz.paragraph_format.space_after = Pt(0)
    absatz.paragraph_format.tab_stops.add_tab_stop(TEXTBREITE, WD_TAB_ALIGNMENT.RIGHT)
    _lauf(absatz, bericht.projektname, fett=True, groesse=GR_KOPF)
    _lauf(absatz, "\t", groesse=GR_KOPF)
    _lauf(absatz, bericht.buero, fett=True, groesse=GR_KOPF)
    _linie(absatz)

    folge = abschnitt.header
    folge.is_linked_to_previous = False
    weiter = folge.paragraphs[0]
    _ohne_stilstopps(weiter)
    weiter.paragraph_format.space_after = Pt(0)
    _lauf(weiter, "", groesse=GR_KOPF)
    _linie(weiter)


def _fusszeile(fuss, bericht: Projektbericht) -> None:
    fuss.is_linked_to_previous = False
    absatz = fuss.paragraphs[0]
    absatz.paragraph_format.space_before = Pt(0)
    absatz.paragraph_format.space_after = Pt(0)
    _ohne_stilstopps(absatz)
    absatz.paragraph_format.tab_stops.add_tab_stop(TEXTBREITE, WD_TAB_ALIGNMENT.RIGHT)
    _linie(absatz, oben=True)

    _lauf(absatz, dateiname(bericht, "doc"), fett=True, groesse=GR_FUSS_DATEI)
    _lauf(absatz, "\t", groesse=GR_FUSS_TEXT, schrift=SCHRIFT_SCHMAL)
    _lauf(absatz, "Seite ", groesse=GR_FUSS_TEXT, schrift=SCHRIFT_SCHMAL)
    _feld(absatz, " PAGE ", "1", groesse=GR_SEITENZAHL,
          schrift=SCHRIFT_SEITENZAHL)
    _lauf(absatz, "/", groesse=GR_FUSS_TEXT, schrift=SCHRIFT_SCHMAL)
    _feld(absatz, " NUMPAGES ", "1", groesse=GR_FUSS_TEXT,
          schrift=SCHRIFT_SCHMAL)


def _fusszeilen(abschnitt, bericht: Projektbericht) -> None:
    abschnitt.footer_distance = FUSS_ABSTAND
    _fusszeile(abschnitt.footer, bericht)
    _fusszeile(abschnitt.first_page_footer, bericht)


def _ueberschrift(dokument, kapitel, marke: str, nummer: int):
    """Kapitelüberschrift: Nummer fett und größer, Titel fett daneben."""
    absatz = _absatz(dokument, abstand_vor=Pt(14), abstand_nach=Pt(4),
                     zusammenhalten=True)
    absatz.paragraph_format.tab_stops.add_tab_stop(EINZUG, WD_TAB_ALIGNMENT.LEFT)
    groesse = GR_NUMMER_H1 if kapitel.ebene == 1 else GR_NUMMER_H2
    _lauf(absatz, kapitel.nummer, fett=True, groesse=groesse)
    _lauf(absatz, "\t", fett=True, groesse=groesse)
    _lauf(absatz, kapitel.titel, fett=True, groesse=GR_TEXT)
    _textmarke(absatz, marke, nummer)
    return absatz


def _fliesstext(dokument, text: str) -> None:
    """Text in Absätze zerlegen; „!“ am Zeilenanfang setzt die Zeile rot.

    Die Markierung ist mit Absicht so schlicht: Auf der Baustelle wird in ein
    Textfeld getippt, nicht in einen Editor mit Werkzeugleiste. Ein Ausrufe-
    zeichen am Zeilenanfang ist schnell gesetzt und im Formular erklärt.
    """
    for zeile in (text or "").replace("\r\n", "\n").split("\n"):
        absatz = _absatz(dokument, einzug=EINZUG, abstand_nach=Pt(2))
        inhalt = zeile.strip()
        if inhalt.startswith(ROT_MARKE):
            _lauf(absatz, inhalt[len(ROT_MARKE):].strip(), farbe=ROT)
        elif inhalt:
            _lauf(absatz, inhalt)


def _baubegehungen(dokument, eintraege: list[Baubegehung]) -> None:
    for eintrag in eintraege:
        absatz = _absatz(dokument, einzug=EINZUG, abstand_nach=Pt(2))
        _lauf(absatz, eintrag.zeile())


def _besprechungen(dokument, eintraege: list[Besprechung]) -> None:
    """Bezeichnung links, Rhythmus und Uhrzeit auf einem Tabstopp.

    Im Original stehen „Planungsbesprechungen“ und „Baubesprechung“
    untereinander und die Angaben dahinter auf gleicher Höhe — deshalb ein
    Tabstopp und keine Tabelle.
    """
    for eintrag in eintraege:
        absatz = _absatz(dokument, einzug=EINZUG, abstand_nach=Pt(2))
        absatz.paragraph_format.tab_stops.add_tab_stop(Cm(5.0), WD_TAB_ALIGNMENT.LEFT)
        _lauf(absatz, eintrag.bezeichnung)
        rest = " ".join(t for t in (eintrag.rhythmus, eintrag.uhrzeit) if t.strip())
        if rest:
            _lauf(absatz, "\t" + rest)


def _soll_ist(dokument, zeilen: list[SollIstZeile]) -> None:
    """SOLL-IST-Vergleich als randlose Tabelle mit fetten Spaltenköpfen."""
    tabelle = dokument.add_table(rows=1, cols=4)
    tabelle.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabelle.autofit = False
    _rahmen_aus(tabelle)

    breiten = (Cm(6.2), Cm(3.6), Cm(3.6), Cm(3.1))
    kopf = tabelle.rows[0]
    for spalte, (breite, titel) in enumerate(
        zip(breiten, ("", "SOLL", "IST (Starttermin)", "Verzug"))
    ):
        zelle = kopf.cells[spalte]
        zelle.width = breite
        absatz = zelle.paragraphs[0]
        absatz.paragraph_format.space_after = Pt(0)
        _lauf(absatz, titel, fett=True)

    for zeile in zeilen:
        reihe = tabelle.add_row()
        werte = (zeile.bezeichnung, zeile.soll, zeile.ist, zeile.verzug)
        for spalte, (breite, wert) in enumerate(zip(breiten, werte)):
            zelle = reihe.cells[spalte]
            zelle.width = breite
            absatz = zelle.paragraphs[0]
            absatz.paragraph_format.space_after = Pt(0)
            _lauf(absatz, wert, fett=True)

    _absatz(dokument, abstand_nach=Pt(4))


def _fotos(dokument, fotos: list[Berichtsfoto]) -> None:
    """Ein Foto je Absatz, Bildunterschrift direkt darunter."""
    for foto in fotos:
        bild, breite, hoehe = _foto_aufbereiten(foto.daten)
        absatz = _absatz(dokument, einzug=EINZUG, abstand_vor=Pt(8),
                         zusammenhalten=True)
        lauf = absatz.add_run()
        lauf.add_picture(io.BytesIO(bild), width=breite, height=hoehe)

        unterschrift = _absatz(dokument, einzug=EINZUG, abstand_vor=Pt(4),
                               abstand_nach=Pt(10))
        _lauf(unterschrift, foto.bildunterschrift or "")


def _verzeichnis(dokument, eintraege) -> None:
    """„Weiterer Inhalt“: Nummer, Titel, Punktführung, Seitenzahl als Feld."""
    ueberschrift = _absatz(dokument, abstand_vor=Pt(24), abstand_nach=Pt(8),
                           ausrichtung=WD_ALIGN_PARAGRAPH.CENTER)
    _lauf(ueberschrift, VERZEICHNIS_UEBERSCHRIFT)

    for lauf_nummer, kapitel in enumerate(eintraege):
        vorher = Pt(4) if kapitel.ebene == 1 and lauf_nummer else Pt(0)
        absatz = _absatz(dokument, abstand_vor=vorher, abstand_nach=Pt(0))
        stopps = absatz.paragraph_format.tab_stops
        stopps.add_tab_stop(EINZUG, WD_TAB_ALIGNMENT.LEFT)
        # Punktführung bis zur Seitenzahl — im Original ein rechter Tabstopp
        # mit Punkten, genau wie Word ein Inhaltsverzeichnis setzt.
        eintrag = stopps.add_tab_stop(TEXTBREITE, WD_TAB_ALIGNMENT.RIGHT)
        eintrag.leader = WD_TAB_LEADER.DOTS

        _lauf(absatz, kapitel.nummer)
        _lauf(absatz, "\t" + kapitel.titel)
        _lauf(absatz, "\t")
        _feld(absatz, f" PAGEREF {_marke(kapitel)} \\h ", "1")


def _marke(kapitel) -> str:
    """Name der Textmarke eines Kapitels — muss ohne Punkte auskommen."""
    return "kap_" + kapitel.schluessel.replace("-", "_")


# ─────────────────────────────────────────────────────────────────────────────
# Der Bericht
# ─────────────────────────────────────────────────────────────────────────────


def erzeuge_bericht(bericht: Projektbericht) -> bytes:
    """Baut das Word-Dokument und gibt es als Bytes zurück."""
    pruefe(bericht)

    kapitel = gliederung.nummeriere(bericht.inhalte())

    dokument = Document()
    _grundschrift(dokument)
    _felder_aktualisieren(dokument)

    abschnitt = dokument.sections[0]
    abschnitt.page_width = Cm(21)
    abschnitt.page_height = Cm(29.7)
    abschnitt.left_margin = RAND_LINKS
    abschnitt.right_margin = RAND_RECHTS
    abschnitt.top_margin = RAND_OBEN
    abschnitt.bottom_margin = RAND_UNTEN
    _kopfzeilen(abschnitt, bericht)
    _fusszeilen(abschnitt, bericht)

    # ── Titelzeile ──
    titel = _absatz(dokument, einzug=EINZUG, abstand_nach=Pt(12))
    _lauf(titel, TITEL_VORLAGE.format(nummer=bericht.nummer), groesse=GR_TITEL)

    # ── Kapitel 1 (Zusammenfassende Bewertung) ──
    erste_gruppe = [k for k in kapitel if k.nummer.split(".")[0] == "1"]
    for nummer, eintrag in enumerate(erste_gruppe, start=1):
        _ueberschrift(dokument, eintrag, _marke(eintrag), nummer)
        _inhalt(dokument, eintrag)

    # ── Verzeichnis ──
    weitere = gliederung.inhaltsverzeichnis(kapitel)
    if weitere:
        _verzeichnis(dokument, weitere)
        dokument.add_page_break()

    # ── Restliche Kapitel ──
    for nummer, eintrag in enumerate(weitere, start=len(erste_gruppe) + 1):
        _ueberschrift(dokument, eintrag, _marke(eintrag), nummer)
        _inhalt(dokument, eintrag)

    puffer = io.BytesIO()
    dokument.save(puffer)
    return puffer.getvalue()


def _inhalt(dokument, eintrag) -> None:
    """Setzt den Inhalt eines Kapitels — je nach Art."""
    if eintrag.nur_ueberschrift:
        return
    art = eintrag.art
    inhalt = eintrag.inhalt
    if gliederung.ist_leer(inhalt):
        return

    if art == gliederung.ART_BAUBEGEHUNGEN:
        _baubegehungen(dokument, inhalt)
    elif art == gliederung.ART_BESPRECHUNGEN:
        _besprechungen(dokument, inhalt)
    elif art == gliederung.ART_SOLLIST:
        _soll_ist(dokument, inhalt)
    elif art == gliederung.ART_FOTOS:
        _fotos(dokument, inhalt)
    else:
        _fliesstext(dokument, str(inhalt))
