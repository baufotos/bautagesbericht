"""Eingegangene Mehrkosten- und Behinderungsanzeigen auslesen.

WAS HIER HEREINKOMMT
====================
Ein Schreiben einer Baufirma — PDF (mit Textebene), Word oder Text. Es gibt
keine Norm dafür: Jede Firma hat ihren eigenen Briefbogen. Was alle gemeinsam
haben, ist die Anordnung eines Geschäftsbriefs nach DIN 5008, und genau darauf
setzt dieses Modul auf.

DER TRICK: SPALTEN STATT ZEILEN
===============================
Ein PDF hat keine Zeilen, nur Wörter mit Koordinaten. Liest man es Zeile für
Zeile, klebt der rechte Briefkopf an der Anschrift:

    "10163 Berlin Ansprechpartner"

— zwei Angaben, die zwölf Zentimeter auseinanderstehen, in einer Zeichenkette.
Deshalb wird jede Zeile zusätzlich an ihren *Lücken* zerlegt: Ein Abstand von
mehr als gut zwei Schriftgrößen ist keine Wortlücke mehr, sondern ein
Spaltenwechsel. Erst danach ist "10163 Berlin" wieder eine Anschriftzeile und
"Ansprechpartner" die Überschrift der rechten Spalte.

WER ABSENDER IST UND WER EMPFÄNGER
==================================
Diese Unterscheidung entscheidet über den ganzen Brief, denn HPP antwortet an
den **Absender** — die Firma, die die Anzeige geschrieben hat. Im Schreiben
selbst steht ihre Anschrift zweimal:

    klein über dem Anschriftfeld   die Rücksendeangabe (DIN 5008)
    ganz unten                     der Impressumsfuß

Groß im Anschriftfeld steht dagegen der *Bauherr*. Die Größe der Schrift ist
damit das verlässlichste Merkmal, und sie wird hier auch so benutzt:
Rücksendeangabe zuerst, Impressumsfuß als Rückfallebene, das Anschriftfeld
gehört immer der Gegenseite.

WAS NICHT GERATEN WIRD
======================
Die Anrede. Ob "Herr" oder "Frau" steht in keinem dieser Schreiben verlässlich
drin, und ein falsch angeredeter Bauleiter ist schlimmer als eine neutrale
Anrede. Findet sich keine ausdrückliche Angabe, bleibt das Feld leer und der
Brief beginnt mit "Sehr geehrte Damen und Herren," — ändern kann man es im
Formular immer.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from app.services import dokumenttext

# ─────────────────────────────────────────────────────────────────────────────
# Maße und Schwellen
# ─────────────────────────────────────────────────────────────────────────────


def _spaltenluecke(groesse: float) -> float:
    """Ab dieser Lücke (in Punkt) beginnt eine neue Spalte.

    Nicht absolut, sondern an der Schriftgröße gemessen — in einem 6-pt-Fuß
    sind 14 pt eine Welt, in einer 11-pt-Überschrift nicht einmal zwei Zeichen.
    """
    return max(14.0, 2.2 * groesse)


#: Zwei Wörter gehören zur selben Zeile, wenn ihre Oberkanten so nah liegen.
ZEILEN_TOLERANZ = 3.0

#: Schrift bis zu dieser Größe gilt als "klein": Rücksendeangabe, Impressum,
#: Spaltenüberschriften des Briefkopfs.
KLEIN_BIS = 7.6

#: Links dieser Grenze (Anteil der Seitenbreite) liegt die Anschriftspalte.
LINKE_SPALTE_BIS = 0.55

#: Nur der obere Teil der ersten Seite kommt für Anschriften in Frage.
ANSCHRIFT_ZONE_BIS = 0.45

MONATE = {
    "januar": 1, "februar": 2, "märz": 3, "maerz": 3, "april": 4, "mai": 5,
    "juni": 6, "juli": 7, "august": 8, "september": 9, "oktober": 10,
    "november": 11, "dezember": 12,
}

#: Rechtsformen und Firmenkennzeichen — woran eine Firmenzeile erkennbar ist.
_RECHTSFORM = re.compile(
    r"\b(?:GmbH|mbH|AG|KG|OHG|GbR|UG|SE|e\.\s?K\.|e\.\s?V\.|Co\.|KGaA|"
    r"Ltd\.?|B\.\s?V\.|Bau|Gruppe|Gesellschaft|Unternehmen|Ingenieure?|"
    r"Architekten|Partner(?:schaft)?)\b",
    re.IGNORECASE,
)

#: Straßenzeile: endet auf eine Hausnummer, oder enthält ein Straßenwort.
_STRASSENWORT = re.compile(
    r"(?:stra(?:ß|ss)e|str\.|weg|allee|platz|ring|damm|gasse|chaussee|ufer|"
    r"kamp|kämpe|hof|berg|feld|graben|markt|wall|steig|pfad|zeile|winkel|"
    r"kehre|redder|stieg|br(?:ü|ue)cke)\b",
    re.IGNORECASE,
)
_HAUSNUMMER_ENDE = re.compile(r"\d+\s*[a-zA-Z]?(?:\s*[-–/]\s*\d+\s*[a-zA-Z]?)?$")

_PLZ_ORT = re.compile(r"^(?:D\s*-\s*)?(\d{5})\s+([A-ZÄÖÜ][^,;]{1,60})$")
_EMAIL = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_TELEFON = re.compile(
    r"(?:Telefon|Tel\.?|Fon)\s*[:\-]?\s*((?:\+|00)?[\d()/\s.\-]{7,25}\d)",
    re.IGNORECASE,
)

#: Die Anzeigearten, die im Büro vorkommen — die längste Schreibweise zuerst,
#: sonst gewinnt "Mehrkostenanzeige" gegen "Behinderungs- und
#: Mehrkostenanzeige" und das Schreiben wäre falsch einsortiert.
ARTEN: tuple[tuple[str, str], ...] = (
    (r"Behinderungs-\s*(?:und|u\.|/)\s*Mehrkostenanzeige",
     "Behinderungs- und Mehrkostenanzeige"),
    (r"Mehrkosten(?:anzeige|anmeldung|ank(?:ü|ue)ndigung)", "Mehrkostenanzeige"),
    (r"Behinderungs(?:anzeige|meldung)", "Behinderungsanzeige"),
    (r"Bedenken(?:anzeige|anmeldung|hinweis)", "Bedenkenanmeldung"),
    (r"Nachtrags?(?:angebot|anmeldung|ank(?:ü|ue)ndigung)", "Nachtragsangebot"),
    (r"Verg(?:ü|ue)tung zus(?:ä|ae)tzlicher Leistungen",
     "Vergütung zusätzlicher Leistungen"),
    (r"Stundenlohnanzeige", "Stundenlohnanzeige"),
    (r"St(?:ö|oe)rungsanzeige", "Störungsanzeige"),
)

#: Kürzel, mit denen die Firmen ihre Schreiben durchzählen (MKA 01, BMA 07).
_KUERZEL_NUMMER = re.compile(
    r"\b(MKA|MKN|MEKO|MK|BMA|BMH|BEH|VZL|NT|NA|SLA)\s*[-–]?\s*"
    r"(?:Nr\.?\s*)?(\d{1,3})\b"
)
_NUMMER_HINTER_ART = re.compile(r"(?:Nr\.?|Nummer)\s*(\d{1,3})")

#: Welche Art zu welchem Kürzel gehört — für Schreiben, die sich selbst nur
#: über ihr Kürzel benennen ("BEH 01", "MEKO 11").
KUERZEL_ART: dict[str, str] = {
    "MKA": "Mehrkostenanzeige",
    "MKN": "Mehrkostenanzeige",
    "MEKO": "Mehrkostenanzeige",
    "MK": "Mehrkostenanzeige",
    "BMA": "Behinderungs- und Mehrkostenanzeige",
    "BMH": "Behinderungs- und Mehrkostenanzeige",
    "BEH": "Behinderungsanzeige",
    "VZL": "Vergütung zusätzlicher Leistungen",
    "NT": "Nachtragsangebot",
    "NA": "Nachtragsangebot",
    "SLA": "Stundenlohnanzeige",
}

#: Umgekehrter Weg: Welches Kürzel das Büro für eine Art benutzt, wenn die
#: Firma selbst keines vergeben hat. Es steht im Dateinamen des Antwort-
#: schreibens ("260507 WEI-VE300.01-BEH 01.docx").
#:
#: Bewusst unvollständig: Aufgenommen ist nur, was in den Schreiben des Büros
#: nachweisbar so heißt. Für Bedenkenanmeldung und Störungsanzeige gibt es
#: dort kein Beispiel — ein erfundenes Kürzel im Dateinamen wäre schlimmer
#: als der ausgeschriebene Name, weil danach niemand sucht.
ART_KUERZEL: dict[str, str] = {
    "Mehrkostenanzeige": "MKA",
    "Behinderungsanzeige": "BEH",
    "Behinderungs- und Mehrkostenanzeige": "BMA",
    "Vergütung zusätzlicher Leistungen": "VZL",
    "Nachtragsangebot": "NT",
}

_VOB = re.compile(
    r"§+\s*\d+(?:\s*Abs(?:atz|\.)?\s*\d+(?:\s*(?:,|und|u\.|bis|-)\s*\d+)*)?"
    r"(?:\s*(?:Nr\.?|Ziff\.?)\s*\d+)?\s*(?:VOB/B|VOB/A|BGB|HOAI)",
    re.IGNORECASE,
)

_LV_POSITION = re.compile(
    r"(?:Pos(?:ition)?\.?|OZ|LV-?Pos\.?)\s*[:\-]?\s*"
    r"(\d{2}(?:[.\-]\d{2,4}){1,3})",
    re.IGNORECASE,
)

_BAUZEIT_WORT = re.compile(
    r"(Verl(?:ä|ae)ngerung der Ausf(?:ü|ue)hrungsdauer|"
    r"Bauzeit(?:en)?verl(?:ä|ae)ngerung|Verl(?:ä|ae)ngerung der Bauzeit|"
    r"Fristverl(?:ä|ae)ngerung|Verz(?:ö|oe)gerung des Bauablaufs|"
    r"Terminverschiebung|Baustellenstillstand)",
    re.IGNORECASE,
)

_FORDERUNG_WORT = re.compile(
    r"(Wir bitten (?:um|Sie)[^.]{0,200}\.|Wir fordern[^.]{0,200}\.|"
    r"Wir melden[^.]{0,200}an\.|Wir behalten uns[^.]{0,200}\.)",
    re.IGNORECASE,
)

_GRUSS = re.compile(
    r"^(?:Mit\s+)?(?:freundlichen?\s+Gr(?:ü|ue)(?:ß|ss)en?|"
    r"Freundliche\s+Gr(?:ü|ue)(?:ß|ss)e|Mit besten Gr(?:ü|ue)(?:ß|ss)en)\s*$",
    re.IGNORECASE,
)
_FUNKTION_ZEILE = re.compile(
    r"^-?\s*(Bauleitung|Bauleiter(?:in)?|Projektleitung|Projektleiter(?:in)?|"
    r"Gesch(?:ä|ae)ftsf(?:ü|ue)hr(?:er|erin|ung)|Partner(?:in)?|"
    r"Prokurist(?:in)?|Kalkulation|Technische Leitung|Oberbauleitung|"
    r"Baumanagement)\s*-?\s*$",
    re.IGNORECASE,
)

#: Sammelpostfächer. Sie stehen im Briefkopf jeder Firma und sind **nicht**
#: die Adresse des Bauleiters — eine Antwort dorthin landet in der Zentrale.
SAMMELPOSTFAECHER = {
    "info", "kontakt", "mail", "email", "office", "zentrale", "buero",
    "büro", "post", "poststelle", "service", "verwaltung", "empfang",
    "anfrage", "anfragen", "sekretariat", "team", "hallo", "willkommen",
}

#: Beschriftungen im Fließtext, hinter denen genau eine Angabe steht.
BESCHRIFTUNGEN: dict[str, tuple[str, ...]] = {
    "betreff": ("Hier", "Betreff", "Betr", "Bezug", "Thema", "Gegenstand"),
    "leistungsort": ("Leistungsort", "Bauvorhaben", "Baustelle", "Objekt",
                     "Bauort", "Projekt", "BV"),
    "gewerk": ("Gewerk", "Leistung", "Leistungsbereich", "Los",
               "Vergabeeinheit", "VE"),
    "projektnummer": ("Projektnummer", "Projekt-Nr", "Projektnr", "Kommission",
                      "Kommissionsnummer", "Auftragsnummer", "Auftrags-Nr",
                      "BV-Nr"),
    "ansprechpartner": ("Ansprechpartner", "Ihr Ansprechpartner", "Bearbeiter",
                        "Sachbearbeiter"),
    "datum": ("Datum",),
}


# ─────────────────────────────────────────────────────────────────────────────
# Ergebnis
# ─────────────────────────────────────────────────────────────────────────────


@dataclass
class Anschrift:
    """Eine Postanschrift, so wie sie in einen Adressblock gehört."""

    firma: str = ""
    zusatz: str = ""
    strasse: str = ""
    plz: str = ""
    ort: str = ""
    land: str = ""

    def leer(self) -> bool:
        return not (self.firma or self.strasse or self.ort)

    def zeilen(self) -> list[str]:
        """Die Anschrift als Briefzeilen.

        "Deutschland" bleibt weg: Inlandspost braucht kein Land, und alle
        Referenzbriefe des Büros lassen es weg.
        """
        raus = [self.firma, self.zusatz, self.strasse]
        if self.plz or self.ort:
            raus.append(f"{self.plz} {self.ort}".strip())
        if self.land and self.land.strip().lower() not in (
            "deutschland", "germany", "de"
        ):
            raus.append(self.land)
        return [z.strip() for z in raus if z and z.strip()]


@dataclass
class Punkt:
    """Ein nummerierter Abschnitt der Anzeige."""

    nummer: str
    titel: str
    text: str = ""


@dataclass
class GelesenesSchreiben:
    """Alles, was aus einem eingegangenen Schreiben herauszuholen war."""

    quelle: str = ""
    seiten: int = 0

    art: str = ""
    nummer: str = ""
    kennung: str = ""            # "MKA 01" — wie die Firma selbst zählt
    datum: date | None = None
    betreff: str = ""
    kurzbezeichnung: str = ""    # Betreff ohne "Mehrkostenanzeige Nr. 01,"

    absender: Anschrift = field(default_factory=Anschrift)
    absender_email: str = ""
    absender_telefon: str = ""
    ansprechpartner: str = ""
    ansprechpartner_email: str = ""

    empfaenger: Anschrift = field(default_factory=Anschrift)

    projektnummer: str = ""
    leistungsort: str = ""
    gewerk: str = ""

    rechtsgrundlage: str = ""
    punkte: list[Punkt] = field(default_factory=list)
    lv_positionen: list[str] = field(default_factory=list)
    bauzeit: str = ""
    forderung: str = ""

    unterzeichner: str = ""
    unterzeichner_funktion: str = ""

    volltext: str = ""
    hinweise: list[str] = field(default_factory=list)


class Lesefehler(Exception):
    """Das Schreiben ließ sich nicht öffnen oder enthielt keinen Text."""


# ─────────────────────────────────────────────────────────────────────────────
# Schritt 1 — Wörter zu Spaltenstücken
# ─────────────────────────────────────────────────────────────────────────────


@dataclass
class Stueck:
    """Ein zusammenhängendes Textstück einer Zeile, mit Ort und Schriftgröße."""

    text: str
    x0: float
    x1: float
    oben: float
    groesse: float
    seite: int


def _stuecke_einer_seite(seite, nummer: int) -> list[Stueck]:
    """Die Wörter einer PDF-Seite als Spaltenstücke, von oben nach unten."""
    try:
        worte = seite.extract_words(extra_attrs=["size", "upright"])
    except Exception:
        return []

    # Gedrehter Text ist bei diesen Briefbögen die senkrechte Randbeschriftung
    # ("Sie befinden sich auf Seite 1 von 2"). Sie kommt verkehrt herum aus der
    # Textebene und würde jede Suche verunreinigen.
    worte = [w for w in worte if w.get("upright", True)]
    if not worte:
        return []

    zeilen: list[list[dict]] = []
    for wort in sorted(worte, key=lambda w: (w["top"], w["x0"])):
        if zeilen and abs(wort["top"] - zeilen[-1][0]["top"]) <= ZEILEN_TOLERANZ:
            zeilen[-1].append(wort)
        else:
            zeilen.append([wort])

    stuecke: list[Stueck] = []
    for zeile in zeilen:
        zeile.sort(key=lambda w: w["x0"])
        gruppe = [zeile[0]]
        for wort in zeile[1:]:
            luecke = wort["x0"] - gruppe[-1]["x1"]
            if luecke > _spaltenluecke(max(float(wort.get("size", 10.0)), 6.0)):
                stuecke.append(_zu_stueck(gruppe, nummer))
                gruppe = [wort]
            else:
                gruppe.append(wort)
        stuecke.append(_zu_stueck(gruppe, nummer))
    return stuecke


def _zu_stueck(worte: list[dict], seite: int) -> Stueck:
    return Stueck(
        text=" ".join(w["text"] for w in worte).strip(),
        x0=min(w["x0"] for w in worte),
        x1=max(w["x1"] for w in worte),
        oben=min(w["top"] for w in worte),
        groesse=max(float(w.get("size", 10.0)) for w in worte),
        seite=seite,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Schritt 2 — Datei einlesen
# ─────────────────────────────────────────────────────────────────────────────


def lies(pfad: Path, quelle: str = "") -> GelesenesSchreiben:
    """Ein Schreiben auslesen. Erkennt PDF, Word und Text an der Endung."""
    name = quelle or pfad.name
    endung = pfad.suffix.lower()
    if endung == ".pdf":
        return _lies_pdf(pfad, name)
    if endung in (".docx", ".docm", ".dotx"):
        return _lies_docx(pfad, name)
    if endung in (".txt", ".text", ".md"):
        return _aus_text(
            pfad.read_text(encoding="utf-8", errors="replace"), name, seiten=1
        )
    raise Lesefehler(
        f"„{name}“ kann nicht gelesen werden. Möglich sind PDF, Word (.docx) "
        f"und Textdateien."
    )


def _lies_pdf(pfad: Path, quelle: str) -> GelesenesSchreiben:
    import pdfplumber

    try:
        with pdfplumber.open(pfad) as pdf:
            seiten = len(pdf.pages)
            breite = float(pdf.pages[0].width) if seiten else 595.0
            hoehe = float(pdf.pages[0].height) if seiten else 842.0
            stuecke: list[Stueck] = []
            for nr, seite in enumerate(pdf.pages, 1):
                stuecke += _stuecke_einer_seite(seite, nr)
    except Exception as fehler:                     # pragma: no cover
        raise Lesefehler(
            f"„{quelle}“ ließ sich nicht öffnen ({fehler.__class__.__name__}). "
            f"Ist die Datei beschädigt?"
        ) from fehler

    if not stuecke:
        raise Lesefehler(
            f"In „{quelle}“ steckt kein auslesbarer Text — das ist ein Scan "
            f"ohne Textebene. Bitte das Schreiben als Original-PDF oder als "
            f"Word-Datei hochladen."
        )

    return _aus_stuecken(stuecke, quelle, seiten, breite, hoehe)


def _lies_docx(pfad: Path, quelle: str) -> GelesenesSchreiben:
    from docx import Document

    try:
        dok = Document(str(pfad))
    except Exception as fehler:
        raise Lesefehler(
            f"„{quelle}“ ließ sich nicht öffnen ({fehler.__class__.__name__})."
        ) from fehler

    teile: list[str] = [p.text for p in dok.paragraphs]
    for tabelle in dok.tables:
        for reihe in tabelle.rows:
            # Verbundene Zellen liefert python-docx mehrfach — das Doppelte
            # würde jede Anschrift verdreifachen.
            entdoppelt: list[str] = []
            for zelle in (z.text.strip() for z in reihe.cells):
                if zelle and (not entdoppelt or entdoppelt[-1] != zelle):
                    entdoppelt.append(zelle)
            if entdoppelt:
                teile.append("    ".join(entdoppelt))

    text = "\n".join(teile)
    if not text.strip():
        raise Lesefehler(f"„{quelle}“ enthält keinen Text.")
    return _aus_text(text, quelle, seiten=1)


# ─────────────────────────────────────────────────────────────────────────────
# Schritt 3 — Auswerten
# ─────────────────────────────────────────────────────────────────────────────


def _aus_stuecken(
    stuecke: list[Stueck],
    quelle: str,
    seiten: int,
    breite: float,
    hoehe: float,
) -> GelesenesSchreiben:
    """Der Weg mit Koordinaten — für PDFs."""
    daten = _aus_text(
        _volltext(stuecke, breite), quelle, seiten, mit_geometrie=True
    )

    # Anschriften und rechte Briefkopfspalte kommen aus der Geometrie; dort
    # sind sie viel verlässlicher als in der Textreihenfolge.
    _anschriften_aus_geometrie(daten, stuecke, breite, hoehe)
    _rechte_spalte(daten, stuecke, breite, hoehe)
    _impressum(daten, stuecke, hoehe)
    _nachbessern(daten)
    return daten


def _volltext(stuecke: list[Stueck], breite: float) -> str:
    """Die Spaltenstücke zurück zu Zeilen — spaltenweise, nicht quer.

    Das Zerlegen an den Lücken (siehe ``_stuecke_einer_seite``) trennt zwei
    Dinge, die verschieden behandelt werden müssen:

        "Hier:"        "Mehrkostenanzeige Nr. 01, …"     eine Angabe, links
        "10163 Berlin" "Ansprechpartner"                 zwei Spalten

    Der Unterschied ist die Spalte, in der das zweite Stück steht. Deshalb
    werden Stücke derselben Zeile nur dann wieder zusammengefügt, wenn beide
    links stehen — der Tabulator hinter "Hier:" verschwindet, der halbe Brief
    zwischen "Berlin" und "Ansprechpartner" bleibt ein Zeilenumbruch.
    """
    grenze = breite * LINKE_SPALTE_BIS
    zeilen: list[str] = []
    for seite in sorted({s.seite for s in stuecke}):
        auf_seite = [s for s in stuecke if s.seite == seite and s.text.strip()]
        gruppen: list[list[Stueck]] = []
        for stueck in sorted(auf_seite, key=lambda s: (s.oben, s.x0)):
            if gruppen and abs(stueck.oben - gruppen[-1][0].oben) <= ZEILEN_TOLERANZ:
                gruppen[-1].append(stueck)
            else:
                gruppen.append([stueck])
        for gruppe in gruppen:
            # Nach x sortieren, nicht nach der Fundreihenfolge: "Gewerk:" und
            # sein Wert liegen im PDF ein Zehntel Punkt auseinander, und nach
            # Oberkante sortiert stünde die Beschriftung hinter ihrem Wert
            # ("Abbruch-, Rückbau- und Entsorgungsmaßnahme Gewerk:").
            gruppe.sort(key=lambda s: s.x0)
            links = [s.text for s in gruppe if s.x0 < grenze]
            rechts = [s.text for s in gruppe if s.x0 >= grenze]
            if links:
                zeilen.append(" ".join(links))
            zeilen += rechts
    return "\n".join(zeilen)


def _aus_text(
    text: str, quelle: str, seiten: int, *, mit_geometrie: bool = False
) -> GelesenesSchreiben:
    """Der Weg ohne Koordinaten — für Word, Text und als Rückfallebene.

    ``mit_geometrie`` sagt: Der Aufrufer hat Koordinaten und wertet die rechte
    Briefkopfspalte selbst aus (siehe ``_rechte_spalte``). Dann bleibt die
    Suche nach allein stehenden Beschriftungen aus, und der Feinschliff
    ebenfalls — beides erledigt ``_aus_stuecken`` danach mit den besseren
    Angaben.

    Der Unterschied ist kein Schönheitsfehler: Im flach gelesenen PDF stehen
    die Spaltenüberschriften zwischen den Anschriftzeilen, und "Projektnummer"
    bekäme die Straße des Absenders als Wert.
    """
    sauber = dokumenttext.xml_sicher(text)
    zeilen = [z.strip() for z in sauber.split("\n")]
    # Umgedrehte Randbeschriftung ("1 etieS") wegwerfen — sie entsteht beim
    # Auslesen gedrehter Textebenen und ergibt nie ein Wort.
    zeilen = [z for z in zeilen if not _ist_umgedreht(z)]

    daten = GelesenesSchreiben(quelle=quelle, seiten=seiten, volltext=sauber)

    _art_und_nummer(daten, sauber)
    _beschriftungen(daten, zeilen)
    if not mit_geometrie:
        _stehende_beschriftungen(daten, zeilen)
    _datum(daten, zeilen, sauber)
    _rechtsgrundlage(daten, sauber)
    _punkte(daten, zeilen)
    _lv_positionen(daten, sauber)
    _bauzeit_und_forderung(daten, sauber)
    _unterschrift(daten, zeilen)
    if daten.absender.leer():
        _anschriften_aus_zeilen(daten, zeilen)
    if not mit_geometrie:
        _nachbessern(daten)
    return daten


def _stehende_beschriftungen(daten: GelesenesSchreiben,
                             zeilen: list[str]) -> None:
    """Beschriftung allein auf einer Zeile, ihr Wert darunter.

    So sieht die rechte Briefkopfspalte aus, wenn man sie ohne Koordinaten
    liest — bei Word- und Textdateien ist das der Normalfall:

        Ansprechpartner
        Qasem Ashgarzada
        Qasem.Ashgarzada@muster-rueckbau.de

    Gelesen wird bis zur nächsten Beschriftung oder zur Leerzeile, höchstens
    drei Zeilen. Mehr steht in einem Briefkopf nie, und ohne Grenze zöge eine
    Beschriftung den halben Brief in ihr Feld.
    """
    for nr, zeile in enumerate(zeilen):
        feld = _welche_beschriftung(zeile)
        if feld is None:
            continue
        werte: list[str] = []
        for folge in zeilen[nr + 1: nr + 4]:
            text = folge.strip()
            if not text or _welche_beschriftung(text) is not None:
                break
            werte.append(text)
        if not werte:
            continue

        if feld == "datum":
            if daten.datum is None:
                daten.datum = _erstes_datum(" ".join(werte))
            continue
        if feld == "ansprechpartner":
            for text in werte:
                mail = _EMAIL.search(text)
                if mail:
                    if not daten.ansprechpartner_email:
                        daten.ansprechpartner_email = mail.group(0)
                elif not daten.ansprechpartner:
                    daten.ansprechpartner = text
            continue
        if not getattr(daten, feld, ""):
            setattr(daten, feld, " ".join(werte))


#: Wörter, die gespiegelt in gedrehten Randbeschriftungen auftauchen.
_UMGEDREHT = {"seite", "von", "sie", "befinden", "sich", "auf", "blatt", "seiten"}


def _ist_umgedreht(zeile: str) -> bool:
    """Erkennt gespiegelte Wörter wie „etieS“, „nednifeb“, „hcis“."""
    text = zeile.strip()
    if not text:
        return False
    worte = [w for w in text.split() if any(c.isalpha() for c in w)]
    if not worte:
        return False
    return all(w.strip(".,").lower()[::-1] in _UMGEDREHT for w in worte)


# ── Art, Nummer, Betreff ─────────────────────────────────────────────────────


def _art_und_nummer(daten: GelesenesSchreiben, text: str) -> None:
    for muster, name in ARTEN:
        treffer = re.search(muster, text, re.IGNORECASE)
        if treffer:
            daten.art = name
            # Nummer möglichst aus derselben Stelle — "Mehrkostenanzeige Nr. 01".
            nummer = _NUMMER_HINTER_ART.search(text[treffer.end(): treffer.end() + 40])
            if nummer:
                daten.nummer = nummer.group(1)
                daten.kennung = f"{name} Nr. {nummer.group(1)}"
            break

    kuerzel = _KUERZEL_NUMMER.search(text)
    if kuerzel:
        if not daten.nummer:
            daten.nummer = kuerzel.group(2)
        daten.kennung = f"{kuerzel.group(1).upper()} {kuerzel.group(2)}"
        if not daten.art:
            daten.art = KUERZEL_ART.get(kuerzel.group(1).upper(), "")

    if not daten.art:
        daten.hinweise.append(
            "Die Art des Schreibens war nicht erkennbar — bitte im Formular "
            "wählen (Mehrkostenanzeige, Behinderungsanzeige, …)."
        )


def _beschriftungen(daten: GelesenesSchreiben, zeilen: list[str]) -> None:
    """Alles, was hinter einer Beschriftung wie „Hier:“ oder „Gewerk:“ steht.

    Mehrzeilige Betreffs kommen mit: Läuft der Betreff über den Zeilenrand,
    steht die Fortsetzung als eigene Zeile darunter — genau wie in
    "Mehrkostenanzeige Nr. 02, Zusätzlicher Rückbauaufwand aufgrund
    abweichender / Konstruktion der Gipskarton-Ständerwände".
    """
    for feld, woerter in BESCHRIFTUNGEN.items():
        if feld == "datum" or getattr(daten, feld, ""):
            continue
        muster = re.compile(
            r"^(?:" + "|".join(re.escape(w) for w in woerter) + r")\s*[:\-]\s*(.+)$",
            re.IGNORECASE,
        )
        for nr, zeile in enumerate(zeilen):
            treffer = muster.match(zeile)
            if not treffer:
                continue
            wert = treffer.group(1).strip()
            if feld == "betreff":
                wert = _betreff_fortsetzen(wert, zeilen, nr)
            if wert:
                setattr(daten, feld, wert)
            break

    if not daten.betreff:
        daten.betreff = _betreff_ueber_der_anrede(zeilen)


def _betreff_ueber_der_anrede(zeilen: list[str]) -> str:
    """Der Betreff ohne Beschriftung: die Zeile direkt über der Anrede.

    Nicht jede Firma schreibt "Hier:" oder "Betreff:" davor — in den
    HPP-Briefen selbst steht der Betreff schlicht fett über der Anrede, und
    genauso machen es viele Baufirmen. Nach DIN 5008 ist das der einzige Ort,
    an dem er stehen darf, also ist die Zeile davor eine verlässliche Quelle.
    """
    for nr, zeile in enumerate(zeilen):
        if not re.match(r"^(?:Sehr geehrte|Guten Tag|Hallo)", zeile, re.IGNORECASE):
            continue
        for rueckwaerts in range(nr - 1, max(nr - 4, -1), -1):
            text = zeilen[rueckwaerts].strip()
            if not text:
                continue
            if not (12 <= len(text) <= 140) or "@" in text:
                return ""
            # Datums- und Zeichenzeile, Anschriftzeilen: keine Betreffs.
            if _PLZ_ORT.match(text) or _ist_strasse(text):
                return ""
            if re.fullmatch(r"[\d.\s/–\-:]+", text):
                return ""
            return text
        return ""
    return ""


def _betreff_fortsetzen(wert: str, zeilen: list[str], nr: int) -> str:
    """Hängt Fortsetzungszeilen an einen umgebrochenen Betreff."""
    for folge in zeilen[nr + 1: nr + 4]:
        text = folge.strip()
        if not text:
            break
        # Eine neue Beschriftung, die Anrede oder ein ganzer Satz beenden ihn.
        if re.match(r"^[A-ZÄÖÜ][A-Za-zÄÖÜäöüß\-. ]{0,20}\s*:", text):
            break
        if re.match(r"^(?:Sehr geehrte|Guten Tag|Hallo)", text, re.IGNORECASE):
            break
        if text.endswith((".", "!", "?", ":")) or len(text) > 110:
            break
        if wert.endswith((",", "-", "–")) or text[:1].islower() or len(text) < 70:
            wert = f"{wert} {text}".strip()
            continue
        break
    return re.sub(r"\s{2,}", " ", wert)


# ── Datum ────────────────────────────────────────────────────────────────────


def _datum(daten: GelesenesSchreiben, zeilen: list[str], text: str) -> None:
    if daten.datum is not None:
        return
    # Zuerst hinter der Beschriftung "Datum" (rechte Briefkopfspalte).
    for nr, zeile in enumerate(zeilen):
        if re.fullmatch(r"Datum\s*:?", zeile, re.IGNORECASE):
            for folge in zeilen[nr + 1: nr + 4]:
                gefunden = _erstes_datum(folge)
                if gefunden:
                    daten.datum = gefunden
                    return
    gefunden = _erstes_datum(text)
    if gefunden:
        daten.datum = gefunden
    else:
        daten.hinweise.append(
            "Im Schreiben stand kein lesbares Datum — bitte im Formular "
            "eintragen. Es steht später im Betreff des HPP-Briefs."
        )


def _erstes_datum(text: str) -> date | None:
    """Erstes Datum im Text, deutsche Schreibweisen inbegriffen."""
    lang = re.search(
        r"(\d{1,2})\.\s*(" + "|".join(MONATE) + r")\s*(\d{4})", text, re.IGNORECASE
    )
    if lang:
        tag, monat, jahr = lang.groups()
        try:
            return date(int(jahr), MONATE[monat.lower()], int(tag))
        except ValueError:
            pass
    kurz = re.search(r"\b(\d{1,2})\.(\d{1,2})\.(\d{2,4})\b", text)
    if kurz:
        tag, monat, jahr = (int(t) for t in kurz.groups())
        # Zweistellige Jahre: 26 heißt 2026, nicht 1926.
        if jahr < 100:
            jahr += 2000
        try:
            return date(jahr, monat, tag)
        except ValueError:
            return None
    return None


# ── Rechtsgrundlage, Punkte, Positionen ──────────────────────────────────────


def _rechtsgrundlage(daten: GelesenesSchreiben, text: str) -> None:
    treffer = _VOB.search(text)
    if treffer:
        daten.rechtsgrundlage = re.sub(r"\s+", " ", treffer.group(0)).strip()


def _punkte(daten: GelesenesSchreiben, zeilen: list[str]) -> None:
    """Die nummerierten Abschnitte der Anzeige, mit Titel und Text.

    Als Überschrift gilt nur eine kurze, für sich stehende Zeile ohne
    Satzpunkt — sonst wäre jede Aufzählung im Fließtext ("1. Wir bitten …")
    plötzlich ein Kapitel.

    Danach bleibt nur der längste **fortlaufende** Lauf ab 1 übrig. Das ist
    nicht Kosmetik: Ein Briefdatum wie "22. Juli 2026" sieht wie eine
    Überschrift Nr. 22 aus, und ohne diese Prüfung stünde der Monat Juli als
    erster Beanstandungspunkt im Antwortschreiben.
    """
    kopf = re.compile(r"^(\d{1,2})([.)])\s+(\S.{2,90})$")
    aktuell: Punkt | None = None
    zeichen: list[str] = []
    gefunden: list[Punkt] = []
    for zeile in zeilen:
        treffer = kopf.match(zeile)
        if treffer and not treffer.group(3).rstrip().endswith((".", ";")):
            if aktuell is not None:
                aktuell.text = aktuell.text.strip()
                gefunden.append(aktuell)
            aktuell = Punkt(nummer=treffer.group(1), titel=treffer.group(3).strip())
            zeichen.append(treffer.group(2))
            continue
        if aktuell is not None and zeile:
            if _GRUSS.match(zeile):
                break
            aktuell.text = f"{aktuell.text} {zeile}".strip()
    if aktuell is not None:
        aktuell.text = aktuell.text.strip()
        gefunden.append(aktuell)

    # Nur Überschriften mit demselben Gliederungszeichen gehören zusammen. In
    # den Referenzbriefen stehen beide Formen im selben Schreiben: die
    # Kapitel als "1." und die Unterpunkte darin als "1)". Ohne diese Trennung
    # würde ein Unterpunkt eine neue Kapitelfolge beginnen.
    lauf: list[Punkt] = []
    for wahl in (".", ")"):
        kandidaten = [p for p, z in zip(gefunden, zeichen) if z == wahl]
        gereiht = _fortlaufend(kandidaten)
        if len(gereiht) > len(lauf):
            lauf = gereiht

    # Ein einziger "Abschnitt" ist fast immer ein Fehlalarm.
    daten.punkte = lauf if len(lauf) >= 2 else []


def _fortlaufend(punkte: list[Punkt]) -> list[Punkt]:
    """Der längste Lauf, der bei 1 beginnt und um 1 weitersteigt."""
    beste: list[Punkt] = []
    stelle = 0
    while stelle < len(punkte):
        if punkte[stelle].nummer.lstrip("0") != "1":
            stelle += 1
            continue
        lauf = [punkte[stelle]]
        erwartet = 2
        weiter = stelle + 1
        while weiter < len(punkte) and \
                punkte[weiter].nummer.lstrip("0") == str(erwartet):
            lauf.append(punkte[weiter])
            erwartet += 1
            weiter += 1
        if len(lauf) > len(beste):
            beste = lauf
        stelle = max(weiter, stelle + 1)
    return beste


def _lv_positionen(daten: GelesenesSchreiben, text: str) -> None:
    gesehen: list[str] = []
    for treffer in _LV_POSITION.finditer(text):
        wert = treffer.group(1)
        if wert not in gesehen:
            gesehen.append(wert)
    daten.lv_positionen = gesehen


def _bauzeit_und_forderung(daten: GelesenesSchreiben, text: str) -> None:
    fluss = re.sub(r"\s+", " ", text)
    treffer = _BAUZEIT_WORT.search(fluss)
    if treffer:
        # Der ganze Satz, in dem das Wort steht — die Dauer steht darin
        # ("… um zwei Wochen").
        anfang = fluss.rfind(".", 0, treffer.start()) + 1
        ende = fluss.find(".", treffer.end())
        satz = fluss[anfang: ende + 1 if ende != -1 else len(fluss)]
        daten.bauzeit = satz.strip()

    forderung = _FORDERUNG_WORT.search(fluss)
    if forderung:
        daten.forderung = forderung.group(1).strip()


def _unterschrift(daten: GelesenesSchreiben, zeilen: list[str]) -> None:
    """Name und Funktion unter dem Grußwort."""
    for nr, zeile in enumerate(zeilen):
        if not _GRUSS.match(zeile):
            continue
        for folge in zeilen[nr + 1: nr + 7]:
            text = folge.strip(" _\t")
            if not text:
                continue
            if _FUNKTION_ZEILE.match(text):
                daten.unterzeichner_funktion = text.strip("- ").strip()
                continue
            if daten.unterzeichner:
                continue
            # Eine Namenszeile: zwei bis vier Wörter, keine Ziffern, kein Satz.
            worte = text.split()
            if (1 < len(worte) <= 4
                    and not any(c.isdigit() for c in text)
                    and not text.endswith((".", ":"))):
                daten.unterzeichner = text
        break


# ── Anschriften ──────────────────────────────────────────────────────────────


def _anschriften_aus_geometrie(
    daten: GelesenesSchreiben,
    stuecke: list[Stueck],
    breite: float,
    hoehe: float,
) -> None:
    """Rücksendeangabe und Anschriftfeld aus dem Kopf der ersten Seite.

    Der Aufbau ist bei jedem Geschäftsbrief derselbe (DIN 5008): links oben
    zuerst die klein gesetzte Rücksendeangabe des Absenders, direkt darunter
    das Anschriftfeld des Empfängers in Lesegröße.
    """
    zone = [
        s for s in stuecke
        if s.seite == 1
        and s.x0 < breite * LINKE_SPALTE_BIS
        and s.oben < hoehe * ANSCHRIFT_ZONE_BIS
        and s.text.strip()
    ]
    zone.sort(key=lambda s: s.oben)
    if not zone:
        return

    klein = [s for s in zone if s.groesse <= KLEIN_BIS]
    gross = [s for s in zone if s.groesse > KLEIN_BIS]

    # Rücksendeangabe: die kleinen Stücke, die ÜBER dem Anschriftfeld stehen.
    erste_grosse = gross[0].oben if gross else hoehe
    ruecksende = [s.text for s in klein if s.oben < erste_grosse]
    if ruecksende:
        gelesen = _anschrift_aus_zeilen(ruecksende)
        if not gelesen.leer():
            daten.absender = gelesen

    # Anschriftfeld: die großen Stücke, bis der Fließtext beginnt.
    felder: list[str] = []
    for stueck in gross:
        text = stueck.text.strip()
        if re.match(r"^(?:Sehr geehrte|Guten Tag|Hallo)", text, re.IGNORECASE):
            break
        if re.match(r"^[A-ZÄÖÜ][A-Za-zÄÖÜäöüß\-. ]{2,24}\s*:", text):
            break
        felder.append(text)
    if felder:
        gelesen = _anschrift_aus_zeilen(felder)
        if not gelesen.leer():
            daten.empfaenger = gelesen


def _trennen(zeilen: list[str]) -> list[str]:
    """Zerlegt "Firma · Straße · PLZ Ort · Land" in einzelne Zeilen."""
    raus: list[str] = []
    for zeile in zeilen:
        teile = re.split(r"\s*[·•|]\s*|,\s(?=\d{5}\s)", zeile)
        raus += [t.strip() for t in teile if t.strip()]
    return raus


def _anschrift_aus_zeilen(zeilen: list[str]) -> Anschrift:
    """Firma / Zusatz / Straße / PLZ Ort aus einer Handvoll Zeilen."""
    anschrift = Anschrift()
    rest: list[str] = []

    for zeile in _trennen(zeilen):
        text = zeile.strip().rstrip(",")
        if not text or "@" in text or text.lower().startswith("www."):
            continue
        plzort = _PLZ_ORT.match(text)
        if plzort and not anschrift.plz:
            anschrift.plz, anschrift.ort = plzort.group(1), plzort.group(2).strip()
            continue
        if _ist_land(text):
            anschrift.land = text
            continue
        if not anschrift.strasse and _ist_strasse(text):
            anschrift.strasse = text
            continue
        rest.append(text)

    for text in rest:
        if not anschrift.firma:
            anschrift.firma = text
        elif not anschrift.zusatz and _ist_person(text):
            anschrift.zusatz = text
    return anschrift


def _ist_strasse(text: str) -> bool:
    if len(text) > 70:
        return False
    if _STRASSENWORT.search(text) and _HAUSNUMMER_ENDE.search(text):
        return True
    # Auch ohne Straßenwort: "Krumbäken Kämpe 2", "Am Wall 5".
    return (bool(_HAUSNUMMER_ENDE.search(text))
            and len(text.split()) <= 4
            and not _RECHTSFORM.search(text))


def _ist_land(text: str) -> bool:
    return text.strip().lower() in {
        "deutschland", "germany", "österreich", "oesterreich", "austria",
        "schweiz", "switzerland", "niederlande", "netherlands", "polen",
        "dänemark", "daenemark", "frankreich", "belgien", "luxemburg",
    }


def _ist_person(text: str) -> bool:
    """Ist das eine Personenzeile („Herrn Steffen Wegner“, „Frau Braun“)?"""
    if re.match(r"^(?:Herrn?|Frau|Fr\.|Hr\.|z\.\s?Hd\.|Familie)\b", text,
                re.IGNORECASE):
        return True
    worte = text.split()
    return (1 < len(worte) <= 4
            and not _RECHTSFORM.search(text)
            and not any(c.isdigit() for c in text)
            and all(w[:1].isupper() for w in worte if w[:1].isalpha()))


def _anschriften_aus_zeilen(daten: GelesenesSchreiben, zeilen: list[str]) -> None:
    """Rückfallebene ohne Koordinaten: Anschrift aus dem Impressumsfuß.

    Bei Word- und Textdateien gibt es keine Schriftgrößen. Dann ist der
    Impressumsfuß die beste Quelle für den Absender — er steht in jedem
    Geschäftsbrief und nennt Firma, Straße und Ort in einer Zeile.
    """
    for zeile in zeilen:
        if (zeile.count("·") + zeile.count("|") >= 2
                and re.search(r"\b\d{5}\b", zeile)):
            gelesen = _anschrift_aus_zeilen([zeile])
            if not gelesen.leer():
                daten.absender = gelesen
                return

    # Sonst die ersten Zeilen: ein Briefbogen beginnt mit dem Firmennamen.
    kopf = [z for z in zeilen[:12] if z.strip()]
    for nr, zeile in enumerate(kopf):
        if _RECHTSFORM.search(zeile) and len(zeile) < 70:
            gelesen = _anschrift_aus_zeilen(kopf[nr: nr + 4])
            if not gelesen.leer():
                daten.absender = gelesen
                return


def _rechte_spalte(
    daten: GelesenesSchreiben,
    stuecke: list[Stueck],
    breite: float,
    hoehe: float,
) -> None:
    """Ansprechpartner, Datum und Projektnummer aus der rechten Briefkopfspalte.

    Dort steht die Beschriftung klein über ihrem Wert. Deshalb wird jeder
    kleinen Beschriftung zugeordnet, was in derselben Spalte darunter folgt,
    bis die nächste Beschriftung kommt.
    """
    spalte = [
        s for s in stuecke
        if s.seite == 1
        and s.x0 >= breite * LINKE_SPALTE_BIS
        and s.oben < hoehe * ANSCHRIFT_ZONE_BIS
        and s.text.strip()
    ]
    spalte.sort(key=lambda s: s.oben)

    beschriftung: str | None = None
    werte: dict[str, list[str]] = {}
    for stueck in spalte:
        text = stueck.text.strip()
        passend = _welche_beschriftung(text)
        if passend and stueck.groesse <= KLEIN_BIS:
            beschriftung = passend
            werte.setdefault(passend, [])
            continue
        if beschriftung:
            werte[beschriftung].append(text)

    for feld, zeilen in werte.items():
        eintraege = [z for z in zeilen if z]
        if not eintraege:
            continue
        if feld == "datum":
            if daten.datum is None:
                daten.datum = _erstes_datum(" ".join(eintraege))
            continue
        if feld == "ansprechpartner":
            for zeile in eintraege:
                mail = _EMAIL.search(zeile)
                if mail:
                    if not daten.ansprechpartner_email:
                        daten.ansprechpartner_email = mail.group(0)
                elif not daten.ansprechpartner:
                    daten.ansprechpartner = zeile
            continue
        if not getattr(daten, feld, ""):
            setattr(daten, feld, " ".join(eintraege))


def _welche_beschriftung(text: str) -> str | None:
    schlicht = text.strip().rstrip(":").strip().lower()
    for feld, woerter in BESCHRIFTUNGEN.items():
        if schlicht in {w.lower() for w in woerter}:
            return feld
    return None


def _impressum(daten: GelesenesSchreiben, stuecke: list[Stueck],
               hoehe: float) -> None:
    """Telefon, E-Mail und — falls noch offen — die Anschrift aus dem Fuß."""
    fuss = " · ".join(
        s.text for s in stuecke
        if s.oben > hoehe * 0.75 and s.groesse <= KLEIN_BIS and s.text.strip()
    )
    if not fuss:
        return
    telefon = _TELEFON.search(fuss)
    if telefon and not daten.absender_telefon:
        daten.absender_telefon = re.sub(r"\s{2,}", " ", telefon.group(1)).strip()
    mail = _EMAIL.search(fuss)
    if mail and not daten.absender_email:
        daten.absender_email = mail.group(0)
    if daten.absender.leer():
        gelesen = _anschrift_aus_zeilen([fuss])
        if not gelesen.leer():
            daten.absender = gelesen


# ── Feinschliff ──────────────────────────────────────────────────────────────


def _nachbessern(daten: GelesenesSchreiben) -> None:
    """Lücken schließen, Kurzbezeichnung bilden, Hinweise sammeln."""
    _mailadressen_zuordnen(daten)

    # Name des Ansprechpartners aus der Unterschrift, wenn die Spalte fehlte.
    if not daten.ansprechpartner and daten.unterzeichner:
        daten.ansprechpartner = daten.unterzeichner

    # Kennung: Hat die Firma kein Kürzel vergeben, nimmt das Büro sein eigenes
    # ("BEH 02"). Es landet im Dateinamen, und dort ist "Behinderungsanzeige
    # Nr. 02" nur lang, nicht klarer.
    kuerzel = ART_KUERZEL.get(daten.art)
    if kuerzel and daten.nummer and not _KUERZEL_NUMMER.search(daten.kennung):
        daten.kennung = f"{kuerzel} {daten.nummer}"

    # Kurzbezeichnung: der Betreff ohne die Zählung davor.
    if daten.betreff:
        kurz = re.sub(
            r"^(?:Behinderungs-\s*(?:und|u\.)\s*)?[A-Za-zÄÖÜäöüß\- ]*"
            r"(?:anzeige|anmeldung|angebot|ank(?:ü|ue)ndigung|Leistungen)\s*"
            r"(?:Nr\.?\s*\d+)?\s*[,:–-]?\s*",
            "",
            daten.betreff,
            count=1,
            flags=re.IGNORECASE,
        ).strip()
        daten.kurzbezeichnung = kurz or daten.betreff

    if daten.absender.leer():
        daten.hinweise.append(
            "Die Anschrift der Firma war nicht sicher zu erkennen — bitte im "
            "Formular ergänzen. Sie stünde sonst leer im Adressfeld."
        )
    elif not daten.absender.strasse or not daten.absender.ort:
        daten.hinweise.append(
            "Von der Anschrift der Firma fehlt ein Teil (Straße oder Ort) — "
            "bitte im Formular nachsehen."
        )
    elif not _passt_zur_maildomain(daten):
        # Der einzige Fehler, der wirklich weh tut: Steht im Adressfeld der
        # Bauherr statt der Firma, geht die Antwort an den falschen Empfänger
        # — und liest sich, als hätte HPP sich selbst geantwortet. Firmenname
        # und Maildomäne müssen deshalb zueinander passen.
        daten.hinweise.append(
            f"Achtung: Die erkannte Firma „{daten.absender.firma}“ passt nicht "
            f"zur E-Mail-Adresse „{daten.ansprechpartner_email}“. Bitte prüfen, "
            f"ob im Adressfeld die anzeigende Firma steht und nicht der Bauherr."
        )
    if not daten.betreff:
        daten.hinweise.append(
            "Es war kein Betreff („Hier:“ / „Betreff:“) zu finden — bitte im "
            "Formular eintragen."
        )
    if not daten.ansprechpartner_email:
        daten.hinweise.append(
            "Im Schreiben stand keine E-Mail-Adresse der Firma — für den "
            "Outlook-Entwurf bitte eine eintragen."
        )


def _mailadressen_zuordnen(daten: GelesenesSchreiben) -> None:
    """Welche der Adressen im Schreiben die des Bauleiters ist.

    In jedem Briefbogen stehen mindestens zwei: das Sammelpostfach der Firma
    im Kopf oder Fuß und die persönliche Adresse des Ansprechpartners. Die
    Antwort gehört an die persönliche — im Sammelpostfach landet sie in der
    Zentrale und braucht Tage bis zur Bauleitung.

    Die Reihenfolge der Vorlieben ist deshalb: eine Adresse, die zum Namen des
    Ansprechpartners passt; sonst irgendeine persönliche; erst zuletzt das
    Sammelpostfach. Fremde Maildomänen bleiben außen vor, denn die Adresse des
    Bauherrn steht in diesen Schreiben ebenfalls oft mit drin.
    """
    adressen: list[str] = []
    for treffer in _EMAIL.finditer(daten.volltext or ""):
        adresse = treffer.group(0)
        if adresse.lower() not in {a.lower() for a in adressen}:
            adressen.append(adresse)
    if not adressen:
        return

    if not daten.absender_email:
        sammel = [a for a in adressen if _ist_sammelpostfach(a)]
        if sammel:
            daten.absender_email = sammel[0]

    if daten.ansprechpartner_email:
        return

    eigene = _hausdomain(daten)
    in_frage = [a for a in adressen
                if not eigene or a.lower().endswith(eigene)] or adressen
    persoenlich = [a for a in in_frage if not _ist_sammelpostfach(a)]
    name = daten.ansprechpartner or daten.unterzeichner
    passend = [a for a in persoenlich if _passt_zum_namen(a, name)]

    for wahl in (passend, persoenlich, in_frage):
        if wahl:
            daten.ansprechpartner_email = wahl[0]
            return


def _ist_sammelpostfach(adresse: str) -> bool:
    return adresse.split("@", 1)[0].strip().lower() in SAMMELPOSTFAECHER


def _passt_zum_namen(adresse: str, name: str) -> bool:
    """Steckt der Nachname des Ansprechpartners in der Adresse?"""
    if not name.strip():
        return False
    nachname = _entumlauten(name.split()[-1])
    if len(nachname) < 3:
        return False
    return nachname in _entumlauten(adresse.split("@", 1)[0])


def _entumlauten(text: str) -> str:
    schlicht = (text.lower()
                .replace("ä", "ae").replace("ö", "oe").replace("ü", "ue")
                .replace("ß", "ss"))
    return re.sub(r"[^a-z0-9]", "", schlicht)


def _hausdomain(daten: GelesenesSchreiben) -> str:
    """Die Maildomäne der schreibenden Firma, z. B. „@servisa-gruppe.de“."""
    if daten.absender_email and "@" in daten.absender_email:
        return "@" + daten.absender_email.split("@", 1)[1].lower()
    return ""


def _passt_zur_maildomain(daten: GelesenesSchreiben) -> bool:
    """Steckt der Firmenname in der Maildomäne (oder umgekehrt)?

    Verglichen wird auf der Ebene der Namensbestandteile, nicht der ganzen
    Zeichenkette: "SERVISA Rückbau & Sanierung GmbH" gegen
    "servisa-gruppe.de" passt über "servisa", "Riedel Bau AG" gegen
    "riedelbau.de" über "riedel". Gibt es keine Adresse zum Vergleichen,
    gilt die Anschrift als in Ordnung — geprüft wird nur, was prüfbar ist.
    """
    adresse = daten.ansprechpartner_email or daten.absender_email
    if not adresse or "@" not in adresse or not daten.absender.firma:
        return True
    domain = adresse.split("@", 1)[1].lower()
    # Die Firmenkennung der Domäne ohne Landesendung ("riedelbau", "servisa").
    kern = re.sub(r"[^a-z0-9]", "", domain.rsplit(".", 1)[0])
    if not kern:
        return True
    for wort in re.findall(r"[A-Za-zÄÖÜäöüß]{4,}", daten.absender.firma):
        schlicht = (wort.lower()
                    .replace("ä", "a").replace("ö", "o")
                    .replace("ü", "u").replace("ß", "ss"))
        if _RECHTSFORM.fullmatch(wort):
            continue
        if schlicht in kern or kern.startswith(schlicht[:5]):
            return True
    return False
