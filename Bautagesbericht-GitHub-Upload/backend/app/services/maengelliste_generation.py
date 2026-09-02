"""Erzeugt die Mängelliste als Word-Dokument aus einer Blanko-Vorlage.

Aufbau wie bei ``app.services.docx_generation``: Vorlage laden, Platzhalter
füllen, Inhalte einsetzen, speichern.

╔══════════════════════════════════════════════════════════════════════════╗
║  HIER WIRD SPÄTER DIE ECHTE HPP-VORLAGE EINGESETZT                       ║
║                                                                          ║
║  Die reale Bürovorlage liegt noch nicht vor. Bis dahin arbeitet der Code  ║
║  mit ``backend/templates/Maengelliste_leer.docx`` — einer schlichten      ║
║  Platzhalter-Vorlage im Stil von ``Bautagesbericht_HPP_leer.docx``.       ║
║                                                                          ║
║  Austausch später, OHNE Codeänderung an anderer Stelle:                   ║
║    1. Die echte Vorlage als ``backend/templates/Maengelliste_leer.docx``  ║
║       ablegen (gleicher Dateiname) — oder einen anderen Pfad über den     ║
║       Parameter ``template_path`` übergeben.                             ║
║    2. In der Vorlage die Platzhalter aus ``PLATZHALTER`` (siehe unten)    ║
║       an die gewünschten Stellen schreiben, z. B. {{PROJEKT}} in den      ║
║       Briefkopf und {{FUSSZEILE}} in die Word-Fußzeile.                   ║
║    3. Genau einen Absatz mit {{MAENGEL}} an die Stelle setzen, an der die ║
║       Mängelblöcke erscheinen sollen. Dieser Absatz wird beim Erzeugen    ║
║       durch je eine Tabelle pro Mangel ersetzt und selbst entfernt.       ║
║                                                                          ║
║  Nicht in der Vorlage vorhandene Platzhalter werden schlicht ignoriert;   ║
║  fehlt {{MAENGEL}}, hängen die Blöcke am Dokumentende an. Der Austausch    ║
║  kann also nichts kaputt machen.                                          ║
╚══════════════════════════════════════════════════════════════════════════╝

Die Mängelblöcke werden bewusst mit der python-docx-Hochsprache gebaut (nicht
mit rohem XML wie in ``docx_generation``): Hier entsteht eine *neue* Tabelle
statt einer Reparatur an vorhandenen Vorlagenzeilen, und eingebettete Fotos
brauchen ohnehin ``add_picture``. Rahmen und Schattierung werden über kleine
XML-Helfer gesetzt, damit die Tabellen ohne einen bestimmten Word-Tabellenstil
funktionieren — die echte HPP-Vorlage muss also keinen "Table Grid"-Stil
mitbringen.

Datenschutz-Regel dieses Moduls: Die interne Bemerkung ("Für Firmen nicht
sichtbar") wird nur ausgegeben, wenn ``data.intern`` gesetzt ist. Sie kommt
gar nicht erst im DTO an, wenn der Export für eine Firma bestimmt ist — siehe
``MangelExportEintrag`` in app.schemas und den Export-Endpunkt in
app.routers.maengel.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.config import settings
from app.services import dokumenttext
from app.schemas import MaengellisteJSON, MangelExportEintrag

# Name der Vorlage im Ordner ``backend/templates`` — beim Austausch gegen die
# echte HPP-Vorlage bleibt der Dateiname gleich.
TEMPLATE_NAME = "Maengelliste_leer.docx"

# Textplatzhalter, die überall im Dokument (Text, Tabellen, Kopf- und Fußzeile)
# ersetzt werden.
PLATZHALTER_PROJEKT = "{{PROJEKT}}"
PLATZHALTER_STAND = "{{STAND}}"
PLATZHALTER_ANZAHL = "{{ANZAHL}}"
PLATZHALTER_FILTER = "{{FILTER}}"
PLATZHALTER_HINWEIS = "{{HINWEIS}}"
PLATZHALTER_FUSSZEILE = "{{FUSSZEILE}}"
PLATZHALTER = (
    PLATZHALTER_PROJEKT,
    PLATZHALTER_STAND,
    PLATZHALTER_ANZAHL,
    PLATZHALTER_FILTER,
    PLATZHALTER_HINWEIS,
    PLATZHALTER_FUSSZEILE,
)

# Absatz-Marker, der durch die Mängelblöcke ersetzt wird.
MARKER_MAENGEL = "{{MAENGEL}}"

FONT = "Arial"
SZ_WERT = Pt(9)
SZ_LABEL = Pt(7.5)
SZ_KOPF = Pt(10.5)
GRAU = RGBColor(0x77, 0x74, 0x6B)
SCHWARZ = RGBColor(0x20, 0x1F, 0x1D)
ROT = RGBColor(0x8A, 0x2E, 0x1F)

FUELLUNG_KOPFZEILE = "F1F0EC"
FUELLUNG_INTERN = "FBF4E1"
RAHMENFARBE = "D8D5CD"

# Nutzbare Textbreite bei A4 und den Rändern der Vorlage.
TABELLE_BREITE = Cm(16.5)
LABEL_BREITE = Cm(3.6)
WERT_BREITE = Cm(12.9)

# Fotos pro Mangel im Dokument. Mehr Fotos bläht die Datei auf; wie viele
# zusätzlich in der App liegen, wird im Dokument vermerkt (siehe
# ``_fotos_zeile``) — die Liste soll nie stillschweigend kürzen.
MAX_FOTOS_PRO_MANGEL = 3
FOTO_BREITE = Cm(5.2)


def _fmt_datum(wert: date | None) -> str:
    return wert.strftime("%d.%m.%Y") if wert else ""


# ───────── XML-Helfer (Rahmen, Schattierung, Spaltenbreiten) ─────────


def _setze_rahmen(tabelle) -> None:
    """Dünne, hellgraue Linien — passend zum ruhigen Layout der Bürovorlagen."""
    tbl_pr = tabelle._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for kante in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{kante}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), RAHMENFARBE)
        borders.append(element)
    tbl_pr.append(borders)


def _setze_fuellung(zelle, farbe: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), farbe)
    zelle._tc.get_or_add_tcPr().append(shd)


def _setze_breite(zelle, breite) -> None:
    """Spaltenbreite zuverlässig setzen (Word ignoriert sonst gern die Vorgabe)."""
    zelle.width = breite
    tc_pr = zelle._tc.get_or_add_tcPr()
    for alt in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(alt)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(breite.twips)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _schreibe(absatz, text: str, *, bold: bool = False, size=SZ_WERT,
              farbe: RGBColor = SCHWARZ, klein_grau: bool = False):
    """Setzt Text mit dem Schriftbild der Bürovorlagen in einen Absatz."""
    absatz.paragraph_format.space_before = Pt(1)
    absatz.paragraph_format.space_after = Pt(1)
    lauf = absatz.add_run(dokumenttext.xml_sicher(text))
    lauf.font.name = FONT
    lauf.font.size = SZ_LABEL if klein_grau else size
    lauf.font.bold = bold
    lauf.font.color.rgb = GRAU if klein_grau else farbe
    return lauf


def _leere_zelle(zelle) -> None:
    """Zelle auf einen einzigen leeren Absatz zurücksetzen.

    Eine neu angelegte Zelle bringt bereits einen Leerabsatz mit; nach einem
    ``merge`` sind es mehrere. Ohne dieses Aufräumen entstehen im Dokument
    unschöne Leerzeilen.
    """
    for absatz in list(zelle.paragraphs)[1:]:
        absatz._p.getparent().remove(absatz._p)
    for lauf in list(zelle.paragraphs[0].runs):
        lauf._r.getparent().remove(lauf._r)


# ───────── Platzhalter ─────────


def _alle_absaetze(doc):
    """Jeden Absatz des Dokuments — Text, Tabellen, Kopf- und Fußzeilen."""
    yield from doc.paragraphs
    for tabelle in doc.tables:
        for zeile in tabelle.rows:
            for zelle in zeile.cells:
                yield from zelle.paragraphs
    for abschnitt in doc.sections:
        for teil in (abschnitt.header, abschnitt.footer,
                     abschnitt.first_page_header, abschnitt.first_page_footer):
            if teil is None:
                continue
            yield from teil.paragraphs
            for tabelle in teil.tables:
                for zeile in tabelle.rows:
                    for zelle in zeile.cells:
                        yield from zelle.paragraphs


def _ersetze_in_absatz(absatz, werte: dict[str, str]) -> None:
    if not any(schluessel in absatz.text for schluessel in werte):
        return

    # Erst versuchen, innerhalb der einzelnen Runs zu ersetzen — so bleibt die
    # Formatierung der Vorlage (Schrift, Größe, Farbe) erhalten.
    for lauf in absatz.runs:
        for schluessel, wert in werte.items():
            if schluessel in lauf.text:
                lauf.text = lauf.text.replace(
                    schluessel, dokumenttext.einzeilig(wert))

    # Ist ein Platzhalter in der Vorlage über mehrere Runs verteilt (kommt bei
    # von Hand bearbeiteten Word-Dateien vor), wird der Absatztext als Ganzes
    # ersetzt und in den ersten Run geschrieben.
    if not any(schluessel in absatz.text for schluessel in werte):
        return
    text = absatz.text
    for schluessel, wert in werte.items():
        text = text.replace(schluessel, dokumenttext.einzeilig(wert))
    if absatz.runs:
        absatz.runs[0].text = text
        for lauf in absatz.runs[1:]:
            lauf.text = ""
    else:
        _schreibe(absatz, text)


def _fuelle_platzhalter(doc, werte: dict[str, str]) -> None:
    for absatz in _alle_absaetze(doc):
        _ersetze_in_absatz(absatz, werte)


def _finde_marker(doc):
    """Der Absatz mit {{MAENGEL}} — oder ``None``, wenn die Vorlage keinen hat."""
    for absatz in doc.paragraphs:
        if MARKER_MAENGEL in absatz.text:
            return absatz
    return None


# ───────── Ein Mangel als Tabelle ─────────


def _zeilen_paare(eintrag: MangelExportEintrag,
                  intern: bool) -> list[tuple[str, str]]:
    """Label/Wert-Paare eines Mangels; leere Felder werden weggelassen."""
    ort = " · ".join(t for t in (eintrag.raumnummer, eintrag.ort) if t)
    fristen = []
    if eintrag.frist_bis:
        fristen.append(f"1. Frist: {_fmt_datum(eintrag.frist_bis)}")
    if eintrag.nachfrist_bis:
        fristen.append(f"Nachfrist: {_fmt_datum(eintrag.nachfrist_bis)}")
    if eintrag.erledigt_am:
        fristen.append(f"erledigt am {_fmt_datum(eintrag.erledigt_am)}")

    paare = [
        ("Typ / Status", " · ".join(t for t in (eintrag.typ, eintrag.status) if t)),
        ("Firma / Büro", eintrag.firma),
        ("Ort", ort),
        ("Priorität", eintrag.prioritaet),
        ("Aufgenommen", _fmt_datum(eintrag.erstellt_am)),
        ("Fristen", "   ".join(fristen)),
        ("Rückmeldung", eintrag.rueckmeldung_status),
        ("Plan-Markierung", eintrag.plan_markierung),
        ("Beschreibung", eintrag.beschreibung),
    ]
    if intern and eintrag.interne_bemerkung:
        paare.append(("Interne Bemerkung", eintrag.interne_bemerkung))
    return [(label, wert) for label, wert in paare if wert]


def _kopfzeile(tabelle, eintrag: MangelExportEintrag) -> None:
    zeile = tabelle.add_row()
    zelle = zeile.cells[0].merge(zeile.cells[1])
    _leere_zelle(zelle)
    _setze_fuellung(zelle, FUELLUNG_KOPFZEILE)
    absatz = zelle.paragraphs[0]
    _schreibe(absatz, f"{eintrag.nummer}   {eintrag.kurzbezeichnung}",
              bold=True, size=SZ_KOPF)
    if eintrag.ist_ueberfaellig:
        _schreibe(absatz, "    Frist überschritten", bold=True, farbe=ROT)


def _wertzeile(tabelle, label: str, wert: str, *, intern_hinweis: bool = False) -> None:
    zeile = tabelle.add_row()
    label_zelle, wert_zelle = zeile.cells[0], zeile.cells[1]
    _setze_breite(label_zelle, LABEL_BREITE)
    _setze_breite(wert_zelle, WERT_BREITE)

    _leere_zelle(label_zelle)
    _schreibe(label_zelle.paragraphs[0], label.upper(), klein_grau=True)

    _leere_zelle(wert_zelle)
    if intern_hinweis:
        _setze_fuellung(wert_zelle, FUELLUNG_INTERN)
    # Mehrzeilige Beschreibungen als eigene Absätze, damit Zeilenumbrüche aus
    # dem Erfassungsformular im Dokument erhalten bleiben.
    zeilen = [z for z in wert.splitlines() if z.strip()] or [""]
    _schreibe(wert_zelle.paragraphs[0], zeilen[0])
    for weitere in zeilen[1:]:
        _schreibe(wert_zelle.add_paragraph(), weitere)


def _fotos_zeile(tabelle, eintrag: MangelExportEintrag) -> None:
    """Bis zu ``MAX_FOTOS_PRO_MANGEL`` Fotos nebeneinander einbetten."""
    if not eintrag.foto_pfade:
        return

    zeile = tabelle.add_row()
    label_zelle, wert_zelle = zeile.cells[0], zeile.cells[1]
    _setze_breite(label_zelle, LABEL_BREITE)
    _setze_breite(wert_zelle, WERT_BREITE)
    _leere_zelle(label_zelle)
    _schreibe(label_zelle.paragraphs[0], "FOTOS", klein_grau=True)

    _leere_zelle(wert_zelle)
    absatz = wert_zelle.paragraphs[0]
    for pfad in eintrag.foto_pfade[:MAX_FOTOS_PRO_MANGEL]:
        try:
            lauf = absatz.add_run()
            lauf.add_picture(str(pfad), width=FOTO_BREITE)
            absatz.add_run("  ")
        except Exception:
            # Fehlende oder beschädigte Datei darf das Dokument nicht kippen.
            _schreibe(absatz, "[Foto nicht lesbar] ", klein_grau=True)

    # Nie stillschweigend kürzen: Wenn mehr Fotos vorliegen als ins Dokument
    # passen, steht das im Dokument.
    weitere = len(eintrag.foto_pfade) - MAX_FOTOS_PRO_MANGEL
    if weitere > 0:
        _schreibe(wert_zelle.add_paragraph(),
                  f"{weitere} weitere(s) Foto(s) nur in der App", klein_grau=True)


def _mangel_tabelle(doc, eintrag: MangelExportEintrag, intern: bool):
    """Baut die Tabelle eines Mangels und gibt sie zurück (noch am Dokumentende)."""
    tabelle = doc.add_table(rows=0, cols=2)
    tabelle.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabelle.autofit = False
    _setze_rahmen(tabelle)

    _kopfzeile(tabelle, eintrag)
    for label, wert in _zeilen_paare(eintrag, intern):
        _wertzeile(tabelle, label, wert,
                   intern_hinweis=(label == "Interne Bemerkung"))
    _fotos_zeile(tabelle, eintrag)
    return tabelle


# ───────── Öffentliche Funktion ─────────


def generate_maengelliste(
    data: MaengellisteJSON,
    template_path: Path | None = None,
) -> Path:
    """Erzeugt das Word-Dokument und gibt den Pfad zur Datei zurück.

    ``template_path`` überschreibt die Standardvorlage — praktisch, um die
    echte HPP-Vorlage zu testen, ohne die Datei im Repo zu tauschen.
    """
    if template_path is None:
        template_path = settings.template_dir / TEMPLATE_NAME

    doc = Document(str(template_path))

    hinweis = (
        "Interne Fassung — enthält interne Bemerkungen. Nicht an Firmen weitergeben."
        if data.intern
        else "Fassung für die ausführende Firma."
    )
    _fuelle_platzhalter(doc, {
        PLATZHALTER_PROJEKT: data.projekt,
        PLATZHALTER_STAND: _fmt_datum(data.stand),
        PLATZHALTER_ANZAHL: str(len(data.maengel)),
        PLATZHALTER_FILTER: data.filter_beschreibung or "alle Mängel",
        PLATZHALTER_HINWEIS: hinweis,
        PLATZHALTER_FUSSZEILE: f"Mängelliste - {data.projekt} - {_fmt_datum(data.stand)}",
    })

    marker = _finde_marker(doc)

    if not data.maengel:
        absatz = marker or doc.add_paragraph()
        for lauf in list(absatz.runs):
            lauf._r.getparent().remove(lauf._r)
        _schreibe(absatz, "Zu dieser Auswahl liegen keine Mängel vor.", klein_grau=True)
        marker = None

    for eintrag in data.maengel:
        tabelle = _mangel_tabelle(doc, eintrag, data.intern)
        abstand = doc.add_paragraph()
        abstand.paragraph_format.space_after = Pt(8)
        if marker is not None:
            # add_table/add_paragraph hängen am Dokumentende an; hier wandern
            # Tabelle und Abstandsabsatz an die Stelle des Markers.
            marker._p.addprevious(tabelle._tbl)
            marker._p.addprevious(abstand._p)

    if marker is not None:
        marker._p.getparent().remove(marker._p)

    sicher_projekt = "".join(
        c if c.isalnum() or c in " -_" else "_" for c in data.projekt
    ).strip().replace(" ", "_")[:40] or "Projekt"
    zusatz = "_intern" if data.intern else ""
    output_path = (
        settings.output_dir
        / f"Maengelliste_{data.stand.isoformat()}_{sicher_projekt}{zusatz}.docx"
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return output_path
