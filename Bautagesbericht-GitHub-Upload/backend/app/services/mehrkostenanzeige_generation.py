"""Erzeugt das HPP-Antwortschreiben auf eine eingegangene Anzeige.

DIE VORLAGE WIRD BEFÜLLT, NICHT NACHGEBAUT
==========================================
Anders als die Mängelanzeige (``maengelanzeige_generation``, die den Briefbogen
aus gemessenen Maßen und einer PNG-Grafik nachbaut) arbeitet dieses Modul mit
der echten Bürovorlage:

    templates/Mehrkostenanzeige_HPP_leer.docx     ("hPP Blanco.docx")

Damit sind Briefkopf, Schriften, Ränder, Falzmarke, Fußzeile und der Rahmen des
Anschriftfelds von Anfang an richtig — nichts davon steht im Code, nichts davon
kann davon abweichen. Wer den Briefbogen ändert (neue Partner, neue Anschrift),
tauscht die Vorlagendatei aus und muss hier keine Zeile anfassen.

Die Vorlage enthält an den auszufüllenden Stellen Platzhaltertexte:

    "Projektnr_Projektname"           die fette Projektzeile
    "Betreff…."                       die fette Betreffzeile
    "Sehr geehrte Damen und Herren,"  die Anrede
    "Götz Gagelmann" / "-Partner-"    Unterzeichner und Funktion

Gefunden werden sie über genau diese Texte (siehe ``PLATZHALTER``). Das ist
Absicht: Findet sich einer nicht mehr, bricht das Erzeugen mit einer klaren
Meldung ab, statt ein Schreiben ohne Betreff auszuliefern.

WARUM DIE ABSCHNITTE UMGEHÄNGT WERDEN
=====================================
Die Vorlage hat zwei Abschnitte: den Brief und dahinter einen für den
Verteiler, der den verkleinerten Briefkopf trägt (nur Wortmarke statt des
vollen Firmenblocks). Ungeändert stimmt das nur bei einem einseitigen Brief —
läuft der Text auf Seite 2, stünde dort der **volle** Briefkopf ein zweites
Mal. Deshalb:

    Brief-Abschnitt   erste Seite = voller Briefkopf (``titlePg``),
                      Folgeseiten = verkleinerter Briefkopf
    Verteiler         nur wenn es einen gibt, dann auf neuer Seite

Ohne Verteiler wird der zweite Abschnitt vollständig entfernt. Er würde sonst
eine leere Seite erzeugen — jedes Schreiben des Büros hätte eine.

DIE UNTERSCHRIFT IST EIN BILD
=============================
In der Vorlage steckt eine eingescannte Unterschrift
(``word/media/image1.jpeg``). Sie wird immer entfernt — jeder Brief geht
zum Unterschreiben aus dem Haus. Siehe ``_unterschrift_entfernen``.
"""

from __future__ import annotations

import copy
import io
import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.config import settings
from app.services import dokumenttext

VORLAGE = "Mehrkostenanzeige_HPP_leer.docx"

#: Die Platzhaltertexte der Vorlage. Werte sind die Feldnamen für Meldungen.
PLATZHALTER = {
    "projekt": "Projektnr_Projektname",
    "betreff": "Betreff….",
    "anrede": "Sehr geehrte Damen und Herren,",
    "gruss": "Mit freundlichen Grüßen",
    "firma": "HPP Architekten GmbH",
    "name": "Götz Gagelmann",
    "funktion": "-Partner-",
}

#: Einzug der LV-Zitate in Twips (1.25 cm) — wie in den Referenzbriefen.
ZITAT_EINZUG = 709

#: Zeile, die einen eingerückten Auszug einleitet. Alles danach wird zitiert,
#: bis eine Leerzeile kommt.
ZITAT_MARKEN = ("auszug lv:", "auszug lv", "auszug aus dem lv:", "lv-auszug:",
                "auszug leistungsverzeichnis:", "zitat:")


class MehrkostenanzeigeFehler(Exception):
    """Die Angaben reichen nicht für ein Schreiben."""


# ─────────────────────────────────────────────────────────────────────────────
# Feste Textbausteine
#
# Der Wortlaut ist aus den acht Referenzschreiben des Büros übernommen. Wer ihn
# ändern will, ändert ihn hier — an einer Stelle, für alle Schreiben.
# ─────────────────────────────────────────────────────────────────────────────

#: Die Haltung des Büros zur Anzeige. Sie bestimmt den Schlusssatz.
HALTUNGEN: dict[str, str] = {
    "ablehnung": "Ablehnung",
    "teilweise": "Teilweise Anerkennung",
    "pruefung": "In Prüfung",
    "anerkennung": "Anerkennung",
    "kenntnisnahme": "Nur Kenntnisnahme",
}

#: Schlusssätze je Haltung. Wortgleich zu den Referenzbriefen.
SCHLUSSSATZ: dict[str, tuple[str, ...]] = {
    "ablehnung": (
        "Aus den vorgenannten Gründen lehnen wir Ihre {art} ab.",
        "Wir bitten um Kenntnisnahme.",
    ),
    "teilweise": (
        "Über die vorstehend anerkannten Positionen hinaus sehen wir keinen "
        "Mehrvergütungsanspruch.",
        "Wir bitten um Kenntnisnahme.",
    ),
    "pruefung": (
        "Ihre {art} befindet sich derzeit in Prüfung. Wir kommen hierauf "
        "unaufgefordert zurück.",
        "Bis dahin bitten wir, die Leistungen vertragsgemäß fortzuführen.",
    ),
    "anerkennung": (
        "Wir bitten um Vorlage eines prüffähigen Nachtragsangebots auf "
        "Grundlage der vertraglichen Preise.",
    ),
    "kenntnisnahme": (
        "Wir bitten um Kenntnisnahme.",
    ),
}

#: Zusatzsatz, wenn die Anzeige eine Bauzeitverlängerung fordert und das Büro
#: ablehnt. Aus „230828 G.100-Desyum-VE300.01-MEKO 11“.
BAUZEIT_ABLEHNUNG = (
    "Die hieraus resultierende Bauzeitverlängerung kann seitens AG nicht "
    "akzeptiert werden und wird vollumfänglich abgelehnt."
)

#: Standard-Einleitung. „Nr. …“ fällt weg, wenn die Anzeige keine Nummer trug.
EINLEITUNG_MIT_NUMMER = (
    "wir haben Ihre {art} Nr. {nummer} vom {datum} erhalten und nehmen hierzu "
    "wie folgt Stellung:"
)
EINLEITUNG_OHNE_NUMMER = (
    "wir haben Ihr Schreiben vom {datum} erhalten und nehmen hierzu wie folgt "
    "Stellung:"
)

#: Betreff der E-Mail. Vorgabe des Büros: Projektname, Art, Datum der Anzeige.
MAIL_BETREFF = "{projekt}_{art} vom {datum}"

#: Text der E-Mail. Wortlaut wie vom Büro vorgegeben; das Datum ist das des
#: HPP-Schreibens, nicht das der Anzeige.
MAIL_TEXT = (
    "Sehr geehrte Damen und Herren,\n"
    "\n"
    "in der Anlage senden wir unsere Schreiben vom {datum}.\n"
)

#: Schrift der E-Mail. Outlook nimmt für reinen Text die Schrift aus den
#: eigenen Einstellungen — die ist auf jedem Rechner anders. Damit die Mail
#: überall gleich aussieht wie das Schreiben, bekommt sie einen HTML-Teil in
#: Arial 10; der Textteil bleibt als Rückfallebene daneben stehen.
MAIL_SCHRIFT = "Arial, Helvetica, sans-serif"
MAIL_SCHRIFTGROESSE = "10pt"


# ─────────────────────────────────────────────────────────────────────────────
# Eingabedaten
# ─────────────────────────────────────────────────────────────────────────────


@dataclass
class Empfaenger:
    """Die Firma, die die Anzeige geschrieben hat — sie bekommt die Antwort."""

    firma: str = ""
    #: "Herr", "Frau" oder leer. Leer heißt "Sehr geehrte Damen und Herren".
    anrede: str = ""
    ansprechpartner: str = ""
    strasse: str = ""
    plz: str = ""
    ort: str = ""
    email: str = ""


@dataclass
class Sachbearbeiter:
    """Wer im Büro unterzeichnet — füllt Datumszeile und Unterschriftsblock."""

    name: str = ""
    funktion: str = "-Baumanagement-"
    zeichen: str = ""
    durchwahl: str = ""
    email: str = ""


@dataclass
class Anzeige:
    """Die eingegangene Anzeige, soweit sie für die Antwort gebraucht wird."""

    art: str = "Mehrkostenanzeige"
    nummer: str = ""
    kennung: str = ""
    datum: date | None = None
    kurzbezeichnung: str = ""
    bauzeit: str = ""


@dataclass
class Schreiben:
    """Alles, was in das Antwortschreiben kommt."""

    empfaenger: Empfaenger
    sachbearbeiter: Sachbearbeiter
    anzeige: Anzeige

    #: Fette Projektzeile, z. B. "G.100-DESYUM-Neubau Besucherzentrum DESY".
    projektzeile: str = ""
    #: Zweite fette Zeile, z. B. "VE300.01- Erweiterter Rohbau". Darf leer sein.
    vergabeeinheit: str = ""
    #: Fette Betreffzeile. Leer = aus der Anzeige gebildet.
    betreff: str = ""

    briefdatum: date | None = None

    #: Die Stellungnahme des Büros — das Info-/Prompt-Feld der Oberfläche.
    stellungnahme: str = ""
    #: Einleitungssatz. Leer = Standard aus ``EINLEITUNG_*``.
    einleitung: str = ""
    haltung: str = "kenntnisnahme"
    #: Schlusssatz. Leer = Standard aus ``SCHLUSSSATZ``.
    schlusssatz: str = ""
    #: Die Bauzeitverlängerung ausdrücklich zurückweisen.
    bauzeit_ablehnen: bool = False

    anlagen: str = ""
    verteiler: str = ""

    #: Kürzel des Projekts für den Dateinamen, z. B. "G.100-DESYUM".
    dateikuerzel: str = ""

    hinweise: list[str] = field(default_factory=list)


# ─────────────────────────────────────────────────────────────────────────────
# Prüfen
# ─────────────────────────────────────────────────────────────────────────────


def pruefe(daten: Schreiben) -> None:
    """Wirft, wenn ein Schreiben daraus nicht verschickbar wäre."""
    fehlt: list[str] = []
    if not daten.empfaenger.firma.strip():
        fehlt.append("die Firma im Adressfeld")
    if not (daten.empfaenger.strasse.strip() or daten.empfaenger.ort.strip()):
        fehlt.append("die Anschrift der Firma")
    if not daten.projektzeile.strip():
        fehlt.append("die Projektzeile")
    if not betreffzeile(daten).strip():
        fehlt.append("der Betreff")
    if not daten.sachbearbeiter.name.strip():
        fehlt.append("der Unterzeichner")
    if not daten.stellungnahme.strip():
        fehlt.append("die Stellungnahme (Infofeld)")
    if fehlt:
        raise MehrkostenanzeigeFehler(
            "Für das Schreiben fehlt noch: " + ", ".join(fehlt) + "."
        )
    if daten.haltung not in HALTUNGEN:
        raise MehrkostenanzeigeFehler(
            f"„{daten.haltung}“ ist keine bekannte Haltung. Möglich sind: "
            + ", ".join(HALTUNGEN)
        )


# ─────────────────────────────────────────────────────────────────────────────
# Texte zusammenstellen
# ─────────────────────────────────────────────────────────────────────────────


def _dat(tag: date | None) -> str:
    return tag.strftime("%d.%m.%Y") if tag else ""


def _dat_kurz(tag: date | None) -> str:
    return tag.strftime("%d.%m.%y") if tag else ""


def anredezeile(empfaenger: Empfaenger) -> str:
    """Die Anrede. Ohne ausdrückliche Angabe bleibt sie neutral.

    Geraten wird nichts: Ob hinter „Q. Ashgarzada“ ein Herr oder eine Frau
    steht, sagt kein Schreiben verlässlich. Ein falsch angeredeter Bauleiter
    ist schlimmer als eine neutrale Anrede.
    """
    anrede = (empfaenger.anrede or "").strip().rstrip(".")
    name = (empfaenger.ansprechpartner or "").strip()
    if not anrede or not name:
        return "Sehr geehrte Damen und Herren,"
    nachname = name.split()[-1]
    if anrede.lower() in ("herr", "herrn", "hr"):
        return f"Sehr geehrter Herr {nachname},"
    if anrede.lower() in ("frau", "fr"):
        return f"Sehr geehrte Frau {nachname},"
    return "Sehr geehrte Damen und Herren,"


def betreffzeile(daten: Schreiben) -> str:
    """Die fette Betreffzeile — von Hand gesetzt oder aus der Anzeige gebildet.

    Aufbau wie in den Referenzbriefen: die **Art** der Anzeige (also
    Mehrkostenanzeige, Behinderungsanzeige, …), ihre Nummer, ihr Datum und —
    wenn bekannt — worum es geht.
    """
    if daten.betreff.strip():
        return daten.betreff.strip()

    anzeige = daten.anzeige
    art = (anzeige.art or "Schreiben").strip()
    teile = [f"Ihre {art}" if art.lower() != "schreiben" else "Ihr Schreiben"]
    if anzeige.nummer.strip():
        teile.append(f"Nr. {anzeige.nummer.strip()}")
    if anzeige.datum:
        teile.append(f"vom {_dat(anzeige.datum)}")
    zeile = " ".join(teile)
    if anzeige.kurzbezeichnung.strip():
        zeile = f"{zeile} – {anzeige.kurzbezeichnung.strip()}"
    return zeile


def einleitungssatz(daten: Schreiben) -> str:
    if daten.einleitung.strip():
        return daten.einleitung.strip()
    anzeige = daten.anzeige
    datum = _dat(anzeige.datum) or _dat(daten.briefdatum)
    if anzeige.nummer.strip():
        return EINLEITUNG_MIT_NUMMER.format(
            art=anzeige.art or "Anzeige", nummer=anzeige.nummer.strip(), datum=datum
        )
    return EINLEITUNG_OHNE_NUMMER.format(datum=datum)


def schlusssaetze(daten: Schreiben) -> list[str]:
    if daten.schlusssatz.strip():
        return [z.strip() for z in daten.schlusssatz.strip().split("\n") if z.strip()]
    art = daten.anzeige.art or "Anzeige"
    return [satz.format(art=art) for satz in SCHLUSSSATZ[daten.haltung]]


# ─────────────────────────────────────────────────────────────────────────────
# Angaben im Infofeld
#
# Das Infofeld ist die Stelle, an der man beim Schreiben ist. Wer dort merkt,
# dass noch eine Anlage fehlt, soll nicht nach oben ins Formular springen
# müssen — er schreibt "Anlage: Auszug LV" in eine eigene Zeile, und die
# Angabe landet auf der Verteilerseite statt im Brieftext.
#
# Erkannt wird nur eine Zeile, die **mit** der Beschriftung beginnt und einen
# Doppelpunkt hat. Ein Satz wie "Die Anlage 3 haben wir geprüft." bleibt
# deshalb Brieftext — sonst verschwänden ganze Sätze aus dem Schreiben.
# ─────────────────────────────────────────────────────────────────────────────

#: Feldname -> Beschriftungen, die dorthin führen. Die längste Schreibweise
#: zuerst, damit "PLZ Ort" nicht an "Ort" hängen bleibt.
ANGABEN: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("anlagen", ("Anlagen", "Anlage")),
    ("verteiler", ("Verteiler", "Verteilung", "Kopie")),
    ("betreff", ("Betreff", "Betr")),
    ("einleitung", ("Einleitung",)),
    ("schlusssatz", ("Schlusssatz", "Schluss")),
    ("haltung", ("Haltung",)),
    ("projektzeile", ("Projektzeile", "Projekt")),
    ("vergabeeinheit", ("Vergabeeinheit", "VE")),
    ("briefdatum", ("Briefdatum", "Datum")),
    # Auch ohne Umlaut: Nicht jeder tippt "Kürzel", und ein nicht
    # erkanntes Kürzel stuende als Zeile mitten im Brief.
    ("dateikuerzel", ("Dateikürzel", "Dateikuerzel", "Kürzel", "Kuerzel")),
    ("firma", ("Firma",)),
    ("anrede", ("Anrede",)),
    ("ansprechpartner", ("Ansprechpartner", "z. Hd.", "z.Hd.", "zHd")),
    # Zwei Mailadressen kommen in diesem Brief vor, und sie zu verwechseln
    # wäre der schlimmste Fehler dieser Liste: Die Adresse der Firma steht im
    # Adressfeld und ist der Empfänger des Outlook-Entwurfs, die des Büros in
    # der Datumszeile. Deshalb muss die eigene ausdrücklich "HPP" heißen —
    # der schlichte Fall gehört der Gegenseite, denn an die schreibt man.
    ("email", ("E-Mail Firma", "Mail Firma", "E-Mail", "EMail", "Mail")),
    ("email_hpp", ("E-Mail HPP", "Mail HPP", "E-Mail Absender")),
    ("strasse", ("Straße", "Strasse")),
    ("plz_ort", ("PLZ Ort", "PLZ/Ort", "PLZ und Ort", "PLZ", "Ort")),
    ("unterzeichner", ("Unterzeichner",)),
    ("funktion", ("Funktion",)),
    ("zeichen", ("Zeichen", "Ze")),
    ("durchwahl", ("Durchwahl", "Telefon")),
)

#: Klartext der Felder für die Rückmeldung. ``feld.capitalize()`` ergäbe
#: "Strasse" und "Email" — in einer Meldung an den Anwender ist das schlampig.
FELD_KLARTEXT: dict[str, str] = {
    "anlagen": "Anlage",
    "verteiler": "Verteiler",
    "betreff": "Betreff",
    "einleitung": "Einleitung",
    "schlusssatz": "Schlusssatz",
    "haltung": "Haltung",
    "projektzeile": "Projektzeile",
    "vergabeeinheit": "Vergabeeinheit",
    "briefdatum": "Datum des Schreibens",
    "dateikuerzel": "Dateikürzel",
    "firma": "Firma",
    "anrede": "Anrede",
    "ansprechpartner": "Ansprechpartner",
    "email": "E-Mail der Firma",
    "email_hpp": "E-Mail HPP",
    "strasse": "Straße",
    "plz_ort": "PLZ und Ort",
    "unterzeichner": "Unterzeichner",
    "funktion": "Funktion",
    "zeichen": "Zeichen",
    "durchwahl": "Durchwahl",
}

#: Felder, bei denen mehrere Zeilen sich sammeln statt sich zu überschreiben.
#: Anlagen und Verteiler sind Listen; alles andere hat genau einen Wert.
SAMMELND = {"anlagen", "verteiler"}

#: Zeilen ohne Doppelpunkt, die einen Schalter umlegen.
SCHALTER: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("bauzeit_ablehnen", (
        "bauzeit ablehnen", "bauzeitverlängerung ablehnen",
        "bauzeitverlaengerung ablehnen", "bauzeit zurückweisen",
    )),
)

#: Wie ein hingeschriebenes Wort zur Haltung wird.
HALTUNG_WORTE: dict[str, str] = {
    "ablehnung": "ablehnung", "ablehnen": "ablehnung", "abgelehnt": "ablehnung",
    "teilweise": "teilweise", "teilweise anerkennung": "teilweise",
    "prüfung": "pruefung", "pruefung": "pruefung", "in prüfung": "pruefung",
    "in pruefung": "pruefung", "prüfen": "pruefung",
    "anerkennung": "anerkennung", "anerkennen": "anerkennung",
    "anerkannt": "anerkennung",
    "kenntnisnahme": "kenntnisnahme", "kenntnis": "kenntnisnahme",
}


def vorbereiten(daten: Schreiben) -> list[str]:
    """Angaben aus dem Infofeld in die Felder des Schreibens übernehmen.

    Gibt zurück, was übernommen wurde — die Oberfläche zeigt das an. Ohne
    diese Rückmeldung wäre es Zauberei: Eine Zeile verschwindet aus dem
    Infofeld und taucht irgendwo im Brief wieder auf, und niemand weiß, ob sie
    angekommen ist.

    **Idempotent.** Die erkannten Zeilen werden aus der Stellungnahme
    entfernt; ein zweiter Aufruf findet nichts mehr. Deshalb dürfen Vorschau
    und Erzeugen beide aufrufen, ohne sich abzustimmen — und beide sehen
    garantiert dasselbe Schreiben.
    """
    if not daten.stellungnahme.strip():
        return []

    uebernommen: list[str] = []
    rest: list[str] = []
    gesammelt: dict[str, list[str]] = {}
    schon_gesetzt: dict[str, str] = {}

    for zeile in dokumenttext.zeilen(daten.stellungnahme):
        schalter = _welcher_schalter(zeile)
        if schalter is not None:
            setattr(daten, schalter, True)
            uebernommen.append(_schalter_klartext(schalter))
            continue

        erkannt = _welche_angabe(zeile)
        if erkannt is None:
            rest.append(zeile)
            continue

        feld, wert = erkannt
        if not wert:
            # "Anlage:" ohne Wert ist ein Tippfehler, kein Befehl.
            rest.append(zeile)
            continue
        if feld in SAMMELND:
            gesammelt.setdefault(feld, []).append(wert)
            continue

        # Zweimal dieselbe Angabe: Der letzte Wert gewinnt — aber still darf
        # das nicht passieren. Wer "Betreff:" zweimal schreibt, hat sich
        # vertippt oder eine Zeile vergessen zu löschen.
        if feld in schon_gesetzt and schon_gesetzt[feld] != wert:
            uebernommen.append(
                f"{FELD_KLARTEXT.get(feld, feld)} stand zweimal im Infofeld "
                f"(„{schon_gesetzt[feld]}“ und „{wert}“) — genommen wurde das "
                f"Letzte."
            )
        schon_gesetzt[feld] = wert
        uebernommen.append(_setze_angabe(daten, feld, wert))

    for feld, werte in gesammelt.items():
        vorhandene = [z for z in (getattr(daten, feld) or "").split("\n") if z.strip()]
        setattr(daten, feld, "\n".join(vorhandene + werte))
        uebernommen.append(
            f"{FELD_KLARTEXT[feld]}: " + ", ".join(f"„{w}“" for w in werte)
        )

    daten.stellungnahme = re.sub(r"\n{3,}", "\n\n", "\n".join(rest)).strip("\n")
    return uebernommen


def _welche_angabe(zeile: str) -> tuple[str, str] | None:
    """Feld und Wert einer Beschriftungszeile, sonst ``None``."""
    for feld, woerter in ANGABEN:
        for wort in woerter:
            muster = re.compile(
                r"^\s*" + re.escape(wort) + r"\s*:\s*(.*)$", re.IGNORECASE
            )
            treffer = muster.match(zeile)
            if treffer:
                return feld, treffer.group(1).strip()
    return None


def _welcher_schalter(zeile: str) -> str | None:
    schlicht = zeile.strip().rstrip(".!").lower()
    for feld, woerter in SCHALTER:
        if schlicht in woerter:
            return feld
    return None


def _schalter_klartext(feld: str) -> str:
    return {
        "bauzeit_ablehnen": "Bauzeitverlängerung wird zurückgewiesen",
    }.get(feld, feld)


def _setze_angabe(daten: Schreiben, feld: str, wert: str) -> str:
    """Einen einzelnen Wert an seinen Platz schreiben, und melden wohin."""
    sauber = dokumenttext.einzeilig(wert)

    if feld == "haltung":
        gewaehlt = HALTUNG_WORTE.get(sauber.lower())
        if gewaehlt is None:
            return (
                f"Haltung „{sauber}“ nicht erkannt — es bleibt bei "
                f"„{HALTUNGEN[daten.haltung]}“. Möglich: "
                + ", ".join(sorted(set(HALTUNG_WORTE.values())))
            )
        daten.haltung = gewaehlt
        return f"Haltung: {HALTUNGEN[gewaehlt]}"

    if feld == "briefdatum":
        from app.services.mehrkostenanzeige_lesen import _erstes_datum

        gelesen = _erstes_datum(sauber)
        if gelesen is None:
            return f"Datum „{sauber}“ nicht lesbar — es bleibt beim bisherigen."
        daten.briefdatum = gelesen
        return f"Datum des Schreibens: {_dat(gelesen)}"

    if feld == "plz_ort":
        teile = sauber.split(None, 1)
        if len(teile) == 2 and teile[0].isdigit():
            daten.empfaenger.plz, daten.empfaenger.ort = teile[0], teile[1]
        else:
            daten.empfaenger.ort = sauber
        return f"{FELD_KLARTEXT[feld]}: {sauber}"

    if feld == "anrede":
        kurz = sauber.rstrip(".").lower()
        if kurz in ("herr", "herrn", "hr"):
            daten.empfaenger.anrede = "Herr"
        elif kurz in ("frau", "fr"):
            daten.empfaenger.anrede = "Frau"
        else:
            daten.empfaenger.anrede = ""
            return f"Anrede „{sauber}“ nicht erkannt — es bleibt bei neutral."
        return f"Anrede: {daten.empfaenger.anrede}"

    # Die übrigen Felder liegen entweder am Schreiben, am Empfänger oder am
    # Sachbearbeiter. Diese Zuordnung steht hier, damit es nur eine gibt.
    klartext = FELD_KLARTEXT.get(feld, feld)
    if feld in ("firma", "ansprechpartner", "strasse", "email"):
        setattr(daten.empfaenger, feld, sauber)
        return f"{klartext}: {sauber}"
    if feld in ("zeichen", "durchwahl"):
        setattr(daten.sachbearbeiter, feld, sauber)
        return f"{klartext}: {sauber}"
    if feld == "email_hpp":
        daten.sachbearbeiter.email = sauber
        return f"{klartext}: {sauber}"
    if feld == "unterzeichner":
        daten.sachbearbeiter.name = sauber
        return f"{klartext}: {sauber}"
    if feld == "funktion":
        daten.sachbearbeiter.funktion = sauber
        return f"{klartext}: {sauber}"

    setattr(daten, feld, sauber)
    return f"{klartext}: {sauber}"


@dataclass
class Block:
    """Ein Absatz des Briefkörpers — normal oder eingerückt (LV-Zitat)."""

    text: str
    #: Eingerückt, weil es ein zitierter LV-Text ist.
    zitat: bool = False
    #: Die einleitende Zeile eines Zitats ("Auszug LV:"). Sie ist nicht
    #: eingerückt, gehört aber ohne Leerzeile zum Zitat darunter.
    marke: bool = False


def koerper(daten: Schreiben) -> list[Block]:
    """Der Briefkörper: Einleitung, Stellungnahme, Bauzeit, Schluss.

    Die Stellungnahme kommt Zeile für Zeile aus dem Infofeld. Zwei Dinge
    werden dabei erkannt, weil das Büro sie in jedem Schreiben so setzt:

        "1) …" / "1. …" / "- …"   bleibt eine eigene Zeile (Absatz)
        "Auszug LV:"              rückt alles Folgende ein, bis zur Leerzeile

    Sonst wird nichts umgeschrieben. Was im Infofeld steht, steht im Brief —
    Wort für Wort. Ein Textbaustein, der den Sinn verschiebt, wäre in einem
    rechtserheblichen Schreiben nicht zu verantworten.
    """
    bloecke: list[Block] = [Block(einleitungssatz(daten))]
    bloecke += _stellungnahme_bloecke(daten.stellungnahme)

    if daten.bauzeit_ablehnen:
        bloecke.append(Block(BAUZEIT_ABLEHNUNG))
    for satz in schlusssaetze(daten):
        bloecke.append(Block(satz))
    return bloecke


def _stellungnahme_bloecke(text: str) -> list[Block]:
    zeilen = dokumenttext.zeilen(text)
    bloecke: list[Block] = []
    puffer: list[str] = []
    im_zitat = False

    def puffer_ablegen() -> None:
        if puffer:
            bloecke.append(Block(" ".join(puffer).strip(), zitat=False))
            puffer.clear()

    for rohzeile in zeilen:
        zeile = rohzeile.strip()

        if not zeile:
            puffer_ablegen()
            im_zitat = False
            continue

        if zeile.lower().rstrip(":") + ":" in ZITAT_MARKEN or \
                zeile.lower() in ZITAT_MARKEN:
            puffer_ablegen()
            bloecke.append(Block(zeile, marke=True))
            im_zitat = True
            continue

        if im_zitat:
            bloecke.append(Block(zeile, zitat=True))
            continue

        # Eine Aufzählung beginnt einen neuen Absatz, läuft aber weiter, wenn
        # der Text der Firma über mehrere Zeilen geht.
        if re.match(r"^(?:\d{1,2}[.)]|[-–•*])\s+", zeile):
            puffer_ablegen()
            puffer.append(zeile)
            continue

        puffer.append(zeile)

    puffer_ablegen()
    return [b for b in bloecke if b.text]


# ─────────────────────────────────────────────────────────────────────────────
# Dateiname, Mail
# ─────────────────────────────────────────────────────────────────────────────


def _sauber(text: str) -> str:
    """Für Dateinamen: alles weg, womit Windows nicht umgehen kann."""
    ohne = re.sub(r'[\\/:*?"<>|\r\n\t]', "-", text)
    return re.sub(r"\s{2,}", " ", ohne).strip(" .-")


def dateiname(daten: Schreiben) -> str:
    """Wie im Büro: "260903 G.100-DESYUM-VE300.01-MKA 01.docx".

    Aufbau JJMMTT, Projektkürzel, Vergabeeinheit, Kennung der Anzeige —
    dieselbe Reihenfolge wie in allen Referenzschreiben. Fehlt ein Teil,
    fällt er samt Trennstrich weg.
    """
    tag = daten.briefdatum or date.today()
    teile = [
        _sauber(daten.dateikuerzel),
        _sauber(_ve_code(daten.vergabeeinheit)),
        _sauber(daten.anzeige.kennung or daten.anzeige.art),
    ]
    rumpf = "-".join(t for t in teile if t)
    return f"{tag.strftime('%y%m%d')} {rumpf}.docx" if rumpf \
        else f"{tag.strftime('%y%m%d')} Anzeige.docx"


def _ve_code(vergabeeinheit: str) -> str:
    """Nur die Kennung der Vergabeeinheit, ohne ihre Bezeichnung.

    "VE300.01- Erweiterter Rohbau" wird zu "VE300.01" — im Dateinamen steht im
    Büro nur die Kennung.
    """
    text = vergabeeinheit.strip()
    if not text:
        return ""
    treffer = re.match(r"^([A-Za-z]{0,3}[\d.\-_/]*\d)", text)
    return treffer.group(1).rstrip("-.") if treffer else text.split()[0]


def mail_betreff(daten: Schreiben) -> str:
    """Betreff der E-Mail: Projektname, Art der Anzeige, Datum der Anzeige."""
    projekt = _projektname(daten.projektzeile)
    art = (daten.anzeige.art or "Anzeige").strip()
    datum = _dat_kurz(daten.anzeige.datum) or _dat_kurz(daten.briefdatum)
    return MAIL_BETREFF.format(projekt=projekt, art=art, datum=datum).strip()


def _projektname(projektzeile: str) -> str:
    """Der Projektname aus der fetten Projektzeile.

    Die Zeile ist im Büro "Nummer_Name" oder "Nummer-Name" aufgebaut
    ("G.100-DESYUM-Neubau Besucherzentrum DESY"). Für den Mailbetreff ist der
    Name gemeint; ohne erkennbare Trennung bleibt die ganze Zeile stehen.
    """
    zeile = projektzeile.strip()
    if "_" in zeile:
        return zeile.split("_", 1)[1].strip() or zeile
    return zeile


def mail_text(daten: Schreiben) -> str:
    return MAIL_TEXT.format(datum=_dat_kurz(daten.briefdatum or date.today()))


def mail_html(text: str) -> str:
    """Der Mailtext als HTML in Arial 10.

    Warum überhaupt HTML: Eine reine Textmail hat keine Schrift — Outlook
    nimmt die aus den Einstellungen des Absenders, und die ist auf jedem
    Rechner eine andere. Ein HTML-Teil legt Arial 10 fest, so wie im
    Schreiben. Der Textteil bleibt daneben stehen, damit die Mail auch dort
    lesbar ist, wo HTML abgeschaltet ist.

    Escaped wird alles: Ein Firmenname mit "&" oder ein Betreff mit "<" darf
    die Mail nicht zerlegen.
    """
    from html import escape

    absaetze = [a for a in text.replace("\r\n", "\n").split("\n\n")]
    rumpf = "".join(
        f'<p style="margin:0 0 {MAIL_SCHRIFTGROESSE} 0">'
        + "<br>".join(escape(zeile) for zeile in absatz.split("\n"))
        + "</p>"
        for absatz in absaetze
    )
    return (
        f'<html><body><div style="font-family:{MAIL_SCHRIFT};'
        f'font-size:{MAIL_SCHRIFTGROESSE}">{rumpf}</div></body></html>'
    )


# ─────────────────────────────────────────────────────────────────────────────
# Word: kleine Werkzeuge am XML
# ─────────────────────────────────────────────────────────────────────────────


def _absatz_mit(dok, text: str):
    """Der erste Absatz mit genau diesem Text (getrimmt), sonst ``None``."""
    for absatz in dok.paragraphs:
        if absatz.text.strip() == text.strip():
            return absatz
    return None


def _setze_text(absatz, text: str) -> None:
    """Text eines Absatzes ersetzen, Absatz- und Zeichenformat behalten.

    Behalten wird das ``rPr`` des ersten Laufs — daran hängen Fettschrift und
    Kursivstellung der Vorlage. Alle anderen Läufe fallen weg; in der Vorlage
    ist ein Satz oft in fünf Läufe zerlegt (Bearbeitungsspuren), und nur der
    erste trägt das Format.
    """
    element = absatz._p
    rpr = None
    for lauf in element.findall(qn("w:r")):
        if rpr is None:
            gefunden = lauf.find(qn("w:rPr"))
            if gefunden is not None:
                rpr = copy.deepcopy(gefunden)
        element.remove(lauf)
    # Korrekturmarken der Rechtschreibprüfung zeigen sonst auf gelöschte Läufe.
    for marke in element.findall(qn("w:proofErr")):
        element.remove(marke)

    if not text:
        return
    lauf = element.makeelement(qn("w:r"), {})
    if rpr is not None:
        lauf.append(rpr)
    knoten = element.makeelement(qn("w:t"), {})
    knoten.text = text
    knoten.set(qn("xml:space"), "preserve")
    lauf.append(knoten)
    element.append(lauf)


def _absatz_nach(hinter, text: str = "", *, format_von=None,
                 einzug: int | None = None):
    """Neuen Absatz hinter ``hinter`` einfügen.

    Das Kopieren des ``pPr`` ist der Kern der Sache: Am Anschriftfeld hängt
    dort der Textrahmen (``framePr``). Ein Absatz ohne dieses ``pPr`` würde
    aus dem Rahmen fallen und mitten im Brief landen.

    ``format_von`` trennt die *Stelle* von der *Form*. Ohne diese Trennung
    erbt jeder Absatz die Form seines Vorgängers, und ein einziges
    eingerücktes LV-Zitat zieht den ganzen Rest des Briefs mit in den Einzug —
    genau das ist beim ersten Probedruck passiert.
    """
    stelle = hinter._p if hasattr(hinter, "_p") else hinter
    vorbild = format_von if format_von is not None else hinter
    vorbild_p = vorbild._p if hasattr(vorbild, "_p") else vorbild

    neu = stelle.makeelement(qn("w:p"), {})
    ppr = vorbild_p.find(qn("w:pPr"))
    if ppr is not None:
        kopie = copy.deepcopy(ppr)
        # Ein Abschnittsumbruch gehört genau einem Absatz — mitkopiert würde
        # er das Dokument in zwei Hälften schneiden.
        for umbruch in kopie.findall(qn("w:sectPr")):
            kopie.remove(umbruch)
        neu.append(kopie)
    stelle.addnext(neu)

    from docx.text.paragraph import Paragraph

    absatz = Paragraph(neu, getattr(hinter, "_parent", None))
    if einzug is not None:
        _einzug_setzen(absatz, einzug)
    if text:
        _setze_text(absatz, text)
    return absatz


def _einzug_setzen(absatz, twips: int) -> None:
    element = absatz._p
    ppr = element.find(qn("w:pPr"))
    if ppr is None:
        ppr = element.makeelement(qn("w:pPr"), {})
        element.insert(0, ppr)
    for alt in ppr.findall(qn("w:ind")):
        ppr.remove(alt)
    ind = ppr.makeelement(qn("w:ind"), {})
    ind.set(qn("w:left"), str(twips))
    # Nach dem pStyle, vor dem rPr — die Reihenfolge im pPr ist im Schema fest.
    _in_ppr_einsortieren(ppr, ind, "w:ind")


#: Reihenfolge der Kindelemente eines ``w:pPr`` nach dem OOXML-Schema, soweit
#: hier gebraucht. Wird sie verletzt, öffnet Word die Datei nicht.
_PPR_REIHENFOLGE = (
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
    "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
    "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku",
    "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
    "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
    "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
    "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
    "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
    "w:rPr", "w:sectPr", "w:pPrChange",
)


def _in_ppr_einsortieren(ppr, element, name: str) -> None:
    stelle = _PPR_REIHENFOLGE.index(name)
    for kind in ppr:
        kurz = _kurzname(kind)
        if kurz in _PPR_REIHENFOLGE and _PPR_REIHENFOLGE.index(kurz) > stelle:
            kind.addprevious(element)
            return
    ppr.append(element)


def _kurzname(element) -> str:
    text = element.tag
    if text.startswith("{"):
        raum, _, rest = text[1:].partition("}")
        if raum.endswith("wordprocessingml/2006/main"):
            return f"w:{rest}"
    return text


def _loeschen(absatz) -> None:
    element = absatz._p if hasattr(absatz, "_p") else absatz
    eltern = element.getparent()
    if eltern is not None:
        eltern.remove(element)


def _zelle_setzen(zelle, text: str) -> None:
    """Text in eine Tabellenzelle, Format der Vorlage behalten."""
    absaetze = zelle.paragraphs
    _setze_text(absaetze[0], text)
    for weiterer in absaetze[1:]:
        _loeschen(weiterer)


# ─────────────────────────────────────────────────────────────────────────────
# Word: Abschnitte und Briefkopf
# ─────────────────────────────────────────────────────────────────────────────

#: Beziehungskennungen der Vorlage. Sie stehen fest, weil die Vorlage
#: mitgeliefert wird — siehe word/_rels/document.xml.rels.
KOPF_VOLL = "rId13"        # header2.xml — voller Briefkopf
KOPF_FOLGE = "rId18"       # header4.xml — nur Wortmarke
FUSS_SEITEN = "rId15"      # footer2.xml — Dateiname und Seitenzahl


def _briefkopf_regeln(dok, mit_verteiler: bool) -> None:
    """Erste Seite voller Briefkopf, Folgeseiten verkleinert.

    Ohne diesen Eingriff trägt jede Folgeseite des Briefs den vollen
    Firmenblock ein zweites Mal — in der Vorlage ist er der Standardkopf des
    Briefabschnitts, und ``titlePg`` ist nicht gesetzt.
    """
    abschnitte = dok.sections
    brief = abschnitte[0]._sectPr

    _bezug_setzen(brief, "w:headerReference", "first", KOPF_VOLL)
    _bezug_setzen(brief, "w:headerReference", "default", KOPF_FOLGE)
    # Die Seitenzahl gehört auch auf Seite 1 — die Vorlage hat dort sonst die
    # leere Fußzeile footer3.
    _bezug_setzen(brief, "w:footerReference", "first", FUSS_SEITEN)
    _titelseite_setzen(brief)

    if len(abschnitte) < 2:
        return

    if mit_verteiler:
        # Der Verteiler beginnt auf einer neuen Seite. "continuous" hinge
        # davon ab, wie voll Seite 1 gerade ist — bei kurzen Briefen stünde
        # der Verteiler unter der Unterschrift, bei langen auf Seite 2.
        _abschnittstyp(abschnitte[1]._sectPr, "nextPage")
        return

    # Ohne Verteiler wird der zweite Abschnitt aufgelöst: Seitenaufbau und
    # Briefkopf des Briefs gelten dann für das ganze Dokument. Sonst bliebe
    # eine leere Seite übrig.
    letzter = abschnitte[1]._sectPr
    letzter.getparent().replace(letzter, copy.deepcopy(brief))
    traeger = _umbruch_absatz(dok)
    if traeger is not None:
        traeger.getparent().remove(traeger)


def _bezug_setzen(sectpr, name: str, art: str, kennung: str) -> None:
    """Kopf-/Fußzeilenbezug einer Art setzen oder anlegen."""
    for bezug in sectpr.findall(qn(name)):
        if bezug.get(qn("w:type")) == art:
            bezug.set(qn("r:id"), kennung)
            return
    neu = sectpr.makeelement(qn(name), {})
    neu.set(qn("w:type"), art)
    neu.set(qn("r:id"), kennung)
    # Kopf- und Fußzeilenbezüge stehen im Schema ganz vorn im sectPr.
    letzter = sectpr.findall(qn("w:headerReference")) + \
        sectpr.findall(qn("w:footerReference"))
    if letzter:
        letzter[-1].addnext(neu)
    else:
        sectpr.insert(0, neu)


def _titelseite_setzen(sectpr) -> None:
    if sectpr.find(qn("w:titlePg")) is not None:
        return
    marke = sectpr.makeelement(qn("w:titlePg"), {})
    # titlePg steht im Schema hinter formProt/vAlign und vor docGrid; am Ende
    # anzuhängen ist bei dieser Vorlage die richtige Stelle.
    sectpr.append(marke)


def _abschnittstyp(sectpr, wert: str) -> None:
    for alt in sectpr.findall(qn("w:type")):
        sectpr.remove(alt)
    typ = sectpr.makeelement(qn("w:type"), {})
    typ.set(qn("w:val"), wert)
    kopfbezuege = sectpr.findall(qn("w:headerReference")) + \
        sectpr.findall(qn("w:footerReference"))
    if kopfbezuege:
        kopfbezuege[-1].addnext(typ)
    else:
        sectpr.insert(0, typ)


def _unterschrift_entfernen(absatz) -> None:
    """Die eingescannte Unterschrift der Vorlage aus dem Brief nehmen.

    In der Vorlage steckt ein Unterschriftsbild (``word/media/image1.jpeg``).
    Es kommt in kein erzeugtes Schreiben: Jeder Brief geht zum Unterschreiben
    aus dem Haus — eine gedruckte Unterschrift unter einem Text, der
    Vergütungsansprüche abwehrt, setzt niemand ungefragt.
    """
    for lauf in absatz._p.findall(qn("w:r")):
        for zeichnung in lauf.findall(qn("w:drawing")):
            lauf.remove(zeichnung)


def _fussnote_dateiname(dok, name: str) -> None:
    """Den Dateinamen in die Fußzeile setzen — wie in jedem Bürodokument.

    Die Vorlage hat dort nur die Seitenzahl. Der Dateiname davor ist die
    Wiedererkennung: Auf jedem Ausdruck im Ordner steht, welche Datei es war.
    """
    fuss = dok.sections[0].footer
    if not fuss.paragraphs:
        return
    absatz = fuss.paragraphs[0]
    element = absatz._p
    vorhandene = element.findall(qn("w:r"))
    rpr = None
    if vorhandene:
        gefunden = vorhandene[0].find(qn("w:rPr"))
        if gefunden is not None:
            rpr = copy.deepcopy(gefunden)

    lauf = element.makeelement(qn("w:r"), {})
    if rpr is not None:
        lauf.append(rpr)
    knoten = element.makeelement(qn("w:t"), {})
    knoten.text = name
    knoten.set(qn("xml:space"), "preserve")
    lauf.append(knoten)
    if vorhandene:
        vorhandene[0].addprevious(lauf)
    else:
        element.append(lauf)


# ─────────────────────────────────────────────────────────────────────────────
# Erzeugen
# ─────────────────────────────────────────────────────────────────────────────


def adressblock(daten: Schreiben) -> list[str]:
    """Das Anschriftfeld, Zeile für Zeile — Aufbau der neueren Briefe.

        Rolfes Bau GmbH & Co. KG
        Frau Braun
        Krumbäken Kämpe 2
        29424 Goldenstedt
        (leer)
        per E-Mail:
        R.Braun@rolfes-gruppe.de
    """
    empfaenger = daten.empfaenger
    zeilen = [
        empfaenger.firma.strip(),
        _person(empfaenger),
        empfaenger.strasse.strip(),
    ]
    plz_ort = f"{empfaenger.plz.strip()} {empfaenger.ort.strip()}".strip()
    if plz_ort:
        zeilen.append(plz_ort)
    zeilen = [z for z in zeilen if z]

    if empfaenger.email.strip():
        zeilen += ["", "per E-Mail:", empfaenger.email.strip()]
    return zeilen


def _person(empfaenger: Empfaenger) -> str:
    """"Herrn Steffen Wegner" / "Frau Braun" / der Name allein."""
    name = (empfaenger.ansprechpartner or "").strip()
    if not name:
        return ""
    anrede = (empfaenger.anrede or "").strip().rstrip(".").lower()
    if anrede in ("herr", "herrn", "hr"):
        return f"Herrn {name}"
    if anrede in ("frau", "fr"):
        return f"Frau {name}"
    return name


def erzeuge(daten: Schreiben, vorlage: Path | None = None) -> bytes:
    """Das fertige Word-Dokument als Bytes."""
    # Zuerst die Angaben aus dem Infofeld an ihren Platz — sonst prüfte die
    # Prüfung Felder, die gleich danach gefüllt werden, und ein Schreiben mit
    # "Projekt: ..." im Infofeld gälte als unvollständig.
    daten.hinweise += vorbereiten(daten)
    pruefe(daten)

    pfad = vorlage or (settings.template_dir / VORLAGE)
    if not pfad.is_file():
        raise MehrkostenanzeigeFehler(
            f"Die Briefvorlage „{pfad.name}“ fehlt im Vorlagenordner "
            f"({pfad.parent}). Ohne sie kann kein Schreiben erzeugt werden."
        )
    dok = Document(str(pfad))

    _fuelle_adressblock(dok, daten)
    _fuelle_datumszeile(dok, daten)
    _fuelle_kopfzeilen(dok, daten)
    _fuelle_koerper(dok, daten)
    _fuelle_unterschrift(dok, daten)

    anlagen = dokumenttext.zeilen(daten.anlagen)
    verteiler = dokumenttext.zeilen(daten.verteiler)
    _fuelle_verteiler(dok, anlagen, verteiler)

    _briefkopf_regeln(dok, mit_verteiler=bool(anlagen or verteiler))
    _fussnote_dateiname(dok, dateiname(daten))

    puffer = io.BytesIO()
    dok.save(puffer)
    return puffer.getvalue()


def _fuelle_adressblock(dok, daten: Schreiben) -> None:
    zeilen = [dokumenttext.einzeilig(z) for z in adressblock(daten)]
    vorhandene = [p for p in dok.paragraphs if p.style.name == "Adressblock"]
    if not vorhandene:
        raise MehrkostenanzeigeFehler(
            "In der Briefvorlage fehlt das Anschriftfeld (Absatzformat "
            "„Adressblock“). Wurde die Vorlage ausgetauscht?"
        )

    letzter = vorhandene[0]
    _setze_text(letzter, zeilen[0] if zeilen else "")
    for zeile in zeilen[1:]:
        letzter = _absatz_nach(letzter, zeile)
    # Die Vorlage hat zwei leere Zeilen im Rahmen; überzählige weg.
    for ueberzaehlig in vorhandene[1:]:
        _loeschen(ueberzaehlig)


def _fuelle_datumszeile(dok, daten: Schreiben) -> None:
    """Datum, Zeichen, Durchwahl und Mailadresse des Bearbeiters."""
    if not dok.tables:
        raise MehrkostenanzeigeFehler(
            "In der Briefvorlage fehlt die Datumszeile. Wurde die Vorlage "
            "ausgetauscht?"
        )
    zellen = dok.tables[0].rows[0].cells
    bearbeiter = daten.sachbearbeiter
    werte = [
        _dat(daten.briefdatum or date.today()),
        f"Ze: {bearbeiter.zeichen.strip()}" if bearbeiter.zeichen.strip() else "",
        f"T - {bearbeiter.durchwahl.strip()}" if bearbeiter.durchwahl.strip() else "",
        bearbeiter.email.strip(),
    ]
    for zelle, wert in zip(zellen, werte):
        _zelle_setzen(zelle, dokumenttext.einzeilig(wert))


def _fuelle_kopfzeilen(dok, daten: Schreiben) -> None:
    """Projektzeile, Vergabeeinheit und Betreff — die drei fetten Zeilen."""
    projekt = _absatz_mit(dok, PLATZHALTER["projekt"])
    betreff = _absatz_mit(dok, PLATZHALTER["betreff"])
    if projekt is None or betreff is None:
        raise MehrkostenanzeigeFehler(
            f"In der Briefvorlage fehlen die Platzhalter "
            f"„{PLATZHALTER['projekt']}“ und „{PLATZHALTER['betreff']}“. "
            f"Wurde die Vorlage ausgetauscht?"
        )

    _setze_text(projekt, dokumenttext.einzeilig(daten.projektzeile))
    if daten.vergabeeinheit.strip():
        _absatz_nach(projekt, dokumenttext.einzeilig(daten.vergabeeinheit))
    _setze_text(betreff, dokumenttext.einzeilig(betreffzeile(daten)))


def _fuelle_koerper(dok, daten: Schreiben) -> None:
    """Anrede und Briefkörper zwischen Anrede und Grußformel.

    Die Vorlage hat dort vier leere Absätze als Platzhalter. Sie fliegen
    heraus und werden durch die Absätze mit je einer Leerzeile dazwischen
    ersetzt — so, wie es in den Referenzbriefen aussieht.
    """
    anrede = _absatz_mit(dok, PLATZHALTER["anrede"])
    gruss = _absatz_mit(dok, PLATZHALTER["gruss"])
    if anrede is None or gruss is None:
        raise MehrkostenanzeigeFehler(
            "In der Briefvorlage fehlen Anrede oder Grußformel. Wurde die "
            "Vorlage ausgetauscht?"
        )

    _setze_text(anrede, anredezeile(daten.empfaenger))

    # Alles zwischen Anrede und Grußformel entfernen. Verglichen wird auf den
    # XML-Elementen: ``dok.paragraphs`` baut bei jedem Zugriff neue
    # Python-Objekte, ein Vergleich der Absätze selbst schlägt deshalb fehl.
    elemente = [p._p for p in dok.paragraphs]
    von = elemente.index(anrede._p)
    bis = elemente.index(gruss._p)
    for dazwischen in elemente[von + 1: bis]:
        _loeschen(dazwischen)

    # Alle Briefabsätze bekommen die Form der Anrede — also den normalen
    # Fließtext der Vorlage — und ihren Einzug ausdrücklich gesetzt.
    letzter = _absatz_nach(anrede, format_von=anrede, einzug=0)
    bloecke = koerper(daten)
    for stelle, block in enumerate(bloecke):
        letzter = _absatz_nach(
            letzter, block.text, format_von=anrede,
            einzug=ZITAT_EINZUG if block.zitat else 0,
        )
        # Zwischen "Auszug LV:" und dem Zitat und innerhalb des Zitats bleibt
        # der Zeilenabstand eng — ein LV-Auszug ist ein Block, keine Folge
        # einzelner Absätze.
        folgt = bloecke[stelle + 1] if stelle + 1 < len(bloecke) else None
        zusammen = (folgt is not None and folgt.zitat
                    and (block.marke or block.zitat))
        if not zusammen:
            letzter = _absatz_nach(letzter, format_von=anrede, einzug=0)


def _fuelle_unterschrift(dok, daten: Schreiben) -> None:
    name = _absatz_mit(dok, PLATZHALTER["name"])
    funktion = _absatz_mit(dok, PLATZHALTER["funktion"])
    firma = _absatz_mit(dok, PLATZHALTER["firma"])
    if name is None or funktion is None:
        raise MehrkostenanzeigeFehler(
            "In der Briefvorlage fehlt der Unterschriftsblock. Wurde die "
            "Vorlage ausgetauscht?"
        )

    bearbeiter = daten.sachbearbeiter
    _setze_text(name, dokumenttext.einzeilig(bearbeiter.name))
    if bearbeiter.funktion.strip():
        _setze_text(funktion, dokumenttext.einzeilig(bearbeiter.funktion))
    else:
        _loeschen(funktion)

    if firma is not None:
        _unterschrift_entfernen(firma)


def _fuelle_verteiler(dok, anlagen: list[str], verteiler: list[str]) -> None:
    """Anlagen- und Verteilerblock in den zweiten Abschnitt der Vorlage.

    Er trägt den verkleinerten Briefkopf — genau dafür ist er da. Gibt es
    nichts zu verteilen, werden seine Absätze geleert und der Abschnitt
    später aufgelöst (siehe ``_briefkopf_regeln``).
    """
    hinten = _abschnitt_zwei(dok)
    if not hinten:
        return

    zeilen: list[str] = []
    if anlagen:
        zeilen.append("Anlage:" if len(anlagen) == 1 else "Anlagen:")
        zeilen += anlagen
    if anlagen and verteiler:
        zeilen.append("")
    if verteiler:
        zeilen.append("Verteiler:")
        zeilen += verteiler

    if not zeilen:
        for absatz in hinten:
            _loeschen(absatz)
        return

    letzter = hinten[0]
    _setze_text(letzter, dokumenttext.einzeilig(zeilen[0]))
    for zeile in zeilen[1:]:
        letzter = _absatz_nach(letzter, dokumenttext.einzeilig(zeile))
    for ueberzaehlig in hinten[1:]:
        _loeschen(ueberzaehlig)


def _umbruch_absatz(dok):
    """Der Absatz, in dem der Abschnittsumbruch des Briefs steckt.

    Zwei Ebenen tiefer als man denkt: Ein Abschnittsumbruch mitten im Text
    steht als ``w:p/w:pPr/w:sectPr``. ``_sectPr.getparent()`` liefert deshalb
    das ``w:pPr`` und nicht den Absatz — mit dieser Verwechslung fand der
    Verteilerblock seinen Platz nicht und die letzte Seite blieb leer.
    """
    if len(dok.sections) < 2:
        return None
    ppr = dok.sections[0]._sectPr.getparent()
    if ppr is None:
        return None
    absatz = ppr.getparent()
    return absatz if absatz is not None and _kurzname(absatz) == "w:p" else None


def _abschnitt_zwei(dok) -> list:
    """Die Absätze hinter dem Abschnittsumbruch des Briefs."""
    umbruch = _umbruch_absatz(dok)
    if umbruch is None:
        return []
    gefunden: list = []
    dahinter = False
    for absatz in dok.paragraphs:
        if absatz._p is umbruch:
            dahinter = True
            continue
        if dahinter:
            gefunden.append(absatz)
    return gefunden
