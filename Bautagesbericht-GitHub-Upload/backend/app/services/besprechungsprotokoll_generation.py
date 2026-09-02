"""Erzeugt das Baubesprechungsprotokoll als Word-Dokument.

WAS HIER NACHGEBAUT WIRD
========================
Die Bürovorlage ``{JJMMTT}_BB_{Nr}.xlsm`` mit ihren vier Blättern. Alle Maße
in diesem Modul sind an einem echten Ausdruck gemessen (Protokoll 16 vom
25.08.2026, vier Seiten) und zusätzlich gegen die Excel-Datei selbst geprüft —
Spaltenbreiten, Zeilenhöhen und Füllfarben stammen aus deren XML und sind
nicht geschätzt.

    Seite 1   Deckblatt "protokoll"    Kopfzeile, Label/Wert-Block, fester
                                       Verteilungshinweis
    Seite 2   Themen-Tabelle           schwarze Kopfzeile mit gedrehter
                                       Schrift, graue Kapitelbalken, je Thema
                                       eine Zeile
    Seite 3   Legende + Abkürzungen    fester Text, dazu die Projektbeteiligten
    Seite 4   "teilnehmerliste"        eigenes Deckblatt, Teilnehmertabelle
    danach    Anlagen                  hochgeladene Scans, ganzseitig

WARUM KEINE PLATZHALTER-VORLAGE
===============================
Der Auftrag schlug den Weg von ``maengelliste_generation`` vor: Blanko-docx
mit ``{{PLATZHALTER}}``. Der passt dort, weil die echte Bürovorlage fehlte und
das Layout schlicht ist. Hier ist es umgekehrt — das Ziel liegt auf den Punkt
vor, und fast alles daran ist Geometrie, die eine Vorlage gar nicht tragen
könnte: um 90 Grad gedrehte weiße Schrift auf schwarzem Grund, sieben Spalten
mit festen Breiten, je Status eine andere Zellfüllung, ein frei positionierter
Briefkopf, vier verschiedene Seitentypen.

Deshalb ist dieses Modul nach dem Vorbild von ``maengelanzeige_generation``
gebaut: Das Dokument entsteht im Code, und alles, was in jedem Protokoll
gleich ist, steht als benannte Konstante ganz oben — einmal, an einer Stelle,
ohne dass jemand eine Word-Datei öffnen muss, um es zu ändern.

FARBEN UND REGELN DER STATUS-SPALTE
===================================
Die Färbung ist **keine** Handarbeit, auch wenn sie im PDF so aussieht. Sie
steht als bedingte Formatierung in der Excel-Vorlage:

    Status "n"  ganze Zeile hellblau (DDEBF7)
    Status "e"  nur die Status-Zelle grün  (C6EFCE, Schrift 006100)
    Status "b"  nur die Status-Zelle gelb  (FFEB9C, Schrift 9C5700)
    Status "k"  nur die Status-Zelle rot   (FFC7CE, Schrift 9C0006)
    Status "i"  keine Füllung

Von Hand gesetzt wird in der Vorlage nur eines: einzelne Zellen der Spalte
"Bearb. bis" werden hellblau hinterlegt, wenn die Bearbeiterin eine Frist im
Blick behalten will. Dafür gibt es keine Regel, deshalb hier auch keine —
das ist das Feld ``hervorheben`` am Themen-Update.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

from app.config import settings
from app.services import dokumenttext


class ProtokollFehler(RuntimeError):
    """Etwas fehlt, ohne das kein gültiges Protokoll entstehen kann."""


# ─────────────────────────────────────────────────────────────────────────────
# Feste Textbausteine
#
# Wortlaut aus dem Referenzprotokoll, Zeichen für Zeichen. Geändert wird er
# nur bewusst und dann genau hier.
# ─────────────────────────────────────────────────────────────────────────────

TITEL_PROTOKOLL = "protokoll"
#: Die Teilnehmerliste setzt "teilnehmer" mager und "liste" fett — so steht es
#: auf dem Ausdruck.
TITEL_TEILNEHMER = ("teilnehmer", "liste")

VERTEILUNGSHINWEIS = (
    "Die Verteilung dieses Protokolls erfolgt per E-Mail an den aufgeführten "
    "Empfängerkreis. Sofern erforderlich, wird um interne Weiterleitung der "
    "Aktennotiz gebeten. Die Empfänger dieses Protokolls werden gebeten, den "
    "Inhalt sowie eventuelle beiliegende Anlagen zu prüfen. Ergeben sich "
    "Einwände, Ergänzungen oder Änderungen, so sind diese dem Verfasser "
    "umgehend mitzuteilen."
)

ANWESENDE_VERWEIS = "siehe Teilnehmerliste"

#: Legende Seite 3, Reihenfolge wie im Original.
LEGENDE = (
    ("k", "kritisch"),
    ("b", "in Bearbeitung"),
    ("e", "erledigt"),
    ("n", "neu"),
    ("i", "informativ"),
)

UEBERSCHRIFT_LEGENDE = "Legende:"
UEBERSCHRIFT_ABKUERZUNGEN = "Abkürzungen Projektbeteiligte:"
UEBERSCHRIFT_AUFGESTELLT = "aufgestellt durch"
BUERO_NAME = "HPP Architekten"

#: Spaltenköpfe der Themen-Tabelle. "Bearb. \nbis" bricht im Original um.
SPALTEN_THEMEN = ("Kapitel", "Inhalt", "BB", "Thema", "Zuständig", "Bearb. \nbis", "Status")
SPALTEN_TEILNEHMER = ("Teilnehmer", "Firma", "Telefon", "Unterschrift/Anwesend")

DECKBLATT_LABELS = (
    "Protokoll-Nr.:",
    "Projekt:",
    "Projekt-Nr.:",
    "Leistung:",
    "Bauherr:",
    "Besprechungsort:",
    "Besprechungsdatum:",
    "Anwesende:",
)


# ─────────────────────────────────────────────────────────────────────────────
# Maße
#
# Alles in Punkt, gemessen am Original-PDF (A4, 595.2 x 841.7 pt). Der
# Nullpunkt ist die linke obere Ecke der Seite.
# ─────────────────────────────────────────────────────────────────────────────

SCHRIFT = "Arial"

SEITE_BREITE = Cm(21)
SEITE_HOEHE = Cm(29.7)

#: Druckrand der Excel-Vorlage: 0.7 Zoll links und rechts.
RAND_LINKS = Pt(50.4)
RAND_RECHTS = Pt(50.4)
#: Oben so knapp, dass die Titelzeile bei y 43.9 pt steht — der Rest der Seite
#: wird über Abstandsabsätze auf die gemessenen Höhen gebracht.
RAND_OBEN = Pt(43.9)
RAND_UNTEN = Pt(39.7)
KOPF_ABSTAND = Pt(10)
#: Fußzeile beginnt bei y 811.2 pt.
FUSS_ABSTAND = Pt(30.5)

#: Nutzbare Textbreite: 595.2 - 50.4 - 50.4.
TEXT_BREITE = Pt(494.4)

# ── Briefkopfgrafiken ────────────────────────────────────────────────────────
#
# Die Werte sind ausgerechnet, nicht geraten: Von der Grafik im Original-PDF
# wurde der Kasten des sichtbaren Inhalts bestimmt und daraus die Platzierung
# abgeleitet, mit der die vorhandene PNG-Datei denselben Kasten trifft.
BRIEFKOPF_DATEI = "hpp_briefkopf.png"
FOLGESEITE_DATEI = "hpp_briefkopf_folgeseite.png"

# Der obere Versatz ist negativ und das ist richtig: Oberhalb der Wortmarke
# hat die PNG-Datei leere Flaeche. Gemessen wird an der Wortmarke selbst —
# im Original steht sie auf Seite 1 bei y 30.0 pt, auf den Folgeseiten bei
# 21.5 pt und weiter rechts.
KOPF_LINKS = Pt(332.1)
KOPF_OBEN = Pt(-4.7)
KOPF_BREITE = Pt(212.9)
KOPF_HOEHE = Pt(559.1)

#: Auf den Folgeseiten steht nur Wortmarke und "Architekten", weiter rechts
#: und etwas größer als auf Seite 1. Der negative obere Versatz ist richtig:
#: Die PNG-Datei hat oberhalb der Wortmarke leere Fläche.
FOLGE_LINKS = Pt(390.9)
FOLGE_OBEN = Pt(-16.6)
FOLGE_BREITE = Pt(235.7)
FOLGE_HOEHE = Pt(619.1)

# ── Schriftgrößen ────────────────────────────────────────────────────────────
GR_TITEL = Pt(24)
GR_KOPFZEILE = Pt(7)
GR_DECKBLATT = Pt(11)
GR_FLIESSTEXT = Pt(9)
GR_TABELLE = Pt(10)
GR_SPALTENKOPF = Pt(9)
GR_FUSS = Pt(7)

SCHWARZ = RGBColor(0x00, 0x00, 0x00)
WEISS = RGBColor(0xFF, 0xFF, 0xFF)

# ── Deckblatt ────────────────────────────────────────────────────────────────
#: Zeilenhöhe der Titelzeile ("protokoll"). Der Titel sitzt im Original bei
#: y 43.9 pt, also direkt am oberen Rand.
TITEL_ZEILE = Pt(30)
#: Erste Zeile (Datum / Zeichen / Durchwahl / E-Mail) bei y 117.7 pt.
#: Darunter folgt der Label-Block, dessen erste Zeile bei 142.1 pt steht.
DECK_KOPFZEILE_OBEN = Pt(40.5)
#: Abstand zwischen Kopfzeile und Label-Block.
DECK_NACH_KOPFZEILE = Pt(14.4)
#: Linker Versatz des Deckblatt-Inhalts gegenüber dem Seitenrand: Text bei
#: 52.9 pt, Rand bei 50.4 pt.
DECK_EINZUG = Pt(2.5)
#: Breite der Textspalte des Verteilungshinweises. Der Wert ist nicht frei
#: waehlbar: Bei 304 pt bricht der Absatz an genau denselben sechs Stellen um
#: wie auf dem Ausdruck — schon 4 pt weniger schieben ein Wort in die naechste
#: Zeile.
HINWEIS_BREITE = Pt(304)
#: Tabstopps der Kopfzeile, gemessen ab Textbeginn (52.3 pt).
DECK_TABS = (Pt(64.5), Pt(128.9), Pt(193.4))
#: Label-Spalte 52.9 pt, Wert-Spalte 181.8 pt.
DECK_LABEL_BREITE = Pt(128.9)
DECK_WERT_BREITE = Pt(365.5)
#: Zeilenabstand der Label-Zeilen (27.6 pt) und innerhalb eines mehrzeiligen
#: Werts (13.8 pt).
DECK_ZEILE = Pt(27.6)
DECK_ZEILE_ENG = Pt(13.8)
#: Der Verteilungshinweis beginnt bei y 709.0 pt.
DECK_HINWEIS_OBEN = Pt(318.9)
DECK_HINWEIS_ZEILE = Pt(11.2)

# ── Themen-Tabelle ───────────────────────────────────────────────────────────
#: Die Tabelle beginnt bei y 82.08 pt, also 38.2 pt unter dem oberen Rand.
THEMEN_OBEN = Pt(38.2)
#: Linker Rand der Tabelle: 58.56 pt, also 8.2 pt eingerückt.
THEMEN_EINZUG = Pt(8.2)
#: Spaltenbreiten aus den senkrechten Linien des Originals:
#: 58.56 | 79.1 | 100.1 | 121.2 | 400.8 | 441.2 | 508.0 | 535.92
SPALTEN_BREITEN = (
    Pt(20.5),   # Kapitel
    Pt(21.0),   # Inhalt
    Pt(21.1),   # BB
    Pt(279.6),  # Thema
    Pt(40.4),   # Zuständig
    Pt(66.8),   # Bearb. bis
    Pt(27.4),   # Status
)
#: Höhe der schwarzen Kopfzeile und der grauen Kapitelbalken.
THEMEN_KOPF_HOEHE = Pt(60.4)
THEMEN_KAPITEL_HOEHE = Pt(19.4)
#: Eine Textzeile in der Tabelle ist genau 12.36 pt hoch.
THEMEN_ZEILE = Pt(12.36)
#: Abstand vom Zellenrand zum Text: 2.3 pt links.
ZELLE_INNEN = Pt(2.3)

FUELLUNG_KOPF = "000000"
FUELLUNG_KAPITEL = "F2F2F2"
FUELLUNG_NEU = "DDEBF7"
RAHMEN_FARBE = "000000"

#: Status -> (Füllung der Status-Zelle, Schriftfarbe). ``None`` = keine Füllung.
STATUS_FARBEN: dict[str, tuple[str | None, RGBColor | None]] = {
    "k": ("FFC7CE", RGBColor(0x9C, 0x00, 0x06)),
    "b": ("FFEB9C", RGBColor(0x9C, 0x57, 0x00)),
    "e": ("C6EFCE", RGBColor(0x00, 0x61, 0x00)),
    "n": (None, None),
    "i": (None, None),
}
#: Status, bei dem die ganze Zeile hinterlegt wird.
STATUS_ZEILE_BLAU = "n"

# ── Seite 3 ──────────────────────────────────────────────────────────────────
#: "Legende:" bei y 84.4 pt.
LEGENDE_OBEN = Pt(37.6)
#: Einzüge: Überschrift 65.8, Kürzel 130.2, Text 161.7, Rolle 338.9 pt.
SP3_UEBERSCHRIFT = Pt(15.4)
SP3_KUERZEL = Pt(79.8)
SP3_TEXT = Pt(111.3)
SP3_ROLLE = Pt(288.5)
SP3_ZEILE = Pt(13.55)
#: Die erste Legendenzeile sitzt 12.3 pt unter der Ueberschrift, nicht 13.55.
SP3_ZEILE_KOPF = Pt(12.3)
#: "aufgestellt durch" bei y 734.1 pt.
AUFGESTELLT_OBEN = Pt(405.7)

# ── Teilnehmerliste ──────────────────────────────────────────────────────────
#: Spaltenbreiten der Teilnehmertabelle, gemessen: 105 | 55 | 110 | 125 pt.
TEILNEHMER_BREITEN = (Pt(122), Pt(60), Pt(120), Pt(140))
#: Der Kopfblock der Teilnehmerliste steht enger als der auf dem Deckblatt:
#: In der Excel-Vorlage sind die Zeilen dort 19.5 pt statt 27.6 pt hoch.
TEILNEHMER_KOPF_ZUSATZ = Pt(5.7)
#: So viele leere Zeilen hängen unter den bekannten Teilnehmern — Platz für
#: alle, die unangemeldet dazukommen und trotzdem unterschreiben sollen.
TEILNEHMER_LEERZEILEN = 14
TEILNEHMER_ZEILE = Pt(15.5)

# ── Anlagen ──────────────────────────────────────────────────────────────────
#: Größte Abbildung einer Anlagenseite (Textbreite x verbleibende Höhe).
ANLAGE_BREITE = Pt(494.4)
ANLAGE_HOEHE = Pt(700)
#: Auflösung, mit der PDF-Anlagen zu Bildern gerendert werden.
ANLAGE_DPI = 150


# ─────────────────────────────────────────────────────────────────────────────
# Eingabedaten
# ─────────────────────────────────────────────────────────────────────────────


@dataclass
class ThemaZeile:
    """Eine Druckzeile der Themen-Tabelle."""

    kapitel_nr: str
    inhalt_nr: str
    bb_nr: str
    thema: str
    zustaendig: str = ""
    bearb_bis: str = ""
    status: str = "n"
    hervorheben: bool = False


@dataclass
class KapitelBlock:
    nummer: str
    titel: str
    zeilen: list[ThemaZeile] = field(default_factory=list)


@dataclass
class TeilnehmerZeile:
    name: str
    firma: str = ""
    telefon: str = ""


@dataclass
class BeteiligterZeile:
    kuerzel: str
    name: str
    rolle: str = ""


@dataclass
class AnlagenSeite:
    pfad: Path
    bezeichnung: str = ""


@dataclass
class ProtokollDaten:
    """Alles, was auf den vier Seiten steht — fertig aufbereitet."""

    nummer: int
    projektname: str
    projekt_adresse_zeilen: list[str]
    projekt_nummer: str
    leistung: str
    bauherr: str
    besprechungsort: str
    besprechungsdatum: date
    #: Datum der Kopfzeile — im Original der Tag, an dem gedruckt wurde.
    erstellt_datum: date
    ersteller_name: str
    ersteller_kuerzel: str
    ersteller_durchwahl: str
    ersteller_email: str
    kapitel: list[KapitelBlock] = field(default_factory=list)
    teilnehmer: list[TeilnehmerZeile] = field(default_factory=list)
    beteiligte: list[BeteiligterZeile] = field(default_factory=list)
    anlagen: list[AnlagenSeite] = field(default_factory=list)


# ─────────────────────────────────────────────────────────────────────────────
# XML-Werkzeuge, die python-docx nicht mitbringt
# ─────────────────────────────────────────────────────────────────────────────


def _marke(name: str) -> Path:
    pfad = settings.template_dir / "marke" / name
    if not pfad.is_file():
        raise ProtokollFehler(
            f"Die Briefkopfgrafik {pfad} fehlt. Ohne sie entsteht kein "
            f"HPP-Protokollbogen."
        )
    return pfad


def _bild_freistellen(absatz, bilddatei: str, *, links, oben, breite, hoehe=None,
                      name: str = "Briefkopf") -> None:
    """Bild an fester Seitenposition, hinter dem Text.

    Gleiches Vorgehen wie in ``maengelanzeige_generation``: python-docx kann
    nur Bilder im Textfluss, deshalb wird das erzeugte ``wp:inline`` zu einem
    ``wp:anchor`` mit absoluter Position umgeschrieben.
    """
    lauf = absatz.add_run()
    if hoehe is not None:
        lauf.add_picture(bilddatei, width=breite, height=hoehe)
    else:
        lauf.add_picture(bilddatei, width=breite)

    inline = lauf._r.find(qn("w:drawing"))[0]
    anker = OxmlElement("wp:anchor")
    for schluessel, wert in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("simplePos", "0"), ("relativeHeight", "251658240"),
        ("behindDoc", "1"), ("locked", "0"), ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anker.set(schluessel, wert)

    einfach = OxmlElement("wp:simplePos")
    einfach.set("x", "0")
    einfach.set("y", "0")
    anker.append(einfach)

    for richtung, versatz in (("wp:positionH", int(links)), ("wp:positionV", int(oben))):
        block = OxmlElement(richtung)
        block.set("relativeFrom", "page")
        wert = OxmlElement("wp:posOffset")
        wert.text = str(versatz)
        block.append(wert)
        anker.append(block)

    for tag in ("wp:extent", "wp:effectExtent"):
        vorhanden = inline.find(qn(tag))
        if vorhanden is not None:
            anker.append(vorhanden)
    anker.append(OxmlElement("wp:wrapNone"))
    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        vorhanden = inline.find(qn(tag))
        if vorhanden is not None:
            anker.append(vorhanden)

    beschreibung = anker.find(qn("wp:docPr"))
    if beschreibung is not None:
        beschreibung.set("name", name)

    inline.getparent().replace(inline, anker)


def _setze_fuellung(zelle, farbe: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), farbe)
    zelle._tc.get_or_add_tcPr().append(shd)


def _setze_breite(zelle, breite) -> None:
    """Spaltenbreite hart setzen — Word ignoriert sonst gern die Vorgabe."""
    zelle.width = breite
    tc_pr = zelle._tc.get_or_add_tcPr()
    for alt in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(alt)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(Emu(int(breite)).twips)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _zellen_innenrand(zelle, links=ZELLE_INNEN, rechts=ZELLE_INNEN) -> None:
    """Innenabstand einer Zelle; oben/unten auf null, damit die Zeilenhöhe
    genau der Zahl der Textzeilen entspricht.

    ``w:left``/``w:right`` und nicht ``w:start``/``w:end``: Word wertet in
    ``w:tcMar`` nur die alten Namen aus, die neuen werden stillschweigend
    ignoriert — und dann steht der Text am Zellenrand.
    """
    tc_pr = zelle._tc.get_or_add_tcPr()
    rand = OxmlElement("w:tcMar")
    for kante, wert in (("top", 0), ("left", int(Emu(int(links)).twips)),
                        ("bottom", 0), ("right", int(Emu(int(rechts)).twips))):
        el = OxmlElement(f"w:{kante}")
        el.set(qn("w:w"), str(wert))
        el.set(qn("w:type"), "dxa")
        rand.append(el)
    tc_pr.append(rand)


def _dreh_zelle(zelle) -> None:
    """Text um 90 Grad drehen (von unten nach oben) — die Spaltenköpfe."""
    tc_pr = zelle._tc.get_or_add_tcPr()
    richtung = OxmlElement("w:textDirection")
    richtung.set(qn("w:val"), "btLr")
    tc_pr.append(richtung)


def _zeilen_hoehe(zeile, hoehe, genau: bool = True) -> None:
    tr_pr = zeile._tr.get_or_add_trPr()
    ht = OxmlElement("w:trHeight")
    ht.set(qn("w:val"), str(int(Emu(int(hoehe)).twips)))
    ht.set(qn("w:hRule"), "exact" if genau else "atLeast")
    tr_pr.append(ht)


def _zeile_zusammenhalten(zeile) -> None:
    """Verbietet, dass eine Themenzeile über den Seitenumbruch reißt."""
    tr_pr = zeile._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _kopfzeile_wiederholen(zeile) -> None:
    """Die schwarze Spaltenkopfzeile auf jeder Folgeseite wiederholen."""
    tr_pr = zeile._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def _tabellen_rahmen(tabelle, farbe: str = RAHMEN_FARBE, staerke: str = "4") -> None:
    tbl_pr = tabelle._tbl.tblPr
    rahmen = OxmlElement("w:tblBorders")
    for kante in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{kante}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), staerke)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), farbe)
        rahmen.append(el)
    tbl_pr.append(rahmen)


def _tabelle_einruecken(tabelle, einzug) -> None:
    tbl_pr = tabelle._tbl.tblPr
    el = OxmlElement("w:tblInd")
    el.set(qn("w:w"), str(int(Emu(int(einzug)).twips)))
    el.set(qn("w:type"), "dxa")
    tbl_pr.append(el)


def _tabelle_feste_breiten(tabelle) -> None:
    """Feste Spaltenbreiten statt Word-Automatik."""
    tbl_pr = tabelle._tbl.tblPr
    for alt in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(alt)
    el = OxmlElement("w:tblLayout")
    el.set(qn("w:type"), "fixed")
    tbl_pr.append(el)


def _schrift_setzen(dokument) -> None:
    stil = dokument.styles["Normal"]
    stil.font.name = SCHRIFT
    stil.font.size = GR_TABELLE
    rpr = stil.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for schluessel in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(schluessel), SCHRIFT)
    sprache = rpr.find(qn("w:lang"))
    if sprache is None:
        sprache = OxmlElement("w:lang")
        rpr.append(sprache)
    for schluessel in ("w:val", "w:eastAsia", "w:bidi"):
        sprache.set(qn(schluessel), "de-DE")


def _feld(absatz, instr: str, platzhalter: str, groesse=GR_FUSS) -> None:
    """Word-Feld (PAGE, NUMPAGES) in einen Absatz setzen."""
    feld = OxmlElement("w:fldSimple")
    feld.set(qn("w:instr"), instr)
    lauf = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), SCHRIFT)
    fonts.set(qn("w:hAnsi"), SCHRIFT)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(Emu(int(groesse)).pt * 2)))
    rpr.append(fonts)
    rpr.append(sz)
    text = OxmlElement("w:t")
    text.text = platzhalter
    lauf.append(rpr)
    lauf.append(text)
    feld.append(lauf)
    absatz._p.append(feld)


# ─────────────────────────────────────────────────────────────────────────────
# Absatz-Bausteine
# ─────────────────────────────────────────────────────────────────────────────


def _leere_zelle(zelle) -> None:
    for absatz in list(zelle.paragraphs)[1:]:
        absatz._p.getparent().remove(absatz._p)
    for lauf in list(zelle.paragraphs[0].runs):
        lauf._r.getparent().remove(lauf._r)


def _schreibe(absatz, text: str, *, groesse=GR_TABELLE, bold: bool = False,
              farbe: RGBColor = SCHWARZ):
    lauf = absatz.add_run(dokumenttext.xml_sicher(text))
    lauf.font.name = SCHRIFT
    lauf.font.size = groesse
    lauf.font.bold = bold
    lauf.font.color.rgb = farbe
    return lauf


def _absatz(behaelter, text: str = "", *, groesse=GR_TABELLE, bold: bool = False,
            farbe: RGBColor = SCHWARZ, zeile=None, vor=None, nach=Pt(0),
            einzug=None, ausrichtung=None):
    absatz = behaelter.add_paragraph()
    fmt = absatz.paragraph_format
    fmt.space_before = vor if vor is not None else Pt(0)
    fmt.space_after = nach
    if zeile is not None:
        fmt.line_spacing = zeile
    if einzug is not None:
        fmt.left_indent = einzug
    if ausrichtung is not None:
        absatz.alignment = ausrichtung
    if text:
        _schreibe(absatz, text, groesse=groesse, bold=bold, farbe=farbe)
    return absatz


def _abstandshalter(dokument, hoehe) -> None:
    """Leerabsatz mit exakt vorgegebener Höhe.

    Genauer als mehrere Leerzeilen: Der Abstand ist eine Zahl und keine
    Rechnung mit der Zeilenhöhe, und er verschiebt sich nicht, wenn jemand
    die Standardschrift ändert.
    """
    absatz = dokument.add_paragraph()
    fmt = absatz.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = hoehe
    lauf = absatz.add_run("")
    lauf.font.size = Pt(1)
    lauf.font.name = SCHRIFT


def _mehrzeilig(zelle, text: str, *, groesse=GR_TABELLE, bold: bool = False,
                farbe: RGBColor = SCHWARZ, zeile=THEMEN_ZEILE,
                ausrichtung=None) -> None:
    """Text mit Zeilenumbrüchen als eigene Absätze in eine Zelle schreiben.

    Zeilenumbrüche sind im Protokoll Inhalt und keine Formatierung: In der
    Spalte "Thema" stehen Aufzählungen mit "-", in "Zuständig" mehrere Kürzel
    untereinander. Beides muss so ankommen, wie es eingegeben wurde.
    """
    _leere_zelle(zelle)
    zeilen = str(text or "").splitlines() or [""]
    ziel = zelle.paragraphs[0]
    for i, inhalt in enumerate(zeilen):
        if i > 0:
            ziel = zelle.add_paragraph()
        fmt = ziel.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = zeile
        if ausrichtung is not None:
            ziel.alignment = ausrichtung
        if inhalt:
            _schreibe(ziel, inhalt, groesse=groesse, bold=bold, farbe=farbe)


def _fmt_datum(wert: date | None) -> str:
    return wert.strftime("%d.%m.%Y") if wert else ""


# ─────────────────────────────────────────────────────────────────────────────
# Seitengerüst: Ränder, Briefkopf, Fußzeile
# ─────────────────────────────────────────────────────────────────────────────


def _richte_seite_ein(dokument, fusszeile_text: str) -> None:
    abschnitt = dokument.sections[0]
    abschnitt.page_width = SEITE_BREITE
    abschnitt.page_height = SEITE_HOEHE
    abschnitt.left_margin = RAND_LINKS
    abschnitt.right_margin = RAND_RECHTS
    abschnitt.top_margin = RAND_OBEN
    abschnitt.bottom_margin = RAND_UNTEN
    abschnitt.header_distance = KOPF_ABSTAND
    abschnitt.footer_distance = FUSS_ABSTAND
    abschnitt.start_type = WD_SECTION_START.NEW_PAGE

    # Seite 1 trägt den vollen Briefkopf, alle weiteren nur Wortmarke und
    # "Architekten" — genau wie im Original.
    abschnitt.different_first_page_header_footer = True

    # Die Kopfzeile trägt ausschließlich die freigestellte Grafik. Ihr Absatz
    # wird auf 1 pt gedrückt, weil Word sonst den ganzen Seiteninhalt um die
    # Zeilenhöhe der Kopfzeile nach unten schiebt — die gemessenen Höhen
    # stimmten dann auf keiner Seite mehr. Die Grafik selbst ist am Seitenrand
    # verankert und von der Absatzhöhe unabhängig.
    for teil, datei, geometrie, name in (
        (abschnitt.first_page_header, BRIEFKOPF_DATEI,
         (KOPF_LINKS, KOPF_OBEN, KOPF_BREITE, KOPF_HOEHE), "HPP Briefkopf"),
        (abschnitt.header, FOLGESEITE_DATEI,
         (FOLGE_LINKS, FOLGE_OBEN, FOLGE_BREITE, FOLGE_HOEHE), "HPP Wortmarke"),
    ):
        teil.is_linked_to_previous = False
        absatz = teil.paragraphs[0]
        absatz.paragraph_format.space_after = Pt(0)
        absatz.paragraph_format.space_before = Pt(0)
        absatz.paragraph_format.line_spacing = Pt(1)
        links, oben, breite, hoehe = geometrie
        _bild_freistellen(absatz, str(_marke(datei)), links=links, oben=oben,
                          breite=breite, hoehe=hoehe, name=name)
        for lauf in absatz.runs:
            lauf.font.size = Pt(1)

    for fuss in (abschnitt.first_page_footer, abschnitt.footer):
        fuss.is_linked_to_previous = False
        absatz = fuss.paragraphs[0]
        # Die eingebaute Formatvorlage "Footer" bringt eigene Tabstopps mit
        # (Mitte 4536, rechts 9072 Twips). Der Tabulator liefe dann dorthin
        # statt an den rechten Rand, und die Seitenzahl stünde mitten auf der
        # Seite. Deshalb zurück auf "Normal", das keine mitbringt.
        absatz.style = dokument.styles["Normal"]
        absatz.paragraph_format.space_before = Pt(0)
        absatz.paragraph_format.space_after = Pt(0)
        absatz.paragraph_format.line_spacing = Pt(10)
        for alt in list(absatz.runs):
            alt._r.getparent().remove(alt._r)
        # Einen Punkt vor dem rechten Rand: genau auf dem Rand verwirft Word
        # den Tabstopp.
        absatz.paragraph_format.tab_stops.add_tab_stop(
            Pt(493.4), WD_TAB_ALIGNMENT.RIGHT
        )
        _schreibe(absatz, fusszeile_text, groesse=GR_FUSS)
        _schreibe(absatz, "\t", groesse=GR_FUSS)
        _feld(absatz, "PAGE", "1")
        _schreibe(absatz, "/", groesse=GR_FUSS)
        _feld(absatz, "NUMPAGES", "1")


def _seitenumbruch(dokument) -> None:
    absatz = dokument.add_paragraph()
    absatz.paragraph_format.space_before = Pt(0)
    absatz.paragraph_format.space_after = Pt(0)
    absatz.paragraph_format.line_spacing = Pt(1)
    lauf = absatz.add_run()
    lauf.font.size = Pt(1)
    umbruch = OxmlElement("w:br")
    umbruch.set(qn("w:type"), "page")
    lauf._r.append(umbruch)


# ─────────────────────────────────────────────────────────────────────────────
# Seite 1 — Deckblatt
# ─────────────────────────────────────────────────────────────────────────────


def _kopfzeile_person(dokument, daten: ProtokollDaten, datum: date) -> None:
    """Datum, Zeichen, Durchwahl, E-Mail auf vier Tabstopps."""
    absatz = _absatz(dokument, "", groesse=GR_KOPFZEILE, vor=DECK_KOPFZEILE_OBEN,
                     nach=DECK_NACH_KOPFZEILE, zeile=Pt(11), einzug=Pt(1.9))
    for stop in DECK_TABS:
        absatz.paragraph_format.tab_stops.add_tab_stop(stop, WD_TAB_ALIGNMENT.LEFT)
    teile = [
        _fmt_datum(datum),
        f"Ze: {daten.ersteller_kuerzel}" if daten.ersteller_kuerzel else "",
        f"T - {daten.ersteller_durchwahl}" if daten.ersteller_durchwahl else "",
        daten.ersteller_email,
    ]
    for i, teil in enumerate(teile):
        if i:
            _schreibe(absatz, "\t", groesse=GR_KOPFZEILE)
        _schreibe(absatz, teil, groesse=GR_KOPFZEILE)


def _deckblatt_block(dokument, daten: ProtokollDaten) -> None:
    """Label/Wert-Zeilen ohne sichtbares Gitter."""
    werte: list[list[str]] = [
        [str(daten.nummer)],
        [daten.projektname, *daten.projekt_adresse_zeilen],
        [daten.projekt_nummer],
        [daten.leistung],
        [daten.bauherr],
        [daten.besprechungsort],
        [_fmt_datum(daten.besprechungsdatum)],
        [ANWESENDE_VERWEIS],
    ]

    _kopfblock(dokument, list(zip(DECKBLATT_LABELS, werte)))


def _kopfblock(dokument, paare: list[tuple[str, list[str]]],
               zusatz=None) -> None:
    """Label/Wert-Zeilen ohne Gitter — Deckblatt und Teilnehmerliste.

    Jede Zeile ist im Original ein Vielfaches von 13.8 pt hoch: eine Zeile für
    den Text, eine als Abstand darunter. Eine dreizeilige Projektangabe
    braucht deshalb 4 x 13.8 = 55.2 pt und nicht 41.4 — genau so steht es auf
    dem Ausdruck.
    """
    tabelle = dokument.add_table(rows=0, cols=2)
    tabelle.autofit = False
    _tabelle_feste_breiten(tabelle)
    _tabelle_rahmen_aus(tabelle)
    _tabelle_einruecken(tabelle, DECK_EINZUG)

    if zusatz is None:
        zusatz = DECK_ZEILE_ENG
    for label, zeilen in paare:
        inhalt = [z for z in zeilen if z]
        zeile = tabelle.add_row()
        _zeilen_hoehe(
            zeile, Pt(DECK_ZEILE_ENG.pt * max(1, len(inhalt)) + zusatz.pt)
        )
        links, rechts = zeile.cells[0], zeile.cells[1]
        _setze_breite(links, DECK_LABEL_BREITE)
        _setze_breite(rechts, DECK_WERT_BREITE)
        _zellen_innenrand(links, links=Pt(0), rechts=Pt(0))
        _zellen_innenrand(rechts, links=Pt(0), rechts=Pt(0))
        links.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        rechts.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _mehrzeilig(links, label, groesse=GR_DECKBLATT, zeile=DECK_ZEILE_ENG)
        _mehrzeilig(rechts, "\n".join(inhalt),
                    groesse=GR_DECKBLATT, zeile=DECK_ZEILE_ENG)


def _tabelle_rahmen_aus(tabelle) -> None:
    tbl_pr = tabelle._tbl.tblPr
    rahmen = OxmlElement("w:tblBorders")
    for kante in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{kante}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        rahmen.append(el)
    tbl_pr.append(rahmen)


def _titelzeile(dokument, teile: tuple[str, ...]) -> None:
    """Die große Wortmarke oben links — "protokoll" bzw. "teilnehmerliste".

    Sie steht im Fließtext und nicht in der Kopfzeile: Eine Kopfzeile mit
    Inhalt schiebt in Word den gesamten Seiteninhalt nach unten, und dann
    stimmt keine der gemessenen Höhen mehr.
    """
    absatz = _absatz(dokument, "", zeile=TITEL_ZEILE)
    for i, teil in enumerate(teile):
        _schreibe(absatz, teil, groesse=GR_TITEL, bold=(i > 0))


def _seite_deckblatt(dokument, daten: ProtokollDaten) -> None:
    _titelzeile(dokument, (TITEL_PROTOKOLL,))
    _kopfzeile_person(dokument, daten, daten.erstellt_datum)
    _deckblatt_block(dokument, daten)

    absatz = _absatz(dokument, "", groesse=GR_FLIESSTEXT, vor=DECK_HINWEIS_OBEN,
                     zeile=DECK_HINWEIS_ZEILE, einzug=Pt(2.3))
    # Der Hinweis läuft im Original nur über gut zwei Drittel der Seitenbreite;
    # sonst bricht er an anderen Stellen um als auf dem Ausdruck.
    absatz.paragraph_format.right_indent = TEXT_BREITE - HINWEIS_BREITE
    _schreibe(absatz, VERTEILUNGSHINWEIS, groesse=GR_FLIESSTEXT)


# ─────────────────────────────────────────────────────────────────────────────
# Seite 2 — Themen-Tabelle
# ─────────────────────────────────────────────────────────────────────────────


def _themen_kopfzeile(tabelle) -> None:
    zeile = tabelle.add_row()
    _zeilen_hoehe(zeile, THEMEN_KOPF_HOEHE)
    _kopfzeile_wiederholen(zeile)
    for zelle, beschriftung, breite in zip(zeile.cells, SPALTEN_THEMEN, SPALTEN_BREITEN):
        _setze_breite(zelle, breite)
        _setze_fuellung(zelle, FUELLUNG_KOPF)
        _zellen_innenrand(zelle, links=Pt(1.5), rechts=Pt(1.5))
        _dreh_zelle(zelle)
        # Bei gedrehtem Text vertauschen sich die Achsen: TOP schiebt die
        # Beschriftung an den *linken* Rand ihrer Spalte, und genau dort steht
        # sie im Original — auch in den breiten Spalten.
        zelle.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _mehrzeilig(zelle, beschriftung, groesse=GR_SPALTENKOPF, bold=True,
                    farbe=WEISS, zeile=Pt(10))


def _kapitel_balken(tabelle, nummer: str, titel: str) -> None:
    zeile = tabelle.add_row()
    _zeilen_hoehe(zeile, THEMEN_KAPITEL_HOEHE)
    _zeile_zusammenhalten(zeile)
    zellen = zeile.cells
    for zelle, breite in zip(zellen, SPALTEN_BREITEN):
        _setze_breite(zelle, breite)
        _setze_fuellung(zelle, FUELLUNG_KAPITEL)
        _zellen_innenrand(zelle)
        zelle.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _leere_zelle(zelle)
    _mehrzeilig(zellen[0], nummer, bold=True, zeile=THEMEN_ZEILE)
    _mehrzeilig(zellen[3], titel, bold=True, zeile=THEMEN_ZEILE)


def _themen_zeile(tabelle, eintrag: ThemaZeile) -> None:
    zeile = tabelle.add_row()
    _zeile_zusammenhalten(zeile)
    zellen = zeile.cells

    fuellung, schriftfarbe = STATUS_FARBEN.get(eintrag.status, (None, None))
    zeile_blau = eintrag.status == STATUS_ZEILE_BLAU

    werte = (
        eintrag.kapitel_nr,
        eintrag.inhalt_nr,
        eintrag.bb_nr,
        eintrag.thema,
        eintrag.zustaendig,
        eintrag.bearb_bis,
        eintrag.status,
    )
    ausrichtungen = (
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    )

    for i, (zelle, wert, breite, ausrichtung) in enumerate(
        zip(zellen, werte, SPALTEN_BREITEN, ausrichtungen)
    ):
        _setze_breite(zelle, breite)
        _zellen_innenrand(zelle)
        # Im Original stehen Nummern, Zuständige und Frist auf Höhe der
        # ersten Textzeile; nur der Statusbuchstabe sitzt mittig in der
        # Zeile — bei einem fünfzeiligen Thema ein deutlicher Unterschied.
        zelle.vertical_alignment = (
            WD_ALIGN_VERTICAL.CENTER if i == 6 else WD_ALIGN_VERTICAL.TOP
        )

        if zeile_blau:
            _setze_fuellung(zelle, FUELLUNG_NEU)
        # Spalte "Bearb. bis": von Hand gesetzte Hervorhebung einer Frist.
        if i == 5 and eintrag.hervorheben and not zeile_blau:
            _setze_fuellung(zelle, FUELLUNG_NEU)
        # Statusspalte: eigene Farbe je Status, auch auf blauer Zeile.
        if i == 6 and fuellung:
            _setze_fuellung(zelle, fuellung)

        farbe = schriftfarbe if (i == 6 and schriftfarbe) else SCHWARZ
        _mehrzeilig(zelle, wert, zeile=THEMEN_ZEILE, farbe=farbe,
                    ausrichtung=ausrichtung)


def _seite_themen(dokument, daten: ProtokollDaten) -> None:
    _abstandshalter(dokument, THEMEN_OBEN)

    tabelle = dokument.add_table(rows=0, cols=len(SPALTEN_THEMEN))
    tabelle.autofit = False
    _tabelle_feste_breiten(tabelle)
    _tabellen_rahmen(tabelle)
    _tabelle_einruecken(tabelle, THEMEN_EINZUG)

    _themen_kopfzeile(tabelle)
    for block in daten.kapitel:
        _kapitel_balken(tabelle, block.nummer, block.titel)
        for eintrag in block.zeilen:
            _themen_zeile(tabelle, eintrag)

    if not daten.kapitel:
        _kapitel_balken(tabelle, "", "Zu dieser Besprechung wurden keine Themen erfasst.")


# ─────────────────────────────────────────────────────────────────────────────
# Seite 3 — Legende und Abkürzungen
# ─────────────────────────────────────────────────────────────────────────────


def _dreispalter(dokument, zeilen: list[tuple[str, str, str]], vor=None) -> None:
    """Kürzel / Text / Rolle auf festen Einzügen — ohne Tabelle.

    Tabstopps statt Tabelle, weil das Original genau so aussieht: drei feste
    Positionen, kein Gitter, keine Zellen.
    """
    for i, (kuerzel, text, rolle) in enumerate(zeilen):
        absatz = _absatz(dokument, "", zeile=SP3_ZEILE, einzug=SP3_KUERZEL,
                         vor=(vor if i == 0 else None))
        absatz.paragraph_format.tab_stops.add_tab_stop(
            SP3_TEXT, WD_TAB_ALIGNMENT.LEFT
        )
        if rolle:
            absatz.paragraph_format.tab_stops.add_tab_stop(
                SP3_ROLLE, WD_TAB_ALIGNMENT.LEFT
            )
        _schreibe(absatz, kuerzel)
        _schreibe(absatz, "\t")
        _schreibe(absatz, text)
        if rolle:
            _schreibe(absatz, "\t")
            _schreibe(absatz, rolle)


def _seite_legende(dokument, daten: ProtokollDaten) -> None:
    _abstandshalter(dokument, LEGENDE_OBEN)
    _absatz(dokument, UEBERSCHRIFT_LEGENDE, zeile=SP3_ZEILE_KOPF, einzug=SP3_UEBERSCHRIFT)
    _dreispalter(dokument, [(k, t, "") for k, t in LEGENDE])

    _absatz(dokument, UEBERSCHRIFT_ABKUERZUNGEN, zeile=SP3_ZEILE,
            einzug=SP3_UEBERSCHRIFT, vor=SP3_ZEILE)
    if daten.beteiligte:
        _dreispalter(
            dokument,
            [(b.kuerzel, b.name, b.rolle) for b in daten.beteiligte],
            vor=SP3_ZEILE,
        )
    else:
        _absatz(dokument,
                "Für dieses Projekt sind noch keine Projektbeteiligten hinterlegt.",
                zeile=SP3_ZEILE, einzug=SP3_KUERZEL)

    _absatz(dokument, UEBERSCHRIFT_AUFGESTELLT, zeile=SP3_ZEILE,
            einzug=SP3_UEBERSCHRIFT, vor=AUFGESTELLT_OBEN)
    absatz = _absatz(dokument, "", zeile=SP3_ZEILE, einzug=SP3_UEBERSCHRIFT,
                     vor=SP3_ZEILE)
    absatz.paragraph_format.tab_stops.add_tab_stop(SP3_TEXT, WD_TAB_ALIGNMENT.LEFT)
    _schreibe(absatz, BUERO_NAME)
    if daten.ersteller_name:
        _schreibe(absatz, "\t")
        _schreibe(absatz, daten.ersteller_name)


# ─────────────────────────────────────────────────────────────────────────────
# Seite 4 — Teilnehmerliste
# ─────────────────────────────────────────────────────────────────────────────


def _seite_teilnehmer(dokument, daten: ProtokollDaten) -> None:
    _titelzeile(dokument, TITEL_TEILNEHMER)
    _kopfzeile_person(dokument, daten, daten.erstellt_datum)

    # Kopfangaben wie auf dem Deckblatt, aber ohne Protokoll-Nr. und
    # "Anwesende" — die Liste ist ja selbst die Antwort darauf. Im Original
    # stehen sie hier enger als auf Seite 1.
    _kopfblock(dokument, [
        ("Projekt:", [daten.projektname, *daten.projekt_adresse_zeilen]),
        ("Projekt-Nr.:", [daten.projekt_nummer]),
        ("Leistung:", [daten.leistung]),
        ("Bauherr:", [daten.bauherr]),
        ("Besprechungsort:", [daten.besprechungsort]),
        ("Besprechungsdatum:", [_fmt_datum(daten.besprechungsdatum)]),
    ], zusatz=TEILNEHMER_KOPF_ZUSATZ)

    _abstandshalter(dokument, Pt(18))

    liste = dokument.add_table(rows=0, cols=len(SPALTEN_TEILNEHMER))
    liste.autofit = False
    _tabelle_feste_breiten(liste)
    _tabellen_rahmen(liste)

    kopf = liste.add_row()
    _zeilen_hoehe(kopf, TEILNEHMER_ZEILE)
    _kopfzeile_wiederholen(kopf)
    for zelle, beschriftung, breite in zip(kopf.cells, SPALTEN_TEILNEHMER,
                                           TEILNEHMER_BREITEN):
        _setze_breite(zelle, breite)
        _zellen_innenrand(zelle, links=Pt(3), rechts=Pt(3))
        zelle.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _mehrzeilig(zelle, beschriftung, zeile=THEMEN_ZEILE)
    # Der Spaltenkopf steht im Original ohne Rahmen über der Tabelle.
    for zelle in kopf.cells:
        _zellen_rahmen_aus(zelle)

    zeilen = list(daten.teilnehmer) + [TeilnehmerZeile("")] * TEILNEHMER_LEERZEILEN
    for eintrag in zeilen:
        reihe = liste.add_row()
        _zeilen_hoehe(reihe, TEILNEHMER_ZEILE)
        werte = (eintrag.name, eintrag.firma, eintrag.telefon, "")
        for zelle, wert, breite in zip(reihe.cells, werte, TEILNEHMER_BREITEN):
            _setze_breite(zelle, breite)
            _zellen_innenrand(zelle, links=Pt(3), rechts=Pt(3))
            zelle.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _mehrzeilig(zelle, wert, zeile=THEMEN_ZEILE)


def _zellen_rahmen_aus(zelle) -> None:
    tc_pr = zelle._tc.get_or_add_tcPr()
    rahmen = OxmlElement("w:tcBorders")
    for kante in ("top", "start", "bottom", "end"):
        el = OxmlElement(f"w:{kante}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        rahmen.append(el)
    tc_pr.append(rahmen)


# ─────────────────────────────────────────────────────────────────────────────
# Anlagen — hochgeladene Scans
# ─────────────────────────────────────────────────────────────────────────────


def _anlagen_bilder(pfad: Path) -> list[bytes]:
    """Eine Anlage in ganzseitige Bilder umwandeln.

    PDF wird Seite für Seite gerendert (``pypdfium2`` liegt ohnehin im Paket,
    es rendert schon die Planvorschau), Bilder werden durchgereicht. Ein
    Format, das keiner von beiden lesen kann, führt nicht zum Abbruch — das
    Protokoll darf an einer unlesbaren Anlage nicht scheitern.
    """
    endung = pfad.suffix.lower()
    if endung == ".pdf":
        try:
            import pypdfium2 as pdfium
        except ImportError:
            return []
        seiten: list[bytes] = []
        dokument = pdfium.PdfDocument(str(pfad))
        try:
            for i in range(len(dokument)):
                bild = dokument[i].render(scale=ANLAGE_DPI / 72).to_pil()
                puffer = io.BytesIO()
                bild.convert("RGB").save(puffer, format="JPEG", quality=85)
                seiten.append(puffer.getvalue())
        finally:
            dokument.close()
        return seiten

    try:
        from PIL import Image

        with Image.open(pfad) as bild:
            puffer = io.BytesIO()
            bild.convert("RGB").save(puffer, format="JPEG", quality=85)
            return [puffer.getvalue()]
    except Exception:
        return []


def _anlagen_abschnitt(dokument):
    """Eigener Abschnitt für die Anlagen — ohne den HPP-Briefkopf.

    Eine Anlage ist ein fremdes Dokument. Der Briefkopf dieses Protokolls
    gehört nicht darüber: Beim eingescannten Ausdruck der Teilnehmerliste
    stünde die Wortmarke sonst zweimal auf der Seite, einmal aus dem Scan und
    einmal aus der Kopfzeile. Die Fußzeile bleibt — die Seitenzählung soll
    durchlaufen.
    """
    abschnitt = dokument.add_section(WD_SECTION_START.NEW_PAGE)
    abschnitt.page_width = SEITE_BREITE
    abschnitt.page_height = SEITE_HOEHE
    abschnitt.left_margin = RAND_LINKS
    abschnitt.right_margin = RAND_RECHTS
    abschnitt.top_margin = RAND_OBEN
    abschnitt.bottom_margin = RAND_UNTEN
    abschnitt.header_distance = KOPF_ABSTAND
    abschnitt.footer_distance = FUSS_ABSTAND
    abschnitt.different_first_page_header_footer = False

    kopf = abschnitt.header
    kopf.is_linked_to_previous = False
    for absatz in list(kopf.paragraphs)[1:]:
        absatz._p.getparent().remove(absatz._p)
    leer = kopf.paragraphs[0]
    for lauf in list(leer.runs):
        lauf._r.getparent().remove(lauf._r)
    leer.paragraph_format.space_before = Pt(0)
    leer.paragraph_format.space_after = Pt(0)
    leer.paragraph_format.line_spacing = Pt(1)

    # Fußzeile weiterführen: is_linked_to_previous bleibt gesetzt, damit
    # Dateiname und Seitenzahl auch auf den Anlagenseiten stehen.
    abschnitt.footer.is_linked_to_previous = True
    return abschnitt


def _seite_anlagen(dokument, daten: ProtokollDaten) -> list[str]:
    """Hängt die Anlagen an. Gibt zurück, was nicht eingebunden werden konnte."""
    probleme: list[str] = []
    bereit = False
    for anlage in daten.anlagen:
        bilder = _anlagen_bilder(anlage.pfad)
        if not bilder:
            probleme.append(anlage.pfad.name)
            continue
        for daten_bild in bilder:
            if not bereit:
                _anlagen_abschnitt(dokument)
                bereit = True
            else:
                _seitenumbruch(dokument)
            if anlage.bezeichnung:
                _absatz(dokument, anlage.bezeichnung, groesse=GR_FLIESSTEXT,
                        zeile=Pt(12), nach=Pt(6))
            absatz = dokument.add_paragraph()
            absatz.paragraph_format.space_before = Pt(0)
            absatz.paragraph_format.space_after = Pt(0)
            lauf = absatz.add_run()
            lauf.add_picture(io.BytesIO(daten_bild), width=ANLAGE_BREITE)
            _bild_hoehe_begrenzen(lauf, ANLAGE_BREITE, ANLAGE_HOEHE)
    return probleme


def _bild_hoehe_begrenzen(lauf, breite, max_hoehe) -> None:
    """Ein hohes Hochformat auf die Seitenhöhe bringen, Seitenverhältnis wahren."""
    bild = lauf._r.find(qn("w:drawing"))
    if bild is None:
        return
    inline = bild[0]
    extent = inline.find(qn("wp:extent"))
    if extent is None:
        return
    cx, cy = int(extent.get("cx")), int(extent.get("cy"))
    if cy <= int(max_hoehe):
        return
    faktor = int(max_hoehe) / cy
    neu_cx, neu_cy = int(cx * faktor), int(max_hoehe)
    extent.set("cx", str(neu_cx))
    extent.set("cy", str(neu_cy))
    for ext in inline.iter(qn("a:ext")):
        ext.set("cx", str(neu_cx))
        ext.set("cy", str(neu_cy))


# ─────────────────────────────────────────────────────────────────────────────
# Öffentliche Funktion
# ─────────────────────────────────────────────────────────────────────────────


def dateiname(daten: ProtokollDaten) -> str:
    """``Protokoll_16_2026-08-25_Neubau_Institutsgebaeude.docx``.

    Die Fußzeile des Originals zeigt den Dateinamen der Excel-Mappe
    (``260825_BB_16.xlsm``). Dieselbe Idee, nur mit dem Namen, den die App
    vergibt — so lässt sich ein Ausdruck weiterhin einer Datei zuordnen.
    """
    sicher = "".join(
        c if c.isalnum() or c in " -_" else "_" for c in daten.projektname
    ).strip().replace(" ", "_")[:40] or "Projekt"
    return (
        f"Protokoll_{daten.nummer:02d}_"
        f"{daten.besprechungsdatum.isoformat()}_{sicher}.docx"
    )


def generate_besprechungsprotokoll(daten: ProtokollDaten) -> tuple[Path, list[str]]:
    """Erzeugt das Word-Dokument.

    Gibt den Pfad zurück und eine Liste der Anlagen, die nicht eingebunden
    werden konnten — die verschweigt das Modul nicht, sondern reicht sie an
    den Router weiter, der sie in die Antwort schreibt.
    """
    if not daten.projektname:
        raise ProtokollFehler("Ohne Projektname kein Protokoll.")

    dokument = Document()
    _schrift_setzen(dokument)

    name = dateiname(daten)
    _richte_seite_ein(dokument, name)

    # Der leere Startabsatz von python-docx würde sonst als Leerzeile ganz
    # oben stehen.
    for absatz in list(dokument.paragraphs):
        absatz._p.getparent().remove(absatz._p)

    _seite_deckblatt(dokument, daten)
    _seitenumbruch(dokument)
    _seite_themen(dokument, daten)
    _seitenumbruch(dokument)
    _seite_legende(dokument, daten)
    _seitenumbruch(dokument)
    _seite_teilnehmer(dokument, daten)
    probleme = _seite_anlagen(dokument, daten)

    ziel = settings.output_dir / name
    ziel.parent.mkdir(parents=True, exist_ok=True)
    dokument.save(str(ziel))
    return ziel, probleme
