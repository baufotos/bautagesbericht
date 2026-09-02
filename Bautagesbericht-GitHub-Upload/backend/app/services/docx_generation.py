"""Erzeugt das HPP-Bautagebuch aus der Blanko-Word-Vorlage.

Die Vorlage (``Bautagesbericht_HPP_leer.docx``) enthält nur das Gerüst:
Logo, Projekt-Label und eine Tabelle mit fünf Zeilen
(Datum, Haupteintrag + Wetterlabels, Firmen, Firmen, Unterschrift).
Wetterwerte, Stundentabelle, Temperaturdiagramm und weitere Firmenblöcke
existieren dort noch nicht — sie werden hier eingefügt.

Aufbau und Formatierung sind an den ausgefüllten HPP-Referenzberichten
(``reference/HPP_Vorlagen``) ausgemessen:

  Projektname   Arial bold 11.5 pt, rechts neben dem Label "Projekt"
  Datum         "Fr 07.08.2026" — Wochentagskürzel + Datum
  Wetterwerte   Station linksbündig, Zahlenwerte rechtsbündig unter ihren Labels
  Stundenwerte  Uhrzeit / Temperaturdiagramm / Temp. (°C) / Wind (m/s) /
                Wind (Grad) / Bewölkung, jeweils zwölf Zweistundenwerte
  Firmenblock   zweispaltig: Labels links, Werte rechts; Firmenname fett,
                leere Felder werden weggelassen
  Fußzeile      "Bautagebuch - {Projekt} - {TT.MM.JJJJ}"
"""

import threading
from datetime import date
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Twips

from app.config import settings
from app.schemas import BautagesberichtJSON
from app.services import dokumenttext

FONT = "Arial"
SYMBOL_FONT = "Segoe UI Symbol"

# Schriftgrößen in Halbpunkten. Etwas großzügiger als die ursprünglich an den
# Referenzberichten ausgemessenen Werte (7 pt / 5.5 pt), da der reine
# Tabellentext sonst gedrängt und schwer lesbar wirkt.
SZ_NORMAL = 18      # 9 pt — Standardtext der Tabelle
SZ_SMALL = 14       # 7 pt — graue Hilfslabels
SZ_PROJEKT = 23     # 11.5 pt — Projektname in der Kopfzeile
SZ_SYMBOL = 22      # 11 pt — Wettersymbole der Bewölkungszeile
GRAY = "999999"

# Spaltenbreiten der Vorlagentabelle in Twips.
COL_WIDTHS = [1091, 2148, 909, 850, 549, 828, 3073]
LABEL_COL = COL_WIDTHS[0]
VALUE_SPAN = sum(COL_WIDTHS[1:])     # 8357 — Breite der sechs Wertspalten
TAB_VALUE_COL = 1050                 # Tabstopp für den Projektnamen

WEEKDAY_ABBR = ["Mo", "Di", "Mi", "Do", "Fr", "Sa", "So"]

# Stundentabelle: Beschriftungsspalte + zwölf Wertspalten.
HOURLY_LABEL_W = 1157
HOURLY_VALUE_W = (VALUE_SPAN - HOURLY_LABEL_W) // 12

# Temperaturdiagramm.
BAR_COLOR = "FFA500"
BAR_MAX_PT = 47
BAR_WIDTH_TWIPS = 150
BAR_INDENT = (HOURLY_VALUE_W - BAR_WIDTH_TWIPS) // 2

# Firmenblock: Label- und Wertspalte.
FIRMA_LABEL_W = 1440
FIRMA_VALUE_W = VALUE_SPAN - FIRMA_LABEL_W

# ── Luft im Dokument ────────────────────────────────────────────────────────
# Alle Werte in Twips (1/20 pt). Die erste Fassung setzte praktisch überall 0
# und entfernte zusätzlich die Zeilenhöhen der Vorlage — der fertige Bericht
# drängte sich dadurch ins obere Seitendrittel, während die untere Hälfte leer
# blieb, und "Firma / Personen / Leistung" klebten aneinander.
#
# Die Werte hier sind bewusst auf Lesbarkeit gestellt, nicht auf möglichst
# wenig Papier: Ein Tag mit einer Firma bleibt einseitig, ab zwei bis drei
# Firmen wird der Bericht zweiseitig. Das ist gewollt — ein Wetterdiagramm und
# fünf Angaben je Firma passen bei ordentlichem Abstand nicht zusammen auf ein
# Blatt, und gequetscht liest sie auf der Baustelle niemand.
#
# ANGEPASST WIRD DAS PRO BERICHT
# ------------------------------
# Ein Tag mit zwei Firmen darf Luft haben. Ein Tag mit sechs Firmen bekommt
# sonst einen Seitenumbruch mitten im Firmenteil, und weil eine Zeile nicht
# zerschnitten werden darf, bleibt unten auf Seite 1 eine handbreite Lücke —
# genau das, was am fertigen Bericht als "komisch" auffällt.
#
# Deshalb wird vor dem Bau geschätzt, wie hoch der Inhalt wird, und daraus
# eines von drei Abstandsmaßen gewählt (siehe ``_luft_waehlen``). Der Bericht
# bleibt so öfter einseitig, ohne dass etwas gequetscht aussieht.
LUFT_MASSE = {
    # Der bisherige Wert: großzügig, für ein bis drei Firmen.
    "weit": {
        "zelle": 90,        # Innenabstand je Zelle der Haupttabelle (4,5 pt)
        "feld": 170,        # zwischen Firma, Ort, Personen, Leistung (8,5 pt)
        "block": 220,       # zwischen zwei Firmenblöcken (11 pt)
        "notiz": 150,       # zwischen den Zeilen des Haupteintrags (7,5 pt)
        "mindesthoehe": 320,
        "zeilenabstand": 288,   # 1,2-fach
    },
    # Spürbar dichter, immer noch ruhig zu lesen.
    "normal": {
        "zelle": 70, "feld": 120, "block": 170, "notiz": 120,
        "mindesthoehe": 260, "zeilenabstand": 264,
    },
    # Für lange Tage. Enger als das wird nicht gesetzt — darunter klebt
    # "Leistung" wieder an "Personen", und auf der Baustelle liest das niemand.
    "eng": {
        "zelle": 50, "feld": 90, "block": 130, "notiz": 90,
        "mindesthoehe": 200, "zeilenabstand": 250,
    },
}

#: Das gerade gültige Maß. Wird von ``generate_bautagesbericht`` gesetzt.
#:
#: Je Thread getrennt: Ein Bericht wird ohne Unterbrechung gebaut, aber zwei
#: Berichte können in verschiedenen Threads gleichzeitig entstehen (eine Woche
#: mit fünf Tagen, zwei Einreichungen parallel). Ein gemeinsamer Wert würde
#: dann mitten im Bau umspringen und einen Bericht mit gemischten Abständen
#: hinterlassen.
_luft_je_thread = threading.local()


class _LuftZugriff:
    """Verhält sich wie ein dict, liest aber den Wert des eigenen Threads."""

    def _mass(self) -> dict:
        vorhanden = getattr(_luft_je_thread, "mass", None)
        if vorhanden is None:
            vorhanden = dict(LUFT_MASSE["weit"])
            _luft_je_thread.mass = vorhanden
        return vorhanden

    def __getitem__(self, name: str) -> int:
        return self._mass()[name]


_luft = _LuftZugriff()


def _luft_setzen(mass: dict) -> None:
    _luft_je_thread.mass = dict(mass)

ZELL_LUFT = LUFT_MASSE["weit"]["zelle"]
FELD_ABSTAND = LUFT_MASSE["weit"]["feld"]
BLOCK_ABSTAND = LUFT_MASSE["weit"]["block"]
NOTIZ_ABSTAND = LUFT_MASSE["weit"]["notiz"]
ZEILEN_MINDESTHOEHE = LUFT_MASSE["weit"]["mindesthoehe"]
ZEILENABSTAND_TEXT = LUFT_MASSE["weit"]["zeilenabstand"]

# Bright-Sky-Icons auf darstellbare Symbole abbilden.
WETTER_SYMBOLE = {
    "clear-day": "☀",
    "clear-night": "☾",
    "partly-cloudy-day": "⛅",
    "partly-cloudy-night": "☁",
    "cloudy": "☁",
    "fog": "░",
    "rain": "☔",
    "sleet": "☔",
    "snow": "❄",
    "hail": "❄",
    "thunderstorm": "⚡",
    "wind": "≋",
}


def _fmt_date(d: date) -> str:
    return d.strftime("%d.%m.%Y")


def _fmt_date_weekday(d: date) -> str:
    return f"{WEEKDAY_ABBR[d.weekday()]} {_fmt_date(d)}"


def _num(value, unit: str, decimals: int = 1) -> str:
    """Zahl mit Einheit; leerer String wenn kein Wert vorliegt."""
    if value is None:
        return ""
    if decimals == 0:
        return f"{round(value):.0f} {unit}".strip()
    return f"{value:.{decimals}f} {unit}".strip()


def _run(text: str, *, bold: bool = False, size: int = SZ_NORMAL,
         color: str | None = None, font: str = FONT) -> str:
    """Ein Textlauf.

    Zwei Dinge passieren hier, die im Aufrufer nicht mehr auftauchen sollen:

    * **Steuerzeichen weg.** XML verbietet sie, und ein einziges davon —
      etwa der Seitenvorschub aus einem kopierten PDF — ließ bisher das
      ganze Dokument nicht entstehen (siehe services/dokumenttext).
    * **Zeilenumbrüche werden welche.** Ein ``\\n`` mitten in ``<w:t>`` ist
      für Word bloßer Leerraum: Aus drei Leistungszeilen wurde eine lange.
      Mit ``<w:br/>`` steht wieder das da, was auf dem Blatt stand.
    """
    rpr = f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
    if bold:
        rpr += "<w:b/>"
    if color:
        rpr += f'<w:color w:val="{color}"/>'
    rpr += f'<w:sz w:val="{size}"/><w:szCs w:val="{size}"/>'

    sauber = dokumenttext.xml_sicher(text)
    stuecke = "<w:br/>".join(
        f'<w:t xml:space="preserve">{escape(zeile)}</w:t>'
        for zeile in sauber.split("\n")
    )
    return f"<w:r><w:rPr>{rpr}</w:rPr>{stuecke}</w:r>"


def _para(runs: str, *, align: str | None = None, after: int = 0,
          shading: str | None = None, indent: int = 0,
          exact_height: int | None = None, zeilenabstand: int = 240) -> str:
    ppr = ""
    if shading:
        ppr += f'<w:shd w:val="clear" w:color="auto" w:fill="{shading}"/>'
    if exact_height:
        ppr += f'<w:spacing w:after="0" w:line="{exact_height}" w:lineRule="exact"/>'
    else:
        ppr += (f'<w:spacing w:after="{after}" w:line="{zeilenabstand}" '
                f'w:lineRule="auto"/>')
    if indent:
        ppr += f'<w:ind w:left="{indent}" w:right="{indent}"/>'
    if align:
        ppr += f'<w:jc w:val="{align}"/>'
    return f"<w:p><w:pPr>{ppr}</w:pPr>{runs}</w:p>"


def _borders(top: bool, bottom: bool) -> str:
    def side(name: str, on: bool) -> str:
        if on:
            return f'<w:{name} w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
        return f'<w:{name} w:val="nil"/>'

    return (
        "<w:tcBorders>"
        + side("top", top)
        + side("left", False)
        + side("bottom", bottom)
        + side("right", False)
        + "</w:tcBorders>"
    )


def _cell(width: int, body: str, *, gridspan: int | None = None,
          top: bool = False, bottom: bool = False,
          valign: str | None = None, luft: int | None = None) -> str:
    if luft is None:
        luft = _luft["zelle"]
    tcpr = f'<w:tcW w:w="{width}" w:type="dxa"/>'
    if gridspan:
        tcpr += f'<w:gridSpan w:val="{gridspan}"/>'
    tcpr += _borders(top, bottom)
    if luft:
        # Innenabstand der Zelle: Ohne ihn sitzt der Text direkt auf der
        # Trennlinie zur nächsten Zeile.
        tcpr += (
            f'<w:tcMar><w:top w:w="{luft}" w:type="dxa"/>'
            f'<w:bottom w:w="{luft}" w:type="dxa"/></w:tcMar>'
        )
    if valign:
        tcpr += f'<w:vAlign w:val="{valign}"/>'
    return f"<w:tc><w:tcPr>{tcpr}</w:tcPr>{body or _para('')}</w:tc>"


def _row(cells: str, *, height: int | None = None,
         mindesthoehe: int | None = -1):
    """Eine Tabellenzeile.

    ``height`` erzwingt eine exakte Höhe (für die Balkengrafik nötig).
    Sonst gilt eine Mindesthöhe: Die Zeile wächst mit dem Inhalt, fällt aber
    nicht auf Schriftgröße zusammen.
    """
    # cantSplit: Eine Zeile wandert ganz auf die nächste Seite, statt am
    # Seitenrand zerschnitten zu werden. Ohne das stand von der
    # Unterschriftszeile nur das Wort "Datum" auf Seite zwei.
    if mindesthoehe == -1:
        mindesthoehe = _luft["mindesthoehe"]
    inhalt = "<w:cantSplit/>"
    if height:
        inhalt += f'<w:trHeight w:val="{height}" w:hRule="exact"/>'
    elif mindesthoehe:
        inhalt += f'<w:trHeight w:val="{mindesthoehe}" w:hRule="atLeast"/>'
    return parse_xml(f"<w:tr {nsdecls('w')}><w:trPr>{inhalt}</w:trPr>{cells}</w:tr>")


def _nested_table_open(width: int, grid: str) -> str:
    return (
        f"<w:tbl><w:tblPr>"
        f'<w:tblW w:w="{width}" w:type="dxa"/>'
        f"<w:tblBorders><w:top w:val=\"nil\"/><w:left w:val=\"nil\"/>"
        f"<w:bottom w:val=\"nil\"/><w:right w:val=\"nil\"/>"
        f"<w:insideH w:val=\"nil\"/><w:insideV w:val=\"nil\"/></w:tblBorders>"
        f'<w:tblLayout w:type="fixed"/>'
        f'<w:tblCellMar><w:top w:w="70" w:type="dxa"/><w:left w:w="0" w:type="dxa"/>'
        f'<w:bottom w:w="70" w:type="dxa"/><w:right w:w="28" w:type="dxa"/></w:tblCellMar>'
        f"</w:tblPr><w:tblGrid>{grid}</w:tblGrid>"
    )


def _weather_value_row(w):
    """Werteszeile unter den Wetter-Labels.

    Station linksbündig, die vier schmalen Zahlenspalten zentriert unter ihrem
    Label — so wie in den Referenzberichten ausgemessen. Die Schnee-Spalte ist
    in der Vorlage bis zum Seitenrand breit; ihr Wert bleibt daher linksbündig
    direkt unter dem Label stehen.
    """
    numeric = [
        (_num(w.temp_max_c, "°C"), "center"),
        (_num(w.temp_min_c, "°C"), "center"),
        (_num(w.regen_mm, "mm"), "center"),
        (_num(w.wind_max_ms, "m/s"), "center"),
        (_num(w.schnee_cm, "cm", decimals=0), None),
    ]
    cells = _cell(LABEL_COL, _para(""))
    cells += _cell(COL_WIDTHS[1], _para(_run(w.station or ""), after=100))
    for width, (value, align) in zip(COL_WIDTHS[2:], numeric):
        cells += _cell(width, _para(_run(value), align=align, after=100))
    return _row(cells)


def _diagramm_bereich(temps: list[float]) -> tuple[float, float]:
    """Untere und obere Kante des Temperaturdiagramms.

    Der Boden ist 0 °C, solange der Tag darüber bleibt — damit sind die
    Diagramme zweier Tage derselben Woche miteinander vergleichbar.

    Bei Frost sinkt der Boden mit. Vorher war er fest auf null: An einem Tag
    zwischen -5 und -1 Grad war deshalb kein einziger Balken zu sehen, das
    Diagramm verschwand im Winter stillschweigend, und im Bericht an den
    Bauherrn blieb an dieser Stelle eine leere Fläche. Ausgerechnet an
    Frosttagen ist die Temperatur aber die Angabe, auf die es ankommt.
    """
    if not temps:
        return 0.0, 1.0
    unten = min(0.0, min(temps))
    # Mindestspanne, damit ein Tag mit durchweg gleicher Temperatur nicht
    # durch Null geteilt wird und nicht als volle Balkenreihe erscheint.
    oben = max(max(temps), unten + 1.0)
    return unten, oben


#: Mindesthöhe eines Balkens in Punkt. Der kälteste Wert des Tages liegt auf
#: dem Boden des Diagramms; ganz ohne Balken sähe diese Stunde aus wie eine,
#: für die es überhaupt keine Messung gab.
BAR_MIN_PT = 3


def _bar_cell(temp: float | None, unten: float, oben: float) -> str:
    """Ein Balken des Temperaturdiagramms, unten ausgerichtet."""
    if temp is None:
        return _cell(HOURLY_VALUE_W, _para(""), valign="bottom")
    anteil = (temp - unten) / (oben - unten)
    height_pt = BAR_MAX_PT * min(1.0, max(0.0, anteil))
    body = _para(
        "",
        shading=BAR_COLOR,
        indent=BAR_INDENT,
        exact_height=max(int(height_pt * 20), BAR_MIN_PT * 20),
    )
    return _cell(HOURLY_VALUE_W, body, valign="bottom")


def _hourly_table(stundenwerte: list) -> str:
    """Verschachtelte Tabelle: Uhrzeiten, Temperaturdiagramm, Messwerte, Symbole."""
    values = stundenwerte[:12]
    temps = [s.temperatur_c for s in values if s.temperatur_c is not None]
    unten, oben = _diagramm_bereich(temps)

    grid = f'<w:gridCol w:w="{HOURLY_LABEL_W}"/>' + (
        f'<w:gridCol w:w="{HOURLY_VALUE_W}"/>' * 12
    )
    tbl = _nested_table_open(VALUE_SPAN, grid)

    def text_row(label: str, render, *, font: str = FONT, size: int = SZ_NORMAL,
                 after: int = 0) -> str:
        row = (
            f'<w:tc><w:tcPr><w:tcW w:w="{HOURLY_LABEL_W}" w:type="dxa"/>'
            f"{_borders(False, False)}</w:tcPr>"
            f"{_para(_run(label, bold=True), after=after)}</w:tc>"
        )
        for s in values:
            row += (
                f'<w:tc><w:tcPr><w:tcW w:w="{HOURLY_VALUE_W}" w:type="dxa"/>'
                f"{_borders(False, False)}</w:tcPr>"
                f"{_para(_run(render(s), font=font, size=size), align='center', after=after)}</w:tc>"
            )
        return f"<w:tr>{row}</w:tr>"

    tbl += text_row("Uhrzeit", lambda s: f"{s.stunde:02d}", after=60)

    # Diagrammzeile: feste Höhe, Balken wachsen von unten.
    if temps:
        chart = (
            f'<w:tc><w:tcPr><w:tcW w:w="{HOURLY_LABEL_W}" w:type="dxa"/>'
            f"{_borders(False, False)}</w:tcPr>{_para('')}</w:tc>"
        )
        for s in values:
            chart += _bar_cell(s.temperatur_c, unten, oben)
        tbl += (
            f'<w:tr><w:trPr><w:trHeight w:val="{BAR_MAX_PT * 20}" '
            f'w:hRule="exact"/></w:trPr>{chart}</w:tr>'
        )

    tbl += text_row("Temp. (°C)", lambda s: "" if s.temperatur_c is None else f"{s.temperatur_c:.1f}", after=40)
    tbl += text_row("Wind (m/s)", lambda s: "" if s.wind_ms is None else f"{s.wind_ms:.1f}", after=40)
    tbl += text_row("Wind (Grad)", lambda s: "" if s.wind_grad is None else f"{s.wind_grad:.0f}", after=40)
    tbl += text_row(
        "Bewölkung",
        lambda s: WETTER_SYMBOLE.get(s.icon or "", ""),
        font=SYMBOL_FONT,
        size=SZ_SYMBOL,
        after=40,
    )

    return tbl + "</w:tbl>"


def _hourly_row(stundenwerte: list):
    body = _hourly_table(stundenwerte) + _para("")
    cells = _cell(LABEL_COL, _para(""))
    cells += _cell(VALUE_SPAN, body, gridspan=6)
    return _row(cells)


def _note_row(text: str, *, bottom: bool):
    paras = "".join(
        _para(_run(line), after=_luft["notiz"])
        for line in text.splitlines() if line.strip()
    )
    cells = _cell(LABEL_COL, _para(""), bottom=bottom)
    cells += _cell(VALUE_SPAN, paras or _para(""), gridspan=6, bottom=bottom)
    return _row(cells)


def _firma_row(firma, *, bottom: bool = True, zeige_label: bool = True):
    """Ein Firmenblock als zweispaltige Tabelle. Leere Felder werden weggelassen."""
    fields: list[tuple[str, str, bool]] = []
    if firma.firma:
        fields.append(("Firma:", firma.firma, True))
    if firma.ort:
        fields.append(("Ort:", firma.ort, False))
    if firma.personen:
        fields.append(("Personen:", str(firma.personen), False))
    if firma.leistung:
        fields.append(("Leistung:", firma.leistung, False))
    if firma.besonderes:
        fields.append(("Besonders:", firma.besonderes, False))

    grid = f'<w:gridCol w:w="{FIRMA_LABEL_W}"/><w:gridCol w:w="{FIRMA_VALUE_W}"/>'
    tbl = _nested_table_open(VALUE_SPAN, grid)
    for i, (label, value, value_bold) in enumerate(fields):
        after = 0 if i == len(fields) - 1 else _luft["feld"]
        tbl += (
            "<w:tr>"
            f'<w:tc><w:tcPr><w:tcW w:w="{FIRMA_LABEL_W}" w:type="dxa"/>'
            f"{_borders(False, False)}</w:tcPr>{_para(_run(label), after=after)}</w:tc>"
            f'<w:tc><w:tcPr><w:tcW w:w="{FIRMA_VALUE_W}" w:type="dxa"/>'
            f"{_borders(False, False)}</w:tcPr>"
            f"{_para(_run(value, bold=value_bold), after=after, zeilenabstand=_luft['zeilenabstand'])}</w:tc>"
            "</w:tr>"
        )
    tbl += "</w:tbl>"

    # Der Abschlussabsatz trennt diesen Firmenblock vom nächsten. Ohne ihn
    # gingen zwei Firmen ineinander über und man müsste die Zeilen zählen,
    # um zu sehen, wo die eine aufhört.
    body = (tbl + _para("", after=_luft["block"])) if fields else _para("")
    # "Firmen" ist die Überschrift des Abschnitts, nicht die jeder einzelnen
    # Firma — bei drei Nachunternehmern stand es sonst dreimal untereinander.
    beschriftung = _para(_run("Firmen", bold=True)) if zeige_label else _para("")
    cells = _cell(LABEL_COL, beschriftung, bottom=bottom)
    cells += _cell(VALUE_SPAN, body, gridspan=6, bottom=bottom)
    return _row(cells)


def _set_cell_paragraphs(cell_tc, paragraphs_xml: str) -> None:
    """Ersetzt den Inhalt einer bestehenden Vorlagenzelle."""
    for child in list(cell_tc):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            cell_tc.remove(child)
    for para in parse_xml(f"<w:root {nsdecls('w')}>{paragraphs_xml}</w:root>"):
        cell_tc.append(para)


def _fill_projekt_header(doc, projekt: str) -> None:
    """Projektname mit Tabstopp rechts neben das Label 'Projekt' setzen."""
    for para in doc.paragraphs:
        if para.text.strip() == "Projekt":
            para.paragraph_format.tab_stops.add_tab_stop(Twips(TAB_VALUE_COL))
            para._p.append(parse_xml(f"<w:r {nsdecls('w')}><w:tab/></w:r>"))
            para._p.append(parse_xml(
                f"<w:r {nsdecls('w')}>"
                f'<w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/><w:b/>'
                f'<w:sz w:val="{SZ_PROJEKT}"/><w:szCs w:val="{SZ_PROJEKT}"/></w:rPr>'
                f'<w:t xml:space="preserve">'
                f'{escape(dokumenttext.einzeilig(projekt))}</w:t></w:r>'
            ))
            return


#: Schriftgrößen der Fußzeile in Halbpunkten — aus der Vorlage übernommen.
SZ_FUSS_ZEILE = 13     # 6,5 pt — "Bautagebuch - Projekt - Datum"
SZ_FUSS_SEITE = 16     # 8 pt — die Seitenzahl

#: Weiches Trennzeichen als Zierstrich links und rechts der Seitenzahl, genau
#: wie in der Blanko-Vorlage ("­ 1 / 1 ­").
ZIERSTRICH = "­"


def _feld(anweisung: str, vorschau: str, *, size: int = SZ_FUSS_SEITE) -> str:
    """Ein Word-Feld wie PAGE oder NUMPAGES.

    Fester Text ginge hier nicht: Die Seitenzahl muss Word beim Umbruch selbst
    ausrechnen. ``vorschau`` ist der Wert, der in Betrachtern ohne
    Feldberechnung stehen bleibt.
    """
    rpr = (f'<w:rPr><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}"/>'
           f'<w:sz w:val="{size}"/><w:szCs w:val="{size}"/></w:rPr>')
    return (
        f"<w:r>{rpr}<w:fldChar w:fldCharType=\"begin\"/></w:r>"
        f"<w:r>{rpr}<w:instrText xml:space=\"preserve\"> {anweisung} </w:instrText></w:r>"
        f"<w:r>{rpr}<w:fldChar w:fldCharType=\"separate\"/></w:r>"
        f'<w:r>{rpr}<w:t>{escape(vorschau)}</w:t></w:r>'
        f"<w:r>{rpr}<w:fldChar w:fldCharType=\"end\"/></w:r>"
    )


def _fill_footer_line(doc, projekt: str, datum: date) -> None:
    """Baut die echte Word-Fußzeile und räumt die Vorlagenzeilen aus dem Text.

    Die Blanko-Vorlage trägt "Bautagebuch" und "­ 1 / 1 ­" als gewöhnliche
    Absätze am Dokumentende — für ein Blatt, das von Hand ausgefüllt wird,
    reicht das. Im erzeugten Bericht ist es falsch: Die Zeilen stehen dann nur
    auf der letzten Seite, und "1 / 1" bleibt "1 / 1", auch wenn der Bericht
    über zwei Seiten läuft.

    Deshalb wandert beides in die Fußzeile des Abschnitts, die Seitenzahl als
    PAGE/NUMPAGES-Feld. Word setzt sie damit auf jeder Seite und rechnet sie
    selbst aus.
    """
    zeile_links = f"Bautagebuch - {projekt} - {_fmt_date(datum)}"

    abschnitt = doc.sections[0]
    fuss = abschnitt.footer
    fuss.is_linked_to_previous = False

    for para in list(fuss.paragraphs):
        para._element.getparent().remove(para._element)

    absaetze = (
        _para(_run(zeile_links, size=SZ_FUSS_ZEILE))
        + _para(
            _run(f"{ZIERSTRICH} ", size=SZ_FUSS_SEITE)
            + _feld("PAGE", "1")
            + _run(" / ", size=SZ_FUSS_SEITE)
            + _feld("NUMPAGES", "1")
            + _run(f" {ZIERSTRICH}", size=SZ_FUSS_SEITE),
            align="right",
        )
    )
    for element in parse_xml(f"<w:root {nsdecls('w')}>{absaetze}</w:root>"):
        fuss._element.append(element)

    # Die Vorlagenzeilen im Text entfernen — sonst stünde alles doppelt.
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text == "Bautagebuch" or text.strip(ZIERSTRICH + " ").replace(
                " ", "") in ("1/1",):
            para._element.getparent().remove(para._element)


def _ist_leer(para) -> bool:
    return not "".join(t.text or "" for t in para.iter(qn("w:t"))).strip()


def _shrink_page_spacer(doc) -> None:
    """Verhindert ein leeres Blatt am Ende des Berichts.

    Die Blanko-Vorlage schiebt die Fußzeile mit einem ~12,7 cm hohen
    ``w:after``-Absatz ans Seitenende — gedacht für den handschriftlichen
    Gebrauch. Im erzeugten Bericht bleibt davon nichts Sichtbares übrig, aber
    genug Höhe, um ein zweites, komplett leeres Blatt zu erzwingen: Bei drei
    Firmen kam genau das heraus, und beim Drucken fällt so ein Blatt jedem
    auf.

    Deshalb werden die leeren Absätze hinter der letzten Tabelle bis auf einen
    entfernt (Word braucht dort einen) und der verbleibende auf minimale Höhe
    gesetzt. Absätze mit Text bleiben unangetastet.
    """
    body = doc.element.body
    for para in body.findall(qn("w:p")):
        ppr = para.find(qn("w:pPr"))
        if ppr is None:
            continue
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            continue
        after = spacing.get(qn("w:after"))
        if after is not None and int(after) > 480:
            spacing.set(qn("w:after"), "240")

    kinder = list(body)
    tabellen = [i for i, k in enumerate(kinder) if k.tag == qn("w:tbl")]
    if not tabellen:
        return

    dahinter = [k for k in kinder[tabellen[-1] + 1:] if k.tag == qn("w:p")]
    leere = [p for p in dahinter if _ist_leer(p)]

    # Alle bis auf den letzten leeren Absatz entfernen.
    for para in leere[:-1]:
        body.remove(para)

    if leere:
        rest = leere[-1]
        ppr = rest.find(qn("w:pPr"))
        if ppr is None:
            ppr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
            rest.insert(0, ppr)
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = parse_xml(f"<w:spacing {nsdecls('w')}/>")
            ppr.append(spacing)
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:line"), "120")
        spacing.set(qn("w:lineRule"), "exact")


def _set_bottom_border(tr, visible: bool) -> None:
    for tc in tr.findall(qn("w:tc")):
        borders = tc.find(f"{qn('w:tcPr')}/{qn('w:tcBorders')}")
        if borders is None:
            continue
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            continue
        if visible:
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "5")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")
        else:
            bottom.set(qn("w:val"), "nil")
            for attr in ("w:sz", "w:color", "w:space"):
                if bottom.get(qn(attr)) is not None:
                    del bottom.attrib[qn(attr)]


def _drop_row_height(tr) -> None:
    """Ersetzt die feste Zeilenhöhe der Vorlage durch eine Mindesthöhe.

    Die Blanko-Vorlage reserviert je Zeile mehrere Zentimeter für die
    handschriftliche Nutzung — das ist im gedruckten Bericht viel zu viel.
    Die Höhe ganz zu entfernen war aber das andere Extrem: Dann fällt die
    Zeile auf Schriftgröße zusammen und der Text sitzt auf der Trennlinie.
    Also: wachsen darf sie, unter ``ZEILEN_MINDESTHOEHE`` fällt sie nicht.
    """
    trpr = tr.find(qn("w:trPr"))
    if trpr is None:
        trpr = parse_xml(f"<w:trPr {nsdecls('w')}/>")
        tr.insert(0, trpr)
    height = trpr.find(qn("w:trHeight"))
    if height is not None:
        height.set(qn("w:val"), str(ZEILEN_MINDESTHOEHE))
        height.set(qn("w:hRule"), "atLeast")
    # Auch die Zeilen aus der Vorlage sollen nicht am Seitenrand zerreißen.
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.insert(0, parse_xml(f"<w:cantSplit {nsdecls('w')}/>"))


def _collapse_cell(tc) -> None:
    """Überzählige Leerabsätze einer Vorlagenzelle entfernen."""
    paragraphs = tc.findall(qn("w:p"))
    for para in paragraphs[1:]:
        tc.remove(para)


# ─────────────────────────────────────────────────────────────────────────────
# Wie hoch wird das?
#
# Word rechnet den Umbruch erst beim Öffnen aus; hier muss vorher entschieden
# werden, wie viel Luft der Bericht verträgt. Die Schätzung ist grob und darf
# es sein: Sie soll nur zwischen "passt bequem", "wird knapp" und "wird lang"
# unterscheiden. Liegt sie daneben, sieht der Bericht etwas luftiger oder
# dichter aus als nötig — es geht nichts verloren.
# ─────────────────────────────────────────────────────────────────────────────

#: Zeichen, die bei Arial 9 pt in eine Wertspalte passen. Aus der Spaltenbreite
#: (6917 Twips) und einer mittleren Zeichenbreite von rund 90 Twips.
ZEICHEN_JE_ZEILE = 76

#: Nutzbare Höhe einer A4-Seite: 297 mm abzüglich der Ränder der Vorlage und
#: abzüglich Kopf- und Fußzeile. In Twips.
SEITENHOEHE_NUTZBAR = 13000

#: Was Datum, Wetterblock und Haupteintrag ungefähr belegen, bevor die erste
#: Firma kommt.
KOPF_HOEHE = 1800


def _zeilen(text: str) -> int:
    """Wie viele Zeilen ein Wert in der Wertspalte braucht."""
    if not text:
        return 0
    gesamt = 0
    for absatz in str(text).splitlines() or [""]:
        gesamt += max(1, -(-len(absatz) // ZEICHEN_JE_ZEILE))
    return max(1, gesamt)


def _hoehe_firma(firma, mass: dict) -> int:
    """Geschätzte Höhe eines Firmenblocks in Twips."""
    felder = [firma.firma, firma.ort, str(firma.personen or ""),
              firma.leistung, firma.besonderes or ""]
    gefuellt = [f for f in felder if str(f).strip()]
    hoehe = 0
    for wert in gefuellt:
        hoehe += _zeilen(wert) * mass["zeilenabstand"] + mass["feld"]
    # Innenabstand der Zelle oben und unten, plus der Trennabsatz zum
    # nächsten Block.
    hoehe += 2 * mass["zelle"] + mass["block"]
    return max(hoehe, mass["mindesthoehe"])


def _geschaetzte_hoehe(data, mass: dict) -> int:
    hoehe = KOPF_HOEHE + 2 * mass["zelle"]
    if data.haupteintrag:
        hoehe += _zeilen(data.haupteintrag) * (mass["zeilenabstand"] + mass["notiz"])
    if data.wetter and data.wetter.stundenwerte:
        # Der Balkenblock hat eine feste Höhe, unabhängig vom Maß.
        hoehe += 1900
    for firma in data.firmen:
        hoehe += _hoehe_firma(firma, mass)
    # Unterschriftszeile.
    hoehe += 700
    return hoehe


def _luft_waehlen(data) -> dict:
    """Das größte Abstandsmaß, mit dem der Bericht noch auf eine Seite passt.

    Passt er auch eng nicht auf eine Seite — bei sechs Firmen mit langen
    Leistungstexten ist das normal —, wird ``normal`` genommen: Dann ist der
    Umbruch ohnehin da, und Enge bringt nichts als schlechtere Lesbarkeit.
    """
    for name in ("weit", "normal", "eng"):
        mass = LUFT_MASSE[name]
        if _geschaetzte_hoehe(data, mass) <= SEITENHOEHE_NUTZBAR:
            return dict(mass)
    return dict(LUFT_MASSE["normal"])


def _mit_naechster_zeile(tr) -> None:
    """Bindet eine Tabellenzeile an die folgende (``keepNext``).

    Gebraucht für die letzte Firma vor der Unterschrift: Ohne das rutscht bei
    einem zweiseitigen Bericht die Unterschriftszeile allein auf Seite 2 und
    steht dort verloren unter einer leeren Fläche.
    """
    for absatz in tr.iter(qn("w:p")):
        eigenschaften = absatz.find(qn("w:pPr"))
        if eigenschaften is None:
            eigenschaften = parse_xml(f"<w:pPr {nsdecls('w')}/>")
            absatz.insert(0, eigenschaften)
        if eigenschaften.find(qn("w:keepNext")) is None:
            eigenschaften.append(parse_xml(f"<w:keepNext {nsdecls('w')}/>"))


def dateiname(projekt: str, datum: date, kennung: str = "") -> str:
    """Der Dateiname des erzeugten Berichts.

    ``kennung`` unterscheidet zwei Berichte für denselben Tag desselben
    Projekts. Ohne sie schrieb der zweite den ersten still über — und weil in
    der Datenbank beide auf denselben Pfad zeigten, lieferte der Download des
    älteren Berichts danach den Inhalt des neueren. Das fällt niemandem auf.
    """
    sauber = "".join(
        c if c.isalnum() or c in " -_" else "_"
        for c in dokumenttext.einzeilig(projekt)
    ).strip().replace(" ", "_")[:40] or "Projekt"
    zusatz = f"_{kennung}" if kennung else ""
    return f"BTB_{datum.isoformat()}_{sauber}{zusatz}.docx"


def anzeigename(projekt: str, datum: date) -> str:
    """Der Name, unter dem der Bericht beim Herunterladen erscheinen soll.

    Getrennt vom Dateinamen auf der Platte: Dort braucht es die Kennung zur
    Unterscheidung, im Download-Ordner des Anwenders wäre sie nur Ballast.
    """
    sauber = dokumenttext.einzeilig(projekt).replace("/", "-") or "Projekt"
    return f"Bautagesbericht {datum.strftime('%Y-%m-%d')} {sauber}.docx"


def generate_bautagesbericht(
    data: BautagesberichtJSON,
    template_path: Path | None = None,
    kennung: str = "",
) -> Path:
    if template_path is None:
        template_path = settings.template_dir / "Bautagesbericht_HPP_leer.docx"

    doc = Document(str(template_path))

    # Muss vor dem ersten Zeilenbau geschehen: Alle Baufunktionen lesen das
    # gültige Abstandsmaß aus ``_luft``.
    _luft_setzen(_luft_waehlen(data))

    _fill_projekt_header(doc, data.projekt)
    _fill_footer_line(doc, data.projekt, data.datum)
    _shrink_page_spacer(doc)

    table = doc.tables[0]
    rows = table._tbl.findall(qn("w:tr"))
    row_datum, row_wetter, row_firma, row_firma_leer, row_unterschrift = rows[:5]

    # Die Vorlage reserviert feste Zeilenhöhen für den handschriftlichen
    # Gebrauch; im generierten Bericht richtet sich die Höhe nach dem Inhalt.
    for tr in rows[:5]:
        _drop_row_height(tr)
    for tc in row_firma_leer.findall(qn("w:tc")):
        _collapse_cell(tc)

    # Datum
    datum_cells = row_datum.findall(qn("w:tc"))
    _set_cell_paragraphs(datum_cells[1], _para(_run(_fmt_date_weekday(data.datum))))

    # Wetterblock unterhalb der bestehenden Label-Zeile aufbauen
    anchor = row_wetter
    block_rows = []

    def append_block_row(new_row) -> None:
        nonlocal anchor
        anchor.addnext(new_row)
        anchor = new_row
        block_rows.append(new_row)

    if data.wetter:
        append_block_row(_weather_value_row(data.wetter))
        if data.wetter.stundenwerte:
            append_block_row(_hourly_row(data.wetter.stundenwerte))

    if data.haupteintrag:
        append_block_row(_note_row(data.haupteintrag, bottom=False))

    if block_rows:
        # Labels und Werte bilden einen Block: nur die letzte Zeile wird
        # nach unten abgegrenzt.
        _set_bottom_border(row_wetter, False)
        _set_bottom_border(block_rows[-1], True)

    # Firmenblöcke: erste Vorlagenzeile ersetzen, weitere dahinter einfügen.
    # Die Vorlage reserviert zwei leere Firmenzeilen für den handschriftlichen
    # Gebrauch (row_firma, row_firma_leer) — sobald echte Firmendaten vorliegen,
    # werden beide Vorlagenzeilen durch die tatsächlichen Blöcke ersetzt, damit
    # keine leere "Firmen"-Zeile übrig bleibt.
    if data.firmen:
        firma_anchor = row_firma
        for i, firma in enumerate(data.firmen):
            new_row = _firma_row(firma, zeige_label=(i == 0))
            firma_anchor.addnext(new_row)
            firma_anchor = new_row
        table._tbl.remove(row_firma)
        table._tbl.remove(row_firma_leer)
        # Die letzte Firma bleibt bei der Unterschrift. Sonst steht die
        # Unterschriftszeile auf einem zweiseitigen Bericht allein oben auf
        # Seite 2 — mit einer leeren Seite darunter.
        if firma_anchor is not row_firma:
            _mit_naechster_zeile(firma_anchor)

    # Unterschriftsdatum über das graue Label setzen
    unterschrift_cells = row_unterschrift.findall(qn("w:tc"))
    signed = data.unterschrift_datum or date.today()
    _set_cell_paragraphs(
        unterschrift_cells[1],
        _para(_run(_fmt_date(signed)), align="center")
        + _para(_run("Datum ", size=SZ_SMALL, color=GRAY), align="center"),
    )
    # In den Referenzberichten steht das Komma vor dem Firmennamen.
    hpp_cell = unterschrift_cells[2]
    first_run = hpp_cell.find(f"{qn('w:p')}/{qn('w:r')}/{qn('w:t')}")
    if first_run is not None and not (first_run.text or "").startswith(","):
        first_run.text = f", {first_run.text or ''}"

    output_path = settings.output_dir / dateiname(
        data.projekt, data.datum, kennung)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return output_path
