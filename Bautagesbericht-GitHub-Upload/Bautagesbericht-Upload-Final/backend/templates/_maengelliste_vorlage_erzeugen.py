"""Erzeugt die Platzhalter-Vorlage ``Maengelliste_leer.docx``.

WARUM DIESES SKRIPT EXISTIERT
=============================
Die echte Mängellisten-Vorlage des Büros liegt noch nicht vor. Damit der
Export trotzdem heute funktioniert, wird hier eine schlichte Blanko-Vorlage im
Stil von ``Bautagesbericht_HPP_leer.docx`` erzeugt: gleiches Seitenformat,
gleiche Ränder, gleiches HPP-Logo (wird aus der Bautagesbericht-Vorlage
übernommen), Arial.

Sobald die echte Vorlage da ist:
  * Datei einfach als ``Maengelliste_leer.docx`` in diesen Ordner legen —
    dieses Skript wird dann nicht mehr gebraucht und kann gelöscht werden.
  * In der echten Vorlage müssen nur die Platzhalter stehen, die
    ``app/services/maengelliste_generation.py`` ersetzt (dort oben im
    Modul-Kommentar aufgeführt): {{PROJEKT}}, {{STAND}}, {{ANZAHL}},
    {{FILTER}}, {{HINWEIS}}, {{FUSSZEILE}} und genau ein Absatz {{MAENGEL}}.

Aufruf (nur bei Bedarf, das Ergebnis liegt im Repo):
    python backend/templates/_maengelliste_vorlage_erzeugen.py
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

ORDNER = Path(__file__).resolve().parent
QUELLE_LOGO = ORDNER / "Bautagesbericht_HPP_leer.docx"
ZIEL = ORDNER / "Maengelliste_leer.docx"

FONT = "Arial"
GRAU = RGBColor(0x77, 0x74, 0x6B)
SCHWARZ = RGBColor(0x20, 0x1F, 0x1D)
LINIE = "D8D5CD"

# Seitenränder der Bautagesbericht-Vorlage (Twips), damit beide Dokumente
# nebeneinander gleich aussehen.
RAND = {"top": 560, "right": 1140, "bottom": 1440, "left": 1420}


def _logo_bytes() -> bytes | None:
    """Holt das HPP-Logo aus der bestehenden Bautagesbericht-Vorlage."""
    if not QUELLE_LOGO.is_file():
        return None
    with zipfile.ZipFile(QUELLE_LOGO) as archiv:
        namen = [n for n in archiv.namelist() if n.startswith("word/media/")]
        if not namen:
            return None
        return archiv.read(sorted(namen)[0])


def _schreibe(absatz, text: str, *, groesse=Pt(9), bold=False, farbe=SCHWARZ):
    lauf = absatz.add_run(text)
    lauf.font.name = FONT
    lauf.font.size = groesse
    lauf.font.bold = bold
    lauf.font.color.rgb = farbe
    return lauf


def _untere_linie(absatz) -> None:
    ppr = absatz._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    unten = OxmlElement("w:bottom")
    unten.set(qn("w:val"), "single")
    unten.set(qn("w:sz"), "4")
    unten.set(qn("w:space"), "2")
    unten.set(qn("w:color"), LINIE)
    borders.append(unten)
    ppr.append(borders)


def _feld(absatz, instr: str, platzhalter: str) -> None:
    """Word-Feld (z. B. PAGE) in einen Absatz setzen."""
    feld = OxmlElement("w:fldSimple")
    feld.set(qn("w:instr"), instr)
    lauf = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    groesse = OxmlElement("w:sz")
    groesse.set(qn("w:val"), "15")
    rpr.append(fonts)
    rpr.append(groesse)
    text = OxmlElement("w:t")
    text.text = platzhalter
    lauf.append(rpr)
    lauf.append(text)
    feld.append(lauf)
    absatz._p.append(feld)


def erzeuge_vorlage() -> Path:
    doc = Document()

    # Standardschrift auf Arial stellen.
    stil = doc.styles["Normal"]
    stil.font.name = FONT
    stil.font.size = Pt(9)

    abschnitt = doc.sections[0]
    abschnitt.top_margin = Twips(RAND["top"])
    abschnitt.right_margin = Twips(RAND["right"])
    abschnitt.bottom_margin = Twips(RAND["bottom"])
    abschnitt.left_margin = Twips(RAND["left"])

    # Kopf: Logo links, Titel darunter.
    logo = _logo_bytes()
    if logo:
        kopf = doc.add_paragraph()
        kopf.paragraph_format.space_after = Pt(10)
        kopf.add_run().add_picture(io.BytesIO(logo), width=Cm(3.4))

    titel = doc.add_paragraph()
    titel.paragraph_format.space_after = Pt(2)
    _schreibe(titel, "Mängelliste", groesse=Pt(16), bold=True)

    untertitel = doc.add_paragraph()
    untertitel.paragraph_format.space_after = Pt(10)
    _schreibe(untertitel, "HPP Architekten Baumanagement", groesse=Pt(8), farbe=GRAU)

    # Kopfangaben als schlanke zweispaltige Tabelle mit Platzhaltern.
    kopfdaten = [
        ("Projekt", "{{PROJEKT}}"),
        ("Stand", "{{STAND}}"),
        ("Auswahl", "{{FILTER}}"),
        ("Anzahl Mängel", "{{ANZAHL}}"),
    ]
    tabelle = doc.add_table(rows=0, cols=2)
    tabelle.autofit = False
    for label, platzhalter in kopfdaten:
        zeile = tabelle.add_row()
        links, rechts = zeile.cells[0], zeile.cells[1]
        links.width = Cm(3.6)
        rechts.width = Cm(12.9)
        for lauf in list(links.paragraphs[0].runs):
            lauf._r.getparent().remove(lauf._r)
        _schreibe(links.paragraphs[0], label.upper(), groesse=Pt(7.5), farbe=GRAU)
        for lauf in list(rechts.paragraphs[0].runs):
            lauf._r.getparent().remove(lauf._r)
        _schreibe(rechts.paragraphs[0], platzhalter, bold=(label == "Projekt"))

    hinweis = doc.add_paragraph()
    hinweis.paragraph_format.space_before = Pt(6)
    hinweis.paragraph_format.space_after = Pt(10)
    _schreibe(hinweis, "{{HINWEIS}}", groesse=Pt(7.5), farbe=GRAU)
    _untere_linie(hinweis)

    # Marker: Hier setzt die Dokumenterzeugung je Mangel eine Tabelle ein.
    marker = doc.add_paragraph()
    _schreibe(marker, "{{MAENGEL}}", groesse=Pt(9), farbe=GRAU)

    # Fußzeile: Beschriftung links, Seitenzahl rechts (wiederholt sich auf
    # jeder Seite — eine Mängelliste wird schnell mehrseitig).
    fuss = abschnitt.footer.paragraphs[0]
    fuss.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _schreibe(fuss, "{{FUSSZEILE}}", groesse=Pt(7.5), farbe=GRAU)
    _schreibe(fuss, "\t\tSeite ", groesse=Pt(7.5), farbe=GRAU)
    _feld(fuss, "PAGE", "1")
    _schreibe(fuss, " von ", groesse=Pt(7.5), farbe=GRAU)
    _feld(fuss, "NUMPAGES", "1")

    doc.save(str(ZIEL))
    return ZIEL


if __name__ == "__main__":
    print(f"Vorlage geschrieben: {erzeuge_vorlage()}")
