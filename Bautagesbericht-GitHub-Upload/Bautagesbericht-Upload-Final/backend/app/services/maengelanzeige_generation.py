"""Erzeugt die zwei Word-Dokumente einer HPP-Mängelanzeige.

ZWEI DATEIEN, IMMER
===================
Ein Vorgang ergibt **zwei** getrennte ``.docx`` — so wie das Büro es versendet:

    1. Anschreiben   Brief mit Fristsetzung gem. § 4 Abs. 7 VOB/B
    2. Anlage        Fotodokumentation, nach Bereichen gegliedert

Nie ein zusammengeführtes Dokument. Die beiden Erzeuger sind deshalb auch zwei
Funktionen (``erzeuge_anschreiben`` / ``erzeuge_anlage``); ``erzeuge_beide``
liefert nur das Paar zurück.

WOHER DIE MASSE KOMMEN
======================
Alle Zahlen unten sind an den beiden Referenz-PDFs **ausgemessen**, nicht
geschätzt — Zeichenpositionen und Bildrahmen wurden aus den PDFs gelesen
(pdfplumber) und hier in Word-Maße übersetzt:

    Seite          A4, 595.3 x 841.9 pt
    Schrift        Arial; Fließtext 11 pt, Bildunterschrift 10 pt,
                   Datumszeile und Fußzeile 7 pt, Funktion 10 pt kursiv
    Textspalte     Brief  x 70.9 … 408.5 pt  (2.50 cm bis 14.41 cm)
                   Anlage x 70.9 … 543.5 pt  (2.50 cm bis 19.17 cm)
    Briefkopf      Grafik x 368.6 … 595.3 pt, y 0 … 594.9 pt
    Fotoraster     zwei Spalten je ~213–226 pt, Abstand ~47 pt
    Fußzeile       Grundlinie y 799.1 pt

DER BRIEFKOPF IST EINE GRAFIK
=============================
Im Original steht rechts kein Text: Logo, Städteliste, Firmenblock, Partner-
und Handelsregisterangaben sind **ein einziges Bild** (629 x 1652 px), das über
dem Textbereich liegt. HPP hat zwei Fassungen davon, gleiche Leinwandgröße:

    hpp_briefkopf.png             alles (für das Anschreiben)
    hpp_briefkopf_folgeseite.png  nur Logo und „Architekten“ (für die Anlage)

Beide werden hier als **freigestellte** Grafik (``wp:anchor``, hinter dem Text,
Position absolut zur Seite) eingefügt — genau an der Stelle des Originals. Das
ist der Grund, warum dieses Modul eigenes XML schreibt: python-docx kann nur
Bilder im Textfluss.

Wer die Angaben im Firmenblock ändern muss (neue Partner, neue Adresse),
tauscht die PNG-Datei aus. Es steht bewusst kein Wort davon im Code — sonst
gäbe es zwei Wahrheiten, und die im Code wäre die veraltete.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt
from PIL import Image, ImageOps

from app.config import settings

# ─────────────────────────────────────────────────────────────────────────────
# Maße und Schrift
# ─────────────────────────────────────────────────────────────────────────────

SCHRIFT = "Arial"

#: Fließtext, Anrede, Betreff, Adressblock.
GR_TEXT = Pt(11)
#: Bildunterschriften in der Anlage.
GR_BILDTEXT = Pt(10)
#: Datumszeile des Briefs und beide Fußzeilen.
GR_KLEIN = Pt(7)
#: Funktionszeile unter dem Namen („-Baumanagement-“).
GR_FUNKTION = Pt(10)

#: Zeilenabstand des Fließtexts: 11 pt Arial, einfach → 12.65 pt.
#: Im Original liegen die Grundlinien 12.65 pt auseinander.
ZEILE = Pt(12.65)

# Seitenränder. Rechts ist beim Brief absichtlich sehr breit: Dort liegt die
# Briefkopfgrafik, und der Text darf nicht darunter laufen.
BRIEF_RAND_LINKS = Cm(2.5)
#: Rechter Rand des Briefs.
#:
#: Am Original gemessen sind es 6.59 cm (Textspalte 337.6 pt). Hier stehen
#: 6.52 cm — 2 pt mehr Spalte. Der Grund ist nachgerechnet: Der erste Satz
#: („bei der Begehung … festgestellt.“) ist in Arial 11 pt 338.9 pt breit und
#: passt damit rechnerisch NICHT in 337.6 pt. Im Original steht er trotzdem
#: auf einer Zeile — Word hat dort die Wortabstände gestaucht. Dieses Stauchen
#: lässt sich nicht ansteuern; mit zwei Punkten mehr Spalte bricht die Zeile
#: an derselben Stelle wie im Original, und alle Blockpositionen des Briefs
#: stimmen. Der Preis: Der rechte Textrand liegt 2 pt weiter außen.
BRIEF_RAND_RECHTS = Cm(6.52)
BRIEF_RAND_OBEN = Cm(4.606)
BRIEF_RAND_UNTEN = Cm(1.9)
BRIEF_FUSS_ABSTAND = Cm(1.24)
#: Rechter Tabstopp der Brieffußzeile (Seitenzahl) — 11.91 cm ab Textbeginn.
BRIEF_FUSS_TAB = Cm(11.91)

ANLAGE_RAND_LINKS = Cm(2.5)
ANLAGE_RAND_RECHTS = Cm(1.78)
ANLAGE_RAND_OBEN = Cm(3.51)
ANLAGE_RAND_UNTEN = Cm(1.9)
ANLAGE_KOPF_ABSTAND = Cm(1.28)
ANLAGE_FUSS_ABSTAND = Cm(1.24)
#: Rechter Tabstopp der Anlagenfußzeile — 15.58 cm ab Textbeginn.
ANLAGE_FUSS_TAB = Cm(15.58)

# Briefkopfgrafik: gemessen x 368.6 pt vom linken Seitenrand, y 0, 226.7 x 594.9 pt.
KOPF_LINKS = Emu(int(368.6 * 12700))
KOPF_OBEN = Emu(0)
KOPF_BREITE = Emu(int(226.7 * 12700))
KOPF_HOEHE = Emu(int(594.9 * 12700))

# Unterschrift (optional): gemessen x 61.6 pt, y 728.8 pt, 150.4 x 34.2 pt.
UNTERSCHRIFT_LINKS = Emu(int(61.6 * 12700))
UNTERSCHRIFT_OBEN = Emu(int(728.8 * 12700))
UNTERSCHRIFT_BREITE = Emu(int(150.4 * 12700))

# Fotoraster der Anlage: zwei Spalten mit Abstand, zusammen die Textbreite.
FOTO_SPALTE = Cm(7.5)
FOTO_ABSTAND = Cm(1.72)
#: Lange Kante der eingebetteten Fotos in Pixeln.
FOTO_MAX_KANTE = 1600
#: JPEG-Qualität der eingebetteten Fotos.
FOTO_QUALITAET = 82
#: Höhe, ab der ein Foto begrenzt wird.
#:
#: Im Original sind die Fotos 282–301 pt hoch (≈ 10 cm) und es stehen zwei
#: Bereichsblöcke auf einer Seite. Beides geht rechnerisch nicht zusammen: Das
#: Original passt nur, weil dort eine Bildunterschrift von Hand *neben* das
#: Foto gerückt wurde. Hier hat der Rhythmus „zwei Blöcke pro Seite“ Vorrang,
#: das Foto ist dafür 6 % kleiner als im Original.
FOTO_MAX_HOEHE = Cm(9.4)

# Tabstopps der Datumszeile, ausgemessen: Datum 71 pt, „Ze: sb“ 155 pt,
# Auftragsnummer 233 pt, E-Mail 311 pt (jeweils ab linkem Seitenrand).
DATUM_TABS = [Cm(2.96), Cm(5.72), Cm(8.48)]

# ── Senkrechtes Raster des Briefs ────────────────────────────────────────────
#
# Der Brief steht im Original auf einem Raster von 12.65 pt (eine Zeile). Die
# Abstände sind hier als *Absatzabstand* hinterlegt und nicht als leere
# Absätze: Leerzeilen zählen sich beim nächsten Umbau falsch, und Word
# behandelt sie am Seitenende anders als echten Abstand. Die Werte sind die am
# Original gemessenen Sprünge zwischen den Blöcken.
#
#   Adresse   131.9   (oberer Rand)
#   Versand   207.9   = Adresse + 4 Zeilen + 25.4
#   Datum     296.1   = Versand + 2 Zeilen + 62.9
#   Betreff   328.6   = Datum + 1 Zeile + 19.85
#   Anrede    411.7   = Betreff + 3 Zeilen + 45.15
#   Absätze   je       + 12.65 (eine Leerzeile)
#
# Wichtig: Word rechnet den Absatzabstand ab dem ENDE des letzten Absatzes.
# Die Werte sind deshalb der Abstand zwischen zwei Blöcken, nicht der Sprung
# zwischen ihren ersten Zeilen.
ABSTAND_VERSAND = Pt(25.4)
ABSTAND_DATUM = Pt(59.4)          # 62.9 minus 3.5: die 7-pt-Zeile sitzt tiefer
ABSTAND_BETREFF = Pt(23.35)       # 19.85 plus dieselben 3.5
ABSTAND_ANREDE = Pt(45.15)
ABSTAND_ABSATZ = Pt(12.65)
ABSTAND_FRIST = Pt(12.75)
ABSTAND_NACH_FRISTABSATZ = Pt(12.7)
ABSTAND_GRUSS = Pt(12.6)
ABSTAND_NAME = Pt(25.3)
#: Abstand zwischen Bereichsüberschrift und Foto in der Anlage (gemessen).
ABSTAND_FOTO = Pt(10.55)

BRIEFKOPF_DATEI = "hpp_briefkopf.png"
FOLGESEITE_DATEI = "hpp_briefkopf_folgeseite.png"

#: Falzmarke am linken Rand — im Original eine 9 pt lange Linie von 0.25 pt
#: Stärke bei x 16.9 pt, y 281.7 pt. Sie zeigt, wo der Brief für den
#: Fensterumschlag gefaltet wird; ohne sie fehlt dem Briefbogen ein Detail,
#: das jeder im Büro kennt.
FALZMARKE_DATEI = "hpp_falzmarke.png"
FALZ_LINKS = Emu(int(16.9 * 12700))
FALZ_OBEN = Emu(int(281.7 * 12700))
FALZ_BREITE = Emu(int(9.0 * 12700))

# ─────────────────────────────────────────────────────────────────────────────
# Feste Textbausteine
#
# Alles, was in jeder Mängelanzeige gleich ist, steht genau hier — einmal.
# Der Wortlaut ist aus dem Referenzbrief übernommen, Zeichen für Zeichen;
# geändert wird er nur bewusst und dann an dieser Stelle.
# ─────────────────────────────────────────────────────────────────────────────

BETREFF_DRITTE_ZEILE = "Mängelanzeige mit Fristsetzung gem. § 4 Abs. 7 VOB/B"

#: Der Zeilenumbruch nach „festgestellt.“ steht so im Original: Der Satz ist
#: 338.9 pt breit und passt damit 1.3 pt nicht in die 337.6 pt breite Spalte —
#: ohne den Umbruch trennt Word „festge-stellt“ und der Brief wird eine Zeile
#: länger. (Das kleine „deren“ am Zeilenanfang ist ebenfalls Original.)
ABSATZ_BEGEHUNG = (
    "bei der Begehung am {begehungsdatum} wurden diverse Mängel festgestellt."
    + chr(10)
    + "deren Einzelheiten Sie bitte der beigefügten Anlage entnehmen."
)

ABSATZ_AUFFORDERUNG = (
    "Gemäß § 4 Abs. 7 VOB/B fordern wir Sie hiermit auf, die festgestellten "
    "Mängel bzw. vertragswidrigen Leistungen unverzüglich zu beseitigen und "
    "durch mangelfreie und vertragsgemäße Leistungen zu ersetzen. Wir setzen "
    "Ihnen hierfür eine Frist zur ordnungsgemäßen Nachbesserung bis zum"
)

ABSATZ_FOLGEN = (
    "Sollten die Mängel bis zu diesem Datum nicht behoben sein, sehen wir uns "
    "gezwungen, die erforderlichen Arbeiten durch ein Drittunternehmen auf "
    "Ihre Kosten durchführen zu lassen. Die durch den Verzug und die "
    "mangelhafte Leistung entstehenden Mehrkosten werden Ihnen in Rechnung "
    "gestellt bzw. von Ihrer Schlussrechnung abgezogen."
)

ABSATZ_ABNAHME = (
    "Bitte beachten Sie, dass eine Abnahme der Arbeiten erst nach "
    "vollständiger und mangelfreier Leistung erfolgen kann."
)

GRUSS = "Mit freundlichen Grüßen"
GRUSS_FIRMA = "HPP Architekten"

ANLAGE_UEBERSCHRIFT = "Anlage"
ANLAGE_KOPF_DRITTE_ZEILE = "Mängelanzeige / Begehung am {begehungsdatum}"


# ─────────────────────────────────────────────────────────────────────────────
# Datenmodell
#
# Reine Datenklassen ohne Web-Abhängigkeit: Das Modul lässt sich damit auch
# aus einem Skript oder einem Test benutzen. Die HTTP-Schicht bringt ihre
# eigenen Pydantic-Modelle mit und füllt diese Klassen (siehe routers).
# ─────────────────────────────────────────────────────────────────────────────


@dataclass
class MangelFoto:
    """Ein Foto mit seiner Bildunterschrift."""

    #: Bilddaten. Bewusst Bytes und kein Pfad: Die Fotos kommen aus dem Upload
    #: oder aus der Datenbank, nicht zwingend aus dem Dateisystem.
    daten: bytes
    beschreibung: str


@dataclass
class MangelBereich:
    """Ein Gebäudebereich mit einem oder mehreren Fotos."""

    bereich: str
    eintraege: list[MangelFoto] = field(default_factory=list)


@dataclass
class Empfaenger:
    firma: str
    ansprechpartner: str = ""
    strasse_hausnummer: str = ""
    plz_ort: str = ""
    versandart: str = "per Mail"
    email: str = ""

    @property
    def anrede(self) -> str:
        """„Herrn Hey“ → „Sehr geehrter Herr Hey,“ — inklusive Fallformen.

        Das Original schreibt den Ansprechpartner im Adressfeld im Akkusativ
        („Herrn Hey“), in der Anrede aber im Nominativ („Sehr geehrter Herr
        Hey"). Ohne diese Umformung stünde im Brief „Sehr geehrter Herrn Hey“.
        """
        name = (self.ansprechpartner or "").strip()
        if not name:
            return "Sehr geehrte Damen und Herren,"
        if name.startswith("Herrn "):
            return f"Sehr geehrter Herr {name[6:].strip()},"
        if name.startswith("Herr "):
            return f"Sehr geehrter Herr {name[5:].strip()},"
        if name.startswith("Frau "):
            return f"Sehr geehrte Frau {name[5:].strip()},"
        return f"Sehr geehrte Damen und Herren,"


@dataclass
class Sachbearbeiter:
    name: str
    funktion: str = "-Baumanagement-"
    zeichen: str = ""
    auftragsnummer: str = ""
    email: str = ""
    #: Optionales Bild der Unterschrift. Ohne Angabe bleibt nur der Platz frei —
    #: eine Unterschrift setzt man nicht ungefragt unter ein Schreiben.
    unterschrift: bytes | None = None


@dataclass
class Maengelanzeige:
    """Alle Angaben eines Vorgangs — Anschreiben und Anlage zusammen."""

    projektbezeichnung: str
    vergabeeinheit: str
    begehungsdatum: date
    dokumentkuerzel: str
    empfaenger: Empfaenger
    sachbearbeiter: Sachbearbeiter
    briefdatum: date
    fristsetzungsdatum: date
    bereiche: list[MangelBereich] = field(default_factory=list)
    #: Datum der Anlage. Ohne Angabe gilt das Begehungsdatum.
    #:
    #: Im Referenzvorgang war die Begehung am 30.07., die Anlage aber vom
    #: 04.08. und der Brief vom 11.08. — die Fotodokumentation entsteht also
    #: irgendwann zwischen Begehung und Schreiben. Deshalb ein eigenes Feld
    #: statt einer Annahme.
    anlagedatum: date | None = None

    @property
    def anlage_stand(self) -> date:
        """Datum, das in Dateiname und Fußzeile der Anlage steht."""
        return self.anlagedatum or self.begehungsdatum


class MaengelanzeigeFehler(ValueError):
    """Eingaben unbrauchbar — die Meldung ist für die Oberfläche gedacht."""


# ─────────────────────────────────────────────────────────────────────────────
# Prüfung
# ─────────────────────────────────────────────────────────────────────────────


def pruefe(daten: Maengelanzeige, *, mit_bereichen: bool = True) -> None:
    """Prüft die Pflichtangaben und wirft mit klarem Text.

    Lieber hier abbrechen als ein Dokument mit „None“ im Betreff erzeugen: Das
    Schreiben geht an eine Firma und setzt eine Rechtsfrist.
    """
    fehlt: list[str] = []
    if not daten.projektbezeichnung.strip():
        fehlt.append("Projektbezeichnung")
    if not daten.vergabeeinheit.strip():
        fehlt.append("Vergabeeinheit")
    if not daten.dokumentkuerzel.strip():
        fehlt.append("Dokumentkürzel (steht in der Fußzeile und im Dateinamen)")
    if not daten.empfaenger.firma.strip():
        fehlt.append("Empfänger: Firma")
    if not daten.sachbearbeiter.name.strip():
        fehlt.append("Sachbearbeiter: Name")
    if daten.begehungsdatum is None:
        fehlt.append("Begehungsdatum")
    if daten.briefdatum is None:
        fehlt.append("Briefdatum")
    if daten.fristsetzungsdatum is None:
        fehlt.append("Fristsetzungsdatum")

    if fehlt:
        raise MaengelanzeigeFehler(
            "Für die Mängelanzeige fehlen: " + ", ".join(fehlt) + "."
        )

    if daten.fristsetzungsdatum < daten.briefdatum:
        raise MaengelanzeigeFehler(
            f"Die Frist ({daten.fristsetzungsdatum:%d.%m.%Y}) liegt vor dem "
            f"Briefdatum ({daten.briefdatum:%d.%m.%Y}). Eine Frist in der "
            f"Vergangenheit ist nicht durchsetzbar."
        )
    if daten.begehungsdatum > daten.briefdatum:
        raise MaengelanzeigeFehler(
            f"Das Begehungsdatum ({daten.begehungsdatum:%d.%m.%Y}) liegt nach "
            f"dem Briefdatum ({daten.briefdatum:%d.%m.%Y})."
        )

    if not mit_bereichen:
        return

    if not daten.bereiche:
        raise MaengelanzeigeFehler(
            "Die Anlage braucht mindestens einen Bereich mit einem Foto."
        )
    for nummer, bereich in enumerate(daten.bereiche, 1):
        if not bereich.bereich.strip():
            raise MaengelanzeigeFehler(
                f"Bereich {nummer} hat keine Überschrift (z. B. „Ostfassade“)."
            )
        if not bereich.eintraege:
            raise MaengelanzeigeFehler(
                f"Der Bereich „{bereich.bereich}“ enthält kein Foto. Ohne Foto "
                f"belegt die Anlage nichts — Bereich löschen oder Foto ergänzen."
            )
        for lauf, eintrag in enumerate(bereich.eintraege, 1):
            if not eintrag.daten:
                raise MaengelanzeigeFehler(
                    f"Bereich „{bereich.bereich}“, Foto {lauf}: Bilddaten fehlen."
                )
            if not eintrag.beschreibung.strip():
                raise MaengelanzeigeFehler(
                    f"Bereich „{bereich.bereich}“, Foto {lauf}: Es fehlt die "
                    f"Bildunterschrift (was ist zu sehen bzw. zu tun?)."
                )


# ─────────────────────────────────────────────────────────────────────────────
# Dateinamen
# ─────────────────────────────────────────────────────────────────────────────


def _jjmmtt(tag: date) -> str:
    return tag.strftime("%y%m%d")


def dateiname_anschreiben(daten: Maengelanzeige) -> str:
    """``{JJMMTT}_{dokumentkuerzel}.docx`` mit dem Briefdatum."""
    return f"{_jjmmtt(daten.briefdatum)}_{daten.dokumentkuerzel}.docx"


def dateiname_anlage(daten: Maengelanzeige) -> str:
    """``{JJMMTT}_Anlage_{dokumentkuerzel}.docx`` mit dem Begehungsdatum.

    Bewusst nicht das Briefdatum: Die Anlage ist ein eigenes Dokument mit
    eigenem Stand (im Original 250804 für die Anlage, 250811 für den Brief).
    Ohne ausdrückliches ``anlagedatum`` gilt der Tag der Begehung.
    """
    return f"{_jjmmtt(daten.anlage_stand)}_Anlage_{daten.dokumentkuerzel}.docx"


# ─────────────────────────────────────────────────────────────────────────────
# Werkzeuge: XML, das python-docx nicht kann
# ─────────────────────────────────────────────────────────────────────────────


def _marke(name: str) -> Path:
    """Pfad einer Briefkopfgrafik im Vorlagenordner."""
    pfad = settings.template_dir / "marke" / name
    if not pfad.is_file():
        raise MaengelanzeigeFehler(
            f"Die Briefkopfgrafik {pfad} fehlt. Ohne sie entsteht kein "
            f"HPP-Briefbogen."
        )
    return pfad


def _schrift_setzen(dokument) -> None:
    """Arial als Standardschrift des Dokuments (auch für Osteuropa/Symbole)."""
    stil = dokument.styles["Normal"]
    stil.font.name = SCHRIFT
    stil.font.size = GR_TEXT
    rpr = stil.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for schluessel in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(schluessel), SCHRIFT)

    # Sprache Deutsch. Word trennt Silben nach der Sprache des Textes; ohne
    # diese Angabe gilt Englisch, und dann entstehen weder „Nachbes-serung“
    # noch die Zeilenumbrüche des Originals.
    sprache = rpr.find(qn("w:lang"))
    if sprache is None:
        sprache = OxmlElement("w:lang")
        rpr.append(sprache)
    for schluessel in ("w:val", "w:eastAsia", "w:bidi"):
        sprache.set(qn(schluessel), "de-DE")

    # Absätze im Original haben keinen Abstand und einfachen Zeilenabstand.
    stil.paragraph_format.space_before = Pt(0)
    stil.paragraph_format.space_after = Pt(0)


def _silbentrennung(dokument) -> None:
    """Automatische Silbentrennung einschalten.

    Im Referenzbrief ist sie aktiv — „Nachbes-serung“, „Drittunter-nehmen“,
    „abgezo-gen“. Ohne sie bricht der Blocksatz an anderen Stellen um, und der
    Brief sieht bei jedem Absatz anders aus als das Original.
    """
    einstellungen = dokument.settings.element
    if einstellungen.find(qn("w:autoHyphenation")) is not None:
        return

    element = OxmlElement("w:autoHyphenation")
    element.set(qn("w:val"), "true")

    # Die Reihenfolge in settings.xml ist im Schema festgelegt. Hinten
    # angehaengt liest Word das Element nicht — der Brief bekommt dann eine
    # Zeile mehr je Absatz, und das ganze Raster verschiebt sich. Richtig
    # sitzt es direkt hinter w:defaultTabStop.
    vorgaenger = einstellungen.find(qn("w:defaultTabStop"))
    if vorgaenger is not None:
        vorgaenger.addnext(element)
    else:
        einstellungen.insert(0, element)


def _absatz(behaelter, text: str = "", *, fett: bool = False,
            kursiv: bool = False, groesse=GR_TEXT,
            ausrichtung=None, abstand_vor=None, abstand_nach=None,
            zusammenhalten: bool = False):
    """Ein Absatz mit genau einem Textlauf — der Normalfall dieses Moduls."""
    absatz = behaelter.add_paragraph()
    if ausrichtung is not None:
        absatz.alignment = ausrichtung
    form = absatz.paragraph_format
    form.space_before = abstand_vor if abstand_vor is not None else Pt(0)
    form.space_after = abstand_nach if abstand_nach is not None else Pt(0)
    form.line_spacing = ZEILE
    if zusammenhalten:
        form.keep_with_next = True
    if text:
        lauf = absatz.add_run(text)
        lauf.font.name = SCHRIFT
        lauf.font.size = groesse
        lauf.bold = fett
        lauf.italic = kursiv
    return absatz


def _leerzeilen(behaelter, anzahl: int, groesse=GR_TEXT) -> None:
    """Leerzeilen wie im Original — der Brief ist auf Zeilenraster gesetzt."""
    for _ in range(anzahl):
        _absatz(behaelter, "", groesse=groesse)


def _feld(absatz, anweisung: str, vorschau: str) -> None:
    """Word-Feld einfügen (``PAGE``, ``NUMPAGES``).

    Als echtes Feld und nicht als Text: Nur dann zählt Word die Seiten selbst,
    und „3/5“ stimmt auch, wenn später ein Bereich dazukommt.

    Gebaut aus den fünf üblichen Läufen (begin, instrText, separate, Vorschau,
    end) und in dieser Reihenfolge angehängt. Die kurze Schreibweise
    ``w:fldSimple`` hatte sich beim Einfügen zwischen die anderen Läufe
    verschoben — dann stand im Brief „12“ statt „1/2“.
    """
    def lauf(inhalt):
        element = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), SCHRIFT)
        rfonts.set(qn("w:hAnsi"), SCHRIFT)
        rpr.append(rfonts)
        groesse = OxmlElement("w:sz")
        groesse.set(qn("w:val"), str(int(GR_KLEIN.pt * 2)))
        rpr.append(groesse)
        element.append(rpr)
        element.append(inhalt)
        absatz._p.append(element)

    for art in ("begin",):
        zeichen = OxmlElement("w:fldChar")
        zeichen.set(qn("w:fldCharType"), art)
        lauf(zeichen)

    anweisungstext = OxmlElement("w:instrText")
    anweisungstext.set(qn("xml:space"), "preserve")
    anweisungstext.text = anweisung
    lauf(anweisungstext)

    zeichen = OxmlElement("w:fldChar")
    zeichen.set(qn("w:fldCharType"), "separate")
    lauf(zeichen)

    text = OxmlElement("w:t")
    text.text = vorschau
    lauf(text)

    zeichen = OxmlElement("w:fldChar")
    zeichen.set(qn("w:fldCharType"), "end")
    lauf(zeichen)


def _bild_freistellen(absatz, bilddaten, *, links: Emu, oben: Emu,
                      breite: Emu, hoehe: Emu | None = None,
                      name: str = "Briefkopf") -> None:
    """Fügt ein Bild an einer festen Seitenposition ein, hinter dem Text.

    python-docx kann nur Bilder im Textfluss. Deshalb wird zuerst ein solches
    eingefügt und sein ``wp:inline`` anschließend zu einem ``wp:anchor``
    umgeschrieben — dieselbe Grafik, aber mit absoluter Position zur Seite.
    So sitzt der Briefkopf auf den Punkt dort, wo er im Original sitzt, und
    der Text läuft darunter hindurch, statt verdrängt zu werden.
    """
    lauf = absatz.add_run()
    if hoehe is not None:
        lauf.add_picture(bilddaten, width=breite, height=hoehe)
    else:
        lauf.add_picture(bilddaten, width=breite)

    inline = lauf._r.find(qn("w:drawing"))[0]
    anchor = OxmlElement("wp:anchor")
    for schluessel, wert in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("simplePos", "0"), ("relativeHeight", "251658240"),
        ("behindDoc", "1"), ("locked", "0"), ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anchor.set(schluessel, wert)

    einfach = OxmlElement("wp:simplePos")
    einfach.set("x", "0")
    einfach.set("y", "0")
    anchor.append(einfach)

    for richtung, bezug, versatz in (
        ("wp:positionH", "page", int(links)),
        ("wp:positionV", "page", int(oben)),
    ):
        block = OxmlElement(richtung)
        block.set("relativeFrom", bezug)
        wert = OxmlElement("wp:posOffset")
        wert.text = str(versatz)
        block.append(wert)
        anchor.append(block)

    # Ausdehnung, Effektrand und Umlauf aus dem Inline-Element übernehmen.
    for tag in ("wp:extent", "wp:effectExtent"):
        vorhanden = inline.find(qn(tag))
        if vorhanden is not None:
            anchor.append(vorhanden)
    anchor.append(OxmlElement("wp:wrapNone"))
    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        vorhanden = inline.find(qn(tag))
        if vorhanden is not None:
            anchor.append(vorhanden)

    beschreibung = anchor.find(qn("wp:docPr"))
    if beschreibung is not None:
        beschreibung.set("name", name)

    inline.getparent().replace(inline, anchor)


def _rahmen_aus(tabelle) -> None:
    """Alle Rahmen einer Tabelle entfernen — sie ist nur Raster, nicht Gitter."""
    eigenschaften = tabelle._tbl.tblPr
    rahmen = OxmlElement("w:tblBorders")
    for kante in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{kante}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        rahmen.append(element)
    eigenschaften.append(rahmen)


def _zellenrand_null(tabelle) -> None:
    """Innenabstand der Zellen entfernen.

    Word gibt jeder Zelle 0.19 cm links und rechts. Damit begänne das erste
    Foto 5 pt hinter dem Seitenrand — im Original sitzt es darauf.
    """
    eigenschaften = tabelle._tbl.tblPr
    raender = OxmlElement("w:tblCellMar")
    for kante in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{kante}")
        element.set(qn("w:w"), "0")
        element.set(qn("w:type"), "dxa")
        raender.append(element)
    eigenschaften.append(raender)


def _zeile_nicht_trennen(zeile) -> None:
    """Verhindert, dass Word eine Tabellenzeile über den Seitenumbruch teilt."""
    eigenschaften = zeile._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    eigenschaften.append(element)


def _foto_aufbereiten(daten: bytes) -> tuple[bytes, Emu, Emu]:
    """Foto drehen, verkleinern und vermessen.

    DREI DINGE, DIE HIER PASSIEREN MÜSSEN
    =====================================
    1. **Drehen.** Handyfotos liegen als Querformat in der Datei und tragen die
       Drehung nur im EXIF-Vermerk. Word beachtet ihn, python-docx nicht —
       ohne ``exif_transpose`` läge jedes Hochformatfoto in der Anlage auf der
       Seite. (Genau diese Drehung steckt auch im Original: Die eingebetteten
       Bilder sind 788 x 591 px, gezeigt werden sie 213 x 284 pt.)
    2. **Verkleinern.** Ein 12-Megapixel-Foto in einer Word-Datei bringt
       nichts außer Dateigröße; bei 13 Fotos wären das 50 MB, die per Mail
       nicht durchgehen. 1600 px an der langen Kante sind bei 7.8 cm
       Druckbreite noch über 400 dpi.
    3. **Vermessen.** Seitenverhältnis bleibt; die Höhenbegrenzung sorgt
       dafür, dass zwei Bereichsblöcke auf eine Seite passen — die Proportion
       des Originals.
    """
    try:
        with Image.open(io.BytesIO(daten)) as bild:
            gedreht = ImageOps.exif_transpose(bild)
            if gedreht.mode not in ("RGB", "L"):
                gedreht = gedreht.convert("RGB")
            gedreht.thumbnail((FOTO_MAX_KANTE, FOTO_MAX_KANTE), Image.LANCZOS)
            breite_px, hoehe_px = gedreht.size
            puffer = io.BytesIO()
            gedreht.save(puffer, format="JPEG", quality=FOTO_QUALITAET,
                         optimize=True)
    except MaengelanzeigeFehler:
        raise
    except Exception as fehler:                       # noqa: BLE001
        raise MaengelanzeigeFehler(
            f"Ein Foto lässt sich nicht lesen ({fehler}). Zulässig sind "
            f"gängige Bildformate wie JPEG, PNG oder HEIC."
        ) from fehler

    if breite_px <= 0 or hoehe_px <= 0:
        raise MaengelanzeigeFehler("Ein Foto hat keine lesbaren Bildmaße.")

    breite = int(FOTO_SPALTE)
    hoehe = int(breite * hoehe_px / breite_px)
    if hoehe > int(FOTO_MAX_HOEHE):
        hoehe = int(FOTO_MAX_HOEHE)
        breite = int(hoehe * breite_px / hoehe_px)
    return puffer.getvalue(), Emu(breite), Emu(hoehe)


# ─────────────────────────────────────────────────────────────────────────────
# Dokument 1 — Anschreiben
# ─────────────────────────────────────────────────────────────────────────────


def erzeuge_anschreiben(daten: Maengelanzeige) -> bytes:
    """Der Brief mit Fristsetzung — eine Seite, Layout wie im Referenz-PDF."""
    pruefe(daten, mit_bereichen=False)

    dokument = Document()
    _schrift_setzen(dokument)
    _silbentrennung(dokument)

    abschnitt = dokument.sections[0]
    abschnitt.page_width = Cm(21)
    abschnitt.page_height = Cm(29.7)
    abschnitt.left_margin = BRIEF_RAND_LINKS
    abschnitt.right_margin = BRIEF_RAND_RECHTS
    abschnitt.top_margin = BRIEF_RAND_OBEN
    abschnitt.bottom_margin = BRIEF_RAND_UNTEN
    abschnitt.footer_distance = BRIEF_FUSS_ABSTAND

    # ── Briefkopfgrafik in die Kopfzeile, freigestellt ──
    kopf = abschnitt.header
    kopf.is_linked_to_previous = False
    kopfabsatz = kopf.paragraphs[0]
    kopfabsatz.paragraph_format.space_after = Pt(0)
    _bild_freistellen(
        kopfabsatz, str(_marke(BRIEFKOPF_DATEI)),
        links=KOPF_LINKS, oben=KOPF_OBEN,
        breite=KOPF_BREITE, hoehe=KOPF_HOEHE,
        name="HPP Briefkopf",
    )
    _bild_freistellen(
        kopfabsatz, str(_marke(FALZMARKE_DATEI)),
        links=FALZ_LINKS, oben=FALZ_OBEN,
        breite=FALZ_BREITE, name="Falzmarke",
    )

    # ── Adressblock ──
    for zeile in (
        daten.empfaenger.firma,
        daten.empfaenger.ansprechpartner,
        daten.empfaenger.strasse_hausnummer,
        daten.empfaenger.plz_ort,
    ):
        if zeile.strip():
            _absatz(dokument, zeile)

    # Abstand zum Versandhinweis: im Original drei Leerzeilen (y 169.9 → 207.9).
    _leerzeilen(dokument, 2)
    if daten.empfaenger.versandart.strip():
        _absatz(dokument, daten.empfaenger.versandart)
    if daten.empfaenger.email.strip():
        _absatz(dokument, daten.empfaenger.email)

    # ── Datumszeile: Datum, Zeichen, Auftragsnummer, E-Mail auf Tabstopps ──
    datumszeile = _absatz(dokument, "", groesse=GR_KLEIN,
                          abstand_vor=ABSTAND_DATUM)
    for stelle in DATUM_TABS:
        datumszeile.paragraph_format.tab_stops.add_tab_stop(
            stelle, WD_TAB_ALIGNMENT.LEFT
        )
    felder = [
        f"{daten.briefdatum:%d.%m.%Y}",
        daten.sachbearbeiter.zeichen,
        daten.sachbearbeiter.auftragsnummer,
        daten.sachbearbeiter.email,
    ]
    lauf = datumszeile.add_run("\t".join(felder))
    lauf.font.name = SCHRIFT
    lauf.font.size = GR_KLEIN

    # ── Betreff, drei Zeilen fett ──
    for lauf_nummer, zeile in enumerate((daten.projektbezeichnung,
                                         daten.vergabeeinheit,
                                         BETREFF_DRITTE_ZEILE)):
        _absatz(dokument, zeile, fett=True,
                abstand_vor=ABSTAND_BETREFF if lauf_nummer == 0 else None)

    # ── Anrede und Fließtext ──
    _absatz(dokument, daten.empfaenger.anrede, abstand_vor=ABSTAND_ANREDE)

    _absatz(
        dokument,
        ABSATZ_BEGEHUNG.format(begehungsdatum=f"{daten.begehungsdatum:%d.%m.%Y}"),
        ausrichtung=WD_ALIGN_PARAGRAPH.JUSTIFY,
        abstand_vor=ABSTAND_ABSATZ,
    )

    _absatz(dokument, ABSATZ_AUFFORDERUNG,
            ausrichtung=WD_ALIGN_PARAGRAPH.JUSTIFY,
            abstand_vor=ABSTAND_ABSATZ)

    # Fristdatum: eigene Zeile, fett, in der Textspalte zentriert.
    _absatz(dokument, f"{daten.fristsetzungsdatum:%d.%m.%Y}.", fett=True,
            ausrichtung=WD_ALIGN_PARAGRAPH.CENTER,
            abstand_vor=ABSTAND_FRIST)

    _absatz(dokument, ABSATZ_FOLGEN, ausrichtung=WD_ALIGN_PARAGRAPH.JUSTIFY,
            abstand_vor=ABSTAND_ABSATZ)

    _absatz(dokument, ABSATZ_ABNAHME, ausrichtung=WD_ALIGN_PARAGRAPH.JUSTIFY,
            abstand_vor=ABSTAND_NACH_FRISTABSATZ)

    # ── Gruß und Unterschrift ──
    _absatz(dokument, GRUSS, abstand_vor=ABSTAND_GRUSS)
    _absatz(dokument, GRUSS_FIRMA)

    # Platz für die Unterschrift — im Original zwei Zeilen (715 → 753 pt).
    namenszeile = _absatz(dokument, daten.sachbearbeiter.name,
                          abstand_vor=ABSTAND_NAME)
    _absatz(dokument, daten.sachbearbeiter.funktion, kursiv=True,
            groesse=GR_FUNKTION)

    if daten.sachbearbeiter.unterschrift:
        _bild_freistellen(
            namenszeile, io.BytesIO(daten.sachbearbeiter.unterschrift),
            links=UNTERSCHRIFT_LINKS, oben=UNTERSCHRIFT_OBEN,
            breite=UNTERSCHRIFT_BREITE, name="Unterschrift",
        )

    _fusszeile(abschnitt, daten, tab=BRIEF_FUSS_TAB, datum=daten.briefdatum)

    puffer = io.BytesIO()
    dokument.save(puffer)
    return puffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Dokument 2 — Anlage
# ─────────────────────────────────────────────────────────────────────────────


def erzeuge_anlage(daten: Maengelanzeige) -> bytes:
    """Die Fotodokumentation — Kopf- und Fußzeile auf jeder Seite."""
    pruefe(daten)

    dokument = Document()
    _schrift_setzen(dokument)

    abschnitt = dokument.sections[0]
    abschnitt.page_width = Cm(21)
    abschnitt.page_height = Cm(29.7)
    abschnitt.left_margin = ANLAGE_RAND_LINKS
    abschnitt.right_margin = ANLAGE_RAND_RECHTS
    abschnitt.top_margin = ANLAGE_RAND_OBEN
    abschnitt.bottom_margin = ANLAGE_RAND_UNTEN
    abschnitt.header_distance = ANLAGE_KOPF_ABSTAND
    abschnitt.footer_distance = ANLAGE_FUSS_ABSTAND

    # ── Kopfzeile: links drei Zeilen, rechts das Logo (freigestellt) ──
    kopf = abschnitt.header
    kopf.is_linked_to_previous = False
    erste = kopf.paragraphs[0]
    erste.paragraph_format.space_after = Pt(0)
    erste.paragraph_format.line_spacing = ZEILE
    lauf = erste.add_run(daten.projektbezeichnung)
    lauf.font.name = SCHRIFT
    lauf.font.size = GR_TEXT
    _absatz(kopf, daten.vergabeeinheit)
    _absatz(
        kopf,
        ANLAGE_KOPF_DRITTE_ZEILE.format(
            begehungsdatum=f"{daten.begehungsdatum:%d.%m.%Y}"
        ),
    )
    _bild_freistellen(
        erste, str(_marke(FOLGESEITE_DATEI)),
        links=KOPF_LINKS, oben=KOPF_OBEN,
        breite=KOPF_BREITE, hoehe=KOPF_HOEHE,
        name="HPP Logo",
    )

    _fusszeile(abschnitt, daten, tab=ANLAGE_FUSS_TAB,
               datum=daten.anlage_stand)

    # ── Überschrift „Anlage“, unterstrichen ──
    ueberschrift = _absatz(dokument, "", zusammenhalten=True)
    lauf = ueberschrift.add_run(ANLAGE_UEBERSCHRIFT)
    lauf.font.name = SCHRIFT
    lauf.font.size = GR_TEXT
    lauf.underline = True

    # ── Bereichsblöcke ──
    for nummer, bereich in enumerate(daten.bereiche):
        # Vor jedem Block eine Leerzeile; vor dem ersten steht schon die
        # Überschrift „Anlage“.
        _leerzeilen(dokument, 1)
        _absatz(dokument, bereich.bereich, fett=True, zusammenhalten=True)
        _bereichsblock(dokument, bereich)

    puffer = io.BytesIO()
    dokument.save(puffer)
    return puffer.getvalue()


def _bereichsblock(dokument, bereich: MangelBereich) -> None:
    """Fotos eines Bereichs im zweispaltigen Raster, Unterschrift je Foto.

    Umgesetzt als randlose Tabelle statt mit freigestellten Bildern: Word
    bricht die Tabelle selbst um, wenn der nächste Block nicht mehr auf die
    Seite passt. Im Original wurde jedes Bild von Hand gesetzt — dabei
    verrutscht auf Dauer alles, und die Bildunterschriften stehen mal unter,
    mal neben dem Bild.
    """
    eintraege = bereich.eintraege
    # Je Zeile höchstens zwei Fotos — so ist das Raster im Original.
    for start in range(0, len(eintraege), 2):
        paar = eintraege[start:start + 2]

        tabelle = dokument.add_table(rows=2, cols=2)
        tabelle.alignment = WD_TABLE_ALIGNMENT.LEFT
        tabelle.autofit = False
        _rahmen_aus(tabelle)
        _zellenrand_null(tabelle)
        _zeile_nicht_trennen(tabelle.rows[0])

        for spalte, breite in enumerate((FOTO_SPALTE + FOTO_ABSTAND, FOTO_SPALTE)):
            for zelle in tabelle.columns[spalte].cells:
                zelle.width = breite

        for spalte in range(2):
            bildzelle = tabelle.cell(0, spalte)
            textzelle = tabelle.cell(1, spalte)
            bildzelle.paragraphs[0].paragraph_format.space_after = Pt(0)
            bildzelle.paragraphs[0].paragraph_format.space_before = ABSTAND_FOTO

            if spalte >= len(paar):
                continue

            eintrag = paar[spalte]
            bild, breite, hoehe = _foto_aufbereiten(eintrag.daten)
            lauf = bildzelle.paragraphs[0].add_run()
            lauf.add_picture(io.BytesIO(bild), width=breite, height=hoehe)

            absatz = textzelle.paragraphs[0]
            absatz.paragraph_format.space_before = Pt(2)
            absatz.paragraph_format.space_after = Pt(0)
            absatz.paragraph_format.line_spacing = Pt(11.5)
            text = absatz.add_run(eintrag.beschreibung)
            text.font.name = SCHRIFT
            text.font.size = GR_BILDTEXT


# ─────────────────────────────────────────────────────────────────────────────
# Gemeinsames
# ─────────────────────────────────────────────────────────────────────────────


def _fusszeile(abschnitt, daten: Maengelanzeige, *, tab, datum: date) -> None:
    """Links „{JJMMTT} {dokumentkuerzel}“, rechts „Seite/Seiten“ als Feld."""
    fuss = abschnitt.footer
    fuss.is_linked_to_previous = False
    absatz = fuss.paragraphs[0]
    # Der mitgelieferte Fußzeilenstil hat Tabstopps in der Mitte und rechts;
    # der Tabulator vor der Seitenzahl würde dort landen statt am gemessenen
    # Platz. Der Absatz bekommt deshalb den Normal-Stil — dann gilt nur der
    # hier gesetzte Stopp. (Der Fußzeilenteil kennt kein ``document``, deshalb
    # wird der Stilname direkt in die Absatzeigenschaften geschrieben.)
    eigenschaften = absatz._p.get_or_add_pPr()
    for vorhandener in eigenschaften.findall(qn("w:pStyle")):
        eigenschaften.remove(vorhandener)
    stil = OxmlElement("w:pStyle")
    stil.set(qn("w:val"), "Normal")
    eigenschaften.insert(0, stil)
    absatz.paragraph_format.space_before = Pt(0)
    absatz.paragraph_format.space_after = Pt(0)
    absatz.paragraph_format.tab_stops.add_tab_stop(tab, WD_TAB_ALIGNMENT.RIGHT)

    lauf = absatz.add_run(f"{_jjmmtt(datum)} {daten.dokumentkuerzel}\t")
    lauf.font.name = SCHRIFT
    lauf.font.size = GR_KLEIN

    _feld(absatz, " PAGE ", "1")
    trenner = absatz.add_run("/")
    trenner.font.name = SCHRIFT
    trenner.font.size = GR_KLEIN
    _feld(absatz, " NUMPAGES ", "1")


def erzeuge_beide(daten: Maengelanzeige) -> dict[str, bytes]:
    """Beide Dokumente auf einmal — Dateiname → Inhalt.

    Der Weg, den die Oberfläche nimmt: Ein Vorgang, zwei Dateien, nie eine
    zusammengeführte. Geprüft wird vorher einmal für beide.
    """
    pruefe(daten)
    return {
        dateiname_anschreiben(daten): erzeuge_anschreiben(daten),
        dateiname_anlage(daten): erzeuge_anlage(daten),
    }
